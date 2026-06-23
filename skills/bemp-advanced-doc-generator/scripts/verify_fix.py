"""验证修复效果：检查'组件内部的模块列表及说明'和'模块2设计说明'"""
from docx import Document

v_new_path = r"d:\code\QJ\BEMP5.0DEV\.trae\skills\bemp-advanced-doc-generator\output\机构管理和管理员管理功能优化-详细设计文档-20260617.docx"
doc = Document(v_new_path)

print("=" * 60)
print("修复验证报告")
print("=" * 60)

# === 验证问题1：组件内部的模块列表及说明 ===
print("\n【问题1验证】'组件内部的模块列表及说明'章节内容")
found_comp = False
for i, para in enumerate(doc.paragraphs):
    text = para.text.strip()
    if '组件' in text and ('列表' in text or '模块列表' in text or '汇总' in text):
        found_comp = True
        print(f"  找到标题: {text}")
        # 打印后续内容
        for j in range(i+1, min(i+15, len(doc.paragraphs))):
            next_text = doc.paragraphs[j].text.strip()
            if next_text:
                print(f"    [{j}] {next_text[:120]}")
            if doc.paragraphs[j].style.name.startswith('Heading'):
                break
        break

if not found_comp:
    # 检查表格
    for table in doc.tables:
        first_row = [cell.text.strip() for cell in table.rows[0].cells]
        if '模块名称' in str(first_row) or '序号' in str(first_row):
            print(f"  找到模块列表表格（{len(table.rows)}行）:")
            for row in table.rows[:8]:
                cells = [cell.text.strip()[:40] for cell in row.cells]
                print(f"    {' | '.join(cells)}")
            break

# === 验证问题2：模块2设计说明下的表格是否有独立子章节 ===
print("\n【问题2验证】'模块2设计说明'下的表格是否有独立子章节标题")
found_module2 = False
for i, para in enumerate(doc.paragraphs):
    if para.style.name == 'Heading 1' and '模块2' in para.text:
        found_module2 = True
        print(f"  找到模块2设计说明")
        # 打印后续H2标题和表格
        h2_count = 0
        table_count = 0
        for j in range(i+1, min(i+50, len(doc.paragraphs))):
            next_para = doc.paragraphs[j]
            if next_para.style.name == 'Heading 1':
                break
            if next_para.style.name == 'Heading 2':
                h2_count += 1
                print(f"    H2[{h2_count}]: {next_para.text.strip()[:80]}")
            if next_para.style.name.startswith('Heading'):
                continue
            text = next_para.text.strip()
            if text and len(text) > 10:
                print(f"      段落: {text[:80]}")
        break

if not found_module2:
    print("  未找到模块2设计说明标题")

# === 检查所有表格的上下文 ===
print("\n【表格上下文检查】")
for i, table in enumerate(doc.tables):
    first_row = [cell.text.strip()[:30] for cell in table.rows[0].cells]
    # 找到表格前最近的标题
    preceding_heading = "无"
    for j in range(len(doc.paragraphs)-1, -1, -1):
        if doc.paragraphs[j].style.name.startswith('Heading'):
            preceding_heading = f"{doc.paragraphs[j].style.name}: {doc.paragraphs[j].text.strip()[:50]}"
            break
    print(f"  表格{i}: {len(table.rows)}行 | 首行: {first_row[:3]} | 前导标题: {preceding_heading}")

print("\n" + "=" * 60)
print("验证完成")
