"""
2026-07-02 新增：模板完整性检查脚本

扫描 assets/ 目录所有 .docx/.json/.xlsx 文件，校验可读性，输出模板清单。
- .docx: 使用 python-docx 打开测试
- .json: 校验 JSON 有效性
- .xlsx: 使用 openpyxl 打开测试
- 检测命名冲突(如 hnnx vs outline-design-hnnx)
- 输出到 stdout + 可选写入文件(JSON)

用法:
  python scripts/check_templates.py
  python scripts/check_templates.py --output d:/code/QJ/BEMP5.0DEV/output/_template-manifest.json
"""
import argparse
import hashlib
import json
import os
import sys
from collections import defaultdict
from pathlib import Path

# 强制 stdout/stderr 使用 UTF-8（Windows GBK 默认编码导致 Unicode 字符输出失败）
try:
    if hasattr(sys.stdout, 'reconfigure'):
        sys.stdout.reconfigure(encoding='utf-8')
        sys.stderr.reconfigure(encoding='utf-8')
except Exception:
    pass

# 复用 paths 模块（确保 SKILL_ROOT 推算正确）
SCRIPT_DIR = Path(__file__).resolve().parent
SKILL_ROOT = SCRIPT_DIR.parent
sys.path.insert(0, str(SCRIPT_DIR))
try:
    from paths import SKILL_ROOT as P_SKILL_ROOT, ASSETS_DIR as P_ASSETS_DIR
    SKILL_ROOT = P_SKILL_ROOT
    ASSETS_DIR = P_ASSETS_DIR
except Exception as e:
    ASSETS_DIR = SKILL_ROOT / 'assets'

# 模板后缀
TEMPLATE_EXTS = {'.docx', '.json', '.xlsx'}


def sha256_file(p: Path) -> str:
    """计算文件 SHA256 哈希（用于检测重复模板）"""
    h = hashlib.sha256()
    with open(p, 'rb') as f:
        for chunk in iter(lambda: f.read(8192), b''):
            h.update(chunk)
    return h.hexdigest()


def check_docx(p: Path) -> dict:
    """校验 .docx 是否能正常打开"""
    result = {'readable': False, 'error': None, 'paragraphs': 0, 'tables': 0}
    try:
        from docx import Document
        doc = Document(str(p))
        result['readable'] = True
        result['paragraphs'] = len(doc.paragraphs)
        result['tables'] = len(doc.tables)
    except ImportError:
        result['error'] = 'python-docx 未安装'
    except Exception as e:
        result['error'] = str(e)
    return result


def check_json(p: Path) -> dict:
    """校验 .json 是否为有效 JSON"""
    result = {'valid': False, 'error': None, 'top_level_keys': 0}
    try:
        with open(p, 'r', encoding='utf-8') as f:
            data = json.load(f)
        result['valid'] = True
        if isinstance(data, dict):
            result['top_level_keys'] = len(data.keys())
    except Exception as e:
        result['error'] = str(e)
    return result


def check_xlsx(p: Path) -> dict:
    """校验 .xlsx 是否能正常打开"""
    result = {'readable': False, 'error': None, 'sheets': 0, 'sheet_names': []}
    try:
        from openpyxl import load_workbook
        wb = load_workbook(str(p), read_only=True, data_only=True)
        result['readable'] = True
        result['sheet_names'] = wb.sheetnames
        result['sheets'] = len(wb.sheetnames)
    except ImportError:
        result['error'] = 'openpyxl 未安装'
    except Exception as e:
        result['error'] = str(e)
    return result


def classify_template(name: str) -> str:
    """根据文件名归类模板用途"""
    n = name.lower()
    if 'outline' in n or '概要' in name:
        return 'outline-design'
    if '差异化' in name or 'detail' in n or '详细' in name:
        return 'detail-design'
    if '测试用例' in name or 'testcase' in n:
        return 'testcase'
    if '测试报告' in name or 'testreport' in n:
        return 'testreport'
    if '单元测试' in name or 'unit' in n:
        return 'unit-test-report'
    if '需求规格' in name or 'srs' in n:
        return 'srs'
    if 'excel' in n or '用例' in name and n.endswith('.xlsx'):
        return 'testcase-xlsx'
    if n.endswith('.json') and ('flowchart' in n or 'mindmap' in n or 'config' in n):
        return 'diagram-config'
    return 'other'


def detect_naming_conflicts(templates: list) -> list:
    """检测命名冲突：
    - hnnx vs outline-design-hnnx（同一银行的概要设计模板命名重复）
    - .doc vs .docx 同名（OLE2 格式 vs OOXML）
    """
    conflicts = []
    by_hash = defaultdict(list)
    for t in templates:
        by_hash[t['sha256']].append(t)

    # 1) 哈希相同的文件（完全重复）
    for h, items in by_hash.items():
        if len(items) > 1:
            conflicts.append({
                'type': 'identical-content',
                'files': [i['relative_path'] for i in items],
                'sha256': h,
                'message': f'文件内容完全相同（SHA256 哈希一致），建议删除重复',
            })

    # 2) 命名近似（hnnx 银行模板重复）
    hnnx_files = [t for t in templates if 'hnnx' in t['name'].lower()]
    if len(hnnx_files) > 1:
        names = [t['name'] for t in hnnx_files]
        # 检查是否有 template-hnnx-* 和 template-*-hnnx 的近似命名
        if any('template-hnnx' in n for n in names) and any('hnnx' in n and 'template-hnnx' not in n for n in names):
            conflicts.append({
                'type': 'hnnx-naming-overlap',
                'files': names,
                'message': '河南农商（hnnx）模板存在多种命名变体（template-hnnx-*.docx vs template-*-hnnx.docx），建议保留一个',
            })

    # 3) .doc vs .docx 同基础名
    by_basename = defaultdict(list)
    for t in templates:
        # 去掉扩展名
        stem = t['name'].rsplit('.', 1)[0] if '.' in t['name'] else t['name']
        by_basename[stem].append(t)
    for stem, items in by_basename.items():
        exts = {i['name'].rsplit('.', 1)[-1].lower() for i in items}
        if 'doc' in exts and 'docx' in exts:
            conflicts.append({
                'type': 'doc-vs-docx',
                'files': [i['relative_path'] for i in items],
                'message': f'.doc（OLE2）与 .docx（OOXML）同名，python-docx 仅能处理 .docx',
            })

    return conflicts


def scan_templates(assets_dir: Path) -> dict:
    """扫描 assets 目录下所有模板文件并校验"""
    templates = []
    for ext in TEMPLATE_EXTS:
        for p in sorted(assets_dir.rglob(f'*{ext}')):
            rel = p.relative_to(assets_dir)
            entry = {
                'name': p.name,
                'relative_path': str(rel).replace('\\', '/'),
                'absolute_path': str(p),
                'extension': ext,
                'size_bytes': p.stat().st_size,
                'sha256': sha256_file(p),
                'category': classify_template(p.name),
                'readable': False,
                'check': {},
            }
            if ext == '.docx':
                entry['check'] = check_docx(p)
                entry['readable'] = entry['check'].get('readable', False)
            elif ext == '.json':
                entry['check'] = check_json(p)
                entry['readable'] = entry['check'].get('valid', False)
            elif ext == '.xlsx':
                entry['check'] = check_xlsx(p)
                entry['readable'] = entry['check'].get('readable', False)
            templates.append(entry)

    # 统计
    by_ext = defaultdict(int)
    by_cat = defaultdict(int)
    readable = 0
    unreadable = 0
    for t in templates:
        by_ext[t['extension']] += 1
        by_cat[t['category']] += 1
        if t['readable']:
            readable += 1
        else:
            unreadable += 1

    conflicts = detect_naming_conflicts(templates)

    return {
        'skill_root': str(SKILL_ROOT),
        'assets_dir': str(assets_dir),
        'total': len(templates),
        'by_extension': dict(by_ext),
        'by_category': dict(by_cat),
        'readable': readable,
        'unreadable': unreadable,
        'conflicts': conflicts,
        'templates': templates,
    }


def print_summary(report: dict, verbose: bool = False):
    """打印模板清单摘要到 stdout"""
    print('=' * 60)
    print(f'模板完整性检查报告')
    print('=' * 60)
    print(f'技能根: {report["skill_root"]}')
    print(f'模板目录: {report["assets_dir"]}')
    print(f'模板总数: {report["total"]}')
    print(f'  - .docx: {report["by_extension"].get(".docx", 0)}')
    print(f'  - .json: {report["by_extension"].get(".json", 0)}')
    print(f'  - .xlsx: {report["by_extension"].get(".xlsx", 0)}')
    print(f'按类别:')
    for cat, count in sorted(report['by_category'].items(), key=lambda x: -x[1]):
        print(f'  - {cat}: {count}')
    print(f'可读: {report["readable"]} / 不可读: {report["unreadable"]}')
    print(f'命名冲突: {len(report["conflicts"])}')
    for c in report['conflicts']:
        print(f'  [WARN] [{c["type"]}] {c["message"]}')
        for f in c.get('files', []):
            print(f'      - {f}')

    if verbose:
        print()
        print('-' * 60)
        print('详细清单:')
        for t in report['templates']:
            status = '[OK]' if t['readable'] else '[FAIL]'
            err = f"  ({t['check'].get('error', '')})" if not t['readable'] else ''
            extra = ''
            if t['extension'] == '.docx' and t['readable']:
                extra = f"  paragraphs={t['check'].get('paragraphs', 0)}, tables={t['check'].get('tables', 0)}"
            elif t['extension'] == '.xlsx' and t['readable']:
                extra = f"  sheets={t['check'].get('sheets', 0)}"
            elif t['extension'] == '.json' and t['readable']:
                extra = f"  top-level keys={t['check'].get('top_level_keys', 0)}"
            print(f'  {status} [{t["category"]:20s}] {t["name"]:60s} ({t["size_bytes"]:>7d} bytes){extra}{err}')


def main():
    parser = argparse.ArgumentParser(description='检查 bemp-advanced-doc-generator 技能模板完整性')
    parser.add_argument('--output', '-o', type=str, help='清单 JSON 输出路径', default=None)
    parser.add_argument('--verbose', '-v', action='store_true', help='输出详细清单')
    parser.add_argument('--assets-dir', type=str, default=None, help='模板目录（默认 SKILL_ROOT/assets）')
    args = parser.parse_args()

    assets_dir = Path(args.assets_dir) if args.assets_dir else ASSETS_DIR
    if not assets_dir.exists():
        print(f'错误: 模板目录不存在: {assets_dir}', file=sys.stderr)
        sys.exit(1)

    report = scan_templates(assets_dir)
    print_summary(report, verbose=args.verbose)

    if args.output:
        out_path = Path(args.output)
        out_path.parent.mkdir(parents=True, exist_ok=True)
        with open(out_path, 'w', encoding='utf-8') as f:
            json.dump(report, f, ensure_ascii=False, indent=2)
        print(f'\n[OK] 清单已写入: {out_path}')

    # 退出码：仅在有"严重"问题时返回非零（命名冲突 - 重复内容）
    has_severe = any(c['type'] == 'identical-content' for c in report['conflicts'])
    sys.exit(2 if has_severe else 0)


if __name__ == '__main__':
    main()
