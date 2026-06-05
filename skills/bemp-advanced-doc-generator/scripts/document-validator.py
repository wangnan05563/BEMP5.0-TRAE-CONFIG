import sys
import json
from docx import Document
from docx.shared import Pt, Emu, RGBColor

from doc_utils import BLUE, is_blue_color, is_blue_paragraph, is_hyperlink_run, has_non_hyperlink_blue, PLACEHOLDER_KEYWORDS

# 关键章节名常量，避免硬编码散落各处
CHAPTER_DB_DESIGN = '数据库设计'
CHAPTER_SYSTEM_INTEGRATION = '系统集成'
CHAPTER_APPENDIX = '附录'

def validate_document(template_path, generated_path):
    template = Document(template_path)
    generated = Document(generated_path)

    result = {
        'passed': True,
        'errors': [],
        'warnings': [],
        'checks': {}
    }

    tpl_styles = {s.name: s for s in template.styles if s.type is not None}
    gen_styles = {s.name: s for s in generated.styles if s.type is not None}

    missing_styles = set(tpl_styles.keys()) - set(gen_styles.keys())
    if missing_styles:
        for ms in missing_styles:
            result['warnings'].append(f'样式缺失: {ms}')
    result['checks']['styles'] = {
        'templateCount': len(tpl_styles),
        'generatedCount': len(gen_styles),
        'missing': list(missing_styles)
    }

    tpl_tables = template.tables
    gen_tables = generated.tables
    table_issues = []
    for i, tpl_table in enumerate(tpl_tables):
        if i >= len(gen_tables):
            table_issues.append(f'表格{i}: 缺失')
            continue
        gen_table = gen_tables[i]
        tpl_header = [cell.text.strip() for cell in tpl_table.rows[0].cells]
        gen_header = [cell.text.strip() for cell in gen_table.rows[0].cells]
        if tpl_header != gen_header:
            table_issues.append(f'表格{i}: 表头不一致 (模板: {tpl_header}, 生成: {gen_header})')
        if len(tpl_table.columns) != len(gen_table.columns):
            table_issues.append(f'表格{i}: 列数不一致 (模板: {len(tpl_table.columns)}, 生成: {len(gen_table.columns)})')
    result['checks']['tables'] = {
        'templateCount': len(tpl_tables),
        'generatedCount': len(gen_tables),
        'issues': table_issues
    }
    if table_issues:
        for issue in table_issues:
            result['errors'].append(issue)

    tpl_headings = [p.text.strip() for p in template.paragraphs if p.style and p.style.name.startswith('Heading')]
    gen_headings = [p.text.strip() for p in generated.paragraphs if p.style and p.style.name.startswith('Heading')]
    result['checks']['headings'] = {
        'templateCount': len(tpl_headings),
        'generatedCount': len(gen_headings)
    }

    blue_paragraphs = []
    blue_table_cells = []
    for i, p in enumerate(generated.paragraphs):
        if has_non_hyperlink_blue(p):
            blue_paragraphs.append(f'段落{i}: {p.text[:60]}')
    for ti, table in enumerate(generated.tables):
        for ri, row in enumerate(table.rows):
            for ci, cell in enumerate(row.cells):
                for para in cell.paragraphs:
                    if has_non_hyperlink_blue(para):
                        blue_table_cells.append(f'表格{ti}行{ri}列{ci}: {para.text[:40]}')
    result['checks']['blueText'] = {
        'paragraphs': len(blue_paragraphs),
        'tableCells': len(blue_table_cells),
        'details': blue_paragraphs[:10] + blue_table_cells[:10]
    }
    if blue_paragraphs:
        for bp in blue_paragraphs[:5]:
            result['errors'].append(f'蓝色文本未清除: {bp}')
    if blue_table_cells:
        for bc in blue_table_cells[:5]:
            result['errors'].append(f'蓝色表格文本未清除: {bc}')

    placeholder_count = 0
    placeholder_details = []
    for i, p in enumerate(generated.paragraphs):
        text = p.text.strip()
        if text.startswith('<') and any(kw in text for kw in PLACEHOLDER_KEYWORDS):
            placeholder_count += 1
            placeholder_details.append(f'段落{i}: {text[:60]}')
    result['checks']['placeholders'] = {
        'remaining': placeholder_count,
        'details': placeholder_details[:10]
    }
    if placeholder_count > 0:
        for pd in placeholder_details[:5]:
            result['errors'].append(f'未替换占位符/注释: {pd}')

    blank_sections = []
    current_h1 = ''
    current_h2 = ''
    for i, p in enumerate(generated.paragraphs):
        style = p.style.name if p.style else ''
        if style == 'Heading 1':
            current_h1 = p.text.strip()
            current_h2 = ''
        elif style == 'Heading 2':
            current_h2 = p.text.strip()
        elif style == 'Heading 3':
            pass
        elif style == 'Normal' and not p.text.strip():
            pass
        elif not is_blue_paragraph(p):
            pass

    blank_headings = []
    for i, p in enumerate(generated.paragraphs):
        style = p.style.name if p.style else ''
        if style.startswith('Heading'):
            heading_level = int(style.split()[-1]) if style.split()[-1].isdigit() else 0
            has_content = False
            for j in range(i + 1, len(generated.paragraphs)):
                np = generated.paragraphs[j]
                np_style = np.style.name if np.style else ''
                if np_style.startswith('Heading'):
                    np_level = int(np_style.split()[-1]) if np_style.split()[-1].isdigit() else 0
                    if np_level <= heading_level:
                        break
                    continue
                if np.text.strip() and len(np.text.strip()) > 20:
                    has_content = True
                    break
            if not has_content:
                blank_headings.append(f'{style}: {p.text.strip()[:50]}')
    result['checks']['blankSections'] = {
        'count': len(blank_headings),
        'details': blank_headings[:10]
    }
    if blank_headings:
        for bh in blank_headings[:5]:
            result['warnings'].append(f'可能空白章节: {bh}')

    tpl_section = template.sections[0] if template.sections else None
    gen_section = generated.sections[0] if generated.sections else None
    page_issues = []
    if tpl_section and gen_section:
        if abs(tpl_section.page_width - gen_section.page_width) > 10000:
            page_issues.append('页面宽度不一致')
        if abs(tpl_section.page_height - gen_section.page_height) > 10000:
            page_issues.append('页面高度不一致')
        if abs(tpl_section.left_margin - gen_section.left_margin) > 10000:
            page_issues.append('左边距不一致')
        if abs(tpl_section.right_margin - gen_section.right_margin) > 10000:
            page_issues.append('右边距不一致')
    result['checks']['pageLayout'] = {'issues': page_issues}
    for pi in page_issues:
        result['errors'].append(pi)

    header_ok = False
    footer_ok = False
    if gen_section and gen_section.header:
        for p in gen_section.header.paragraphs:
            if p.text.strip():
                header_ok = True
    if gen_section and gen_section.footer:
        for p in gen_section.footer.paragraphs:
            if p.text.strip():
                footer_ok = True
    result['checks']['headerFooter'] = {'header': header_ok, 'footer': footer_ok}
    if not header_ok:
        result['warnings'].append('页眉为空')
    if not footer_ok:
        result['warnings'].append('页脚为空')

    filled_sections = []
    for i, p in enumerate(generated.paragraphs):
        style = p.style.name if p.style else ''
        if style.startswith('Heading'):
            for j in range(i + 1, min(i + 5, len(generated.paragraphs))):
                np = generated.paragraphs[j]
                if np.style and np.style.name.startswith('Heading'):
                    break
                if np.text.strip() and len(np.text.strip()) > 20:
                    filled_sections.append({
                        'heading': p.text.strip()[:50],
                        'contentLength': len(np.text.strip())
                    })
                    break
    result['checks']['contentSummary'] = {
        'filledSections': len(filled_sections),
        'totalHeadings': len(gen_headings),
        'details': [s['heading'] for s in filled_sections[:15]]
    }

    if result['errors']:
        result['passed'] = False

    return result

def validate_generated_doc(generated_path):
    """无需模板的文档质量验证，用于生成后自动检查"""
    generated = Document(generated_path)

    result = {
        'passed': True,
        'errors': [],
        'warnings': [],
        'checks': {}
    }

    heading_issues = []
    last_level = 0
    for i, p in enumerate(generated.paragraphs):
        style = p.style.name if p.style else ''
        if style.startswith('Heading'):
            parts = style.split()
            if len(parts) == 2 and parts[1].isdigit():
                level = int(parts[1])
                if level > last_level + 1:
                    heading_issues.append(f'标题跨级: "{p.text.strip()[:40]}" 从H{last_level}跳到H{level}')
                last_level = level
    result['checks']['headingHierarchy'] = {
        'issues': heading_issues
    }
    for issue in heading_issues:
        result['errors'].append(issue)

    blue_count = 0
    for p in generated.paragraphs:
        if has_non_hyperlink_blue(p):
            blue_count += 1
    for table in generated.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    if has_non_hyperlink_blue(para):
                        blue_count += 1
    result['checks']['blueText'] = {'nonHyperlinkBlueCount': blue_count}
    if blue_count > 0:
        result['errors'].append(f'非超链接蓝色字体未清理: {blue_count}处')

    # 3. 表格字体大小检查
    inconsistent_tables = []
    for ti, table in enumerate(generated.tables):
        sizes = set()
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        if run.font.size and run.text.strip():
                            sizes.add(run.font.size)
        if len(sizes) > 2:
            inconsistent_tables.append(f'表格{ti}: 字体大小不统一 {[str(s) for s in sorted(sizes)]}')
    result['checks']['tableFonts'] = {
        'inconsistentCount': len(inconsistent_tables),
        'details': inconsistent_tables[:5]
    }
    for it in inconsistent_tables[:3]:
        result['warnings'].append(it)

    # 4. ER章节位置检查
    er_h1_found = False
    er_after_system_integration = False
    er_before_appendix = False
    h1_order = []
    for p in generated.paragraphs:
        if p.style and p.style.name == 'Heading 1':
            h1_order.append(p.text.strip())
    for idx, title in enumerate(h1_order):
        if CHAPTER_DB_DESIGN in title:
            er_h1_found = True
            if idx > 0 and CHAPTER_SYSTEM_INTEGRATION in h1_order[idx - 1]:
                er_after_system_integration = True
            if idx < len(h1_order) - 1 and CHAPTER_APPENDIX in h1_order[idx + 1]:
                er_before_appendix = True
    result['checks']['erChapter'] = {
        'found': er_h1_found,
        'afterSystemIntegration': er_after_system_integration,
        'beforeAppendix': er_before_appendix
    }
    if not er_h1_found:
        result['warnings'].append(f'未找到"{CHAPTER_DB_DESIGN}"H1章节')
    elif not er_after_system_integration or not er_before_appendix:
        result['warnings'].append(f'ER章节位置不正确，应在"{CHAPTER_SYSTEM_INTEGRATION}"之后、"{CHAPTER_APPENDIX}"之前')

    # 5. 占位符检查
    placeholder_count = 0
    for p in generated.paragraphs:
        text = p.text.strip()
        if text.startswith('<') and any(kw in text for kw in PLACEHOLDER_KEYWORDS):
            placeholder_count += 1
    result['checks']['placeholders'] = {'remaining': placeholder_count}
    if placeholder_count > 0:
        result['errors'].append(f'未替换占位符: {placeholder_count}处')

    if result['errors']:
        result['passed'] = False

    return result

if __name__ == '__main__':
    import sys as _sys
    # Windows 控制台默认 GBK，强制 UTF-8 输出避免编码错误
    _sys.stdout.reconfigure(encoding='utf-8', errors='replace')
    if len(_sys.argv) < 3:
        print('Usage: python document-validator.py <template.docx> <generated.docx>', file=_sys.stderr)
        _sys.exit(1)
    result = validate_document(_sys.argv[1], _sys.argv[2])
    print(json.dumps(result, ensure_ascii=False, indent=2))
    _sys.exit(0 if result['passed'] else 1)