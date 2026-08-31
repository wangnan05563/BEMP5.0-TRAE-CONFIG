# -*- coding: utf-8 -*-
"""
BEMP 河南农信"反洗钱功能优化（HNNS-EB-STD-REQ-002，PRD v3.4.8）全面测试"交付文档生成器
产出：正式测试报告 .docx（7 章节结构：测试概述/测试策略与用例复用/测试执行汇总/
缺陷记录与闭环/已知问题与遗留项/测试结论/附录证据索引）。
生成模式与 gen_aml4_design.py 先例一致（全代码构建，样式工具函数同源）。
数据来源：AML-FT-20260830-01 / AML3-RT-20260830-01 / AML4-FT-20260829-01 /
AML4-RT-20260829-02 / antimoney_e2e_regression_20260830_v1.md / DEF-02-03-变更处置.sql。
"""
import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

OUTPUT_DIR = r"d:\code\QJ\BEMP5.0DEV\docs\delivery\反洗钱功能优化"
OUTPUT_FILE = os.path.join(OUTPUT_DIR, "反洗钱功能优化-测试报告-v1.docx")

# ===================== 样式工具函数（与 gen_aml4_design.py 先例一致） =====================

def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = tcPr.find(qn('w:tcBorders'))
    if tcBorders is None:
        tcBorders = OxmlElement('w:tcBorders')
        tcPr.append(tcBorders)
    for edge in ('top', 'left', 'bottom', 'right'):
        tag = 'w:{}'.format(edge)
        element = tcBorders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            tcBorders.append(element)
        element.set(qn('w:val'), 'single')
        element.set(qn('w:sz'), '4')
        element.set(qn('w:color'), '000000')


def set_cell_background(cell, color="D9D9D9"):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color)
    tcPr.append(shd)


def set_run_font(run, name="宋体", size=10.5, bold=False, color=None):
    run.font.name = name
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    rFonts.set(qn('w:eastAsia'), name)
    rFonts.set(qn('w:ascii'), name)
    rFonts.set(qn('w:hAnsi'), name)


def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        set_run_font(run, name="黑体", size=max(12, 18 - level * 2), bold=True, color="000000")
    return p


def add_para(doc, text, bold=False, indent=True, size=10.5):
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.first_line_indent = Cm(0.74)
    p.paragraph_format.line_spacing = 1.5
    run = p.add_run(text)
    set_run_font(run, name="宋体", size=size, bold=bold)
    return p


def add_bullet(doc, text, size=10.5):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.line_spacing = 1.5
    run = p.add_run(text)
    set_run_font(run, name="宋体", size=size)
    return p


def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr_cells[i].text = ""
        p = hdr_cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(header)
        set_run_font(run, name="宋体", size=10, bold=True)
        set_cell_background(hdr_cells[i])
        set_cell_border(hdr_cells[i])
        hdr_cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    for r_idx, row in enumerate(rows, 1):
        row_cells = table.rows[r_idx].cells
        for c_idx, val in enumerate(row):
            row_cells[c_idx].text = ""
            p = row_cells[c_idx].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(str(val))
            set_run_font(run, name="宋体", size=10)
            set_cell_border(row_cells[c_idx])
            row_cells[c_idx].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    if col_widths:
        for i, width in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(width)
    doc.add_paragraph()
    return table


def add_toc_field(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run()
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = 'TOC \\o "1-3" \\h \\z \\u'
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')
    fldChar3 = OxmlElement('w:t')
    fldChar3.text = '右键此处选择"更新域"以生成目录'
    fldChar4 = OxmlElement('w:fldChar')
    fldChar4.set(qn('w:fldCharType'), 'end')
    r_element = run._r
    r_element.append(fldChar1)
    r_element.append(instrText)
    r_element.append(fldChar2)
    r_element.append(fldChar3)
    r_element.append(fldChar4)


def add_page_break(doc):
    p = doc.add_paragraph()
    run = p.add_run()
    run.add_break(WD_BREAK.PAGE)


def setup_page_and_header_footer(doc):
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(3.0)
        section.right_margin = Cm(2.5)
        header = section.header
        if not header.paragraphs[0].runs:
            hp = header.paragraphs[0]
            hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = hp.add_run("河南农商票据交易管理平台 - 测试报告")
            set_run_font(run, name="宋体", size=9, color="808080")
        footer = section.footer
        fp = footer.paragraphs[0]
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = fp.add_run()
        set_run_font(run, name="宋体", size=9, color="808080")
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')
        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = 'PAGE'
        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'end')
        run._r.append(fldChar1)
        run._r.append(instrText)
        run._r.append(fldChar2)


def inject_update_fields(doc):
    settings = doc.settings.element
    update_fields = settings.find(qn('w:updateFields'))
    if update_fields is None:
        update_fields = OxmlElement('w:updateFields')
        update_fields.set(qn('w:val'), 'true')
        settings.append(update_fields)


# ===================== 文档生成主流程 =====================

def build_test_report_doc():
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = '宋体'
    style.font.size = Pt(10.5)
    rPr = style.element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    rFonts.set(qn('w:eastAsia'), '宋体')

    setup_page_and_header_footer(doc)

    # ===== 封面 =====
    for _ in range(3):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("河南农商银行")
    set_run_font(run, name="黑体", size=26, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("新一代票据系统个性化开发")
    set_run_font(run, name="黑体", size=22, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("反洗钱功能优化")
    set_run_font(run, name="黑体", size=22, bold=True)

    for _ in range(2):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("测 试 报 告")
    set_run_font(run, name="黑体", size=20, bold=True)

    for _ in range(6):
        doc.add_paragraph()

    info_table = doc.add_table(rows=6, cols=2)
    info_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    info_data = [
        ("项目名称", "BEMP5.0 河南农商银行个性化开发"),
        ("需求名称", "反洗钱功能优化——中互金关注名单校验（HNNS-EB-STD-REQ-002）"),
        ("需求依据", "PRD v3.4.8 + 待确认清单 v3.4.8（2026-08-28）"),
        ("测试范围", "市场交易侧 + 四阶段客户维护侧全面测试（53 条用例）"),
        ("文档版本", "V1.0"),
        ("编写日期", "2026-08-30"),
    ]
    for i, (k, v) in enumerate(info_data):
        cells = info_table.rows[i].cells
        cells[0].text = ""
        cells[1].text = ""
        p1 = cells[0].paragraphs[0]
        p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r1 = p1.add_run(k)
        set_run_font(r1, name="宋体", size=12, bold=True)
        p2 = cells[1].paragraphs[0]
        p2.alignment = WD_ALIGN_PARAGRAPH.LEFT
        r2 = p2.add_run(v)
        set_run_font(r2, name="宋体", size=12)
        for c in cells:
            set_cell_border(c)

    add_page_break(doc)

    # ===== 修订记录 =====
    add_heading(doc, "修订记录", level=1)
    add_table(doc,
        ["版本", "日期", "修订内容", "修订人"],
        [
            ["V1.0", "2026-08-30", "初始版本：全面测试汇总——53 条用例一轮执行（AML-FT-20260830-01）、缺陷闭环 DEF-02~05、二轮回归 8 项（AML3-RT-20260830-01）、二轮端到端全场景 PASS、已知问题与遗留项、测试结论与建议、证据文件索引；同步落实用例评审轻微项 m1/m2 回写", "BEMP文档交付工程师"],
        ],
        col_widths=[2, 3, 8, 3]
    )
    add_page_break(doc)

    # ===== 目录 =====
    add_heading(doc, "目录", level=1)
    add_toc_field(doc)
    add_page_break(doc)

    # ===== 第一章 测试概述 =====
    add_heading(doc, "第一章 测试概述", level=1)

    add_heading(doc, "1.1 测试目的与依据", level=2)
    add_para(doc, "本轮测试为反洗钱功能优化需求的全面测试：在前期分侧测试（市场交易侧、四阶段校验侧）基础上，汇总验证市场交易侧与四阶段客户维护侧功能符合 PRD v3.4.8 口径，完成缺陷闭环与二轮回归，形成可交付结论。")
    add_table(doc,
        ["项", "内容"],
        [
            ["需求编号", "HNNS-EB-STD-REQ-002（票据业务系统反洗钱工作对接·日常优化类）"],
            ["需求版本", "PRD v3.4.8（2026-08-28，含公司金融部 docx 对照补充 Q-11/Q-12）"],
            ["需求确认依据", "《反洗钱功能优化-待确认清单》v3.4.8（业务项全部已确认；Q-11 测试假设；Q-12 进展已补记）"],
            ["测试目的", "验证中互金关注名单数据管理（全量/增量抽取、查询/导出/导入）、市场交易侧 5 角色两轮匹配校验、四阶段客户维护侧校验、面客/对内文案分流、独立开关与异常降级符合 PRD v3.4.8 与既有确认口径"],
            ["测试结论预告", "市场交易侧 + 四阶段客户维护侧功能符合 PRD v3.4.8；P0/P1 代码缺陷 0；P2 缺陷 4 项（配置/数据类）全部闭环；建议进入下一阶段/上线评估（详见第六章）"],
        ],
        col_widths=[3, 12.5]
    )

    add_heading(doc, "1.2 测试范围", level=2)
    add_table(doc,
        ["范围域", "覆盖内容"],
        [
            ["名单数据管理", "HNNXTK020113 定时抽取 INCR 增量 upsert（OBJ_ID 主键）/FULL 全量快照双模式、19 列文件解析（GBK/0x03 分隔）、口径过滤（DATA_SRC=18、VLD_ST=1）、名单查询/多条件分页/批量导出/导入、名单基线管理"],
            ["市场交易侧校验", "转贴现买入/卖出、承兑、再贴现（仅 00102006）业务入口；出票人/交易对手/承兑行/贴现行/贴现人 5 角色独立校验；round1 统一社会信用代码优先（certType='22'+certNo）+ round2 名称兜底两轮匹配；预检接口 + commit 兜底双层拦截；对内文案 HNNX0BE320009/010"],
            ["四阶段客户维护侧", "④企业客户维护（insert/update 落库前校验）；面客文案 HNNX0BE320011（角色前缀+96288，不暴露名单性质）；独立开关 hnnx.cust.antiMoneyList.block；异常降级 fail-closed/fail-open"],
            ["网银三分支（审查段）", "①直贴贴现分支（NES.007）、②提示承兑分支（NES.002）、③背书转让分支（NES.006）代码审查段（bean 顶替/取数链/校验时点/空集合短路）；报文应答段 ⏸ 延后网银联调"],
            ["横切能力", "校验开关（市场/客户双开关隔离与实时生效）、异常分层降级（预检/round1/commit 三层）、审计留痕（LOGGER 日志载体）、越权访问拦截"],
        ],
        col_widths=[3.2, 12.3]
    )
    add_para(doc, "范围外（明确不做）：网银报文应答段实测（需网银外围联调环境，⏸ 延后）；灰名单人工审批流（二期规划，一期以提示中断降级承载）；行内页面直贴/承兑/背书入口（原文操作路径均为网银端发起）；业务表直插造数（以完整票据流转为准）。")

    add_heading(doc, "1.3 测试环境", level=2)
    add_table(doc,
        ["项", "终态实测状态"],
        [
            ["服务", "前端 8091 / 后端 8010 / Redis 6379 / ZooKeeper 21811 / Adapter 8090 全 UP；Served 于 DEF-02 配置生效后重启（PID 40256）"],
            ["业务库", "10.20.42.211:1521/orcl，应用实际 schema=BEMP_HNNX（JDBC 通道 run-oracle-jdbc.ps1 -Schema BEMP_HNNX；Oracle MCP 连接的 BEMP schema 为演示遗留库，不承载反洗钱数据）"],
            ["名单表基线", "HNNX_M_CUST_SPECIAL_INFO 共 8 条基线：R1 黑（21/长期有效）、R2 灰（21/20250101~20281231）、R5 已失效（END=20250731）、R9 未生效（START=20300101）、R10 证件类型 01 对照、R22 统一社会信用代码对照（CERT_TYPE='22'/CERT_NO=91861230741970469R）、R-SPACE 空格名称对照、R99 存量保留对照"],
            ["开关参数", "hnnx.market.antiMoneyList.block='1'、hnnx.cust.antiMoneyList.block='1'（均 LEGAL_NO='000000'）；同步模式 hnnx.market.antimoney.sync.mode=INCR"],
            ["数据清洁度", "TM_CUST_CORP 测试客户档案零残留（重名检查 CNT=0）；ecif 目录 2 份标准测试文件（D:/code/home/hnnxbank/ecif/{日期}/ECF_M_CUST_SPECIAL_INFO_CBS_{日期}.txt，GBK 编码、0x03 分隔、19 列）"],
            ["业务日期", "BusiDateService 实际返回 20250828（2025-08-28），全部时效类断言（R5 已失效/R9 未生效/R22 生效时间）以该日期为基准"],
        ],
        col_widths=[2.8, 12.7]
    )

    add_heading(doc, "1.4 测试账号", level=2)
    add_table(doc,
        ["账号", "角色权限", "使用轮次"],
        [
            ["mllzs01", "法人管理员（test_config.json → banks.hnnxbank.login.admin）", "四阶段功能测试一轮/二轮回归（Playwright + Chrome DevTools 双通道）"],
            ["wangnan01", "场内交易 + 中互金名单管理权限", "二轮端到端回归（E2E 全场景）"],
        ],
        col_widths=[2.5, 7.5, 5.5]
    )

    add_heading(doc, "1.5 测试方法与轮次", level=2)
    add_table(doc,
        ["轮次", "日期", "执行通道", "报告编号/留档"],
        [
            ["四阶段功能测试（一轮）", "2026-08-29", "Playwright + DB 断言 + 日志留痕 + 单测", "AML4-FT-20260829-01"],
            ["四阶段二轮回归", "2026-08-29", "Chrome DevTools 真实浏览器（防同通道偶发/共因误判）", "AML4-RT-20260829-02"],
            ["全面测试一轮（53 条汇总）", "2026-08-30", "多通道（页面实证/代码审查/DB 断言/JDBC/日志）", "AML-FT-20260830-01"],
            ["缺陷闭环处置", "2026-08-30", "配置登记 + JDBC 落库 + 重启 Served + 修复验证", "DEF-02-03-变更处置.sql 等留档"],
            ["二轮回归（8 项）", "2026-08-30", "JDBC + 运行时验证（round3_*.sql 轨迹）", "AML3-RT-20260830-01"],
            ["二轮端到端（E2E）", "2026-08-30", "Playwright（Chromium headless）真实浏览器（⚠Chrome DevTools MCP 超时污染降级，降级记录随报告归档）", "antimoney_e2e_regression_20260830_v1.md"],
        ],
        col_widths=[4.2, 2.2, 5.6, 3.5]
    )
    add_page_break(doc)

    # ===== 第二章 测试策略与用例复用说明 =====
    add_heading(doc, "第二章 测试策略与用例复用说明", level=1)

    add_heading(doc, "2.1 用例复用策略", level=2)
    add_para(doc, "本轮全面测试不重新编制用例，完整复用既有用例基线 v3.4.11（经两轮用例评审放行口径），共 53 条：P0=26、P1=22、P2=5。市场交易侧 30 条（P0=14、P1=15、P2=1）与四阶段增量 23 条（P0=12、P1=7、P2=4）合并归集执行。")
    add_table(doc,
        ["用例文件", "优先级构成", "路径"],
        [
            ["反洗钱名单校验-P0用例.md", "P0=26（市场侧 14 + 四阶段增量 P0-015~026 共 12）", "aotutests-devtools/testcases/反洗钱名单校验/"],
            ["反洗钱名单校验-P1用例.md", "P1=22（市场侧 15 + 四阶段增量 P1-026~032 共 7）+ P2=5（市场侧 P2-026 + 四阶段 P2-027~030）", "同上"],
            ["配套文件", "功能地图与优先级矩阵 / 测试数据准备（§2.3 名单文件、§4.3 场景记录）", "同上"],
        ],
        col_widths=[4.5, 6, 5]
    )

    add_heading(doc, "2.2 v3.4.8 增量核对与 P0-022 修订", level=2)
    add_para(doc, "执行前以 PRD v3.4.8 为基准对 53 条用例做覆盖度核对（v3.4.11 修订，2026-08-30），核对结论：覆盖无缺口，不新增用例编号、总数不变。核心修订两项：")
    add_table(doc,
        ["修订项", "内容", "影响"],
        [
            ["P0-022 灰名单全文断言", "补灰名单 HNNX0BE320010 全文现状断言（「客户[{0}]在中互金灰名单中，暂不可办理票据业务，请联系风控人员处理」，hnnx_mt_zh_CN.properties L10 锚点），并追加 Q-11 待确认项标注块（PRD v3.4.8 §3.2.3 三偏差对照：原文面客模板 vs 市场侧对内文案，本用例按现状文案断言，\"按原文案模板调整\"为基于 Q-11 假设的待确认项）", "断言锚点精确化；Q-11 确认后若文案调整，仅需回归该断言段"],
            ["P0-018/019 Q-12 确认结论标注", "关联需求补 Q-12 确认结论标注（④客户信息维护纳入本期实现，网银三分支报文应答段 ⏸ 报文联调阶段执行）", "用例与需求确认进展一致（m1 同步回写待确认清单）"],
        ],
        col_widths=[3.2, 8.3, 4]
    )

    add_heading(doc, "2.3 用例评审闭环与轻微项落实", level=2)
    add_para(doc, "用例评审结论：通过（严重 0 / 主要 0 / 轻微 2）。两项轻微项已于本轮交付前全部落实回写：")
    add_table(doc,
        ["编号", "轻微项", "落实情况"],
        [
            ["m1", "待确认清单 Q-12 处补记确认进展", "已落实：Q-12 汇总表行状态更新为\"进展补记（2026-08-30）\"，详情节追加确认进展块——④客户信息维护已纳入本期实现（对应用例已执行验证）；网银三分支报文段延后报文联调阶段执行（⏸ 待网银联调环境）"],
            ["m2", "P0 用例文件头部\"用例总数\"补注", "已落实：补注\"另含四阶段增量 P0-015~026 共 12 条，合计 P0=26 条\"（v3.4.11 m2 覆盖度核对补注）"],
        ],
        col_widths=[1.5, 4.5, 9.5]
    )

    add_heading(doc, "2.4 测试数据准备", level=2)
    add_para(doc, "数据准备遵循《反洗钱名单校验-测试数据准备.md》（v3.4.10 口径），名单基线 8 条覆盖黑/灰/时效/证件类型/名称格式/存量保留全部对照场景；名单测试文件含 INCR 特殊行（OBJ_ID 空 R11、单条失败 R12）与过滤行（R3/R4/R6），保障双模式与口径过滤断言。执行期经 DEF-02（配置）/DEF-03/DEF-05（数据）缺陷闭环后达到 1.3 节终态。")
    add_page_break(doc)

    # ===== 第三章 测试执行汇总 =====
    add_heading(doc, "第三章 测试执行汇总", level=1)

    add_heading(doc, "3.1 一轮功能测试执行分布（AML-FT-20260830-01）", level=2)
    add_para(doc, "一轮功能测试覆盖 53 条用例，执行结果：33 通过 + 4 部分通过 + 16 ⏸（延后）+ 0 失败，P0/P1 代码缺陷 0。优先级与结果双维度分布如下：")
    add_table(doc,
        ["优先级", "用例数", "占比", "主要覆盖"],
        [
            ["P0", "26", "49.1%", "双模式抽取、5 角色命中拦截、④客户维护拦截、面客/对内文案、独立开关、两轮匹配口径、降级策略"],
            ["P1", "22", "41.5%", "名单查询/导出/导入、越权拦截、审计留痕载体、灰名单面客文案、未命中放行、开关互不影响、预约承兑自动覆盖"],
            ["P2", "5", "9.4%", "来文行不校验、时效边界四象限、多角色命中顺序、贴现申请人回退、承兑行反查降级"],
            ["合计", "53", "100%", "—"],
        ],
        col_widths=[1.8, 1.8, 1.8, 10.1]
    )
    add_table(doc,
        ["执行结果", "数量", "说明"],
        [
            ["通过", "33", "断言段全部执行且全部通过（含代码审查段全过用例）"],
            ["部分通过", "4", "主要断言段通过、个别延后段未执行（如审查段通过 + 页面段/报文段 ⏸）"],
            ["⏸ 延后", "16", "存在明确延后归因（联调环境/数据前置/二期范围），不阻塞门禁"],
            ["失败", "0", "无功能性失败；无阻塞用例"],
        ],
        col_widths=[2.5, 1.8, 11.2]
    )
    add_para(doc, "⏸ 与部分通过的构成归集（不阻塞门禁，逐条明细以一轮报告原始记录为准）：")
    add_table(doc,
        ["构成类别", "关联用例", "延后原因与出路"],
        [
            ["网银报文应答段", "P0-023/024/025、P1-027、P1-030（报文段）", "需网银外围联调环境/报文模拟器；代码审查段已全部通过；列入已知问题 ⏸ 清单待网银联调"],
            ["eDisc 页面段", "P0-020/021", "TB_BILL_INFO 无出票人=名单客户的在途票据（0/21），业务表按原则不直插；审查段（commit 兜底时点/角色组装）已通过；需完整票据流转数据后回归"],
            ["数据前置类", "P0-001/002（cbsFilePath 配置）、P0-026 round1 段、P1-022（R22）", "随 DEF-02~05 缺陷闭环后，二轮回归 8/8 全部补验通过（见 3.3）"],
            ["二期/环境类", "P0-007（灰名单人工审批流二期，一期提示中断降级承载已验证）、P1-028 步骤 5/6（运行时异常注入演练段，降级两分支已由单测+代码审查覆盖）", "随二期交付/环境能力就绪跟踪"],
            ["场景补验类", "场景 A/C/D、19 列解析、round1 信用代码优先命中、轮次互斥", "随数据/配置就绪后二轮回归补验通过（见 3.3）"],
        ],
        col_widths=[2.8, 5.2, 7.5]
    )

    add_heading(doc, "3.2 一轮关键执行证据", level=2)
    add_table(doc,
        ["证据项", "实测内容", "对应证据"],
        [
            ["④路径页面实证", "企业客户维护页面新增/修改黑名单客户：拦截文案「客户名：该客户[河南宏信贸易有限公司]暂不符合相关政策要求，暂不支持提供该服务，详情请咨询开户机构或96288」逐字命中、TM_CUST_CORP 不落库（CNT=0）、命中日志 ERROR 留痕；灰名单场景同构且不暴露名单性质", "results/fill_P0-017~019.png、result_*.json、后端 app.log"],
            ["双重越权拦截", "名单管理页面/按钮级越权访问双重拦截实测：无权限角色访问名单接口被鉴权拦截、按钮级权限不渲染（btnAuth.add.isShow=false 经 Vue 实例通道旁路完成表单验证，权限配置建议已列入已知问题）", "p1_019_auth_result.json、round3_04_auth_probe.sql"],
            ["口径过滤实测", "DATA_SRC=18、VLD_ST=1 口径过滤实测：非中互金来源/无效记录不参与命中（未命中放行）与命中拦截正反对照成立；round1 命中后排除出 round2 索引（轮次互斥）", "round3_13/17/18 SQL 轨迹、单测 HnnxAntiMoneyCustValidateTest 8 条全绿"],
            ["面客文案逐字断言", "E2E 拦截弹窗文案 exactMsg=true 逐字一致；接口响应 HTTP 200 + {\"retCode\":\"HNNX0BE320011\",\"retMsg\":…完整文案}；对内文案 320009/010 与面客 320011 分流无交叉", "antimoney_e2e_regression_20260830_v1.md §二"],
            ["双重 trim 口径", "带首尾空格黑名单客户名提交被拦截：前端表单序列化 trim + 后端交易侧 trim（HnnxAntiMoneyValidateUtil L217-222）双层口径一致，无漏匹配风险", "AML4-RT V2b 变体实测"],
            ["开关实时生效", "开关 UPDATE='0' 提交放行落库且日志 INFO 跳过留痕，恢复='1' 同场景恢复拦截——无需重启/清缓存", "AML4-RT V2c 变体实测、一轮 P0-017 复现"],
        ],
        col_widths=[2.8, 8.2, 4.5]
    )

    add_heading(doc, "3.3 二轮回归 8 项（AML3-RT-20260830-01）", level=2)
    add_para(doc, "缺陷闭环（DEF-02~05）后执行二轮回归，对一轮 ⏸/数据前置用例全部补验，结果 8/8 全过。全程以 round3_01~19 SQL 轨迹落盘为证，执行后环境复位（两开关='1'、测试数据零残留、名单基线 8 条完整）。")
    add_table(doc,
        ["序号", "回归项", "关联用例", "结果", "SQL 轨迹"],
        [
            ["1", "INCR 增量 upsert + 幂等（OBJ_ID 主键 update/insert、不 deleteAll、OBJ_ID 空跳过、单条失败 continue）", "P0-001", "PASS", "round3_05_p0001_verify / round3_06_p0001_rerun"],
            ["2", "FULL 全量快照（deleteAll+insertBatch 全量替换、切回 INCR 恢复）", "P0-002", "PASS", "round3_08~12_p0002（backup/switch_full/verify/switch_back/restore）"],
            ["3", "场景 A 补验（黑名单拦截场景，回归客户 AML3RT0001X）", "P0-003 组", "PASS", "round3_14_cleanup_aml3rt0001x"],
            ["4", "场景 C 补验（时效场景，回归客户 AML3RT0002Y）", "P2-027 组", "PASS", "round3_15_cleanup_aml3rt0002y"],
            ["5", "场景 D 补验（场景化组合验证）", "P0 组", "PASS", "round3_*.sql 轨迹"],
            ["6", "19 列解析验证（GBK/0x03 分隔 19 列名单文件解析落库）", "P0-001/002 前置", "PASS", "round3_05/10（解析统计核对）"],
            ["7", "round1 信用代码优先命中（R22 修复后 CERT_TYPE='22'+certNo 优先匹配）", "P1-022 / P0-013", "PASS", "round3_13_p1022_dbcheck / round3_17_p1022_x3_verify"],
            ["8", "轮次互斥（round1 命中排除出 round2 索引、无信用代码不入索引）", "P0-026", "PASS", "round3_18_p026_y2_verify"],
        ],
        col_widths=[1.2, 6.2, 2.6, 1.5, 4.0]
    )
    add_para(doc, "辅助轨迹：round3_01_baseline（基线 8 条核对）、round3_02_busidate_probe（业务日期 20250828 基准探查）、round3_03_menu_probe / round3_04_auth_probe（菜单与越权探查，支撑双重越权拦截证据）、round3_16_def05_fix_r22（DEF-05 修复执行）、round3_19_final_cleanup（终态清理）。")

    add_heading(doc, "3.4 二轮端到端回归（E2E，antimoney_e2e_regression_20260830_v1.md）", level=2)
    add_para(doc, "真实浏览器端到端回归覆盖登录→名单页→客户维护拦截→修正提交→清理闭环全链路，全场景 PASS，全程 fatal JS=0、无白屏、无 4xx/5xx 业务接口失败。执行通道说明：Chrome DevTools MCP 在场景 1 截图阶段触发 Page.captureScreenshot 协议超时污染（CDP 响应槽位损坏，重启进程未恢复），按降级规则切换 Playwright（Chromium headless）通道继续执行，MCP 阶段已完成断言在降级通道全部复验通过（降级记录随报告归档，证据等价：request/response 事件监听 + performance API 替代 list_network_requests）。")
    add_table(doc,
        ["场景", "内容", "结果", "关键证据"],
        [
            ["S0", "登录 wangnan01（含强制登录确认弹窗处理）", "PASS", "登录链路 queryVersion→getSecurityKey(SM4)→userLogin 全 200"],
            ["S1", "中互金名单页面渲染 + 基线核对 + 控制台健康", "PASS", "DataGrid 渲染；基线客户在列；致命错误=0"],
            ["S2", "名称模糊查询/重置恢复/导出按钮触发", "PASS", "输入\"河南\"精确命中 1 条（10→1 行）；重置恢复；导出可触发且 fatal 新增=0"],
            ["S3-1", "企业客户维护新增黑名单客户拦截", "PASS", "弹窗文案逐字命中（exactMsg=true）；retCode=HNNX0BE320011；TM_CUST_CORP 不落库"],
            ["S3-2", "修正客户名后提交成功", "PASS", "retCode=000000 新增成功（顺带实证：黑名单拦截先于参数格式校验执行，见 NOTE-3-04）"],
            ["S3-3", "列表可见 + 注销清理 + 终态", "PASS", "新客户列表可见（E2E20260830001）；注销后有效残留=0"],
            ["S4", "校验开关开启态（行为级）", "PASS", "黑名单拦截按预期触发，行为级证明开关为开启态"],
            ["S5", "全程监控（console/网络/白屏）", "PASS", "fatal=0；非致命 6 条均为环境噪声（CLodop 打印控件/WS 9080/路由守卫提示）；无白屏"],
            ["S6", "清理恢复终态", "PASS", "测试客户注销、名单页查询重置、基线未变更"],
        ],
        col_widths=[1.5, 5.8, 1.6, 6.6]
    )
    add_page_break(doc)

    # ===== 第四章 缺陷记录与闭环 =====
    add_heading(doc, "第四章 缺陷记录与闭环", level=1)

    add_heading(doc, "4.1 缺陷统计", level=2)
    add_table(doc,
        ["严重度", "数量", "明细", "闭环状态"],
        [
            ["P0", "0", "无", "—"],
            ["P1", "0", "无", "—"],
            ["P2", "4", "DEF-02（配置类）/ DEF-03（数据类）/ DEF-04（数据类）/ DEF-05（数据类）", "全部修复 + 验证双重确认关闭（修复与验证角色分离）"],
            ["P3", "0 缺陷 / 5 项已知问题", "E2E 发现 1~3 等转入第五章已知问题清单", "记录跟踪，不阻塞交付"],
        ],
        col_widths=[1.8, 2.2, 7.0, 4.5]
    )
    add_para(doc, "历史衔接：四阶段侧一轮唯一缺陷 DEF-AML4-001（P2/数据类，灰名单 R2、已失效 R5 名单记录未落库）已于 AML4-FT-20260829-01 轮闭环（SQL 直插 + 重测三重断言通过），本轮 8 条名单基线中 R2/R5 已稳定在库。")

    add_heading(doc, "4.2 DEF-02（P2/配置类）：cbsFile.cbsFilePath 未配置导致抽取任务失败", level=2)
    add_table(doc,
        ["环节", "内容"],
        [
            ["发现", "P0-001/002 执行时，HNNXTK020113 名单抽取任务（及手动触发 func_importFile）运行时异常\"未配置文件根目录 cbsFile.cbsFilePath\"，任务失败"],
            ["根因", "该配置项为 Spring @Value(\"${cbsFile.cbsFilePath}\") 注入（无默认值占位符），读取方为 HnnxAntiMoneyListServiceImpl / BOPC010101MessageConverter / SyncPjgcs*JobServiceImpl，解析源为 bemp_home 下 *.properties（PropertiesAutoLoad）与 jres 配置中心；env.properties 河南农信段未登记该 key（代码引用与配置下发缺口），TM_BUSINESS_PARAMETER 参数表不参与 @Value 解析"],
            ["处置", "D:\\code\\home\\hnnxbank\\bemp_home\\env.properties 河南农信段后追加 cbsFile.cbsFilePath=D:/code/home/hnnxbank/ecif；@Value 在 Bean 创建时绑定、无运行时热加载触发点，重启 Served 后生效（重启后 PID 40256）；测试文件按 {cbsFilePath}/{日期}/ECF_M_CUST_SPECIAL_INFO_CBS_{日期}.txt 规范落位（GBK、0x03 分隔 19 列）"],
            ["验证", "重启后抽取任务正常读取文件；INCR/FULL 双模式回归通过（round3_05~12）；19 列解析回归通过"],
            ["状态/留档", "已关闭（修复+验证双重确认）；变更留档 DEF-02-03-变更处置.sql；生产/验收部署对照说明见《反洗钱缺陷修复-BUG-005配置增量说明-cbsFile.cbsFilePath.md》（该 key 必须纳入配置中心应用配置清单，一处配置四处生效：CBS 报文/机构关系/机构参数/名单抽取）"],
        ],
        col_widths=[2.2, 13.3]
    )

    add_heading(doc, "4.3 DEF-03（P2/数据类）：R22 统一社会信用代码名单记录未落库", level=2)
    add_table(doc,
        ["环节", "内容"],
        [
            ["发现", "round1 信用代码优先命中断言（P1-022）执行时名单库无 CERT_TYPE='22' 记录，round1 无法命中"],
            ["根因", "测试数据准备 §4.3 定义的 R22 记录未落库；且文档示例 CERT_NO=91410000MA3X5K9L9A 在票据主档无对应出票人，直接照搬无法构造 round1 命中链路"],
            ["处置", "按\"SOC_A 需与票据主档 DRWR_SOC_CODE 一致\"原则，改用 TB_BILL_INFO 主流值 91861230741970469R（15 笔票据在用）落库 R22（OBJ_ID=OBJ20260825000022、ID=MAX(ID)+1 取号、字段风格对齐既有 R1、DELETE+INSERT 幂等）；CERT_NO 登记值=91861230741970469R"],
            ["验证", "round1 信用代码优先命中通过（round3_13/17）；轮次互斥通过（round3_18）；OBJ_ID 唯一性自查通过"],
            ["状态/留档", "已关闭（修复+验证双重确认）；变更留档 DEF-02-03-变更处置.sql（含口径差异说明与自查 SQL）"],
        ],
        col_widths=[2.2, 13.3]
    )

    add_heading(doc, "4.4 DEF-04（P2/数据类）：测试文件 CUST_NO 超长致落库失败 ORA-12899", level=2)
    add_table(doc,
        ["环节", "内容"],
        [
            ["发现", "名单测试文件抽取落库时报 ORA-12899（列长度超限），记录未落库"],
            ["根因", "测试文件中个别记录 CUST_NO 值超过目标列定义长度（13 位），属测试数据制作错误，非代码缺陷（生产文件由 ODS 按接口规范供数，超长由单条失败 continue 分支兜底）"],
            ["处置", "修正测试文件 CUST_NO 至规范长度后重新落位"],
            ["验证", "修正后解析落库通过；19 列解析回归覆盖（round3 轨迹）；单条失败 continue 分支行为另有 R12 专项行验证"],
            ["状态/留档", "已关闭（修复+验证双重确认）；测试文件修正记录随 ecif 目录测试文件留档"],
        ],
        col_widths=[2.2, 13.3]
    )

    add_heading(doc, "4.5 DEF-05（P2/数据类）：R22 生效时间晚于业务日期致 round1 不命中", level=2)
    add_table(doc,
        ["环节", "内容"],
        [
            ["发现", "DEF-03 修复后 round1 仍不命中，名单查询直查可命中而校验链路不命中"],
            ["根因", "R22 落库时 CTRL_START_TIME=20260101 晚于系统业务日期 20250828，按时效判定（生效≤当前≤失效）该记录\"未生效\"不参与命中——校验行为正确，属数据时效基准设置错误"],
            ["处置", "round3_16 将 R22 生效时间修正为 20250101（早于业务日期）"],
            ["验证", "round1 信用代码优先命中通过（round3_17_p1022_x3_verify）；轮次互斥通过（round3_18_p026_y2_verify）；R9 未生效记录仍正确放行（时效判定双向对照成立）"],
            ["状态/留档", "已关闭（修复+验证双重确认）；变更留档 round3_16_def05_fix_r22.sql"],
        ],
        col_widths=[2.2, 13.3]
    )
    add_page_break(doc)

    # ===== 第五章 已知问题与遗留项 =====
    add_heading(doc, "第五章 已知问题与遗留项", level=1)
    add_para(doc, "以下事项不阻塞本期交付，均已在各轮测试中如实记录，并给出影响评估与后续计划。")

    add_heading(doc, "5.1 P3 已知问题", level=2)
    add_table(doc,
        ["编号", "描述", "影响评估", "后续计划"],
        [
            ["NOTE-3-01", "FULL 模式返回值 -1 口径：FULL 全量替换任务返回值 -1 的语义口径待统一（快照替换正确性已验证，仅返回值语义存在口径分歧）", "低：不影响名单快照替换正确性与业务功能", "开发侧确认返回值语义口径（成功/计数/异常编码），必要时随下一版本统一"],
            ["NOTE-3-02", "法人 userType=4 名单接口鉴权不放行 + 中互金名单菜单 30308 未绑法人角色：法人管理员经名单接口访问被鉴权拦截，菜单未绑定法人角色", "中：法人侧名单查询入口不可用（功能测试经具备权限账号完成，核心校验链路不受影响）", "执行菜单 30308 角色授权（授权草稿 SQL 已备：V202301.03.081_202608281000_T202608280001，需正式化评审后执行）+ 鉴权口径确认"],
            ["NOTE-3-03", "WebSocket 消息推送服务 9080 端口未启动：登录后 ws://127.0.0.1:9080 持续重连失败（环境噪声）", "低：消息推送功能未验证，主链路与本次反洗钱需求无交集", "需消息推送功能验证时启动对应推送服务；已列入 ignorable_patterns 白名单"],
            ["NOTE-3-04", "参数校验晚于黑名单拦截顺序确认项：法人证件类型后端格式校验（0BE229905099）晚于黑名单拦截执行，提交黑名单客户时非法证件值报错被掩盖", "低：仅影响错误提示顺序（先拦截后参数报错），不影响拦截正确性", "设计确认：如属预期（拦截优先）无需处理；如需参数校验前置由开发评估调整校验顺序"],
            ["NOTE-3-05", "E2E 并行时序快照说明：名单页实测 10 行比 8 条基线多 2 条（反洗钱文件抽取测试客户B/C），为测试与抽取任务并行时序下的快照现象/上轮残留，属数据残留非功能缺陷", "低：不影响基线比对结论（8 条基线客户全部在列）", "测试数据管理员确认来源；如属残留建议清理，避免影响后续基线比对"],
        ],
        col_widths=[2.0, 5.6, 3.9, 4.0]
    )

    add_heading(doc, "5.2 ⏸ 延后测试项", level=2)
    add_table(doc,
        ["延后项", "关联用例", "延后原因/所需条件"],
        [
            ["网银报文应答段（PICE070104/NES.002/NES.006 应答 retCode/retMsg 断言）", "P0-023/024/025、P1-027、P1-030 报文段", "需网银外围联调环境/报文模拟器；代码审查段（bean 顶替/取数链/校验时点/空集合短路）已全部通过，响应承载机制经 func_insertCustCorp retCode/retMsg 实测规范佐证"],
            ["eDisc 页面段（贴现申请批次提交命中/未命中）", "P0-020/021 步骤 1-6 页面操作", "TB_BILL_INFO 无出票人=名单客户的在途票据（0/21），业务表按原则不直插；需经完整票据流转（签发/承兑/背书）构造在途批次后回归；两轮评估均不可构造，代码审查段已通过"],
            ["运行时异常注入演练（停库/超时注入）", "P1-028 步骤 5/6", "测试环境不具备异常注入手段；降级两分支已由单测（入参注入）+代码审查完整覆盖"],
        ],
        col_widths=[4.8, 4.2, 6.5]
    )

    add_heading(doc, "5.3 待需求方确认项", level=2)
    add_table(doc,
        ["编号", "事项", "现状与测试处理", "待确认内容"],
        [
            ["Q-11", "拦截报错文案口径（公司金融部 docx 原文 vs 当前实现）", "测试假设：按原文案模板调整。四阶段面客文案 HNNX0BE320011 已按原文模板实现并逐字验证通过；市场侧对内文案 HNNX0BE320009/010 保持现状并按现状断言（P0-022 v3.4.11 修订）", "市场侧对内文案是否同步按原文模板调整（当前断言基于 Q-11 假设，确认后如调整需回归相应断言段）"],
            ["Q-12", "网银端四阶段校验是否本期范围（进展已补记 m1）", "④客户信息维护已纳入本期实现并经测试验证；网银三分支（①②③）代码审查段通过、报文段 ⏸ 延后报文联调", "网银三分支报文段的联调计划与验收口径确认"],
            ["RK-011", "审计留痕载体（日志级 vs DB 级）", "代码事实：无命中留痕/抽取批次留痕 DB 表，载体=LOGGER 日志（HnnxAntiMoneyValidateUtil L218/L308）；全部留痕断言已按日志载体验证通过", "RK-011 合规口径：日志级留痕是否满足；若需 DB 级请确认二期补建留痕表的交付计划"],
            ["—", "灰名单人工审批流二期交付计划", "一期以提示中断（HNNX0BE320010/320011）降级承载并已验证；Q-05/Q-08 确认口径为\"灰名单提示并需人工审批放行\"，一期存在功能缺口", "确认二期交付计划（范围/时间点），并确认一期\"提示中断\"降级承载是否可接受"],
        ],
        col_widths=[1.5, 3.8, 5.7, 4.5]
    )
    add_page_break(doc)

    # ===== 第六章 测试结论 =====
    add_heading(doc, "第六章 测试结论", level=1)

    add_heading(doc, "6.1 测试结论", level=2)
    add_para(doc, "结论：市场交易侧 + 四阶段客户维护侧功能符合 PRD v3.4.8 口径，具备交付条件。")
    add_bullet(doc, "核心拦截链路（round1 信用代码优先 + round2 名称兜底两轮匹配、时效判定、黑/灰差异化、开关启停、异常降级 fail-closed/fail-open、面客 HNNX0BE320011 与对内 HNNX0BE320009/010 文案分流）经一轮功能测试、二轮回归（双浏览器通道交叉）、端到端回归三个层级验证，结论稳定复现，全程 0 失败。")
    add_bullet(doc, "P0/P1 代码缺陷为 0；P2 缺陷 4 项（DEF-02 配置类、DEF-03/04/05 数据类）全部完成\"修复 + 验证\"双重确认闭环，修复过程与口径差异均已留档。")
    add_bullet(doc, "名单数据管理双模式（INCR 增量 upsert + FULL 全量快照）、19 列文件解析、口径过滤（DATA_SRC=18/VLD_ST=1）经缺陷修复后二轮回归 8/8 全部补验通过。")
    add_bullet(doc, "环境终态清洁：5 服务 UP、名单基线 8 条完整、双开关='1'、sync.mode=INCR、TM_CUST_CORP 零残留、ecif 目录 2 份标准测试文件。")

    add_heading(doc, "6.2 交付门禁核对", level=2)
    add_table(doc,
        ["门禁项", "核对结果"],
        [
            ["缺陷闭环", "通过——P0/P1 缺陷 0；P2 缺陷 4 项全部修复验证关闭，无未关闭缺陷"],
            ["测试完成度", "通过——53 条用例执行段全部落实（33 通过/4 部分通过/16 ⏸/0 失败）；⏸ 项均有明确归因与出路，其中 8 项数据前置类经二轮回归补验通过；P0 用例执行段 100% 落实"],
            ["需求确认闭环", "通过——业务项全部已确认；Q-11 测试假设经面客文案逐字验证（P0-015/018/P1-026）；Q-12 确认进展已补记（m1）；RK-011 与灰名单二期为需求方确认项，已列入遗留清单不阻塞"],
            ["用例评审闭环", "通过——严重 0/主要 0/轻微 2，m1/m2 已落实回写（见 2.3）"],
            ["交付物完整性", "通过——测试用例文档（53 条 v3.4.11）+ 测试报告（本文档）+ 需求文档（PRD v3.4.8）齐备；详细设计文档（反洗钱四阶段校验-详细设计说明书-v1.docx）已交付"],
        ],
        col_widths=[3.0, 12.5]
    )

    add_heading(doc, "6.3 建议与风险评估", level=2)
    add_table(doc,
        ["维度", "建议/评估"],
        [
            ["阶段推进", "建议进入下一阶段：网银报文段 5 条随网银联调环境就绪后补验；四阶段客户维护侧与市场交易侧可提交上线评估"],
            ["上线前检查单", "①cbsFile.cbsFilePath 纳入生产配置中心应用配置清单（BUG-005 增量说明，一处配置四处生效）；②中互金名单菜单 30308 法人角色授权 SQL 正式化执行（NOTE-3-02）；③名单基线初始化与双开关/sync.mode 参数登记纳入环境部署检查单；④名单数据部署流程将测试数据准备 §2.3/§4.3 纳入环境初始化清单"],
            ["风险评估（中）", "网银报文应答段未经实测：响应承载机制经产品 dealException 机制与 retCode/retMsg 实测结构佐证，但报文级断言（PICE070104/NES.002/NES.006）留待联调，存在联调期暴露问题的可能——建议网银联调排期时优先安排"],
            ["风险评估（低）", "5 项 P3 已知问题（5.1）均有影响评估与后续计划，不构成上线阻断；RK-011 若合规要求 DB 级留痕，需二期补建（功能风险已由日志载体覆盖）"],
            ["下轮测试计划", "①网银联调：报文段 5 条补验；②eDisc 在途批次数据构造后 P0-020/021 页面段回归；③Q-11/Q-12/RK-011 确认后对应断言段回归；④R-SPACE 名单侧含空格反向场景随真实 ODS 数据联调验证"],
        ],
        col_widths=[2.6, 12.9]
    )
    add_page_break(doc)

    # ===== 附录 证据文件索引 =====
    add_heading(doc, "附录 证据文件索引", level=1)
    add_para(doc, "除特别注明外，路径根为 d:\\code\\QJ\\BEMP5.0DEV\\。")
    add_table(doc,
        ["类别", "证据", "路径"],
        [
            ["需求文档", "PRD v3.4.8 / 待确认清单 v3.4.8（含 m1 补记）", "docs\\prd\\反洗钱校验功能\\HNNS-EB-STD-REQ-002-反洗钱功能优化-v3.md；HNNS-EB-STD-REQ-002-反洗钱功能优化-待确认-v3.md"],
            ["测试用例", "53 条用例（v3.4.11，含 m2 补注）+ 功能地图 + 测试数据准备", "aotutests-devtools\\testcases\\反洗钱名单校验\\反洗钱名单校验-P0用例.md；反洗钱名单校验-P1用例.md；反洗钱名单校验-功能地图与优先级矩阵.md；反洗钱名单校验-测试数据准备.md"],
            ["分侧报告", "四阶段功能测试报告（AML4-FT-20260829-01）/ 四阶段二轮回归报告（AML4-RT-20260829-02）", "aotutests-devtools\\testcases\\反洗钱名单校验\\四阶段功能测试报告.md；四阶段二轮回归测试报告.md"],
            ["全面测试报告", "一轮 AML-FT-20260830-01 / 二轮回归 AML3-RT-20260830-01（会话报告，关键数据以 SQL 轨迹与统计留档为证）", "本报告第三章为权威汇总；SQL 轨迹见下"],
            ["E2E 报告", "二轮端到端回归报告（含 Chrome DevTools MCP 降级记录）", "aotutests-devtools\\reports\\2026-08-30\\antimoney_e2e_regression_20260830_v1.md"],
            ["缺陷留档", "DEF-02/03 变更处置 SQL / DEF-05 修复 SQL / BUG-005 配置增量说明", "aotutests-devtools\\testcases\\反洗钱名单校验\\DEF-02-03-变更处置.sql；round3_16_def05_fix_r22.sql；docs\\prd\\反洗钱校验功能\\反洗钱缺陷修复-BUG-005配置增量说明-cbsFile.cbsFilePath.md"],
            ["二轮 SQL 轨迹", "round3_01~19（基线/探查/补验/修复/清理）", "aotutests-devtools\\testcases\\反洗钱名单校验\\round3_*.sql"],
            ["四阶段 SQL 轨迹", "round2_01~09（就绪/R1/V2a/V2b/V2c/评估/清理）", "aotutests-devtools\\testcases\\反洗钱名单校验\\round2_*.sql"],
            ["一轮截图", "名单查询/导出/导入、四阶段表单与拦截对照等", "aotutests-devtools\\testcases\\反洗钱名单校验\\results\\（fill_*.png / result_*.png / p1-011_*、p1-012_*、p1-017_* 等）"],
            ["E2E 截图", "名单页/拦截弹窗/修正提交/终态列表等时序快照", "aotutests-devtools\\shots\\hnnxbank-aml-round-e2e\\（step1~step3 系列）"],
            ["四阶段二轮截图", "R1 拦截弹窗/页面终态/Console 日志", "aotutests-devtools\\screenshots\\2026-08-29\\AML4-round2\\；aotutests-devtools\\console-logs\\2026-08-29\\AML4-round2-console.json"],
            ["后端日志", "命中拦截/跳过校验/抽取统计留痕", "D:\\code\\home\\hnnxbank\\bemp_home\\log\\bemp\\app.log（关键字：命中中互金关注名单/命中中互金灰名单/跳过校验/抽取）"],
            ["授权草稿 SQL", "中互金名单菜单 30308 角色授权草稿（NOTE-3-02 后续执行依据）", "docs\\prd\\反洗钱校验功能\\增量SQL\\V202301.03.081_202608281000_T202608280001_中互金名单菜单30308角色授权草稿.dml.sql"],
            ["E2E 降级脚本", "Playwright 降级通道执行脚本与结构化结果", "aotutests-devtools\\scripts\\aml_e2e_phase_*.py；aotutests-devtools\\logs\\aml_e2e_phase_a_result.json"],
            ["详细设计文档", "反洗钱四阶段校验-详细设计说明书-v1（已交付交付物）", "docs\\delivery\\反洗钱四阶段校验\\反洗钱四阶段校验-详细设计说明书-v1.docx"],
        ],
        col_widths=[2.2, 5.0, 8.3]
    )

    inject_update_fields(doc)
    os.makedirs(OUTPUT_DIR, exist_ok=True)
    doc.save(OUTPUT_FILE)

    # ===== 生成统计（字数/页数概况） =====
    total_chars = 0
    table_count = len(doc.tables)
    para_count = len(doc.paragraphs)
    for p in doc.paragraphs:
        total_chars += len(p.text)
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                total_chars += len(cell.text)
    print(f"OK: {OUTPUT_FILE}")
    print(f"STAT: paragraphs={para_count}, tables={table_count}, chars={total_chars}")
    return OUTPUT_FILE


if __name__ == "__main__":
    build_test_report_doc()
