"""验证模块2设计说明章节内容"""
from docx import Document
from docx.oxml.ns import qn

doc_path = r"D:\code\QJ\BEMP5.0DEV\.trae\skills\bemp-advanced-doc-generator\output\机构管理和管理员管理功能优化-详细设计文档-20260618.docx"
doc = Document(doc_path)

print("=" * 70)
print("模块2设计说明验证报告")
print("=" * 70)

# 找到所有H1标题
h1_list = []
for i, para in enumerate(doc.paragraphs):
    style_name = para.style.name if para.style else 'None'
    if 'Heading 1' in style_name or style_name == '1':
        h1_list.append((i, para.text.strip()))

print("\n所有H1标题:")
for idx, (i, text) in enumerate(h1_list):
    print(f"  [{idx}] 位置{i}: {text}")

# 找到模块2的位置
module2_pos = None
module2_idx = None
for idx, (i, text) in enumerate(h1_list):
    if '模块2' in text:
        module2_pos = i
        module2_idx = idx
        break

if module2_pos is None:
    print("\n❌ 未找到模块2!")
else:
    print(f"\n✅ 找到模块2: 位置{module2_pos}")
    
    # 找到下一个H1的位置
    next_h1_pos = None
    if module2_idx + 1 < len(h1_list):
        next_h1_pos = h1_list[module2_idx + 1][0]
        print(f"下一个H1位置: {next_h1_pos} ({h1_list[module2_idx + 1][1]})")
    
    # 统计模块2下的所有内容
    print(f"\n模块2下的详细内容:")
    
    h2_count = 0
    h3_count = 0
    para_count = 0
    table_count = 0
    h2_titles = []
    
    start_pos = module2_pos + 1
    end_pos = next_h1_pos if next_h1_pos else len(doc.paragraphs)
    
    for i in range(start_pos, min(end_pos, start_pos + 100)):
        para = doc.paragraphs[i]
        style_name = para.style.name if para.style else 'None'
        text = para.text.strip()
        
        if 'Heading 2' in style_name or style_name == '2':
            h2_count += 1
            h2_titles.append(text)
            print(f"  [H2] 位置{i}: {text[:80] if text else '(空)'}")
        elif 'Heading 3' in style_name or style_name == '3':
            h3_count += 1
            print(f"  [H3] 位置{i}: {text[:80] if text else '(空)'}")
        elif text:
            para_count += 1
            if para_count <= 15:
                print(f"  [段落] 位置{i}: {text[:80]}")
    
    # 统计表格
    for table in doc.tables:
        table_count += 1
    
    print(f"\n统计汇总:")
    print(f"  H2数量: {h2_count}")
    print(f"  H3数量: {h3_count}")
    print(f"  段落数量: {para_count}")
    print(f"  表格数量: {table_count}")
    
    print(f"\nH2标题列表:")
    for idx, title in enumerate(h2_titles, 1):
        print(f"  {idx}. {title}")
    
    if h2_count == 0:
        print("\n❌ 模块2下没有H2标题!")
    elif h2_count < 5:
        print(f"\n⚠️ 模块2下H2标题较少（{h2_count}个），可能不完整")
    else:
        print(f"\n✅ 模块2下有 {h2_count} 个H2标题，结构完整")
    
    if para_count == 0:
        print("❌ 模块2下没有段落内容!")
    else:
        print(f"✅ 模块2下有 {para_count} 个段落")
