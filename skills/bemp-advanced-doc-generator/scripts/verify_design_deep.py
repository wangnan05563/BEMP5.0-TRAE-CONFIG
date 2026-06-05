"""深入验证详细设计文档内容"""
from docx import Document
import sys

doc_path = sys.argv[1]
doc = Document(doc_path)

# 检查封面标题
print('=== 封面检查 ===')
for i, p in enumerate(doc.paragraphs[:15]):
    text = p.text.strip()
    if text:
        print(f'  [{i}] {p.style.name}: {text[:80]}')

# 检查各标题下的内容
print('\n=== 章节内容检查 ===')
for i, p in enumerate(doc.paragraphs):
    if p.style and p.style.name.startswith('Heading'):
        # 检查该标题后面是否有内容
        has_content = False
        content_preview = ''
        for j in range(i+1, min(i+5, len(doc.paragraphs))):
            next_p = doc.paragraphs[j]
            if next_p.style and next_p.style.name.startswith('Heading'):
                break
            if next_p.text.strip():
                has_content = True
                content_preview = next_p.text.strip()[:60]
                break
        status = '有内容' if has_content else '空'
        print(f'  {p.style.name}: {p.text.strip()[:50]} [{status}] {content_preview}')

# 检查蓝色文本是否为超链接
print('\n=== 蓝色文本分析 ===')
blue_hyperlink = 0
blue_plain = 0
for p in doc.paragraphs:
    for run in p.runs:
        if run.font.color and run.font.color.rgb:
            color = str(run.font.color.rgb)
            if color.startswith('00') and color not in ['000000']:
                # 检查是否在超链接中
                parent = run._element.getparent()
                tag = parent.tag.split('}')[-1] if '}' in parent.tag else parent.tag
                if tag == 'hyperlink':
                    blue_hyperlink += 1
                else:
                    blue_plain += 1
                    if blue_plain <= 5:
                        print(f'  非超链接蓝色: {run.text[:40]}')

print(f'超链接蓝色: {blue_hyperlink}')
print(f'非超链接蓝色: {blue_plain}')

# 检查模板是否被正确使用（对比模板原始结构）
print('\n=== 文档结构对比 ===')
print(f'总段落数: {len(doc.paragraphs)}')
print(f'总表格数: {len(doc.tables)}')
h1_count = sum(1 for p in doc.paragraphs if p.style and p.style.name == 'Heading 1')
h2_count = sum(1 for p in doc.paragraphs if p.style and p.style.name == 'Heading 2')
print(f'Heading 1 数: {h1_count}')
print(f'Heading 2 数: {h2_count}')
