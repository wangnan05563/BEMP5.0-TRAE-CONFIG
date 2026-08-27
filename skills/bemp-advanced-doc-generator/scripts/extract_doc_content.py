from pathlib import Path
SKILL_ROOT = Path(__file__).resolve().parent.parent
PROJECT_ROOT = SKILL_ROOT.parent.parent.parent

from pathlib import Path
"""提取详细设计文档内容用于分析"""
from docx import Document
import json

doc_path = str(SKILL_ROOT / "output" / "机构管理和管理员管理功能优化-详细设计文档-20260617.docx")
doc = Document(doc_path)

result = {
    "headings": [],
    "paragraphs": [],
    "tables": []
}

# 提取标题
for para in doc.paragraphs:
    if para.style.name.startswith('Heading'):
        level = int(para.style.name.replace('Heading ', ''))
        result["headings"].append({
            "level": level,
            "text": para.text.strip()
        })
    elif para.text.strip():
        result["paragraphs"].append(para.text.strip()[:200])  # 限制长度

# 提取表格
for i, table in enumerate(doc.tables):
    table_data = []
    for row in table.rows:
        row_data = [cell.text.strip()[:50] for cell in row.cells]
        table_data.append(row_data)
    result["tables"].append({
        "index": i,
        "rows": len(table.rows),
        "cols": len(table.rows[0].cells) if table.rows else 0,
        "preview": table_data[:3]  # 前3行预览
    })

# 保存到文件
with open(str(SKILL_ROOT / "output" / "doc_analysis.json"), 'w', encoding='utf-8') as f:
    json.dump(result, f, ensure_ascii=False, indent=2)

print(f"提取完成：{len(result['headings'])} 个标题，{len(result['paragraphs'])} 个段落，{len(result['tables'])} 个表格")
print("\n=== 标题结构 ===")
for h in result['headings'][:30]:  # 显示前30个标题
    indent = "  " * (h['level'] - 1)
    print(f"{indent}H{h['level']}: {h['text']}")
