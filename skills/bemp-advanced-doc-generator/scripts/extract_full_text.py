from pathlib import Path
SKILL_ROOT = Path(__file__).resolve().parent.parent
PROJECT_ROOT = SKILL_ROOT.parent.parent.parent

from pathlib import Path
"""提取详细设计文档完整文本内容"""
from docx import Document

doc_path = str(SKILL_ROOT / "output" / "机构管理和管理员管理功能优化-详细设计文档-20260617.docx")
doc = Document(doc_path)

output = []

for para in doc.paragraphs:
    style = para.style.name
    text = para.text.strip()
    if not text:
        continue
    if style.startswith('Heading'):
        level = int(style.replace('Heading ', ''))
        indent = "#" * level
        output.append(f"\n{indent} {text}\n")
    else:
        output.append(text)

# 表格内容
for i, table in enumerate(doc.tables):
    output.append(f"\n--- 表格{i+1} ---")
    for row_idx, row in enumerate(table.rows):
        cells = [cell.text.strip().replace('\n', ' ')[:80] for cell in row.cells]
        output.append(f"  行{row_idx}: {' | '.join(cells)}")

content = "\n".join(output)
with open(str(SKILL_ROOT / "output" / "doc_full_text.md"), 'w', encoding='utf-8') as f:
    f.write(content)

print(f"总字符数: {len(content)}")
print(f"段落数: {len(doc.paragraphs)}")
print(f"表格数: {len(doc.tables)}")
