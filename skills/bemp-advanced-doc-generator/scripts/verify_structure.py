from pathlib import Path
SKILL_ROOT = Path(__file__).resolve().parent.parent
PROJECT_ROOT = SKILL_ROOT.parent.parent.parent

from pathlib import Path
"""详细验证文档结构和内容"""
from docx import Document
from docx.oxml.ns import qn

doc_path = str(SKILL_ROOT / "output" / "机构管理和管理员管理功能优化-详细设计文档-20260617.docx")
doc = Document(doc_path)

print("=" * 80)
print("文档结构详细分析")
print("=" * 80)

# 统计各级标题
h1_list = []
h2_list = []
current_h1 = None

for para in doc.paragraphs:
    style_name = para.style.name if para.style else ''
    text = para.text.strip()
    
    if not text:
        continue
    
    if 'Heading 1' in style_name or style_name == '1':
        h1_list.append(text)
        current_h1 = text
        print(f"\n【H1】{text}")
    elif 'Heading 2' in style_name or style_name == '2':
        h2_list.append((current_h1, text))
        print(f"  【H2】{text}")

print("\n" + "=" * 80)
print(f"H1 章节总数: {len(h1_list)}")
print(f"H2 章节总数: {len(h2_list)}")
print("=" * 80)

# 检查关键章节
print("\n【关键章节检查】")
key_sections = [
    "组件内部的模块列表及说明",
    "模块1设计说明",
    "模块2设计说明",
    "模块3设计说明",
    "模块4设计说明"
]

for key in key_sections:
    found = any(key in h1 for h1 in h1_list)
    print(f"  {key}: {'✓ 找到' if found else '✗ 未找到'}")

# 检查表格分布
print("\n【表格分布统计】")
table_count = 0
for i, para in enumerate(doc.paragraphs):
    if para._element.tag.endswith('tbl'):
        table_count += 1
        # 找到该表格前的最近标题
        prev_heading = None
        for j in range(i - 1, -1, -1):
            prev_para = doc.paragraphs[j]
            if prev_para.style and 'Heading' in prev_para.style.name:
                prev_heading = prev_para.text.strip()
                break
        print(f"  表格 {table_count}: 位于标题 '{prev_heading}' 之后")

print(f"\n总表格数: {table_count}")
