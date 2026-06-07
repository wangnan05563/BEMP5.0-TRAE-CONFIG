"""
文档通用格式化工具：段落缩进、表格样式、空章节占位、TOC域注入等。
所有"先A后B"的串行操作使用临时文件+重命名，避免 zipfile 追加。
"""
import os
import re
import sys
import shutil
import zipfile

from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# 从 doc_utils 导入蓝色检测和超链接检测工具
from doc_utils import is_blue_color, is_hyperlink_run


HEADING_STYLE_VALS = {
    '1', '2', '3', '4', '5', '6',
    'Heading1', 'Heading2', 'Heading3', 'Heading4', 'Heading5', 'Heading6',
    'heading1', 'heading2', 'heading3', 'heading4', 'heading5', 'heading6',
    'heading 1', 'heading 2', 'heading 3', 'heading 4', 'heading 5', 'heading 6',
}


# 2026-06-07 新增：模板 Heading 样式 styleId 解析缓存
# 解决"模板中 Heading 样式是数字 styleId（1/2/3）而非 'Heading 1' 字符串"导致
# 设置 pStyle='Heading1' 后 Word 找不到样式、回退到 Normal 的问题。
# 通用化设计：从 doc.styles 动态扫描，按名称匹配映射到数字 styleId。
_HEADING_STYLE_ID_CACHE = {}


def _resolve_heading_style_id(doc, level):
    """根据 doc.styles 解析 HeadingN 对应的实际 styleId（通用化）

    2026-06-07 新增。模板中 Heading 样式可能使用：
      - 数字 styleId：'1' / '2' / '3'（中文 WPS 常见）
      - 标准 styleId：'Heading1' / 'Heading 1'（Word 默认）
    通过扫描 doc.styles 中所有 style，按 name 包含 "heading N" 或 "标题 N"
    匹配到对应 styleId，并缓存以避免重复 IO。

    Args:
        doc: Document 对象（python-docx）
        level: 1~6，Heading 级别

    Returns:
        str: styleId（例如 '1' / '2' / 'Heading 1'）
    """
    if doc is None or level not in (1, 2, 3, 4, 5, 6):
        return f'Heading {level}'
    cache_key = (id(doc), level)
    if cache_key in _HEADING_STYLE_ID_CACHE:
        return _HEADING_STYLE_ID_CACHE[cache_key]
    candidates = []
    try:
        for style in doc.styles:
            sid = getattr(style, 'style_id', None) or ''
            name = (getattr(style, 'name', None) or '').strip()
            if not name:
                continue
            # 匹配 heading N / 标题 N（不区分大小写）
            nlow = name.lower().replace('  ', ' ')
            if nlow == f'heading {level}' or nlow == f'heading{level}' or nlow == f'标题 {level}':
                candidates.append((sid, name))
    except Exception:
        candidates = []
    # 优先级：完全匹配 'heading N' > 'HeadingN' > 标题 N
    resolved = None
    for sid, name in candidates:
        if name.lower() == f'heading {level}':
            resolved = sid
            break
    if not resolved:
        for sid, name in candidates:
            if name.lower() == f'heading{level}':
                resolved = sid
                break
    if not resolved and candidates:
        resolved = candidates[0][0]
    # 兜底：未找到则用数字 styleId（兼容中文 WPS 模板）
    if not resolved:
        resolved = str(level)
    _HEADING_STYLE_ID_CACHE[cache_key] = resolved
    return resolved


def _is_heading_paragraph(paragraph):
    """判断段落是否为标题样式"""
    if paragraph is None or paragraph.style is None:
        return False
    name = paragraph.style.name or ''
    return name.startswith('Heading')


def apply_paragraph_indent(paragraph, chars=2):
    """为首行缩进设置段落格式（单位：字符数，对应 w:firstLineChars / w:firstLine）

    chars=2 表示首行缩进2字符，符合中文文档常规排版。
    """
    if paragraph is None:
        return
    p_elem = paragraph._element
    pPr = p_elem.find(qn('w:pPr'))
    if pPr is None:
        pPr = OxmlElement('w:pPr')
        p_elem.insert(0, pPr)
    # 清理旧的缩进设置
    existing_ind = pPr.find(qn('w:ind'))
    if existing_ind is not None:
        pPr.remove(existing_ind)
    ind = OxmlElement('w:ind')
    # firstLineChars 以百分之一字符为单位，200 表示 2 字符
    ind.set(qn('w:firstLineChars'), str(int(chars * 100)))
    # 兜底：也设置 firstLine 磅值（按 10.5pt 字号换算近似）
    ind.set(qn('w:firstLine'), str(int(chars * 210)))
    pPr.append(ind)


def apply_body_indent_to_doc(doc, chars=2, skip_headings=True, skip_tables=True):
    """遍历文档中所有正文段落，对非标题段落应用首行缩进

    跳过：标题（Heading* 样式）、表格内段落。
    """
    if doc is None:
        return 0
    count = 0
    for p in doc.paragraphs:
        if _is_heading_paragraph(p):
            if skip_headings:
                continue
        apply_paragraph_indent(p, chars=chars)
        count += 1
    if not skip_tables:
        for tbl in doc.tables:
            for row in tbl.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        apply_paragraph_indent(p, chars=chars)
                        count += 1
    return count


def apply_table_style(table, style_name='Table Grid', font_size=None, header_bold=True):
    """统一表格样式：边框 + 字体 + 表头加粗

    style_name: 内置表格样式名（默认 Table Grid）
    font_size: 字号（docx.shared.Pt），默认 10.5pt
    header_bold: 表头是否加粗
    """
    if table is None:
        return
    if font_size is None:
        font_size = Pt(10.5)
    tbl = table._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    # 设置表格样式
    existing_style = tblPr.find(qn('w:tblStyle'))
    if existing_style is None:
        tbl_style = OxmlElement('w:tblStyle')
        tbl_style.set(qn('w:val'), style_name)
        tblPr.insert(0, tbl_style)
    else:
        existing_style.set(qn('w:val'), style_name)
    # 设置边框（防止模板未定义时缺失）
    existing_borders = tblPr.find(qn('w:tblBorders'))
    if existing_borders is not None:
        tblPr.remove(existing_borders)
    tblBorders = OxmlElement('w:tblBorders')
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement('w:' + border_name)
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), '000000')
        tblBorders.append(border)
    tblPr.append(tblBorders)
    # 单元格字体+加粗
    for ri, row in enumerate(table.rows):
        for cell in row.cells:
            for p in cell.paragraphs:
                for run in p.runs:
                    if run.font.size is None or run.font.size != font_size:
                        run.font.size = font_size
                    if header_bold and ri == 0:
                        run.font.bold = True
            # 表头底色
            if ri == 0:
                tcPr = cell._tc.find(qn('w:tcPr'))
                if tcPr is None:
                    tcPr = OxmlElement('w:tcPr')
                    cell._tc.insert(0, tcPr)
                existing_shd = tcPr.find(qn('w:shd'))
                if existing_shd is not None:
                    tcPr.remove(existing_shd)
                shd = OxmlElement('w:shd')
                shd.set(qn('w:val'), 'clear')
                shd.set(qn('w:color'), 'auto')
                shd.set(qn('w:fill'), 'D9D9D9')
                tcPr.append(shd)


def has_toc_field(doc):
    """检查文档正文是否已包含 TOC 域（fldChar begin 或 instrText 'TOC'）"""
    if doc is None:
        return False
    for el in doc.element.iter():
        tag = el.tag
        if tag.endswith('}instrText') and el.text and 'TOC' in el.text.upper():
            return True
    return False


def insert_toc_field_after(heading_paragraph, title='目录'):
    """在指定标题段落下插入 Word TOC 域（含 begin/separate/end 三段）"""
    if heading_paragraph is None:
        return None
    parent = heading_paragraph._element.getparent()
    if parent is None:
        return None
    # 创建空段落承载 TOC
    p = OxmlElement('w:p')
    # begin
    r1 = OxmlElement('w:r')
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    r1.append(fldChar1)
    p.append(r1)
    # instrText
    r2 = OxmlElement('w:r')
    instr = OxmlElement('w:instrText')
    instr.set(qn('xml:space'), 'preserve')
    instr.text = ' TOC \\o "1-3" \\h \\z \\u '
    r2.append(instr)
    p.append(r2)
    # separate
    r3 = OxmlElement('w:r')
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')
    r3.append(fldChar2)
    p.append(r3)
    # 占位文字（Word 更新前可见）
    r4 = OxmlElement('w:r')
    t4 = OxmlElement('w:t')
    t4.set(qn('xml:space'), 'preserve')
    t4.text = '请在 Word 中按 F9 或右键更新域以生成目录'
    r4.append(t4)
    p.append(r4)
    # end
    r5 = OxmlElement('w:r')
    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'end')
    r5.append(fldChar3)
    p.append(r5)
    heading_paragraph._element.addnext(p)
    return p


def inject_update_fields(output_path):
    """在 docx 的 settings.xml 中注入 updateFields=true（临时文件+重命名模式）

    使用"读全部 → 改 → 写全部"的 ZIP 重写方式，避免 zipfile 追加重复条目。
    先写到临时文件，再原子重命名。
    """
    if not output_path or not os.path.exists(output_path):
        return
    try:
        settings_xml = None
        with zipfile.ZipFile(output_path, 'r') as zf_in:
            if 'word/settings.xml' not in zf_in.namelist():
                return
            settings_xml = zf_in.read('word/settings.xml').decode('utf-8')
        if settings_xml is None or 'updateFields' in settings_xml:
            return
        insert_tag = '<w:updateFields w:val="true"/>'
        settings_xml = settings_xml.replace('</w:settings>', insert_tag + '</w:settings>')
        # 写临时文件后重命名
        tmp_path = output_path + '.tmp'
        with zipfile.ZipFile(output_path, 'r') as zf_in:
            entries = []
            for item in zf_in.infolist():
                data = settings_xml.encode('utf-8') if item.filename == 'word/settings.xml' else zf_in.read(item.filename)
                entries.append((item, data))
        with zipfile.ZipFile(tmp_path, 'w', zipfile.ZIP_DEFLATED) as zf_out:
            for item, data in entries:
                zf_out.writestr(item, data)
        shutil.move(tmp_path, output_path)
    except Exception as e:
        print(f'[WARN] 注入 updateFields 失败: {e}', file=sys.stderr)


def clear_all_blue_runs(paragraph):
    """彻底清除段落中所有蓝色 run 的内容并变黑，仅对非超链接 run 处理

    只清空 r.text = '' 但不修改颜色属性会导致24处蓝色run残留（文本为空但颜色仍为蓝色），
    因此在清空文本后同时将颜色设为黑色。
    """
    if paragraph is None:
        return
    from doc_utils import _is_hyperlink_style_run
    for r in list(paragraph.runs):
        if _is_hyperlink_style_run(r):
            continue
        try:
            color = r.font.color.rgb if r.font.color else None
        except Exception:
            color = None
        if color and is_blue_color(color):
            r.text = ''
            try:
                r.font.color.rgb = RGBColor(0, 0, 0)
            except Exception:
                pass


def enforce_business_alignment(paragraph, business_keywords):
    """检查段落是否与业务模块名相关；不相关时返回 True 表示可删除

    business_keywords: 业务模块关键词集合（如 {"账户", "票据", "贴现"}）
    """
    if paragraph is None:
        return False
    text = (paragraph.text or '').strip()
    if not text:
        return False
    for kw in business_keywords:
        if kw in text:
            return False
    # 命中通用术语关键词也不算业务相关（保留）
    generic_terms = {'设计目标', '设计原则', '设计约束', '非功能性', '可用性', '可靠性', '安全性', '可维护性', '可扩展性'}
    for g in generic_terms:
        if g in text:
            return False
    return True


def remove_design_constraint_irrelevant(paragraphs, business_keywords):
    """在"设计约束"标题下，删除与业务无关的段落

    找到"设计约束"标题，遍历其下到下一标题间的段落，
    若段落内容与业务关键词无关，则标记删除。
    """
    if not paragraphs or not business_keywords:
        return []
    target_idx = None
    for i, p in enumerate(paragraphs):
        text = (p.text or '').strip()
        if text == '设计约束' or (text and '设计约束' in text and len(text) <= 10):
            target_idx = i
            break
    if target_idx is None:
        return []
    next_idx = None
    for j in range(target_idx + 1, len(paragraphs)):
        if _is_heading_paragraph(paragraphs[j]):
            next_idx = j
            break
    end_idx = next_idx if next_idx is not None else len(paragraphs)
    to_remove = []
    for j in range(target_idx + 1, end_idx):
        p = paragraphs[j]
        if _is_heading_paragraph(p):
            break
        if enforce_business_alignment(p, business_keywords):
            to_remove.append(p)
    for p in to_remove:
        try:
            parent = p._element.getparent()
            if parent is not None:
                parent.remove(p._element)
        except Exception:
            pass
    return to_remove


def format_ui_paragraph(paragraph, max_chars=40):
    """对界面段落做软换行处理：每行不超过 max_chars 字符

    改造方式：在 run 内的 w:t 节点之后插入 w:br 元素，达到换行效果。
    避免改变段落数（仍属同一段落）。
    """
    if paragraph is None:
        return
    text = paragraph.text or ''
    if not text:
        return
    if len(text) <= max_chars:
        return
    # 清空所有 run
    for r in paragraph.runs:
        r.text = ''
    if not paragraph.runs:
        paragraph.add_run('')
    first_run = paragraph.runs[0]
    # 按 max_chars 切分
    chunks = []
    i = 0
    while i < len(text):
        chunks.append(text[i:i + max_chars])
        i += max_chars
    first_run.text = chunks[0]
    insert_after = first_run._element
    for chunk in chunks[1:]:
        br = OxmlElement('w:br')
        insert_after.addnext(br)
        new_r = OxmlElement('w:r')
        rPr = first_run._element.find(qn('w:rPr'))
        if rPr is not None:
            from copy import deepcopy
            new_r.append(deepcopy(rPr))
        t = OxmlElement('w:t')
        t.set(qn('xml:space'), 'preserve')
        t.text = chunk
        new_r.append(t)
        br.addnext(new_r)
        insert_after = new_r


# ═══════════════════════════════════════════════════════════════
# 扩展 API（按 outline-design-generator.py 需求新增）
# 目的：与 outline-design-generator.py 中 doc_formatter.* 调用兼容
# ═══════════════════════════════════════════════════════════════

def insert_caption_paragraph(doc, heading_para, caption_text, style_name='Caption',
                              alignment=None):
    """在标题段落后插入图表说明段落

    设计原则：
    - 通用工具函数，不硬编码任何业务
    - style_name 不存在时降级为 Normal
    - alignment 缺省时使用居中

    Args:
        doc: docx Document 对象
        heading_para: 标题段落（Paragraph）
        caption_text: 说明文本
        style_name: 段落样式名（默认 'Caption'）
        alignment: 对齐方式（None → 居中）

    Returns:
        新插入的段落对象
    """
    if heading_para is None:
        return None
    new_p = doc.add_paragraph()
    try:
        new_p.style = doc.styles[style_name]
    except KeyError:
        try:
            new_p.style = doc.styles['Normal']
        except KeyError:
            pass
    if alignment is None:
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        alignment = WD_ALIGN_PARAGRAPH.CENTER
    new_p.alignment = alignment
    run = new_p.add_run(caption_text)
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    heading_para._element.addnext(new_p._element)
    return new_p


def force_insert_toc(doc, after_paragraph=None, levels='1-3'):
    """强制插入/刷新动态目录域

    - 若文档已包含 TOC 域：跳过（避免重复插入）
    - 若 after_paragraph 指定：在其后插入
    - 否则：插入到文档正文开头

    Args:
        doc: docx Document 对象
        after_paragraph: 在该段落后插入；None 则插入到正文开头
        levels: TOC 层级范围（如 '1-3'）

    Returns:
        bool: 是否成功插入（False 表示已存在 TOC）
    """
    if has_toc_field(doc):
        return False

    # 构造 TOC 域段落
    p_elem = OxmlElement('w:p')

    r_begin = OxmlElement('w:r')
    fld_begin = OxmlElement('w:fldChar')
    fld_begin.set(qn('w:fldCharType'), 'begin')
    r_begin.append(fld_begin)
    p_elem.append(r_begin)

    r_instr = OxmlElement('w:r')
    instr = OxmlElement('w:instrText')
    instr.set(qn('xml:space'), 'preserve')
    instr.text = f' TOC \\o "{levels}" \\h \\z \\u '
    r_instr.append(instr)
    p_elem.append(r_instr)

    r_sep = OxmlElement('w:r')
    fld_sep = OxmlElement('w:fldChar')
    fld_sep.set(qn('w:fldCharType'), 'separate')
    r_sep.append(fld_sep)
    p_elem.append(r_sep)

    r_placeholder = OxmlElement('w:r')
    t_placeholder = OxmlElement('w:t')
    t_placeholder.set(qn('xml:space'), 'preserve')
    t_placeholder.text = '请在 Word 中按 F9 或右键更新域以生成目录'
    r_placeholder.append(t_placeholder)
    p_elem.append(r_placeholder)

    r_end = OxmlElement('w:r')
    fld_end = OxmlElement('w:fldChar')
    fld_end.set(qn('w:fldCharType'), 'end')
    r_end.append(fld_end)
    p_elem.append(r_end)

    if after_paragraph is not None:
        after_paragraph._element.addnext(p_elem)
    else:
        body = doc.element.body
        first_child = body[0] if len(body) > 0 else None
        if first_child is not None:
            first_child.addprevious(p_elem)
        else:
            body.append(p_elem)
    return True


# ═══════════════════════════════════════════════════════════════
# 配置化规则层（YAML 驱动）
# ----------------------------------------------------------------
# 设计目标：
#   1. 把 doc_formatter 中硬编码的关键词/字号/填充规则全部抽到
#      doc_rules.yaml，运行时按需加载。
#   2. YAML 缺失或解析失败 → 使用 _DEFAULT_RULES 内置默认值。
#   3. 现有 11 个公共函数签名保持不变（向后兼容），新增的 rules
#      参数全部为可选；调用方无需感知配置层。
#   4. 预留 _register_generator() hook，让 outline-design-generator.py
#      / design-generator.py 注入自定义 generator（按 type 路由）。
# ═══════════════════════════════════════════════════════════════

import copy as _copy
import threading as _threading

try:
    import yaml as _yaml  # PyYAML；缺失时降级为内置 dict
    _YAML_AVAILABLE = True
except ImportError:
    _yaml = None
    _YAML_AVAILABLE = False


# 内置默认规则：当 doc_rules.yaml 缺失/解析失败时使用
_DEFAULT_RULES = {
    'toc': {
        'levels': '1-3',
        'title_size_pt': 14,
        'indent_chars': 0,
        'placeholder_text': '请在 Word 中按 F9 或右键更新域以生成目录',
    },
    'empty_chapter_keywords': [
        '设计目标', '设计原则', '输入项', '输出项', '代码示例',
        '性能优化', '附录', '错误码', '模块复用分析',
        '组件内部的模块列表及说明', 'API 接口清单', '备注',
        '模块划分', '设计约束', '外部接口',
        '术语和缩写', '术语定义', '参考资料', '参考文档',
    ],
    'fill_chapters': {
        '概要设计': [
            {'keywords': ['设计目标', '设计原则'], 'type': 'text',
             'generator': '_build_design_goal_text'},
            {'keywords': ['术语和缩写', '术语定义'], 'type': 'table',
             'generator': '_build_glossary_data'},
            {'keywords': ['参考资料', '参考文献', '参考文档'], 'type': 'table',
             'generator': '_build_references_table'},
            {'keywords': ['类图', '顺序图', '活动图', '状态图'],
             'type': 'uml_placeholder', 'generator': '_build_uml_placeholder'},
        ],
        '详细设计': [
            {'keywords': ['概述', '项目概述', '系统概述'], 'type': 'text',
             'generator': '_build_overview_text'},
            {'keywords': ['系统组件', '组件结构图'], 'type': 'component_renumber',
             'generator': '_renumber_components'},
            {'keywords': ['关键技术'], 'type': 'component_specific',
             'generator': '_build_component_specific_tech'},
        ],
    },
    'table_style': {
        'font_name': '宋体',
        'font_size_pt': 10.5,
        'header_bold': True,
        'border': 'Table Grid',
        'align': 'left',
        'header_fill': 'D9D9D9',
        'border_color': '000000',
        'border_size_eighths': 4,
    },
    'paragraph': {
        'first_line_indent_chars': 2,
        'font_name': '宋体',
        'font_size_pt': 12,
        'line_spacing': 1.5,
        'apply_to_body': True,
        'skip_headings': True,
        'skip_tables': True,
    },
    'chapter_content_correction': [
        {'source_keywords': ['范围说明'],
         'target_keywords': ['设计目标'],
         'type': 'move_table'},
    ],
    'heading_numbering': {
        'skip_h1': False,
        'separator': '.',
        'start_at': 1,
    },
    'design_constraint': {
        'enable': True,
        'generic_terms': [
            '设计目标', '设计原则', '设计约束', '非功能性',
            '可用性', '可靠性', '安全性', '可维护性', '可扩展性',
        ],
    },
    'ui_paragraph': {
        'enable': True,
        'max_chars_per_line': 40,
    },
    'placeholder': {
        'empty_chapter': '不涉及',
        'ui_soft_break_keep': '（保持换行展示）',
    },
    'blue_runs': {
        'enable': True,
        'preserve_hyperlinks': True,
        'preserve_underline_runs': True,
    },
    # 2026-06-07 v7.1：UML 图表配置（硬编码外置）
    'uml': {
        'enable': True,
        'keywords': ['类图', '顺序图', '活动图', '状态图', '组件图',
                   '时序图', '协作图', '部署图'],
        'required_headings': ['类图', '顺序图', '活动图'],
        'diagram_type_priority': ['类图', '顺序图', '活动图', '状态图'],
        'file_matchers': {
            '类图': ['class-diagram.png', 'uml-类图.png', 'uml-类.png'],
            '顺序图': ['sequence-*.png', 'uml-顺序图.png', 'uml-顺序.png'],
            '活动图': ['activity-*.png', 'uml-活动图.png', 'uml-活动.png'],
            '状态图': ['state-*.png', 'uml-状态图.png', 'uml-状态.png'],
            '组件图': ['component-*.png', 'uml-组件图.png', 'uml-组件.png'],
        },
        'placeholder_cleaners': [
            '类图待补充', '顺序图待补充', '活动图待补充',
            '状态图待补充', '组件图待补充',
            '建议使用工具', '请在详细设计阶段补充', '建议在详细设计阶段',
        ],
        'fallback_placeholder_template': '【{diagram_name}待补充】\n建议在详细设计阶段使用 Visio / draw.io / Mermaid 等工具补充专业的{diagram_name}。',
        'min_diagram_size_kb': 10,
        'fallback_generation': True,
        'fallback_class_diagram': {
            'classes': [
                {'name': 'Controller', 'stereotype': 'control'},
                {'name': 'Service', 'stereotype': 'service'},
                {'name': 'Repository', 'stereotype': 'dataAccess'},
                {'name': 'Entity', 'stereotype': 'entity'},
                {'name': 'DTO', 'stereotype': 'data'},
            ],
            'relations': [
                {'from': 'Controller', 'to': 'Service', 'label': '调用'},
                {'from': 'Service', 'to': 'Repository', 'label': '使用'},
                {'from': 'Service', 'to': 'DTO', 'label': '转换'},
                {'from': 'Repository', 'to': 'Entity', 'label': '操作'},
            ],
            'title': '业务实体类图',
        },
        'fallback_sequence_diagram': {
            'actors': ['用户', '前端', '后端服务', '认证中心', '数据库'],
            'messages': [
                {'from': '用户', 'to': '前端', 'label': '输入凭证'},
                {'from': '前端', 'to': '后端服务', 'label': '登录请求'},
                {'from': '后端服务', 'to': '认证中心', 'label': '身份验证'},
                {'from': '认证中心', 'to': '数据库', 'label': '查询用户'},
                {'from': '数据库', 'to': '认证中心', 'label': '返回用户信息'},
                {'from': '认证中心', 'to': '后端服务', 'label': '返回Token'},
                {'from': '后端服务', 'to': '前端', 'label': '返回Token'},
                {'from': '前端', 'to': '用户', 'label': '跳转主页'},
            ],
            'title': '登录鉴权顺序图',
        },
        'fallback_activity_diagram': {
            'nodes': [
                {'id': 'start', 'shape': 'circle', 'label': '开始'},
                {'id': 'req', 'shape': 'box', 'label': '接收请求'},
                {'id': 'validate', 'shape': 'diamond', 'label': '参数校验'},
                {'id': 'business', 'shape': 'box', 'label': '业务处理'},
                {'id': 'persist', 'shape': 'box', 'label': '持久化'},
                {'id': 'end', 'shape': 'circle', 'label': '结束'},
            ],
            'edges': [
                {'from': 'start', 'to': 'req'},
                {'from': 'req', 'to': 'validate'},
                {'from': 'validate', 'to': 'business', 'label': '通过'},
                {'from': 'validate', 'to': 'end', 'label': '失败'},
                {'from': 'business', 'to': 'persist'},
                {'from': 'persist', 'to': 'end'},
            ],
            'title': '业务处理流程活动图',
        },
    },
    # 2026-06-07 v7.1：标题编号规范化
    'title_normalize': {
        'enable': True,
        'heading_style_ids': [
            '1', '2', '3',
            'Heading1', 'Heading2', 'Heading3',
            'heading1', 'heading2', 'heading3',
            'heading 1', 'heading 2', 'heading 3',
        ],
        'number_pattern': r'^(\d+(\.\d+)*)[\s.、，,：:]*',
        'text_strip_patterns': [
            r'^(\d+\.\d+)\D+',
            r'^(\d+)、\s*',
            r'^\([一二三四五六七八九十0-9]+\)\s*',
        ],
        'fallback_h2_body': '子模块',
        'fallback_h3_body': '子章节',
        'override_inherited_numbering': True,
        'strip_style_numbering': True,
    },
    # 2026-06-07 v7.1：空表格处理（增强版）
    'empty_table': {
        'enable': True,
        'delete_if': ['组件列表', '接口定义', '字段说明', '不涉及'],
        'fill_if': {
            '功能说明': '本模块实现{module_name}的核心业务功能，涵盖关键子功能。',
            '业务规则': '1. 业务规则一\n2. 业务规则二\n3. 业务规则三',
            '输入项': '本模块的输入项包括业务参数、交易上下文、用户身份信息等。',
            '输出项': '本模块的输出项包括业务处理结果、状态码、错误信息等。',
        },
        'default_placeholder': '（不涉及）',
        'dual_layer_check': True,
    },
    # 2026-06-07 v7.1：ER 图插入配置
    'er_diagram_migration': {
        'enable': True,
        'source_heading': '适用范围',
        'target_h1': '数据库ER关系图',
        'h2_prefix': 'ER',
        'max_groups': 35,
        'max_other_ratio': 0.3,
        'max_tables_per_group': 50,
        'section_intro': '本章集中展示系统涉及的数据库实体关系图（ER图），用于直观呈现表结构及表间关系。',
        'file_naming_pattern': 'ER_{group_name}.png',
    },
    # 2026-06-07 v7.1：关键技术差异化
    'tech_description': {
        'enable': True,
        'type_keywords': {
            'query': ['查询', '检索', '浏览', '列表', '详情', '统计', '报表'],
            'write': ['新增', '修改', '删除', '保存', '提交', '编辑', '录入'],
            'batch': ['批量', '导入', '导出', '同步', '迁移', '上载', '下发'],
            'approval': ['审批', '审核', '流程', '复核', '签批', '授权'],
            'integration': ['对接', '接入', '集成', '外部', '三方', '前置', '网关'],
        },
        'subset_size': [5, 7],
        'integration_variants': [
            '通过Feign远程调用{bank_name}前置服务',
            '通过Dubbo RPC调用{bank_name}核心服务',
            '通过Kafka消息队列异步通知{bank_name}',
            '通过HTTP+JSON接口调用{bank_name}外围系统',
            '通过SOAP WS调用{bank_name}企业服务总线',
        ],
        'fallback_tech_stack': [
            'Spring Boot', 'MyBatis', 'Redis', 'MySQL',
            'ZooKeeper', 'Logback', 'Swagger',
        ],
        'title_template': '{module_name}关键技术',
        'tech_item_template': '{name}：{role}',
    },
    # 2026-06-07 v7.1：图表生成引擎配置
    'chart_engine': {
        'enable': True,
        'engine_priority': ['antv', 'matplotlib'],
        'uml_engine': 'graphviz',
        'fallback_strategy': 'reuse_then_placeholder',
        'antv_timeout_ms': 60000,
        'min_diagram_size_kb': 10,
        'parallel': True,
        'graphviz_dot_cmd': 'dot',
        'output_dir': 'output/diagrams',
        'fallback_message': '【图表待补充】\n当前环境缺少图表生成工具，请手动补充。',
    },
}


_RULES_LOCK = _threading.RLock()
_RULES_CACHE = {}                  # path -> dict
_GENERATOR_REGISTRY = {}           # (doc_type, type) -> callable
_GENERATOR_BY_NAME = {}            # generator_name -> callable


def _rules_path_default():
    """计算 doc_rules.yaml 的默认路径：与 doc_formatter.py 同目录"""
    return os.path.join(os.path.dirname(os.path.abspath(__file__)), 'doc_rules.yaml')


def _merge_defaults(loaded):
    """把 loaded 字典与 _DEFAULT_RULES 深度合并，缺失键自动补齐"""
    if not isinstance(loaded, dict):
        return _copy.deepcopy(_DEFAULT_RULES)
    result = _copy.deepcopy(_DEFAULT_RULES)
    for k, v in loaded.items():
        if isinstance(v, dict) and isinstance(result.get(k), dict):
            result[k].update(v)
        else:
            result[k] = _copy.deepcopy(v)
    return result


def load_doc_rules(yaml_path=None):
    """加载 doc_rules.yaml；缺失/解析失败时返回 _DEFAULT_RULES 深拷贝

    Args:
        yaml_path: 自定义 YAML 路径；None 时使用与本文件同目录的
                   doc_rules.yaml

    Returns:
        dict: 合并后的规则字典（始终包含 _DEFAULT_RULES 中所有顶层键）
    """
    path = yaml_path or _rules_path_default()
    with _RULES_LOCK:
        if path in _RULES_CACHE:
            return _copy.deepcopy(_RULES_CACHE[path])
        if not os.path.exists(path):
            _RULES_CACHE[path] = _copy.deepcopy(_DEFAULT_RULES)
            return _copy.deepcopy(_DEFAULT_RULES)
        if not _YAML_AVAILABLE:
            print(f'[WARN] PyYAML 未安装，使用内置 _DEFAULT_RULES: {path}',
                  file=sys.stderr)
            _RULES_CACHE[path] = _copy.deepcopy(_DEFAULT_RULES)
            return _copy.deepcopy(_DEFAULT_RULES)
        try:
            with open(path, 'r', encoding='utf-8') as f:
                loaded = _yaml.safe_load(f) or {}
        except Exception as e:
            print(f'[WARN] 解析 {path} 失败，使用内置 _DEFAULT_RULES: {e}',
                  file=sys.stderr)
            _RULES_CACHE[path] = _copy.deepcopy(_DEFAULT_RULES)
            return _copy.deepcopy(_DEFAULT_RULES)
        merged = _merge_defaults(loaded)
        _RULES_CACHE[path] = _copy.deepcopy(merged)
        return _copy.deepcopy(merged)


def clear_rules_cache(yaml_path=None):
    """清空已缓存的规则；测试场景或热重载配置时使用

    Args:
        yaml_path: 仅清空该路径；None 时清空全部缓存
    """
    with _RULES_LOCK:
        if yaml_path is None:
            _RULES_CACHE.clear()
        else:
            _RULES_CACHE.pop(yaml_path, None)


def register_generator(generator_name, callable_obj=None, doc_type=None,
                       gen_type=None):
    """注册自定义 generator（hook 点）

    两种调用方式：
      1. 按名称注册：register_generator('_build_x', fn)
         → 后续 apply_yaml_rules_to_doc 命中该 generator_name 时调用 fn
      2. 按 (doc_type, type) 路由注册：
         register_generator(callable_obj, doc_type='详细设计',
                            gen_type='table')
         → 优先于按名称注册

    Args:
        generator_name: 字符串名（与 YAML 中 generator 字段一致）
        callable_obj:   可调用对象
        doc_type:       章节类型（可选，用于精确路由）
        gen_type:       内容类型（可选，用于精确路由）
    """
    if doc_type is not None and gen_type is not None:
        _GENERATOR_REGISTRY[(doc_type, gen_type)] = callable_obj
        return
    if generator_name is not None and callable_obj is not None:
        _GENERATOR_BY_NAME[generator_name] = callable_obj
    elif generator_name is not None and generator_name in _GENERATOR_BY_NAME:
        # 仅做解绑
        _GENERATOR_BY_NAME.pop(generator_name, None)


def _resolve_generator(generator_name, doc_type, gen_type):
    """根据规则项解析真实可调用对象，遵循 (doc_type,type) > name > None"""
    if doc_type and gen_type:
        fn = _GENERATOR_REGISTRY.get((doc_type, gen_type))
        if fn is not None:
            return fn
    if generator_name:
        fn = _GENERATOR_BY_NAME.get(generator_name)
        if fn is not None:
            return fn
    return None


def get_default_empty_keywords(rules=None):
    """从规则中提取空章节关键词列表（供 fill_empty_chapter_* 默认值）"""
    rules = rules or _DEFAULT_RULES
    return list(rules.get('empty_chapter_keywords', []))


def get_table_style_params(rules=None):
    """从规则中提取表格样式参数，返回 dict（与 apply_table_style 参数对齐）"""
    rules = rules or _DEFAULT_RULES
    ts = rules.get('table_style', {}) or {}
    return {
        'style_name': ts.get('border', 'Table Grid'),
        'font_size': Pt(ts.get('font_size_pt', 10.5)),
        'header_bold': ts.get('header_bold', True),
        'header_fill': ts.get('header_fill', 'D9D9D9'),
        'border_color': ts.get('border_color', '000000'),
        'border_size_eighths': ts.get('border_size_eighths', 4),
    }


def get_paragraph_params(rules=None):
    """从规则中提取段落格式参数"""
    rules = rules or _DEFAULT_RULES
    p = rules.get('paragraph', {}) or {}
    return {
        'chars': p.get('first_line_indent_chars', 2),
        'font_name': p.get('font_name', '宋体'),
        'font_size_pt': p.get('font_size_pt', 12),
        'line_spacing': p.get('line_spacing', 1.5),
        'apply_to_body': p.get('apply_to_body', True),
        'skip_headings': p.get('skip_headings', True),
        'skip_tables': p.get('skip_tables', True),
    }


def get_toc_params(rules=None):
    """从规则中提取 TOC 参数"""
    rules = rules or _DEFAULT_RULES
    t = rules.get('toc', {}) or {}
    return {
        'levels': t.get('levels', '1-3'),
        'placeholder_text': t.get('placeholder_text',
                                  '请在 Word 中按 F9 或右键更新域以生成目录'),
    }


def _is_element_between(elem, start_elem, end_elem):
    """检查 elem 是否在 start_elem 和 end_elem 之间（按文档顺序）"""
    parent = start_elem.getparent()
    if parent is None:
        return False
    siblings = list(parent)
    try:
        start_idx = siblings.index(start_elem)
    except ValueError:
        return False
    if end_elem is not None:
        try:
            end_idx = siblings.index(end_elem)
        except ValueError:
            end_idx = len(siblings)
    else:
        end_idx = len(siblings)
    try:
        elem_idx = siblings.index(elem)
    except ValueError:
        return False
    return start_idx < elem_idx < end_idx


def fill_empty_chapter(doc, keywords=None, placeholder='不涉及', rules=None):
    """[向后兼容] 检查指定章节关键词之间的内容；如为空则插入占位段落

    新增 rules 参数：
        - 当 keywords=None 时，使用 rules['empty_chapter_keywords']
        - 当 placeholder 仍是默认 '不涉及' 且 rules 含
          rules['placeholder']['empty_chapter'] 时，后者优先
    """
    if rules is None:
        rules = _DEFAULT_RULES
    if not keywords:
        kws = list(rules.get('empty_chapter_keywords', []))
    else:
        kws = [k if isinstance(k, str) else k[0] for k in keywords]
    ph_default = rules.get('placeholder', {}).get('empty_chapter', placeholder)
    effective_placeholder = ph_default if placeholder == '不涉及' else placeholder
    if doc is None or not kws:
        return 0
    paragraphs = list(doc.paragraphs)
    filled = 0
    for kw in kws:
        target_idx = None
        for i, p in enumerate(paragraphs):
            if not _is_heading_paragraph(p):
                continue
            text = (p.text or '').strip()
            if not text:
                continue
            if text == kw or kw in text:
                target_idx = i
                break
        if target_idx is None:
            continue
        target_p = paragraphs[target_idx]
        next_heading_idx = None
        for j in range(target_idx + 1, len(paragraphs)):
            if _is_heading_paragraph(paragraphs[j]):
                next_heading_idx = j
                break
        end_idx = next_heading_idx if next_heading_idx is not None else len(paragraphs)
        is_empty = True
        for j in range(target_idx + 1, end_idx):
            t = (paragraphs[j].text or '').strip()
            if t:
                is_empty = False
                break
        # 2026-06-05 修复：段落检查通过后，还需检查是否包含表格内容。
        # fill_empty_chapter 原先只检查段落文本，不检查表格，导致仅有表格的章节被误判为空。
        if is_empty and doc and doc.tables:
            target_elem = target_p._element
            next_heading_elem = paragraphs[next_heading_idx]._element if next_heading_idx is not None else None
            for table in doc.tables:
                tbl_elem = table._tbl
                if _is_element_between(tbl_elem, target_elem, next_heading_elem):
                    if any(cell.text.strip() for row in table.rows for cell in row.cells):
                        is_empty = False
                        break
        if not is_empty:
            continue
        parent = target_p._element.getparent()
        if parent is None:
            continue
        new_p = OxmlElement('w:p')
        r = OxmlElement('w:r')
        rPr = OxmlElement('w:rPr')
        color = OxmlElement('w:color')
        color.set(qn('w:val'), '000000')
        rPr.append(color)
        r.append(rPr)
        t_elem = OxmlElement('w:t')
        t_elem.set(qn('xml:space'), 'preserve')
        t_elem.text = effective_placeholder
        r.append(t_elem)
        new_p.append(r)
        target_elem = target_p._element
        target_elem.addnext(new_p)
        filled += 1
        paragraphs = list(doc.paragraphs)
    return filled


def fill_empty_chapter_compat(doc, keywords=None, placeholder='不涉及',
                              skip_if_has_content=True, rules=None):
    """[向后兼容] fill_empty_chapter_compat 的规则化版本

    行为同 doc_formatter.fill_empty_chapter_compat；新增 rules 参数。
    """
    if rules is None:
        rules = _DEFAULT_RULES
    if not keywords:
        kws = list(rules.get('empty_chapter_keywords', []))
    else:
        kws = list(keywords)
    ph_default = rules.get('placeholder', {}).get('empty_chapter', placeholder)
    effective_placeholder = ph_default if placeholder == '不涉及' else placeholder
    if doc is None or not kws:
        return 0
    paragraphs = list(doc.paragraphs)
    filled = 0
    for kw in kws:
        target_p = None
        for p in paragraphs:
            if not _is_heading_paragraph(p):
                continue
            text = (p.text or '').strip()
            if not text:
                continue
            if text == kw or kw in text:
                target_p = p
                break
        if target_p is None:
            continue
        target_idx = paragraphs.index(target_p)
        next_heading_idx = None
        for j in range(target_idx + 1, len(paragraphs)):
            if _is_heading_paragraph(paragraphs[j]):
                next_heading_idx = j
                break
        end_idx = next_heading_idx if next_heading_idx is not None else len(paragraphs)
        if skip_if_has_content:
            is_empty = True
            for j in range(target_idx + 1, end_idx):
                t = (paragraphs[j].text or '').strip()
                if t:
                    is_empty = False
                    break
            if not is_empty:
                continue
        parent = target_p._element.getparent()
        if parent is None:
            continue
        new_p = OxmlElement('w:p')
        r = OxmlElement('w:r')
        rPr = OxmlElement('w:rPr')
        color = OxmlElement('w:color')
        color.set(qn('w:val'), '000000')
        rPr.append(color)
        r.append(rPr)
        t_elem = OxmlElement('w:t')
        t_elem.set(qn('xml:space'), 'preserve')
        t_elem.text = effective_placeholder
        r.append(t_elem)
        new_p.append(r)
        target_p._element.addnext(new_p)
        filled += 1
        paragraphs = list(doc.paragraphs)
    return filled


def apply_yaml_rules_to_doc(doc, doc_type, rules=None,
                            generators=None, dry_run=False):
    """按 doc_type（'概要设计' / '详细设计' / '单元测试报告' / ...）
    应用 doc_rules.yaml 中定义的 fill_chapters 规则

    流程：
      1. 遍历 rules['fill_chapters'][doc_type] 中每条规则
      2. 找到匹配标题的章节
      3. 若该章节空（无内容）→ 调用对应 generator 生成内容并插入
      4. 若未注册 generator 或 generator 返回 None → 退化为
         fill_empty_chapter（写入"不涉及"占位）

    Args:
        doc:       docx Document
        doc_type:  章节类型，必须在 rules['fill_chapters'] 中
        rules:     规则字典（None 时自动 load_doc_rules()）
        generators:dict，key=generator_name，value=callable；临时覆盖
                    此次调用的 generator 解析表（不影响全局 registry）
        dry_run:   True 时只统计将要操作的章节，不实际改动文档

    Returns:
        dict: {'matched': N, 'filled': M, 'skipped': K, 'dry_run': bool}
    """
    if rules is None:
        rules = load_doc_rules()
    fill_map = rules.get('fill_chapters', {}) or {}
    rule_list = fill_map.get(doc_type, []) or []
    if not rule_list:
        return {'matched': 0, 'filled': 0, 'skipped': 0, 'dry_run': bool(dry_run)}

    result = {'matched': 0, 'filled': 0, 'skipped': 0, 'dry_run': bool(dry_run)}
    paragraphs = list(doc.paragraphs)

    for rule in rule_list:
        keywords = rule.get('keywords', []) or []
        gen_type = rule.get('type', 'text')
        gen_name = rule.get('generator')
        target_p = None
        target_idx = None
        for i, p in enumerate(paragraphs):
            if not _is_heading_paragraph(p):
                continue
            text = (p.text or '').strip()
            if not text:
                continue
            for kw in keywords:
                if text == kw or kw in text:
                    target_p = p
                    target_idx = i
                    break
            if target_p is not None:
                break
        if target_p is None:
            continue
        result['matched'] += 1
        next_idx = None
        for j in range(target_idx + 1, len(paragraphs)):
            if _is_heading_paragraph(paragraphs[j]):
                next_idx = j
                break
        end_idx = next_idx if next_idx is not None else len(paragraphs)
        has_content = any(
            (paragraphs[j].text or '').strip()
            for j in range(target_idx + 1, end_idx)
        )
        if has_content:
            result['skipped'] += 1
            continue
        if dry_run:
            result['filled'] += 1
            continue
        # 解析 generator
        gen_fn = None
        if generators and gen_name and gen_name in generators:
            gen_fn = generators[gen_name]
        if gen_fn is None:
            gen_fn = _resolve_generator(gen_name, doc_type, gen_type)
        if gen_fn is not None:
            try:
                gen_fn(doc, target_p, rule)
                result['filled'] += 1
                paragraphs = list(doc.paragraphs)
                continue
            except Exception as e:
                print(f'[WARN] generator {gen_name} 失败: {e}', file=sys.stderr)
        # 退化：写"不涉及"占位
        ph = rules.get('placeholder', {}).get('empty_chapter', '不涉及')
        new_p = OxmlElement('w:p')
        r = OxmlElement('w:r')
        t = OxmlElement('w:t')
        t.set(qn('xml:space'), 'preserve')
        t.text = ph
        r.append(t)
        new_p.append(r)
        target_p._element.addnext(new_p)
        result['filled'] += 1
        paragraphs = list(doc.paragraphs)
    return result


def apply_global_styles_from_rules(doc, rules=None):
    """应用 rules['paragraph'] + rules['table_style'] 的全局样式

    作用：根据 YAML 规则一次性调用：
      - apply_body_indent_to_doc（缩进）
      - 遍历 doc.tables → apply_table_style（表样式）

    Args:
        doc:   docx Document
        rules: 规则字典（None 时 load_doc_rules()）

    Returns:
        dict: {'paragraphs_indented': N, 'tables_styled': M}
    """
    if rules is None:
        rules = load_doc_rules()
    p_params = get_paragraph_params(rules)
    n_indent = apply_body_indent_to_doc(
        doc,
        chars=p_params.get('chars', 2),
        skip_headings=p_params.get('skip_headings', True),
        skip_tables=p_params.get('skip_tables', True),
    )
    ts = get_table_style_params(rules)
    n_tbl = 0
    for tbl in doc.tables:
        apply_table_style(
            tbl,
            style_name=ts.get('style_name', 'Table Grid'),
            font_size=ts.get('font_size'),
            header_bold=ts.get('header_bold', True),
        )
        n_tbl += 1
    return {'paragraphs_indented': n_indent, 'tables_styled': n_tbl}


# ═══════════════════════════════════════════════════════════════
# 示例内容识别与清除（Template Example Content Clearing）
# ────────────────────────────────────────────────────────────────
# 设计目标：
#   模板中常有"示例内容"用于说明填写方式，这些内容需在填充前清除。
#   识别策略：
#     1. 文本前缀匹配：如"示例："、"样例："、"如："等
#     2. 占位符模式：如"XXX"、"YYY"、"【请填写】"等
#     3. 模板说明段落：如"以下为示例"、"请根据实际情况修改"等
#     4. 灰色/浅色字体段落：模板中用于区分正式内容的说明文字
#   清除策略：
#     - 整个段落匹配示例内容 → 删除段落
#     - 段落部分匹配 → 仅清除匹配部分
#     - 表格中的示例内容 → 清空单元格
# ═══════════════════════════════════════════════════════════════

# 示例内容识别关键词（可配置，通过 doc_rules.yaml 扩展）
_EXAMPLE_CONTENT_PREFIXES = [
    '示例：', '示例:', '样例：', '样例:', '例如：', '例如:', '如：',
    '举例：', '举例:', '说明：', '说明:', '注：', '注:', '备注：',
    '注意：', '注意:', '提示：', '提示:', '※', '★',
]

_EXAMPLE_CONTENT_WHOLE_LINE = [
    '以下为示例', '以下为样例', '示例内容', '样例内容',
    '请根据实际情况', '请根据实际需求', '请修改为实际', '请替换为实际',
    '此处填写', '此处填入', '请在此处', '模板内容', '占位内容',
    '以下内容仅供参考', '实际内容请根据', '请自行补充',
]

_EXAMPLE_PLACEHOLDER_PATTERNS = [
    '【请填写', '【请补充', '【待填写', '【待补充',
    'XXX', 'YYY', 'ZZZ', '...', '……',
    '{{', '}}', '<<', '>>',
]

# 灰色/浅色字体阈值（RGB各分量 > 128 视为浅色，一般用于说明文字）
_LIGHT_COLOR_THRESHOLD = 128


def _is_example_prefix(text):
    """检查文本是否以示例前缀开头"""
    if not text:
        return False
    text_stripped = text.strip()
    for prefix in _EXAMPLE_CONTENT_PREFIXES:
        if text_stripped.startswith(prefix):
            return True
    return False


def _is_example_whole_line(text):
    """检查整行文本是否属于示例内容"""
    if not text:
        return False
    text_stripped = text.strip()
    for keyword in _EXAMPLE_CONTENT_WHOLE_LINE:
        if keyword in text_stripped:
            return True
    return False


def _is_placeholder_text(text):
    """检查文本是否包含占位符模式"""
    if not text:
        return False
    for pattern in _EXAMPLE_PLACEHOLDER_PATTERNS:
        if pattern in text:
            return True
    return False


def _is_light_color_run(run):
    """检查 run 的颜色是否为浅色（说明文字）"""
    try:
        color = run.font.color.rgb if run.font.color else None
    except Exception:
        return False
    if color is None:
        return False
    # 排除蓝色（已由 clear_all_blue_runs 处理）
    if is_blue_color(color):
        return False
    r_val = (color >> 16) & 0xFF
    g_val = (color >> 8) & 0xFF
    b_val = color & 0xFF
    return (r_val > _LIGHT_COLOR_THRESHOLD and
            g_val > _LIGHT_COLOR_THRESHOLD and
            b_val > _LIGHT_COLOR_THRESHOLD)


def clear_example_content(doc, aggressive=False):
    """识别并清除模板中的示例内容

    清除策略：
      1. 整段匹配示例关键词 → 删除整个段落
      2. 段落含占位符模式 → 清除段落内容
      3. 段落含浅色文字 → 若 aggressive=True 则清除
      4. 表格单元格含示例内容 → 清空单元格

    Args:
        doc: docx Document 对象
        aggressive: 是否激进模式（清除浅色文字段落）

    Returns:
        dict: 统计信息
    """
    stats = {
        'paragraphs_removed': 0,
        'paragraphs_cleared': 0,
        'cells_cleared': 0,
        'placeholder_count': 0,
    }

    # 处理段落
    paragraphs_to_remove = []
    for p in doc.paragraphs:
        if _is_heading_paragraph(p):
            continue
        text = (p.text or '').strip()
        if not text:
            continue

        # 整段匹配示例内容 → 删除
        if _is_example_whole_line(text):
            paragraphs_to_remove.append(p)
            stats['paragraphs_removed'] += 1
            continue

        # 占位符文本 → 清除段落
        if _is_placeholder_text(text):
            for run in p.runs:
                run.text = ''
            stats['placeholder_count'] += 1
            stats['paragraphs_cleared'] += 1
            continue

        # 示例前缀 → 清除段落
        if _is_example_prefix(text):
            for run in p.runs:
                run.text = ''
            stats['paragraphs_cleared'] += 1
            continue

        # 激进模式：清除浅色文字段落
        if aggressive:
            has_light = False
            for run in p.runs:
                if _is_light_color_run(run):
                    has_light = True
                    break
            if has_light:
                for run in p.runs:
                    if not is_hyperlink_run(run):
                        run.text = ''
                stats['paragraphs_cleared'] += 1

    for p in paragraphs_to_remove:
        try:
            parent = p._element.getparent()
            if parent is not None:
                parent.remove(p._element)
        except Exception:
            pass

    # 处理表格
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                cell_text = cell.text.strip()
                if not cell_text:
                    continue
                if _is_example_whole_line(cell_text) or _is_placeholder_text(cell_text):
                    for para in cell.paragraphs:
                        for run in para.runs:
                            run.text = ''
                    stats['cells_cleared'] += 1

    return stats


def clean_template_remarks(doc):
    """清理模板中的备注说明段落

    模板中常有以下类型的备注说明：
      - "以下内容为模板说明，请根据实际情况修改"
      - "此处为示例，请替换为实际内容"
      - 蓝色字体的备注说明（已由 clear_all_blue_runs 处理）
      - 括号内的填写说明，如"（请填写项目名称）"

    本函数识别并清除这些备注说明，为内容填充做准备。
    """
    REMARK_KEYWORDS = [
        '模板说明', '模板备注', '填写说明', '填写指南',
        '请根据实际情况', '请根据项目实际', '请替换为',
        '以下为示例', '以下内容仅', '本模板',
        '说明：以下', '备注：以下',
    ]

    stats = {'paragraphs_removed': 0, 'cells_cleared': 0}

    # 清理段落中的备注说明
    paragraphs_to_remove = []
    for p in doc.paragraphs:
        if _is_heading_paragraph(p):
            # 标题中的括号说明 → 清除括号内容
            text = (p.text or '').strip()
            if '（' in text and '）' in text:
                # 保留标题，清除括号内的备注
                new_text = re.sub(r'（[^）]*请[^）]*）', '', text)
                new_text = re.sub(r'【[^】]*请[^】]*】', '', new_text)
                if new_text != text:
                    for run in p.runs:
                        run.text = ''
                    if p.runs:
                        p.runs[0].text = new_text.strip()
            continue

        text = (p.text or '').strip()
        if not text:
            continue

        for kw in REMARK_KEYWORDS:
            if kw in text:
                paragraphs_to_remove.append(p)
                stats['paragraphs_removed'] += 1
                break

    for p in paragraphs_to_remove:
        try:
            parent = p._element.getparent()
            if parent is not None:
                parent.remove(p._element)
        except Exception:
            pass

    # 清理表格中的备注说明
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                cell_text = cell.text.strip()
                if not cell_text:
                    continue
                for kw in REMARK_KEYWORDS:
                    if kw in cell_text:
                        for para in cell.paragraphs:
                            for run in para.runs:
                                run.text = ''
                        stats['cells_cleared'] += 1
                        break

    return stats


def enhanced_blue_cleanup(doc):
    """增强版蓝色文本清理：分阶段处理

    阶段1：段落级蓝色清理（含标题、正文、占位符）
    阶段2：表格级蓝色清理（含单元格）
    阶段3：残留蓝色 run 清理（二次扫描）

    相比 clear_all_blue_runs 的单次处理，此函数更彻底。
    """
    stats = {'paragraphs_cleared': 0, 'cells_cleared': 0, 'runs_cleared': 0}

    # 阶段1：段落级清理
    paragraphs_to_remove = []
    for p in doc.paragraphs:
        text = (p.text or '').strip()
        is_heading = _is_heading_paragraph(p)

        # 蓝色空标题 → 删除
        if is_heading and not text:
            has_blue = False
            for run in p.runs:
                if not (run.font.color and run.font.color.rgb):
                    continue
                if is_blue_color(run.font.color.rgb):
                    has_blue = True
                    break
            if has_blue:
                paragraphs_to_remove.append(p)
                continue

        # 蓝色非标题段落 → 检查是否占位符
        if not is_heading:
            has_blue = False
            for run in p.runs:
                if not (run.font.color and run.font.color.rgb):
                    continue
                if is_blue_color(run.font.color.rgb) and not is_hyperlink_run(run):
                    has_blue = True
                    break
            if has_blue:
                # 占位符文本 → 删除整段
                if _is_placeholder_text(text) or _is_example_prefix(text):
                    paragraphs_to_remove.append(p)
                else:
                    # 非占位符 → 清空蓝色 run 内容
                    clear_all_blue_runs(p)
                    stats['paragraphs_cleared'] += 1

    for p in paragraphs_to_remove:
        try:
            parent = p._element.getparent()
            if parent is not None:
                parent.remove(p._element)
        except Exception:
            pass

    # 阶段2：表格级清理
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    has_blue = False
                    for run in para.runs:
                        if not (run.font.color and run.font.color.rgb):
                            continue
                        if is_blue_color(run.font.color.rgb) and not is_hyperlink_run(run):
                            has_blue = True
                            break
                    if has_blue:
                        for run in para.runs:
                            if is_hyperlink_run(run):
                                continue
                            if run.font.color and run.font.color.rgb and is_blue_color(run.font.color.rgb):
                                run.text = ''
                                try:
                                    run.font.color.rgb = RGBColor(0, 0, 0)
                                except Exception:
                                    pass
                        stats['cells_cleared'] += 1

    # 阶段3：残留检查（二次扫描确保无遗漏）
    for p in doc.paragraphs:
        for run in p.runs:
            if not (run.font.color and run.font.color.rgb):
                continue
            if is_blue_color(run.font.color.rgb) and not is_hyperlink_run(run):
                if run.text.strip():
                    run.text = ''
                    stats['runs_cleared'] += 1

    return stats


def full_template_cleanup(doc, aggressive=False, rules=None):
    """完整的模板清理流程（一键调用）

    按顺序执行：
      1. 蓝色文本清理（enhanced_blue_cleanup）
      2. 示例内容清除（clear_example_content）
      3. 模板备注清除（clean_template_remarks）
      4. 图片后冗余文字清除（remove_redundant_text_after_images，2026-06-06 新增）

    Args:
        doc: docx Document 对象
        aggressive: 是否激进模式
        rules: 自定义规则（None 时 load_doc_rules()）

    Returns:
        dict: 各阶段清理统计
    """
    if rules is None:
        rules = load_doc_rules()
    stats = {}
    stats['blue'] = enhanced_blue_cleanup(doc)
    stats['example'] = clear_example_content(doc, aggressive=aggressive)
    stats['remarks'] = clean_template_remarks(doc)
    # 阶段4：图片后冗余文字（仅在 rules 启用时执行）
    if rules.get('image_redundant_text', {}).get('enable', True):
        stats['image_redundant'] = remove_redundant_text_after_images(doc, rules=rules)
    return stats


# ═══════════════════════════════════════════════════════════════
# 图片后冗余文字检测（2026-06-06 新增）
# ────────────────────────────────────────────────────────────────
# 背景：
#   模板在"类图/顺序图/活动图/架构图"等标题下往往已写了
#   "图1" / "Figure 1" / "（图1）" 等占位说明，图片插入后
#   这些文字仍残留，与自动插入的图说重复。需要扫描整个文档
#   所有 w:drawing 后的非标题段落，按规则去除冗余文字。
#
# 关键设计：
#   - 跨章节统一处理：扫描整个 body（不仅是当前章节）
#   - 三档处理：整段删除 / 剥除前导图号 / 保留（>=12 字说明）
#   - 保护关键词：含"架构图/类图/ER图"等明确图说关键词的段落不处理
# ═══════════════════════════════════════════════════════════════

def _iter_drawing_paragraphs(doc, drawing_tags=None):
    """遍历文档中所有包含图片元素的段落（OXML 层级）

    2026-06-06 新增：跨章节扫描，统一在 doc.element.body 上遍历，
    避免上层段落 API 跳过嵌套元素（表格内图片、文本框等）。
    """
    if doc is None:
        return
    if drawing_tags is None:
        drawing_tags = ['w:drawing', 'w:pict']
    # 末尾统一加上花括号
    suffix_tags = [t.split(':', 1)[-1] for t in drawing_tags]
    body = doc.element.body
    for p_elem in body.iter():
        if not p_elem.tag.endswith('}p'):
            continue
        for t in suffix_tags:
            # OXML tag: {namespace}p
            qn_tag = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}' + t
            if p_elem.find('.//' + qn_tag) is not None:
                yield p_elem
                break


def _get_paragraph_text_from_elem(p_elem):
    """从 OXML 段落元素中提取纯文本（避免 docx Paragraph 包装的层级访问问题）"""
    if p_elem is None:
        return ''
    text_parts = []
    for t in p_elem.findall('.//' + qn('w:t')):
        if t.text:
            text_parts.append(t.text)
    return ''.join(text_parts).strip()


def _is_protected_caption(text, protected_keywords):
    """检查文本是否含保护关键词（含则不处理）"""
    if not text or not protected_keywords:
        return False
    for kw in protected_keywords:
        if kw in text:
            return True
    return False


def _unescape_yaml_regex(pattern):
    """YAML 双引号字符串中的反斜杠会被转义；本函数将 `\\s` 还原为 `\\s` 等

    2026-06-06 修复：doc_rules.yaml 中正则因 YAML 双引号字符串特性
    出现 `\\\\s`，Python re 期望 `\\s`。本函数在加载后自动还原。
    不可逆地处理：是只对 re 模式做 `replace('\\\\\\\\', '\\\\')`。
    """
    if not isinstance(pattern, str):
        return pattern
    # 仅当字符串含 "\\" 时才处理（避免误伤正常文本）
    if '\\' not in pattern:
        return pattern
    # 将四个反斜杠 → 两个反斜杠（YAML 解析 `\\\\s` 后为 `\\\\s`）
    return pattern.replace('\\\\', '\\')


def _match_figure_number_only(text, patterns):
    """检查文本是否仅含图号（无其他说明）"""
    if not text or not patterns:
        return False
    for pat in patterns:
        try:
            pat_unescaped = _unescape_yaml_regex(pat)
            if re.match(pat_unescaped, text):
                return True
        except re.error:
            continue
    return False


def _strip_figure_number(text, pattern, min_length):
    """从文本中剥除前导图号；若剥除后剩余说明 < min_length 则整段清除

    Args:
        text: 原始文本
        pattern: 复合图号正则（含一个捕获组为前导图号）
        min_length: 说明部分最小长度阈值

    Returns:
        (new_text, full_removed)
        - new_text: 剥除图号后的文本（若为 None 表示整段删除）
        - full_removed: 是否整段删除
    """
    if not text or not pattern:
        return text, False
    pat_unescaped = _unescape_yaml_regex(pattern)
    m = re.match(pat_unescaped, text)
    if not m:
        return text, False
    # 关键：使用 m.end(1) 获取 group1（前导图号）的结束位置
    # 因为 (.+)$ 贪婪匹配，m.end() 会指向字符串末尾
    try:
        end_of_prefix = m.end(1)
    except (IndexError, AttributeError):
        end_of_prefix = m.end()
    body = text[end_of_prefix:].strip()
    if not body:
        return None, True
    if len(body) < min_length:
        # 短说明视为冗余，整段删除
        return None, True
    return body, False


def remove_redundant_text_after_images(doc, rules=None, drawing_tags=None, max_offset=1):
    """扫描整篇文档，移除/剥除图片后的冗余图说文字

    处理规则（按 image_redundant_text 节配置）：
      1) 仅图号段落（"图1"）→ 删除整段
      2) 图号+短说明（< caption_description_min_length）→ 删除整段
      3) 图号+长说明（>= 阈值）→ 剥除前导图号，保留说明
      4) 含保护关键词 → 不处理

    Args:
        doc: docx Document 对象
        rules: 规则字典（None 时 load_doc_rules()）
        drawing_tags: 自定义图片元素标签列表
        max_offset: 图片后偏移段落数

    Returns:
        dict: {'paragraphs_removed': N, 'paragraphs_trimmed': M, 'skipped': K}
    """
    if rules is None:
        rules = load_doc_rules()
    rt_config = rules.get('image_redundant_text', {}) or {}
    if not rt_config.get('enable', True):
        return {'paragraphs_removed': 0, 'paragraphs_trimmed': 0, 'skipped': 0}
    if drawing_tags is None:
        drawing_tags = rt_config.get('drawing_tags', ['w:drawing', 'w:pict'])
    if max_offset is None:
        max_offset = rt_config.get('max_offset_after', 1)
    fig_only_patterns = rt_config.get('figure_number_only_patterns', [])
    fig_with_cap_pattern = rt_config.get('figure_with_caption_pattern')
    min_len = rt_config.get('caption_description_min_length', 12)
    protected = rt_config.get('protected_keywords', [])

    stats = {'paragraphs_removed': 0, 'paragraphs_trimmed': 0, 'skipped': 0}
    body = doc.element.body
    # 1) 收集所有图片段落 elem（不修改 doc.paragraphs，避免索引失效）
    image_paragraphs = list(_iter_drawing_paragraphs(doc, drawing_tags))
    if not image_paragraphs:
        return stats
    # 2) 遍历每个图片段落，定位其后 1..max_offset 个段落
    siblings = list(body)
    for img_p in image_paragraphs:
        try:
            img_idx = siblings.index(img_p)
        except ValueError:
            continue
        for offset in range(1, max_offset + 1):
            target_idx = img_idx + offset
            if target_idx >= len(siblings):
                break
            target_elem = siblings[target_idx]
            if not target_elem.tag.endswith('}p'):
                continue
            # 跳过标题段落
            pPr = target_elem.find(qn('w:pPr') + '/' + qn('w:pStyle'))
            if pPr is not None:
                sv = pPr.get(qn('w:val'), '') or ''
                if sv in HEADING_STYLE_VALS or 'Heading' in sv or sv.startswith('heading'):
                    continue
            text = _get_paragraph_text_from_elem(target_elem)
            if not text:
                continue
            # 保护关键词 → 跳过
            if _is_protected_caption(text, protected):
                stats['skipped'] += 1
                continue
            # 仅图号 → 整段删除
            if _match_figure_number_only(text, fig_only_patterns):
                body.remove(target_elem)
                siblings = list(body)
                stats['paragraphs_removed'] += 1
                continue
            # 复合图号 → 剥除前导图号
            if fig_with_cap_pattern:
                new_text, full_removed = _strip_figure_number(
                    text, fig_with_cap_pattern, min_len
                )
                if full_removed:
                    body.remove(target_elem)
                    siblings = list(body)
                    stats['paragraphs_removed'] += 1
                elif new_text and new_text != text:
                    _replace_paragraph_text(target_elem, new_text)
                    stats['paragraphs_trimmed'] += 1
    return stats


def _replace_paragraph_text(p_elem, new_text):
    """将 OXML 段落元素的文本内容替换为 new_text（保留首 run 的样式）"""
    if p_elem is None:
        return
    runs = p_elem.findall(qn('w:r'))
    if not runs:
        # 无 run 时新建一个
        r = OxmlElement('w:r')
        t = OxmlElement('w:t')
        t.set(qn('xml:space'), 'preserve')
        t.text = new_text
        r.append(t)
        p_elem.append(r)
        return
    # 清空所有 run 文本
    for r in runs:
        for t in r.findall(qn('w:t')):
            t.text = ''
    # 首 run 写入新文本
    first_run = runs[0]
    t_elems = first_run.findall(qn('w:t'))
    if t_elems:
        t_elems[0].text = new_text
        t_elems[0].set(qn('xml:space'), 'preserve')
    else:
        t = OxmlElement('w:t')
        t.set(qn('xml:space'), 'preserve')
        t.text = new_text
        first_run.append(t)


# ═══════════════════════════════════════════════════════════════
# 标题编号重构（2026-06-06 新增）
# ────────────────────────────────────────────────────────────────
# 解决 4 类问题：
#   1) 子章节标题序号重复显示（模板硬编码"4.1 4.2..." + 新计算编号叠加）
#   2) 序号定位错位（H1 实际位置在第 7 章但子节编号写"4.1"）
#   3) 序号与章节层级不对应（H2 显示 "1.1" 但实际是 H3）
#   4) 模板 H1 标题含括号"（XXX）"导致正则匹配失败
#
# 关键设计：
#   - 按 OXML 层级动态计算编号（H1→n1, H2→n1.n2, H3→n1.n2.n3）
#   - 先剥除段落开头已存在的硬编码编号，再写入新编号
#   - 支持 skip_h1 / apply_levels 配置
# ═══════════════════════════════════════════════════════════════

def _get_paragraph_level(p_elem):
    """从 OXML 段落元素读取 H 级别（1~6）；非标题返回 0"""
    if p_elem is None or not p_elem.tag.endswith('}p'):
        return 0
    pPr = p_elem.find(qn('w:pPr'))
    if pPr is None:
        return 0
    pStyle = pPr.find(qn('w:pStyle'))
    if pStyle is None:
        return 0
    sv = pStyle.get(qn('w:val'), '') or ''
    # 数字形式 '1' '2' '3'
    try:
        n = int(sv)
        if 1 <= n <= 6:
            return n
    except (TypeError, ValueError):
        pass
    # 名称形式 'Heading1' / 'heading 1' / 'Heading 1'
    sv_lower = sv.lower().replace(' ', '')
    for lvl in range(1, 7):
        if sv_lower == f'heading{lvl}':
            return lvl
    return 0


def _strip_existing_number(text, patterns):
    """从文本开头剥除已存在的编号前缀

    2026-06-06 修复：循环剥除模式（count=None），处理"4.1 4.1"这种重复前缀。
    原因：_renumber_h2_under_h1 写入"4.1 标题"后，若 renumber_headings 又对该
    H2 触发，新一轮剥除应能一次性剥除多个连续编号，避免出现"4.1 4.1 标题"。
    """
    if not text or not patterns:
        return text
    prev = text
    for _ in range(5):  # 最多剥 5 次，防止死循环
        cur = prev
        for pat in patterns:
            try:
                # 2026-06-06 修复：YAML 解析后正则中的反斜杠需 unescape
                pat_unescaped = _unescape_yaml_regex(pat)
                cur = re.sub(pat_unescaped, '', cur, count=1).strip()
            except re.error:
                continue
            if cur != prev:
                break
        if cur == prev:
            return cur
        prev = cur
    return prev


def _format_level_label(level, counters, level_format):
    """根据当前级别与计数器生成新编号标签

    Args:
        level: 1/2/3
        counters: dict，key 为级别，value 为计数
        level_format: dict，{级别: 格式串}

    Returns:
        编号字符串（含分隔符），如 "1.2"
    """
    fmt = (level_format or {}).get(str(level)) or (level_format or {}).get(level)
    if not fmt:
        return ''
    try:
        if level == 1:
            return fmt.format(n=counters.get(1, 0))
        if level == 2:
            return fmt.format(n=counters.get(1, 0), n1=counters.get(1, 0), n2=counters.get(2, 0))
        if level == 3:
            return fmt.format(
                n=counters.get(3, 0),
                n1=counters.get(1, 0), n2=counters.get(2, 0), n3=counters.get(3, 0)
            )
    except (KeyError, IndexError):
        return ''
    return ''


def renumber_headings(doc, rules=None):
    """重构标题编号系统：按 OXML 层级动态重新编号

    处理流程：
      1) 加载 heading_numbering 配置（默认 {1:'{n}', 2:'{n1}.{n2}', 3:'{n1}.{n2}.{n3}'}）
      2) 遍历 body 中所有段落，记录每个段落的 H 级别
      3) 按 H 级别累加计数器
      4) 替换段落文本：若启用 dedupe_existing，先剥除旧编号再写入新编号
      5) 跳过未在 apply_levels 范围内的级别

    Args:
        doc: docx Document 对象
        rules: 规则字典（None 时 load_doc_rules()）

    Returns:
        dict: {'h1': N1, 'h2': N2, 'h3': N3, 'renamed': total}
    """
    if rules is None:
        rules = load_doc_rules()
    hn_config = rules.get('heading_numbering', {}) or {}
    if not hn_config.get('enable', True):
        return {'h1': 0, 'h2': 0, 'h3': 0, 'renamed': 0}
    apply_levels = set(hn_config.get('apply_levels', [1, 2, 3]))
    skip_h1 = hn_config.get('skip_h1', False)
    if skip_h1:
        apply_levels.discard(1)
    level_format = hn_config.get('level_format', {
        1: '{n}', 2: '{n1}.{n2}', 3: '{n1}.{n2}.{n3}'
    })
    dedupe_existing = hn_config.get('dedupe_existing', True)
    patterns = hn_config.get('number_prefix_patterns', [
        r'^\s*\d+(\.\d+)*\s+',
        r'^\s*[一二三四五六七八九十]+[、\.]\s*',
    ])

    counters = {1: 0, 2: 0, 3: 0, 4: 0, 5: 0, 6: 0}
    stats = {'h1': 0, 'h2': 0, 'h3': 0, 'renamed': 0}
    body = doc.element.body
    for p_elem in body.iter():
        if not p_elem.tag.endswith('}p'):
            continue
        # 表格内段落跳过
        ancestor = p_elem.getparent()
        in_table = False
        while ancestor is not None:
            if ancestor.tag.endswith('}tbl'):
                in_table = True
                break
            ancestor = ancestor.getparent()
        if in_table:
            continue
        level = _get_paragraph_level(p_elem)
        if level == 0 or level not in apply_levels:
            continue
        # 累加：低级计数器清零
        counters[level] += 1
        for higher in range(level + 1, 7):
            counters[higher] = 0
        # 生成新编号
        new_label = _format_level_label(level, counters, level_format)
        # 读取现有文本
        old_text = _get_paragraph_text_from_elem(p_elem)
        if dedupe_existing:
            stripped = _strip_existing_number(old_text, patterns)
        else:
            stripped = old_text
        # 写回：新编号 + 空格 + 内容
        new_full = f'{new_label} {stripped}'.strip() if new_label else stripped
        if new_full != old_text:
            _replace_paragraph_text(p_elem, new_full)
            stats['renamed'] += 1
        stats[f'h{level}'] = stats.get(f'h{level}', 0) + 1
    return stats


# ═══════════════════════════════════════════════════════════════
# 功能模块表格内容映射（2026-06-06 新增）
# ────────────────────────────────────────────────────────────────
# 解决表格内容与业务无关、混入技术实现细节的问题：
#   - 仅保留业务相关列（去除"技术实现"/"代码路径"/"类名"等）
#   - 清洗行内容（移除含"@Autowired"/"SQL"等技术关键词的行）
#   - 自动从 businessSubmodules 数据生成业务行
# ═══════════════════════════════════════════════════════════════

def clean_function_module_table(table, rules=None):
    """清洗功能模块表格：去除技术实现列、清洗行内容

    Args:
        table: docx Table 对象
        rules: 规则字典（None 时 load_doc_rules()）

    Returns:
        dict: {'columns_removed': N, 'rows_removed': M, 'columns_kept': K, 'rows_kept': M}
    """
    if rules is None:
        rules = load_doc_rules()
    fmt_config = rules.get('function_module_table', {}) or {}
    if not fmt_config.get('enable', True):
        return {'columns_removed': 0, 'rows_removed': 0, 'columns_kept': 0, 'rows_kept': 0}
    forbidden_cols = fmt_config.get('forbidden_column_keywords', [])
    forbidden_rows = fmt_config.get('forbidden_row_keywords', [])
    fallback_text = fmt_config.get('fallback_row_text', '本模块提供核心业务功能处理能力')
    stats = {'columns_removed': 0, 'rows_removed': 0, 'columns_kept': 0, 'rows_kept': 0}
    if not table.rows:
        return stats
    # 1) 识别并移除禁止列（基于表头）
    header_row = table.rows[0]
    header_texts = [cell.text.strip() for cell in header_row.cells]
    keep_indices = []
    for idx, htxt in enumerate(header_texts):
        drop = False
        for kw in forbidden_cols:
            if kw in htxt:
                drop = True
                break
        if not drop:
            keep_indices.append(idx)
        else:
            stats['columns_removed'] += 1
    stats['columns_kept'] = len(keep_indices)
    if not keep_indices:
        return stats
    # 2) 删除禁止列：按 OXML 层级移除
    tbl = table._tbl
    rows_xml = tbl.findall(qn('w:tr'))
    # 倒序删除以避免索引错位
    drop_indices = [i for i in range(len(header_texts)) if i not in keep_indices]
    for tr in rows_xml:
        tcs = tr.findall(qn('w:tc'))
        for drop_i in sorted(drop_indices, reverse=True):
            if drop_i < len(tcs):
                tr.remove(tcs[drop_i])
    # 3) 清洗行内容
    rows_to_remove = []
    for ri, row in enumerate(table.rows):
        if ri == 0:
            continue  # 跳过表头
        # 拼接该行所有保留单元格文本用于关键词匹配
        row_text = '|'.join(cell.text.strip() for cell in row.cells)
        drop = False
        for kw in forbidden_rows:
            if kw in row_text:
                drop = True
                break
        if drop:
            rows_to_remove.append(row)
            stats['rows_removed'] += 1
        else:
            # 检查是否仅含兜底文本
            if not any(cell.text.strip() for cell in row.cells):
                # 空行：填充兜底文本
                if row.cells:
                    _set_cell_text_safe(row.cells[-1], fallback_text)
                stats['rows_kept'] += 1
            else:
                stats['rows_kept'] += 1
    for row in rows_to_remove:
        tbl.remove(row._tr)
    return stats


def _set_cell_text_safe(cell, text):
    """安全地设置单元格文本（清空后写入新文本）"""
    if cell is None:
        return
    for p in cell.paragraphs:
        for r in p.runs:
            r.text = ''
    if cell.paragraphs:
        first_p = cell.paragraphs[0]
        if first_p.runs:
            first_p.runs[0].text = text
        else:
            first_p.add_run(text)


def clean_all_function_module_tables(doc, rules=None):
    """扫描整篇文档，识别并清洗所有功能模块表格

    识别策略：扫描所有 H2/H3 标题，标题文本命中 table_title_keywords
    后，下一个表格即为功能模块表。

    Args:
        doc: docx Document 对象
        rules: 规则字典（None 时 load_doc_rules()）

    Returns:
        dict: {'tables_scanned': N, 'tables_cleaned': M, 'details': [...]}
    """
    if rules is None:
        rules = load_doc_rules()
    fmt_config = rules.get('function_module_table', {}) or {}
    if not fmt_config.get('enable', True):
        return {'tables_scanned': 0, 'tables_cleaned': 0, 'details': []}
    title_keywords = fmt_config.get('table_title_keywords', [])
    if not title_keywords:
        return {'tables_scanned': 0, 'tables_cleaned': 0, 'details': []}
    body = doc.element.body
    children = list(body)
    result = {'tables_scanned': 0, 'tables_cleaned': 0, 'details': []}
    for i, child in enumerate(children):
        if not child.tag.endswith('}p'):
            continue
        level = _get_paragraph_level(child)
        if level not in (2, 3):
            continue
        text = _get_paragraph_text_from_elem(child)
        if not any(kw in text for kw in title_keywords):
            continue
        # 找下一个表格
        for j in range(i + 1, min(i + 20, len(children))):
            sibling = children[j]
            if sibling.tag.endswith('}tbl'):
                # 找到目标表格
                from docx.table import Table
                tbl = Table(sibling, doc)
                result['tables_scanned'] += 1
                stats = clean_function_module_table(tbl, rules=rules)
                if stats['columns_removed'] > 0 or stats['rows_removed'] > 0:
                    result['tables_cleaned'] += 1
                result['details'].append({
                    'heading_text': text,
                    'stats': stats,
                })
                break
            # 遇到下一个标题则停止
            if sibling.tag.endswith('}p') and _get_paragraph_level(sibling) <= level:
                break
    return result


# ═══════════════════════════════════════════════════════════════
# 标题样式统一（2026-06-06 新增）
# ────────────────────────────────────────────────────────────────
# 解决"4.1功能描述 vs 4.2模块划分字号不一致"问题：
#   - 扫描文档中所有"标准 H2 关键词"段落（如"功能描述"/"模块划分"等）
#   - 若其当前样式不是 Heading 2，强制升级为 Heading 2 并应用统一字号
#   - 同时清理 Normal 段落的"假标题"（仅靠加粗/字体大识别）
# ═══════════════════════════════════════════════════════════════

def unify_heading_styles(doc, rules=None):
    """按 doc_rules.heading_style_unify 配置统一标题样式

    2026-06-06 新增：解决模板中"4.1功能描述"是 Heading 2 样式、4.2"模块划分"是
    Normal 样式导致的字号不一致问题。

    工作流程：
    1) 加载 heading_style_unify 配置（含 h1/h2_keywords、字号、加粗、字体名）
    2) 遍历 body 中所有段落
    3) 对每个段落：剥除编号前缀 → 文本匹配 → 推断 H 级别 → 统一应用 pStyle
    4) 字体属性：run.font.size / run.font.bold / run.font.name

    Returns:
        dict: {'unified_h1': N, 'unified_h2': M, 'total_paragraphs': T}
    """
    if rules is None:
        rules = load_doc_rules()
    hsu_config = rules.get('heading_style_unify', {}) or {}
    if not hsu_config.get('enable', True):
        return {'unified_h1': 0, 'unified_h2': 0, 'unified_h3': 0, 'total_paragraphs': 0}

    h1_keywords = hsu_config.get('h1_keywords', []) or []
    h2_keywords = hsu_config.get('h2_keywords', []) or []
    h3_keywords = hsu_config.get('h3_keywords', []) or []  # 2026-06-07 新增
    h1_size = hsu_config.get('h1_size_pt', 16)
    h2_size = hsu_config.get('h2_size_pt', 14)
    h3_size = hsu_config.get('h3_size_pt', 12)
    h1_bold = hsu_config.get('h1_bold', True)
    h2_bold = hsu_config.get('h2_bold', True)
    h3_bold = hsu_config.get('h3_bold', True)
    font_name = hsu_config.get('font_name', '黑体')
    # 2026-06-07 新增：中文/英文章节前缀清理配置
    ccp_config = hsu_config.get('chinese_chapter_prefix', {}) or {}
    ccp_enabled = ccp_config.get('enable', True)
    ccp_strip_patterns = ccp_config.get('strip_patterns', []) or []

    body = doc.element.body
    stats = {'unified_h1': 0, 'unified_h2': 0, 'unified_h3': 0, 'total_paragraphs': 0}

    for p_elem in list(body.iter()):
        if not p_elem.tag.endswith('}p'):
            continue
        # 表格内段落跳过
        ancestor = p_elem.getparent()
        in_table = False
        while ancestor is not None:
            if ancestor.tag.endswith('}tbl'):
                in_table = True
                break
            ancestor = ancestor.getparent()
        if in_table:
            continue
        stats['total_paragraphs'] += 1
        # 当前段落样式（先读取，可能已被 _renumber_h2_under_h1 升级）
        current_level = _get_paragraph_level(p_elem)
        # 读取文本并剥除编号
        text = _get_paragraph_text_from_elem(p_elem)
        text_normalized = re.sub(r'^\s*\d+(\.\d+)*\s*', '', text).strip()
        if not text_normalized:
            continue
        # 2026-06-07 新增：清理硬编码的"第N章/第N节/Chapter N/Section N"等章节前缀
        # 解决"第八章 技术实现细节"与自动编号叠加导致"第八章 5.1 技术实现细节"的问题
        if ccp_enabled and ccp_strip_patterns:
            text_stripped = text_normalized
            changed = False
            for pat in ccp_strip_patterns:
                try:
                    new_text = re.sub(pat, '', text_stripped, count=1)
                except re.error:
                    continue
                if new_text != text_stripped:
                    text_stripped = new_text
                    changed = True
            if changed and text_stripped.strip():
                # 写回剥除前缀后的文本（避免与 renumber 阶段的编号叠加）
                _replace_paragraph_text(p_elem, text_stripped.strip())
                text_normalized = text_stripped.strip()
        # 推断 H 级别
        inferred_level = 0
        # H1 优先（"系统概述"等比 H2 关键词更具体）
        for kw in h1_keywords:
            if text_normalized == kw or text_normalized.endswith(kw):
                inferred_level = 1
                break
        if inferred_level == 0:
            for kw in h2_keywords:
                if text_normalized == kw or text_normalized.endswith(kw):
                    inferred_level = 2
                    break
        # 2026-06-07 新增：H3 关键词匹配（深度子节"核心算法""代码示例""性能优化"等）
        if inferred_level == 0 and h3_keywords:
            for kw in h3_keywords:
                if text_normalized == kw or text_normalized.endswith(kw):
                    inferred_level = 3
                    break
        # 2026-06-07 修复：若段落已被 _renumber_headings 等前置流程设置为 Heading N，
        # 视为可信的标题（不再要求关键词命中），统一应用视觉属性；否则按推断级别处理
        if inferred_level == 0 and current_level > 0:
            inferred_level = current_level
        if inferred_level == 0:
            continue
        # 2026-06-07 修复：已设标题不升级原则
        # 根因：段落已设为 H2 (current_level=2)，但 h3_keywords 匹配"业务背景"等
        # → inferred_level=3 → 3 > 2 触发升级 → H2 被错误升为 H3
        # 修复：若 current_level > 0 (已有 heading 样式)，永不改为其他级别，
        # 仅应用视觉属性（字号/加粗/字体）确保一致性。
        if current_level > 0:
            if current_level == 1:
                _apply_heading_visual(p_elem, h1_size, h1_bold, font_name)
                stats['unified_h1'] += 1
            elif current_level == 2:
                _apply_heading_visual(p_elem, h2_size, h2_bold, font_name)
                stats['unified_h2'] += 1
            elif current_level == 3:
                _apply_heading_visual(p_elem, h3_size, h3_bold, font_name)
                stats['unified_h3'] += 1
            continue

        # current_level == 0（Normal 段落）→ 按关键词推断并设置级别
        if inferred_level == current_level:
            # 已是正确级别，但仍应用字号/加粗/字体名（确保视觉一致）
            if inferred_level == 1:
                _apply_heading_visual(p_elem, h1_size, h1_bold, font_name)
                stats['unified_h1'] += 1
            elif inferred_level == 2:
                _apply_heading_visual(p_elem, h2_size, h2_bold, font_name)
                stats['unified_h2'] += 1
            elif inferred_level == 3:
                _apply_heading_visual(p_elem, h3_size, h3_bold, font_name)
                stats['unified_h3'] += 1
            continue
        # 升级/降级
        _set_paragraph_style(p_elem, inferred_level, doc=doc)
        # 应用视觉属性
        if inferred_level == 1:
            _apply_heading_visual(p_elem, h1_size, h1_bold, font_name)
            stats['unified_h1'] += 1
        elif inferred_level == 2:
            _apply_heading_visual(p_elem, h2_size, h2_bold, font_name)
            stats['unified_h2'] += 1
        elif inferred_level == 3:
            _apply_heading_visual(p_elem, h3_size, h3_bold, font_name)
            stats['unified_h3'] += 1
    return stats


def _set_paragraph_style(p_elem, level, doc=None):
    """强制设置段落的 pStyle 为 HeadingN（OXML 直接操作）

    2026-06-06 新增：与 _get_paragraph_level 对应，确保样式名与 pStyle 同步。
    2026-06-07 优化：通过 doc.styles 解析模板实际 styleId（如 '1'/'2'/'3'），
    避免对 WPS/中文模板设置 'Heading 1' 后回退到 Normal 的问题。
    """
    if p_elem is None or level not in (1, 2, 3, 4, 5, 6):
        return
    # 2026-06-07 优化：通用化——优先用模板自带的 heading 样式 styleId
    style_name = _resolve_heading_style_id(doc, level) if doc is not None else f'Heading{level}'
    pPr = p_elem.find(qn('w:pPr'))
    if pPr is None:
        pPr = OxmlElement('w:pPr')
        p_elem.insert(0, pPr)
    # 清理已有 pStyle
    existing_pStyle = pPr.find(qn('w:pStyle'))
    if existing_pStyle is not None:
        pPr.remove(existing_pStyle)
    pStyle = OxmlElement('w:pStyle')
    pStyle.set(qn('w:val'), style_name)
    pPr.insert(0, pStyle)


def _wrap_orphaned_text_to_run(p_elem):
    """把段落中挂在 <w:p> 直接子节点的裸文本节点包装为 <w:r><w:t>...</w:t></w:r>

    2026-06-07 新增：解决模板中"无 run 元素"标题的视觉属性无法应用问题。

    OXML 规范要求文本必须放在 <w:r> 内的 <w:t> 中，但部分模板会直接写裸文本
    （如 <w:p>标题文本</w:p>）。这种段落虽然 _get_paragraph_text_from_elem 能读到
    文本，但 _apply_heading_visual 的 findall('w:r') 找不到 run，导致字号/加粗
    设置无效。

    工作流程：
    1) 遍历 <w:p> 的直接子节点
    2) 对 .tag 为 'w:t' 但父节点不是 'w:r' 的节点，包装为新的 <w:r><w:t/></w:r>
    3) 保留 text 内容和 xml:space="preserve" 属性
    4) 保留其他非文本/非 run 节点不动（如 w:pPr, w:bookmarkStart 等）

    Returns:
        int: 实际包装的裸文本节点数
    """
    if p_elem is None or not p_elem.tag.endswith('}p'):
        return 0
    wrapped = 0
    # 用 list() 复制防止迭代时修改
    for child in list(p_elem):
        # 跳过 pPr（段落属性）、bookmark、hyperlink 等容器
        if not child.tag.endswith('}t'):
            continue
        # 如果父节点是 w:r，跳过
        parent = child.getparent()
        if parent is not None and parent.tag.endswith('}r'):
            continue
        # 跳过空文本
        if not (child.text or '').strip():
            # 删除空裸文本节点
            p_elem.remove(child)
            continue
        # 包装为 run
        new_r = OxmlElement('w:r')
        new_t = OxmlElement('w:t')
        # 保留原 xml:space 属性（让首尾空格不被压缩）
        if child.get('{http://www.w3.org/XML/1998/namespace}space') == 'preserve':
            new_t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
        new_t.text = child.text
        new_r.append(new_t)
        # 替换原节点
        p_elem.replace(child, new_r)
        wrapped += 1
    return wrapped


def _apply_heading_visual(p_elem, size_pt, bold, font_name):
    """应用标题视觉属性：字号 + 加粗 + 字体名

    2026-06-07 增强：模板原标题可能没有 run 元素（仅有裸文本节点挂在 <w:p> 下），
    此时 _get_paragraph_text_from_elem 能读到文本但 findall('w:r') 找不到 run，
    导致字号/加粗/字体名设置无效。处理策略：
    1) 先用 _wrap_orphaned_text_to_run 把裸文本节点包装为 run
    2) 再按原逻辑遍历 runs 应用 rPr 属性
    """
    if p_elem is None:
        return
    # 兜底：把裸文本节点包装成 run（处理模板中无 rPr 的标题段落）
    _wrap_orphaned_text_to_run(p_elem)
    runs = p_elem.findall(qn('w:r'))
    for r in runs:
        rPr = r.find(qn('w:rPr'))
        if rPr is None:
            rPr = OxmlElement('w:rPr')
            r.insert(0, rPr)
        # 字号
        existing_sz = rPr.find(qn('w:sz'))
        if existing_sz is not None:
            rPr.remove(existing_sz)
        sz = OxmlElement('w:sz')
        sz.set(qn('w:val'), str(int(size_pt * 2)))  # OXML 单位是半磅
        rPr.append(sz)
        # 加粗
        if bold:
            existing_b = rPr.find(qn('w:b'))
            if existing_b is None:
                b = OxmlElement('w:b')
                rPr.append(b)
        # 字体名（中文）
        existing_rFonts = rPr.find(qn('w:rFonts'))
        if existing_rFonts is not None:
            rPr.remove(existing_rFonts)
        rFonts = OxmlElement('w:rFonts')
        rFonts.set(qn('w:ascii'), font_name)
        rFonts.set(qn('w:eastAsia'), font_name)
        rFonts.set(qn('w:hAnsi'), font_name)
        rPr.append(rFonts)


def auto_promote_bold_subheadings(doc, rules=None):
    """自动提升 Normal/加粗的子标题为对应 Heading 样式

    2026-06-07 新增：解决"子标题未设置带样式标题"问题。

    场景：模板中部分子标题（如"业务背景""设计目标""范围说明"等）在原文档
    中被设置为 Normal 样式 + 加粗格式，未带 Heading 样式，导致：
      1) 在 Word 大纲视图中无法识别为标题
      2) TOC 字段不会包含这些子标题
      3) unify_heading_styles 因为 current_level==0 跳过

    检测策略（避免误判普通加粗段落）：
      1) 段落当前样式 = Normal（无 Heading 样式）
      2) 段落文本非空、长度 <= max_chars（默认 24）
      3) 段落不以 。，；！？：、 等句子级标点结尾
      4) 段落不以 。 结尾
      5) 段落加粗（任何 run.bold == True）或字号 >= body_size_pt
      6) 文本命中 H1/H2/H3 关键词 或 匹配章节前缀清理模式

    处理：
      1) 推断 H 级别（与 unify_heading_styles 一致：H1 > H2 > H3 优先）
      2) 调用 _set_paragraph_style 应用 Heading 样式
      3) 应用视觉属性（字号/加粗/字体）

    Args:
        doc: docx Document 对象
        rules: 规则字典（None 时 load_doc_rules()）

    Returns:
        dict: {'promoted_h1': N, 'promoted_h2': M, 'promoted_h3': K, 'skipped': T}
    """
    if rules is None:
        rules = load_doc_rules()
    hsu_config = rules.get('heading_style_unify', {}) or {}
    if not hsu_config.get('enable', True):
        return {'promoted_h1': 0, 'promoted_h2': 0, 'promoted_h3': 0, 'skipped': 0}

    h1_keywords = hsu_config.get('h1_keywords', []) or []
    h2_keywords = hsu_config.get('h2_keywords', []) or []
    h3_keywords = hsu_config.get('h3_keywords', []) or []
    h1_size = hsu_config.get('h1_size_pt', 16)
    h2_size = hsu_config.get('h2_size_pt', 14)
    h3_size = hsu_config.get('h3_size_pt', 12)
    h1_bold = hsu_config.get('h1_bold', True)
    h2_bold = hsu_config.get('h2_bold', True)
    h3_bold = hsu_config.get('h3_bold', True)
    font_name = hsu_config.get('font_name', '黑体')

    # 2026-06-07 新增：可配置阈值
    auto_cfg = hsu_config.get('auto_promote', {}) or {}
    if not auto_cfg.get('enable', True):
        return {'promoted_h1': 0, 'promoted_h2': 0, 'promoted_h3': 0, 'skipped': 0}
    max_chars = auto_cfg.get('max_chars', 24)
    require_bold = auto_cfg.get('require_bold', True)
    min_size_pt = auto_cfg.get('min_size_pt', 11)  # 比正文字号略大
    body_size_pt = auto_cfg.get('body_size_pt', 10.5)

    # 匹配：仅当文本命中关键词时升级（避免误判普通加粗短句）
    keyword_match_required = auto_cfg.get('keyword_match_required', True)

    # 句子级标点：包含则视为段落正文而非子标题
    sentence_endings = set('。，；！？：、;!?:\u3000 \t\n')

    body = doc.element.body
    stats = {'promoted_h1': 0, 'promoted_h2': 0, 'promoted_h3': 0, 'skipped': 0}

    for p_elem in list(body.iter()):
        if not p_elem.tag.endswith('}p'):
            continue
        # 表格内段落跳过
        ancestor = p_elem.getparent()
        in_table = False
        while ancestor is not None:
            if ancestor.tag.endswith('}tbl'):
                in_table = True
                break
            ancestor = ancestor.getparent()
        if in_table:
            continue
        # 已经是 Heading 样式 → 跳过
        current_level = _get_paragraph_level(p_elem)
        if current_level > 0:
            continue
        # 读取文本
        text = _get_paragraph_text_from_elem(p_elem).strip()
        if not text:
            continue
        # 长度过滤：超过阈值视为正文段落
        if len(text) > max_chars:
            continue
        # 末尾标点过滤：以句子级标点结尾视为正文
        if text[-1] in sentence_endings:
            continue
        # 加粗 / 字号判断
        runs = p_elem.findall(qn('w:r'))
        has_bold = False
        has_large_size = False
        for r in runs:
            rPr = r.find(qn('w:rPr'))
            # 2026-06-07 修复：同时检查 OXML 显式 <w:b> 与 run.bold（处理继承场景）
            if rPr is not None:
                b_elem = rPr.find(qn('w:b'))
                if b_elem is not None:
                    has_bold = True
                sz_elem = rPr.find(qn('w:sz'))
                if sz_elem is not None:
                    try:
                        sz_pt = int(sz_elem.get(qn('w:val'), '21')) / 2
                        if sz_pt >= min_size_pt:
                            has_large_size = True
                    except (ValueError, TypeError):
                        pass
            # python-docx 的 run.bold 会处理样式继承
            try:
                if r.text and r.bold is True:
                    has_bold = True
            except (AttributeError, ValueError):
                pass
        if require_bold and not has_bold and not has_large_size:
            stats['skipped'] += 1
            continue
        # 关键词匹配：必须命中 H1/H2/H3 关键词（避免误判普通加粗短句）
        inferred_level = 0
        text_norm = re.sub(r'^\s*\d+(\.\d+)*\s*', '', text).strip()
        for kw in h1_keywords:
            if text_norm == kw or text_norm.endswith(kw):
                inferred_level = 1
                break
        if inferred_level == 0:
            for kw in h2_keywords:
                if text_norm == kw or text_norm.endswith(kw):
                    inferred_level = 2
                    break
        if inferred_level == 0 and h3_keywords:
            for kw in h3_keywords:
                if text_norm == kw or text_norm.endswith(kw):
                    inferred_level = 3
                    break
        # 若不要求关键词命中且有加粗/大字号，按文本长度推断级别
        if not keyword_match_required and inferred_level == 0:
            if len(text_norm) <= 6 and has_large_size:
                inferred_level = 1
            elif len(text_norm) <= 12:
                inferred_level = 2
            else:
                inferred_level = 3
        if inferred_level == 0:
            stats['skipped'] += 1
            continue
        # 应用 Heading 样式 + 视觉属性
        _set_paragraph_style(p_elem, inferred_level, doc=doc)
        if inferred_level == 1:
            _apply_heading_visual(p_elem, h1_size, h1_bold, font_name)
            stats['promoted_h1'] += 1
        elif inferred_level == 2:
            _apply_heading_visual(p_elem, h2_size, h2_bold, font_name)
            stats['promoted_h2'] += 1
        elif inferred_level == 3:
            _apply_heading_visual(p_elem, h3_size, h3_bold, font_name)
            stats['promoted_h3'] += 1
    return stats


def build_function_module_rows_from_subsystems(subsystems, rules=None):
    """根据业务子模块列表生成功能模块表格行

    Args:
        subsystems: 业务子模块列表，每项 dict 或 str
        rules: 规则字典（None 时 load_doc_rules()）

    Returns:
        list: 表格行数据（每行为 [序号, 名称, 说明]）
    """
    if rules is None:
        rules = load_doc_rules()
    fmt_config = rules.get('function_module_table', {}) or {}
    fallback = fmt_config.get('fallback_row_text', '本模块提供核心业务功能处理能力')
    rows = []
    if not isinstance(subsystems, list):
        return rows
    for idx, sub in enumerate(subsystems, 1):
        if isinstance(sub, dict):
            name = sub.get('name') or sub.get('title') or f'子模块{idx}'
            desc = sub.get('description') or ''
            if isinstance(desc, dict):
                desc = desc.get('description', '')
        elif isinstance(sub, str):
            name = sub
            desc = ''
        else:
            continue
        if not name:
            continue
        rows.append([str(idx), name, desc or fallback])
    return rows

