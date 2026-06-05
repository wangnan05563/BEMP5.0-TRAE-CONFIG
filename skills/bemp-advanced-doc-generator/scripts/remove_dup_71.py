"""删除重复的7.1"""
import paths
from docx import Document

path = str(paths.detail_design_template())
doc = Document(path)

count_71 = 0
to_remove = []
for p in doc.paragraphs:
    if p.style and p.style.name == 'Heading 2' and p.text.strip() == '7.1 功能描述':
        count_71 += 1
        if count_71 > 1:  # 保留第一个
            to_remove.append(p)

for p in to_remove:
    p._element.getparent().remove(p._element)

doc.save(path)
print(f"删除了 {len(to_remove)} 个重复")

# 验证
doc2 = Document(path)
for p in doc2.paragraphs:
    if p.style and p.style.name.startswith('Heading') and ('7.' in p.text or '额度占用' in p.text):
        print(f"  {p.style.name}: {p.text.strip()}")
