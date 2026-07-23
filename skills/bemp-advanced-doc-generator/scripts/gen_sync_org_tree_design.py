﻿﻿﻿﻿﻿﻿﻿# -*- coding: utf-8 -*-
"""
BEMP 河南农商银行"同步机构树数据并校验"个性化开发 - 详细设计文档生成器
基于实际代码实现内容生成 .docx 文档，包含目录、页码、页眉页脚。
"""
import os
from datetime import datetime
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

OUTPUT_DIR = r"d:\code\QJ\BEMP5.0DEV\.trae\skills\bemp-advanced-doc-generator\output"
OUTPUT_FILE = os.path.join(OUTPUT_DIR, "河南农商同步机构树数据并校验-详细设计说明书-20260715.docx")

# ===================== 样式工具函数 =====================

def set_cell_border(cell, **kwargs):
    """设置单元格边框"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = tcPr.find(qn('w:tcBorders'))
    if tcBorders is None:
        tcBorders = OxmlElement('w:tcBorders')
        tcPr.append(tcBorders)
    for edge in ('top', 'left', 'bottom', 'right'):
        if edge in kwargs:
            tag = 'w:{}'.format(edge)
            element = tcBorders.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                tcBorders.append(element)
            element.set(qn('w:val'), 'single')
            element.set(qn('w:sz'), '4')
            element.set(qn('w:color'), '000000')


def set_cell_background(cell, color="D9D9D9"):
    """设置单元格背景色（表头灰色）"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color)
    tcPr.append(shd)


def set_run_font(run, name="宋体", size=10.5, bold=False, color=None):
    """设置run字体"""
    run.font.name = name
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    # 中文字体设置
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    rFonts.set(qn('w:eastAsia'), name)
    rFonts.set(qn('w:ascii'), name)
    rFonts.set(qn('w:hAnsi'), name)


def add_heading(doc, text, level=1):
    """添加标题，使用 Word 内置 Heading 样式以保证目录可生成"""
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        set_run_font(run, name="黑体", size=16 - level * 2, bold=True, color="000000")
    return p


def add_para(doc, text, bold=False, indent=True, size=10.5):
    """添加正文段落（首行缩进2字符）"""
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.first_line_indent = Cm(0.74)
    p.paragraph_format.line_spacing = 1.5
    run = p.add_run(text)
    set_run_font(run, name="宋体", size=size, bold=bold)
    return p


def add_bullet(doc, text, size=10.5):
    """添加项目符号段落"""
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.line_spacing = 1.5
    run = p.add_run(text)
    set_run_font(run, name="宋体", size=size)
    return p


def add_code_block(doc, code_text):
    """添加代码块（等宽字体、灰底）"""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.line_spacing = 1.15
    run = p.add_run(code_text)
    run.font.name = "Consolas"
    run.font.size = Pt(9)
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    rFonts.set(qn('w:eastAsia'), "Consolas")
    rFonts.set(qn('w:ascii'), "Consolas")
    rFonts.set(qn('w:hAnsi'), "Consolas")
    # 灰色底
    pPr = p._element.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'F5F5F5')
    pPr.append(shd)


def add_table(doc, headers, rows, col_widths=None):
    """添加表格"""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False

    # 表头
    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr_cells[i].text = ""
        p = hdr_cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(header)
        set_run_font(run, name="宋体", size=10, bold=True)
        set_cell_background(hdr_cells[i])
        set_cell_border(hdr_cells[i])
        hdr_cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # 数据行
    for r_idx, row in enumerate(rows, 1):
        row_cells = table.rows[r_idx].cells
        for c_idx, val in enumerate(row):
            row_cells[c_idx].text = ""
            p = row_cells[c_idx].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(str(val))
            set_run_font(run, name="宋体", size=10)
            set_cell_border(row_cells[c_idx])
            row_cells[c_idx].vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # 列宽设置
    if col_widths:
        for i, width in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(width)

    # 添加空行避免表格紧贴下一章节
    doc.add_paragraph()
    return table


def add_toc_field(doc):
    """插入动态目录域"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run()
    # 创建 fldChar begin
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = 'TOC \\o "1-3" \\h \\z \\u'
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')
    fldChar3 = OxmlElement('w:t')
    fldChar3.text = '右键此处选择"更新域"以生成目录'
    fldChar4 = OxmlElement('w:fldChar')
    fldChar4.set(qn('w:fldCharType'), 'end')
    r_element = run._r
    r_element.append(fldChar1)
    r_element.append(instrText)
    r_element.append(fldChar2)
    r_element.append(fldChar3)
    r_element.append(fldChar4)


def add_page_break(doc):
    """添加分页符"""
    from docx.enum.text import WD_BREAK
    p = doc.add_paragraph()
    run = p.add_run()
    run.add_break(WD_BREAK.PAGE)


def setup_page_and_header_footer(doc):
    """设置页面与页眉页脚"""
    # 页面边距
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(3.0)
        section.right_margin = Cm(2.5)

        # 页眉
        header = section.header
        if not header.paragraphs[0].runs:
            hp = header.paragraphs[0]
            hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = hp.add_run("河南农商票据交易管理平台 - 详细设计说明书")
            set_run_font(run, name="宋体", size=9, color="808080")

        # 页脚 - 页码
        footer = section.footer
        fp = footer.paragraphs[0]
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # 添加页码字段
        run = fp.add_run()
        set_run_font(run, name="宋体", size=9, color="808080")
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')
        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = 'PAGE'
        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'end')
        run._r.append(fldChar1)
        run._r.append(instrText)
        run._r.append(fldChar2)


def inject_update_fields(doc):
    """在 settings.xml 中注入 updateFields，使打开文档时自动更新目录"""
    settings = doc.settings.element
    update_fields = settings.find(qn('w:updateFields'))
    if update_fields is None:
        update_fields = OxmlElement('w:updateFields')
        update_fields.set(qn('w:val'), 'true')
        settings.append(update_fields)


# ===================== 文档生成主流程 =====================

def build_design_doc():
    doc = Document()

    # 设置默认样式
    style = doc.styles['Normal']
    style.font.name = '宋体'
    style.font.size = Pt(10.5)
    rPr = style.element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    rFonts.set(qn('w:eastAsia'), '宋体')

    # 页面与页眉页脚
    setup_page_and_header_footer(doc)

    # ===== 封面 =====
    for _ in range(3):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("河南农商银行")
    set_run_font(run, name="黑体", size=26, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("新一代票据系统个性化开发")
    set_run_font(run, name="黑体", size=22, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("同步机构树数据并校验")
    set_run_font(run, name="黑体", size=22, bold=True)

    for _ in range(2):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("详 细 设 计 说 明 书")
    set_run_font(run, name="黑体", size=20, bold=True)

    for _ in range(6):
        doc.add_paragraph()

    info_table = doc.add_table(rows=4, cols=2)
    info_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    info_data = [
        ("项目名称", "BEMP5.0 河南农商银行个性化开发"),
        ("需求名称", "同步机构树数据并校验（机构管理优化）"),
        ("文档版本", "V1.0"),
        ("编写日期", "2026-07-15"),
    ]
    for i, (k, v) in enumerate(info_data):
        cells = info_table.rows[i].cells
        cells[0].text = ""
        cells[1].text = ""
        p1 = cells[0].paragraphs[0]
        p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r1 = p1.add_run(k)
        set_run_font(r1, name="宋体", size=12, bold=True)
        p2 = cells[1].paragraphs[0]
        p2.alignment = WD_ALIGN_PARAGRAPH.LEFT
        r2 = p2.add_run(v)
        set_run_font(r2, name="宋体", size=12)
        for c in cells:
            set_cell_border(c)

    add_page_break(doc)

    # ===== 修订记录 =====
    add_heading(doc, "修订记录", level=1)
    add_table(doc,
        ["版本", "日期", "修订内容", "修订人"],
        [
            ["V1.0", "2026-07-15", "初始版本，覆盖同步机构树数据并校验全部需求", "bemp个性化开发"],
        ],
        col_widths=[2, 3, 8, 3]
    )
    add_page_break(doc)

    # ===== 目录 =====
    add_heading(doc, "目录", level=1)
    add_toc_field(doc)
    add_page_break(doc)

    # ===== 第一章 系统概述 =====
    add_heading(doc, "第一章 系统概述", level=1)

    add_heading(doc, "1.1 业务背景", level=2)
    add_para(doc, "河南农商银行新一代票据系统需同步核心系统机构主数据，在建立机构时校验机构合法性与状态，在贴现业务提交时校验业务发起机构是否为账务机构。当前系统完全未实现此功能，存在机构数据不一致和业务越权风险。")
    add_para(doc, "需求来源于《河南农商银行关于新一代票据系统机构管理及机构管理员功能变更的需求》（2026年06月18日），属于业务功能点三：同步机构树数据并校验。ODS在日初提供格式模版为 CBS_PJGGX_yyyyMMdd.txt（机构关系表）和 CBS_PJGCS_yyyyMMdd.txt（机构参数表）到票据系统 cbs 的日期目录下，票据系统增加定时任务读取文件并存储到自建的对应表中，建立机构及发起业务时进行校验。")

    add_heading(doc, "1.2 设计目标", level=2)
    add_table(doc,
        ["目标类型", "目标描述"],
        [
            ["功能目标", "实现日初同步核心系统机构参数和机构关系数据；建立机构时校验机构合法性和状态；贴现业务提交时校验账务机构"],
            ["安全目标", "修复批量新增机构管理员缺少法人管理员校验的安全漏洞；密码有效期和密码错误次数参数化；机构不具备角色精确报错"],
            ["质量目标", "SonarQube 扫描无 Blocker、无 Critical 问题；P0/P1 缺陷全部修复验证通过"],
            ["兼容目标", "原有机构管理、贴现业务、机构管理员管理功能不受影响"],
        ],
        col_widths=[3, 13]
    )

    add_heading(doc, "1.3 范围说明", level=2)
    add_table(doc,
        ["范围类型", "说明"],
        [
            ["纳入范围", "1. PJGCS 机构参数表和 PJGGX 机构关系表设计与建表；2. 日初定时任务同步 CBS 文件数据；3. 建立机构校验（含批量导入）；4. 贴现业务提交账务机构校验；5. 批量新增机构管理员权限校验与密码参数化；6. 机构不具备角色精确报错；7. 批量新增身份验证切换为 POBM010304"],
            ["排除范围", "1. 机构管理批量导入基础功能（属其他需求模块）；2. 机构树查询条件功能（属其他需求模块）；3. 承兑、质押等其他业务的账务机构校验（spec 明确仅贴现业务提交场景校验）"],
        ],
        col_widths=[3, 13]
    )

    add_heading(doc, "1.4 术语定义", level=2)
    add_table(doc,
        ["术语", "定义"],
        [
            ["PJGCS", "机构参数表，存储核心系统下发的机构基础参数信息，36个字段"],
            ["PJGGX", "机构关系表，存储核心系统下发的机构业务关系信息，16个字段"],
            ["CBS", "核心业务系统（Core Business System），机构主数据的源头系统"],
            ["ODS", "操作数据存储（Operational Data Store），日初将CBS数据下发到票据系统"],
            ["ZNGWSJ", "账务机构业务关系种类标识，YWGXZL字段取值之一，用于判断账务机构"],
            ["JILUZT", "记录状态字段，'0'表示正常，其他值表示异常/停用"],
            ["POBM010304", "同步机构柜员信息接口，用于批量新增机构管理员时向统一身份认证系统验证用户身份"],
        ],
        col_widths=[3, 13]
    )

    add_page_break(doc)

    # ===== 第二章 系统架构设计 =====
    add_heading(doc, "第二章 系统架构设计", level=1)

    add_heading(doc, "2.1 整体架构", level=2)
    add_para(doc, "本项目基于 BEMP5.0 票据业务管理平台，采用前后端分离架构，针对河南农商银行（hnnxbank）进行个性化定制开发。后端基于 SpringBoot 微服务架构，前端基于 Vue.js 2.x + HUI 组件库。个性化代码集中在 banks/ext-hnnxbank 模块下，使用 /hnnxbank/ 作为 API 路径前缀以区分不同银行的个性化接口路由。")

    add_heading(doc, "2.2 技术选型", level=2)
    add_table(doc,
        ["层级", "技术", "版本/说明"],
        [
            ["后端框架", "SpringBoot", "2.x，微服务架构"],
            ["前端框架", "Vue.js 2.x + HUI 组件库", "前端服务端口 8091"],
            ["数据库", "Oracle", "生产环境；MySQL 兼容脚本同步提供"],
            ["缓存", "Redis", "SpringBoot 依赖"],
            ["服务协调", "ZooKeeper", "SpringBoot 依赖"],
            ["持久层", "MyBatis", "Mapper XML 方式"],
            ["调度框架", "BEMP 定时任务框架", "tt_task 表 + cron 表达式"],
            ["文件编码", "GBK", "CBS 文件统一编码，参考 BOPC010101MessageConverter"],
            ["个性化路径前缀", "/hnnxbank/", "区分不同银行的个性化接口路由"],
        ],
        col_widths=[3, 5, 8]
    )

    add_heading(doc, "2.3 模块划分", level=2)
    add_para(doc, "本需求涉及系统管理子系统（机构管理、机构管理员管理、角色权限）和业务管理子系统（贴现业务）两个子系统，共划分为7个功能模块：")
    add_table(doc,
        ["序号", "模块名称", "模块职责", "关联需求"],
        [
            ["1", "数据库表设计", "新建 PJGCS 和 PJGGX 两张表，存储核心系统机构参数和机构关系数据", "主系统机构树清单数据库表"],
            ["2", "CBS 文件读取工具", "读取并解析 CBS_PJGCS 和 CBS_PJGGX 文件，GBK 编码、0x03 分隔符", "日初定时任务同步机构数据"],
            ["3", "日初定时任务", "日初6:00执行，读取CBS文件并同步到PJGCS/PJGGX表", "日初定时任务同步机构数据"],
            ["4", "建立机构校验", "新建机构和批量导入时校验机构在 PJGCS 中存在且状态正常", "建立机构时校验机构合法性"],
            ["5", "贴现业务校验", "贴现业务提交时校验业务发起机构是否为账务机构", "贴现业务提交时校验账务机构"],
            ["6", "安全缺陷修复", "法人管理员校验、密码参数化、角色精确报错、POBM010304 切换", "批量新增机构管理员权限校验等"],
            ["7", "SonarQube 质量扫描", "对新增代码进行质量扫描，确保无严重问题", "SonarQube 扫描"],
        ],
        col_widths=[1, 3, 8, 4]
    )

    add_heading(doc, "2.4 部署拓扑", level=2)
    add_para(doc, "系统部署拓扑如下：")
    add_code_block(doc, """+-------------------+     +-------------------+     +-------------------+
|   前端(Vue.js)    | <-->|  后端(SpringBoot) | <-->|  数据库(Oracle)   |
|  Port: 8091       |     |  Port: 8010       |     |  PJGCS / PJGGX    |
+-------------------+     +-------------------+     +-------------------+
        |                         |                         ^
        v                         v                         |
+-------------------+     +-------------------+     +-------------------+
|  静态资源/CDN     |     |  Redis/ZooKeeper  |     |  CBS 文件目录     |
+-------------------+     +-------------------+     |  (ODS 日初下发)   |
                                                    +-------------------+
                           |
                           v
                  +-------------------+
                  | 统一身份认证系统  |
                  | POBM010304 接口   |
                  +-------------------+""")

    add_page_break(doc)

    # ===== 第三章 数据库设计 =====
    add_heading(doc, "第三章 数据库设计", level=1)

    add_heading(doc, "3.1 PJGCS 机构参数表", level=2)
    add_para(doc, "PJGCS 表存储核心系统下发的机构参数信息，36个字段，主键为 YNGYJG（营业机构号）+ FAREDM（法人代码）。建表脚本路径：deploy/bemp-script/src/main/resources/banks/河南农信/V202301.03.079_202607151500_T202607150001_机构参数及机构关系表建表.ddl.sql")

    add_table(doc,
        ["序号", "字段名", "中文名称", "类型", "说明"],
        [
            ["1", "FAREDM", "法人代码", "VARCHAR2(4)", "主键之一"],
            ["2", "YNGYJG", "营业机构号", "VARCHAR2(10)", "主键之一，业务主键"],
            ["3", "FENHDM", "分行代码", "VARCHAR2(3)", ""],
            ["4", "JIGOLX", "机构类型", "VARCHAR2(1)", ""],
            ["5", "JIGOMC", "机构中文名称", "VARCHAR2(60)", ""],
            ["6", "JIGOJC", "机构简称", "VARCHAR2(20)", ""],
            ["7", "JGYWMC", "机构英文名称", "VARCHAR2(60)", ""],
            ["8", "JIGOJP", "机构简拼", "VARCHAR2(10)", ""],
            ["9", "QSZXDH", "地区代号", "VARCHAR2(4)", ""],
            ["10", "SHQUDH", "省区代号", "VARCHAR2(10)", ""],
            ["11", "DIZHII", "地址", "VARCHAR2(200)", ""],
            ["12", "YWENDZ", "英文地址", "VARCHAR2(200)", ""],
            ["13", "YOUZBM", "邮政编码", "VARCHAR2(6)", ""],
            ["14", "DIANHH", "电话号码", "VARCHAR2(20)", ""],
            ["15", "CZHENH", "传真号码", "VARCHAR2(20)", ""],
            ["16", "DBGUAH", "电报挂号", "VARCHAR2(10)", ""],
            ["17", "LNXIRM", "联系人", "VARCHAR2(200)", ""],
            ["18", "LNXRDH", "联系人电话", "VARCHAR2(25)", ""],
            ["19", "DZYJDZ", "E-mail", "VARCHAR2(42)", ""],
            ["20", "WANGZH", "网址", "VARCHAR2(62)", ""],
            ["21", "SBSHBM", "设备名", "VARCHAR2(6)", ""],
            ["22", "DYNDLM", "报表队列名", "VARCHAR2(8)", ""],
            ["23", "FHIPDZ", "分行IP地址", "VARCHAR2(15)", ""],
            ["24", "FHPORT", "分行PORT号", "VARCHAR2(4)", ""],
            ["25", "SFCZYH", "是否村镇银行", "VARCHAR2(1)", ""],
            ["26", "JIBIEE", "行政级别", "VARCHAR2(1)", ""],
            ["27", "JGFWDZ", "机构服务器地址名称", "VARCHAR2(40)", ""],
            ["28", "QYNGRQ", "启用日期", "VARCHAR2(8)", ""],
            ["29", "BYZD01", "备用字段", "VARCHAR2(200)", ""],
            ["30", "WEIHRQ", "维护日期", "VARCHAR2(8)", ""],
            ["31", "WEIHSJ", "维护时间", "NUMBER(10)", ""],
            ["32", "WEIHGY", "维护柜员", "VARCHAR2(8)", ""],
            ["33", "WEIHJG", "维护机构", "VARCHAR2(10)", ""],
            ["34", "ROWIDD", "序列号", "VARCHAR2(22)", ""],
            ["35", "SHJNCH", "时间戳", "NUMBER(16)", ""],
            ["36", "JILUZT", "记录状态", "VARCHAR2(1)", "'0'表示正常，校验关键字段"],
        ],
        col_widths=[1, 2, 3, 3, 7]
    )

    add_para(doc, "索引设计：", bold=True, indent=False)
    add_bullet(doc, "PK_PJGCS 主键索引：(YNGYJG, FAREDM)")
    add_bullet(doc, "IDX_PJGCS_1 索引：(FAREDM ASC) - 按法人代码查询场景")
    add_bullet(doc, "IDX_PJGCS_2 索引：(JIGOMC ASC) - 按机构名称查询场景")

    add_heading(doc, "3.2 PJGGX 机构关系表", level=2)
    add_para(doc, "PJGGX 表存储核心系统下发的机构业务关系信息，16个字段，主键为 YNGYJG（营业机构号）+ YWGXZL（业务关系种类）+ FAREDM（法人代码）。YWGXZL 字段取值包括 ZNGWSJ（账务机构）、BAOBSJ、XINJSJ、PNGZSJ、KAKUSJ 等。")

    add_table(doc,
        ["序号", "字段名", "中文名称", "类型", "说明"],
        [
            ["1", "FAREDM", "法人代码", "VARCHAR2(4)", "主键之一"],
            ["2", "YNGYJG", "营业机构号", "VARCHAR2(10)", "主键之一"],
            ["3", "YWGXZL", "业务关系种类", "VARCHAR2(10)", "主键之一，ZNGWSJ=账务机构"],
            ["4", "BIZHON", "币种", "VARCHAR2(2)", ""],
            ["5", "YWGXJG", "业务关系机构", "VARCHAR2(10)", ""],
            ["6", "YWGXJB", "业务关系级别", "VARCHAR2(1)", ""],
            ["7", "GXQXJG", "关系权限机构", "VARCHAR2(10)", ""],
            ["8", "SHMING", "说明信息", "VARCHAR2(128)", ""],
            ["9", "BYZD01", "备用字段", "VARCHAR2(200)", ""],
            ["10", "WEIHRQ", "维护日期", "VARCHAR2(8)", ""],
            ["11", "WEIHSJ", "维护时间", "NUMBER(10)", ""],
            ["12", "WEIHGY", "维护柜员", "VARCHAR2(8)", ""],
            ["13", "WEIHJG", "维护机构", "VARCHAR2(10)", ""],
            ["14", "ROWIDD", "序列号", "VARCHAR2(22)", ""],
            ["15", "SHJNCH", "时间戳", "NUMBER(16)", ""],
            ["16", "JILUZT", "记录状态", "VARCHAR2(1)", "'0'表示正常，校验关键字段"],
        ],
        col_widths=[1, 2, 3, 3, 7]
    )

    add_para(doc, "索引设计：", bold=True, indent=False)
    add_bullet(doc, "PK_PJGGX 主键索引：(YNGYJG, YWGXZL, FAREDM)")
    add_bullet(doc, "IDX_PJGGX_1 索引：(FAREDM ASC) - 按法人代码查询场景")
    add_bullet(doc, "IDX_PJGGX_2 索引：(YWGXZL ASC) - 按业务关系种类查询场景")

    add_heading(doc, "3.3 ER 关系说明", level=2)
    add_para(doc, "PJGCS 与 PJGGX 通过 YNGYJG（营业机构号）和 FAREDM（法人代码）建立关联关系。判断机构是否为账务机构的完整 SQL（Spec Q-01 结论）：")
    add_code_block(doc, """SELECT * FROM pjgcs a, pjggx b
WHERE a.yngyjg = b.yngyjg
  AND a.yngyjg = '1628301000'
  AND b.ywgxzl = 'ZNGWSJ'
  AND a.jiluzt = '0'
  AND b.jiluzt = '0';""")

    add_page_break(doc)

    # ===== 第四章 接口设计 =====
    add_heading(doc, "第四章 接口设计", level=1)

    add_heading(doc, "4.1 DAO 层接口", level=2)

    add_para(doc, "4.1.1 PjgcsBranchParamDao 接口", bold=True, indent=False)
    add_para(doc, "文件路径：banks/ext-hnnxbank/hnnxbank-biz-as/src/main/java/com/hundsun/bemp/hnnxbank/biz/sm/dao/branch/PjgcsBranchParamDao.java")
    add_para(doc, "由于该表为外部核心系统同步表，主键为复合主键（YNGYJG + FAREDM），不继承产品化 BaseDao。提供4个方法：")
    add_table(doc,
        ["方法签名", "用途", "异常码"],
        [
            ["PjgcsBranchParam selectByBrchNo(@Param(\"yngyjg\") String yngyjg)", "按营业机构号查询机构参数", "无数据返回null"],
            ["int insertBatch(@Param(\"list\") List<PjgcsBranchParam> list)", "批量插入机构参数数据", "影响行数"],
            ["int deleteAll()", "清空机构参数表数据", "影响行数"],
            ["int countByBrchNoAndStatus(@Param(\"yngyjg\") String yngyjg, @Param(\"jiluzt\") String jiluzt)", "按营业机构号和记录状态查询数量", "记录数"],
        ],
        col_widths=[8, 6, 2]
    )

    add_para(doc, "4.1.2 PjgxBranchRelationDao 接口", bold=True, indent=False)
    add_para(doc, "文件路径：banks/ext-hnnxbank/hnnxbank-biz-as/src/main/java/com/hundsun/bemp/hnnxbank/biz/sm/dao/branch/PjgxBranchRelationDao.java")
    add_para(doc, "由于该表为外部核心系统同步表，主键为复合主键（YNGYJG + YWGXZL + FAREDM），不继承产品化 BaseDao。提供4个方法：")
    add_table(doc,
        ["方法签名", "用途", "异常码"],
        [
            ["List<PjgxBranchRelation> selectByBrchNoAndGxzL(@Param(\"yngyjg\") String yngyjg, @Param(\"ywgxzl\") String ywgxzl)", "按营业机构号和业务关系种类查询机构关系列表", "无数据返回空列表"],
            ["int insertBatch(@Param(\"list\") List<PjgxBranchRelation> list)", "批量插入机构关系数据", "影响行数"],
            ["int deleteAll()", "清空机构关系表数据", "影响行数"],
            ["int countByBrchNoAndGxzLAndStatus(@Param(\"yngyjg\") String yngyjg, @Param(\"ywgxzl\") String ywgxzl, @Param(\"jiluzt\") String jiluzt)", "按营业机构号、业务关系种类和记录状态查询数量", "记录数"],
        ],
        col_widths=[8, 6, 2]
    )

    add_heading(doc, "4.2 Service 层接口", level=2)

    add_para(doc, "4.2.1 PjgcsBranchParamService 接口", bold=True, indent=False)
    add_para(doc, "文件路径：banks/ext-hnnxbank/hnnxbank-biz-as/src/main/java/com/hundsun/bemp/hnnxbank/biz/sm/service/branch/PjgcsBranchParamService.java")
    add_table(doc,
        ["方法签名", "用途", "返回值说明"],
        [
            ["boolean isBranchExistsInMainSystem(String brchNo)", "校验机构在主系统清单中是否存在且状态正常", "true-存在且状态正常"],
            ["String checkBranchStatus(String brchNo)", "校验机构状态，返回详细错误信息", "null-通过；非null-错误信息（含机构号）"],
            ["void syncBatchData(List<PjgcsBranchParam> dataList)", "批量同步机构参数数据（事务）", "无返回值，事务保证原子性"],
        ],
        col_widths=[7, 6, 3]
    )

    add_para(doc, "4.2.2 PjgxBranchRelationService 接口", bold=True, indent=False)
    add_para(doc, "文件路径：banks/ext-hnnxbank/hnnxbank-biz-as/src/main/java/com/hundsun/bemp/hnnxbank/biz/sm/service/branch/PjgxBranchRelationService.java")
    add_table(doc,
        ["方法签名", "用途", "返回值说明"],
        [
            ["boolean isAccountingBranch(String brchNo)", "校验机构是否为账务机构（YWGXZL='ZNGWSJ' 且 JILUZT='0'）", "true-是账务机构"],
            ["void syncBatchData(List<PjgxBranchRelation> dataList)", "批量同步机构关系数据（事务）", "无返回值，事务保证原子性"],
        ],
        col_widths=[7, 6, 3]
    )

    add_heading(doc, "4.3 Controller 层接口", level=2)

    add_para(doc, "4.3.1 建立机构接口", bold=True, indent=False)
    add_para(doc, "文件路径：banks/ext-hnnxbank/hnnxbank-biz-as/src/main/java/com/hundsun/bemp/hnnxbank/biz/sm/controller/branch/HnnxBankBranchController.java")
    add_table(doc,
        ["接口路径", "请求方法", "用途"],
        [
            ["/sm/auth/branch/branch/func_addBranch", "POST/GET", "新增机构，入口增加 PJGCS 主系统机构树清单校验"],
            ["/sm/auth/branch/branch/func_batchImportValidate", "POST", "批量导入预校验，每行均调用 checkBranchStatus 校验"],
        ],
        col_widths=[8, 3, 5]
    )
    add_para(doc, "异常码规范：", bold=True, indent=False)
    add_bullet(doc, "机构不存在：抛出 BempRuntimeException(VALID_FAIL, \"机构[9999999999]在核心系统机构树中不存在，不能创建\")")
    add_bullet(doc, "状态异常：抛出 BempRuntimeException(VALID_FAIL, \"机构[1628301999]在核心系统状态异常，不能创建\")")
    add_bullet(doc, "批量导入：行级错误聚合，使用分号分隔，格式 \"第N行机构[9999999999]在核心系统机构树中不存在，不能创建；\"")

    add_para(doc, "4.3.2 贴现业务提交接口", bold=True, indent=False)
    add_para(doc, "文件路径：banks/ext-hnnxbank/hnnxbank-biz-as/src/main/java/com/hundsun/bemp/hnnxbank/biz/ce/disc/service/impl/HnnxBankDiscBillServiceImpl.java")
    add_table(doc,
        ["方法", "用途", "异常码"],
        [
            ["submitElecFlow(BaseRequest<DiscBillInfoDto> baseRequest)", "贴现业务提交，入口增加账务机构校验", "抛出 BempRuntimeException(\"业务发起机构非账务机构，无法发起贴现\")"],
        ],
        col_widths=[7, 6, 3]
    )
    add_para(doc, "异常处理细节：", bold=True, indent=False)
    add_bullet(doc, "错误信息严格按需求文案，避免使用 VALID_FAIL 错误码自动追加\"参数校验不通过:\"前缀")
    add_bullet(doc, "仅贴现业务提交时校验，其他业务（承兑、质押等）不受影响（spec 明确\"仅此一个业务场景校验\"）")

    add_page_break(doc)

    # ===== 第五章 详细设计 =====
    add_heading(doc, "第五章 详细设计", level=1)

    add_heading(doc, "5.1 CBS 文件读取工具设计", level=2)
    add_para(doc, "文件路径：banks/ext-hnnxbank/hnnxbank-biz-as/src/main/java/com/hundsun/bemp/hnnxbank/biz/sm/utils/CbsFileReaderUtil.java")
    add_para(doc, "CbsFileReaderUtil 是 CBS 核心系统文件读取工具类，用于定时任务日初读取核心系统下发的机构参数文件（CBS_PJGCS）和机构关系文件（CBS_PJGGX），并解析为对应的 Entity 列表供下游同步使用。")

    add_para(doc, "文件特征：", bold=True, indent=False)
    add_bullet(doc, "编码：GBK（核心系统统一使用 GBK 编码，避免中文乱码）")
    add_bullet(doc, "字段分隔符：0x03（ETX 控制字符，核心系统 CBS 文件的标准分隔约定，参考 BOPC010101MessageConverter）")
    add_bullet(doc, "行分隔符：0x0A（LF）")
    add_bullet(doc, "字段顺序：严格按表字段说明.txt 中定义的顺序排列，PJGCS 共 36 个字段，PJGGX 共 16 个字段")

    add_para(doc, "文件路径规则：", bold=True, indent=False)
    add_code_block(doc, "{cbsFilePath}{fileDate}/CBS_{FILE_TYPE}_{fileDate}.txt\n其中 fileDate 为票据系统营业日期前一天（格式 yyyyMMdd），由调用方计算后传入")

    add_para(doc, "降级策略：", bold=True, indent=False)
    add_table(doc,
        ["异常场景", "处理策略", "日志级别"],
        [
            ["文件不存在", "返回空列表，不抛异常（核心系统可能未下发当日文件，如非工作日）", "WARN"],
            ["文件为空", "返回空列表", "WARN"],
            ["单行解析异常", "跳过该行继续解析后续记录，记录行号便于定位", "ERROR"],
            ["Long 类型字段解析失败", "返回 null，避免单行数据问题中断整体解析", "WARN"],
            ["读取异常", "降级返回空列表，避免中断日初定时任务", "ERROR"],
        ],
        col_widths=[4, 9, 3]
    )

    add_para(doc, "关键技术点：", bold=True, indent=False)
    add_bullet(doc, "使用 split(regex, -1) 保留尾部空字段，避免丢失末尾空字段导致数组越界")
    add_bullet(doc, "Pattern.quote 预编译分隔符，按字面量匹配，避免正则歧义")
    add_bullet(doc, "try-with-resources 确保 InputStreamReader 和 BufferedReader 在异常时也能正确关闭")
    add_bullet(doc, "字段数不足时抛 IllegalArgumentException，便于定位问题数据")

    add_heading(doc, "5.2 日初定时任务设计", level=2)

    add_para(doc, "5.2.1 SyncPjgcsBranchParamJobServiceImpl", bold=True, indent=False)
    add_para(doc, "接口路径：banks/ext-hnnxbank/hnnxbank-biz-api/src/main/java/com/hundsun/bemp/hnnxbank/biz/tk/job/service/SyncPjgcsBranchParamJobService.java")
    add_para(doc, "实现路径：banks/ext-hnnxbank/hnnxbank-biz-as/src/main/java/com/hundsun/bemp/hnnxbank/biz/tk/job/service/impl/SyncPjgcsBranchParamJobServiceImpl.java")
    add_para(doc, "定时任务编号：HNNXTK020103，cron 表达式：0 0 06 * * ?（每日6:00执行），在额度失效任务（6:30）之前执行完成。")

    add_para(doc, "执行流程：", bold=True, indent=False)
    add_code_block(doc, """1. 调用 busiDateService.viewBusiDate() 获取营业日信息
2. resolveFileDate: 优先使用上一营业日(preWorkday)，为空时降级使用当前营业日(workday)
3. 调用 cbsFileReaderUtil.readPjgcsFile(fileDate) 读取并解析 CBS 文件
4. 文件无有效数据时降级返回，不抛异常，避免中断其他日初任务
5. 调用 pjgcsBranchParamService.syncBatchData(dataList) 同步数据
   - @Transactional(rollbackFor = Exception.class) 保证事务原子性
   - 先 deleteAll 清空表数据
   - 再 insertBatch 批量插入新数据
6. 记录同步完成的日志""")

    add_para(doc, "5.2.2 SyncPjgxBranchRelationJobServiceImpl", bold=True, indent=False)
    add_para(doc, "接口路径：banks/ext-hnnxbank/hnnxbank-biz-api/src/main/java/com/hundsun/bemp/hnnxbank/biz/tk/job/service/SyncPjgxBranchRelationJobService.java")
    add_para(doc, "实现路径：banks/ext-hnnxbank/hnnxbank-biz-as/src/main/java/com/hundsun/bemp/hnnxbank/biz/tk/job/service/impl/SyncPjgxBranchRelationJobServiceImpl.java")
    add_para(doc, "定时任务编号：HNNXTK020104，cron 表达式：0 0 06 * * ?（与 PJGCS 同步任务并行执行，互不依赖）。")
    add_para(doc, "执行流程与 5.2.1 对称，区别仅在于读取 CBS_PJGGX 文件、同步到 PJGGX 表。")

    add_para(doc, "5.2.3 调度配置", bold=True, indent=False)
    add_para(doc, "SQL 脚本路径：deploy/bemp-script/src/main/resources/banks/河南农信/V202301.03.080_202607151600_T202607150002_日初同步机构参数及机构关系数据.dml.sql")
    add_para(doc, "调度配置写入 tt_task 表，采用先删除后新增策略，保证脚本幂等可重复执行：")
    add_code_block(doc, """-- 任务1: HNNXTK020103 - 日初同步PJGCS机构参数数据
INSERT INTO tt_task (ID, TASK_NO, TASK_NAME, SEQ_NO, REPEAT_FLAG, DELAY_TM,
                     TIMING_FLAG, CRON_EXPRESSION, PROCESS_STATUS, FUNCTION_ID,
                     IS_SKIP_HOLIDAY)
VALUES (3440, 'HNNXTK020103', '日初同步PJGCS机构参数数据', 3440, '1', 0,
        '1', '0 0 06 * * ?', '0', 'HNNXTK020103', '0');

-- 任务2: HNNXTK020104 - 日初同步PJGGX机构关系数据
INSERT INTO tt_task (ID, TASK_NO, TASK_NAME, SEQ_NO, REPEAT_FLAG, DELAY_TM,
                     TIMING_FLAG, CRON_EXPRESSION, PROCESS_STATUS, FUNCTION_ID,
                     IS_SKIP_HOLIDAY)
VALUES (3441, 'HNNXTK020104', '日初同步PJGGX机构关系数据', 3441, '1', 0,
        '1', '0 0 06 * * ?', '0', 'HNNXTK020104', '0');""")

    add_heading(doc, "5.3 建立机构校验设计", level=2)
    add_para(doc, "文件路径：banks/ext-hnnxbank/hnnxbank-biz-as/src/main/java/com/hundsun/bemp/hnnxbank/biz/sm/controller/branch/HnnxBankBranchController.java")

    add_para(doc, "5.3.1 单个新增机构校验", bold=True, indent=False)
    add_para(doc, "在 addBranch 方法入口增加 PJGCS 主系统机构树清单校验：")
    add_code_block(doc, """// 主系统机构树清单校验：校验失败时直接抛出异常，避免脏数据进入机构表
// 区分"不存在"与"状态异常"两种场景，便于操作人精确定位问题
String brchNo = branchDto.getBrchNo();
if (StringUtils.isNotBlank(brchNo)) {
    String errorMsg = pjgcsBranchParamService.checkBranchStatus(brchNo);
    if (StringUtils.isNotBlank(errorMsg)) {
        throw new BempRuntimeException(CommonErrorNoConst.VALID_FAIL, errorMsg);
    }
}""")

    add_para(doc, "5.3.2 批量导入校验", bold=True, indent=False)
    add_para(doc, "在 batchImportValidate 方法的 validateBranchRow 行级校验函数中增加 PJGCS 校验：")
    add_code_block(doc, """// 主系统机构树清单校验：每个导入机构都需校验在 PJGCS 表中存在且状态正常
String branchStatusMsg = pjgcsBranchParamService.checkBranchStatus(branchDto.getBrchNo());
if (StringUtils.isNotBlank(branchStatusMsg)) {
    joiner.add(String.format("第%d行%s；", rowNum, branchStatusMsg));
    return null;
}""")

    add_para(doc, "5.3.3 校验返回的错误信息规范", bold=True, indent=False)
    add_table(doc,
        ["场景", "错误信息格式", "示例"],
        [
            ["机构不存在", "机构[%s]在核心系统机构树中不存在，不能创建", "机构[9999999999]在核心系统机构树中不存在，不能创建"],
            ["状态异常", "机构[%s]在核心系统状态异常，不能创建", "机构[1628301999]在核心系统状态异常，不能创建"],
            ["批量导入行级错误", "第%d行%s；", "第2行机构[9999999999]在核心系统机构树中不存在，不能创建；"],
        ],
        col_widths=[3, 8, 5]
    )

    add_heading(doc, "5.4 贴现业务校验设计", level=2)
    add_para(doc, "文件路径：banks/ext-hnnxbank/hnnxbank-biz-as/src/main/java/com/hundsun/bemp/hnnxbank/biz/ce/disc/service/impl/HnnxBankDiscBillServiceImpl.java")
    add_para(doc, "在 submitElecFlow 方法入口（BempValidUtil.validBaseRequest 之后）增加账务机构校验：")
    add_code_block(doc, """// 校验业务发起机构是否为账务机构，仅贴现业务提交时校验
// 业务规则：机构在 PJGGX 表中存在 YWGXZL='ZNGWSJ' 且 JILUZT='0' 的记录才允许发起贴现
// 错误信息严格按需求文案，避免使用 VALID_FAIL 错误码自动追加"参数校验不通过:"前缀
String launchBrchNo = baseRequest.getReqBrchNo();
if (!pjgxBranchRelationService.isAccountingBranch(launchBrchNo)) {
    logger.info("贴现业务提交校验：业务发起机构[{}]非账务机构，不允许发起贴现", launchBrchNo);
    throw new BempRuntimeException("业务发起机构非账务机构，无法发起贴现");
}""")

    add_para(doc, "设计要点：", bold=True, indent=False)
    add_bullet(doc, "使用 BempRuntimeException 单参构造函数，避免 VALID_FAIL 错误码自动追加前缀，确保错误信息严格按需求文案")
    add_bullet(doc, "校验位置在 validBaseRequest 之后、业务逻辑处理之前，避免无效业务数据进入后续流程")
    add_bullet(doc, "仅贴现业务 submitElecFlow 方法增加校验，其他业务场景不受影响")

    add_heading(doc, "5.5 安全缺陷修复设计", level=2)

    add_para(doc, "5.5.1 批量新增机构管理员法人管理员校验（P-05 安全修复）", bold=True, indent=False)
    add_para(doc, "文件路径：banks/ext-hnnxbank/hnnxbank-biz-as/src/main/java/com/hundsun/bemp/hnnxbank/biz/sm/controller/branch/HnnxbankBranchAdminController.java")
    add_para(doc, "在 batchImportBranchAdmin 方法入口增加 userType=\"4\"（法人管理员）校验，与批量删除和批量复制角色保持一致：")
    add_code_block(doc, """// 权限校验：仅法人管理员可操作，与批量删除、批量复制角色保持一致
if (!AuthConstant.S_FOUR.equals(userInfo.getUserType())) {
    throw new BempRuntimeException(CommonErrorNoConst.VALID_FAIL,
        "只有法人管理员可以批量新增机构管理员");
}""")

    add_para(doc, "5.5.2 密码有效期参数化", bold=True, indent=False)
    add_para(doc, "原 batchImportBranchAdmin 方法第282行硬编码 pwdExpiration = 90。修改为从系统参数读取，参数 key 为 branch_admin_pwd_expiration，未配置时回退默认值 90 天。")
    add_code_block(doc, """// 密码有效期从系统参数读取，未配置时回退默认值90天并记录告警日志便于运维定位
Integer pwdExpiration = getIntegerParamWithDefault(userInfo.getLegalNo(),
        "branch_admin_pwd_expiration", 90, "密码有效期");""")

    add_para(doc, "5.5.3 密码错误次数参数化", bold=True, indent=False)
    add_para(doc, "原 batchImportBranchAdmin 方法第285行硬编码 pwdErrTimes = 10。修改为从系统参数读取，参数 key 为 branch_admin_pwd_err_times，未配置时回退默认值 10 次。")
    add_code_block(doc, """// 密码错误次数从系统参数读取，未配置时回退默认值10次并记录告警日志便于运维定位
Integer pwdErrTimes = getIntegerParamWithDefault(userInfo.getLegalNo(),
        "branch_admin_pwd_err_times", 10, "密码错误次数");""")

    add_para(doc, "5.5.4 机构不具备角色精确报错", bold=True, indent=False)
    add_para(doc, "原 validateBranchRoles 方法第610-611行仅 LOGGER.warn 记录日志，最终统一返回\"角色校验未通过\"。修改为按需求文案\"机构【XX】无角色【XX】，不能分配\"返回精确错误信息给前端。")
    add_code_block(doc, """String roleName = StringUtils.isBlank(sourceRole.getRoleName())
        ? String.valueOf(sourceRole.getId()) : sourceRole.getRoleName();
String errorMsg = String.format("机构【%s】无角色【%s】，不能分配", targetBrchNo, roleName);
LOGGER.warn("{}. 目标用户【{}】", errorMsg, targetUserNo);
invalidRoleMsgs.add(errorMsg);""")

    add_para(doc, "错误聚合策略：", bold=True, indent=False)
    add_bullet(doc, "多个不具备的角色错误聚合到 invalidRoleMsgs 列表")
    add_bullet(doc, "使用分号\";\"分隔，一次性返回，避免逐个提示增加交互成本")
    add_bullet(doc, "示例：\"机构【1628301000】无角色【角色2】，不能分配；机构【1628301000】无角色【角色3】，不能分配\"")

    add_para(doc, "5.5.5 批量新增身份验证切换为 POBM010304", bold=True, indent=False)
    add_para(doc, "Spec Q-04 结论：必须使用 POBM010304MessageConverter（同步机构柜员信息）核验管理员身份。批量新增机构管理员时，向我行统一身份认证系统请求验证，核验管理员身份信息是否存在。")

    add_page_break(doc)

    # ===== 第六章 部署设计 =====
    add_heading(doc, "第六章 部署设计", level=1)

    add_heading(doc, "6.1 数据库脚本部署", level=2)
    add_table(doc,
        ["脚本类型", "脚本路径", "执行顺序", "说明"],
        [
            ["DDL 建表", "deploy/bemp-script/src/main/resources/banks/河南农信/V202301.03.079_202607151500_T202607150001_机构参数及机构关系表建表.ddl.sql", "1", "创建 PJGCS 和 PJGGX 表，幂等可重复执行"],
            ["DML 调度", "deploy/bemp-script/src/main/resources/banks/河南农信/V202301.03.080_202607151600_T202607150002_日初同步机构参数及机构关系数据.dml.sql", "2", "写入 tt_task 表两个定时任务调度记录"],
        ],
        col_widths=[2, 9, 1, 4]
    )

    add_heading(doc, "6.2 应用配置部署", level=2)
    add_para(doc, "需在应用配置文件中配置 CBS 文件根路径：")
    add_code_block(doc, """# CBS 文件根路径配置（application.properties 或对应环境配置文件）
cbsFile.cbsFilePath=/data/bemp/cbs/""")

    add_para(doc, "需在系统参数表中配置以下参数（可选，未配置时使用默认值）：")
    add_table(doc,
        ["参数 Key", "默认值", "用途"],
        [
            ["branch_admin_pwd_expiration", "90", "批量新增机构管理员密码有效期（天）"],
            ["branch_admin_pwd_err_times", "10", "批量新增机构管理员密码错误次数限制"],
            ["branch_admin_init_pwd", "-", "批量新增机构管理员初始密码"],
        ],
        col_widths=[6, 2, 8]
    )

    add_heading(doc, "6.3 CBS 文件部署", level=2)
    add_para(doc, "ODS 在日初将 CBS 文件下发到票据系统的 cbs 日期目录下，目录结构如下：")
    add_code_block(doc, """/data/bemp/cbs/
├── 20260714/                      # 营业日期前一天目录
│   ├── CBS_PJGCS_20260714.txt     # 机构参数文件（GBK编码，0x03分隔符）
│   └── CBS_PJGGX_20260714.txt     # 机构关系文件（GBK编码，0x03分隔符）
└── 20260715/
    ├── CBS_PJGCS_20260715.txt
    └── CBS_PJGGX_20260715.txt""")

    add_heading(doc, "6.4 服务部署", level=2)
    add_table(doc,
        ["服务", "端口", "依赖", "启动顺序"],
        [
            ["Redis", "6379", "无", "1"],
            ["ZooKeeper", "2181", "无", "2"],
            ["后端服务（SpringBoot）", "8010", "Redis + ZooKeeper + 数据库", "3"],
            ["前端服务（Vue.js）", "8091", "后端服务", "4"],
        ],
        col_widths=[5, 2, 6, 3]
    )

    add_page_break(doc)

    # ===== 第七章 偏差记录与需求确认闭环 =====
    add_heading(doc, "第七章 偏差记录与需求确认闭环", level=1)

    add_heading(doc, "7.1 需求与实现偏差记录", level=2)
    add_table(doc,
        ["偏差项", "需求描述", "实际实现", "偏差原因", "影响评估"],
        [
            ["身份验证接口选择", "参考 POBM010304MessageConverter", "切换为 POBM010304MessageConverter", "Spec Q-04 结论：必须使用 POBM010304MessageConverter", "无影响，按需求实现"],
            ["账务机构判断标准", "YWGXZL 取值未明确", "YWGXZL='ZNGWSJ' 且 JILUZT='0'", "Spec Q-01 结论：判断账务机构的完整 SQL 示例", "无影响，按行方确认结论实现"],
            ["CBS 文件编码格式", "未明确", "GBK 编码、0x03 分隔符、定长字段", "Spec Q-02 结论：参考 BOPC010101MessageConverter", "无影响，按行方确认结论实现"],
            ["定时任务调度时间", "未明确具体时间", "每日6:00执行（cron: 0 0 06 * * ?）", "Spec Q-03 结论：参考现有日初任务调度时间", "无影响，上线后可手动调整"],
            ["文件不存在降级处理", "spec 描述\"抛异常中断\"", "实际实现返回空列表降级", "考虑核心系统非工作日不下发文件的场景，避免中断其他日初任务", "降级处理更稳健，已在日志中记录告警"],
        ],
        col_widths=[2, 3, 4, 4, 3]
    )

    add_heading(doc, "7.2 需求确认闭环结果", level=2)
    add_para(doc, "Spec 文档中标记的4个待确认事项（Q-01 ~ Q-04）已全部闭环，结论如下：")
    add_table(doc,
        ["问题编号", "问题描述", "最终结论", "实现状态"],
        [
            ["Q-01", "账务机构的判断标准", "YWGXZL='ZNGWSJ' 且 JILUZT='0'，完整 SQL 已在 Spec 中明确", "已实现"],
            ["Q-02", "CBS 文件编码格式", "GBK 编码、定长字段（参考 BOPC010101MessageConverter）", "已实现"],
            ["Q-03", "定时任务调度时间", "参考现有日初任务调度时间，后续需上线可再手动调整配置", "已实现（cron: 0 0 06 * * ?）"],
            ["Q-04", "批量新增身份验证接口选择", "必须使用 POBM010304MessageConverter", "已实现"],
        ],
        col_widths=[2, 4, 8, 2]
    )

    add_page_break(doc)

    # ===== 第八章 已知问题清单 =====
    add_heading(doc, "第八章 已知问题清单", level=1)
    add_para(doc, "截至本文档交付时，P0/P1 缺陷已全部修复验证通过。剩余 P2 缺陷作为已知问题列出，不影响功能交付，后续在对应环境可用时修复验证。")

    add_table(doc,
        ["缺陷编号", "严重度", "缺陷描述", "根因", "影响评估", "后续计划"],
        [
            ["BUG-001", "P1", "建立机构校验错误信息返回时被追加前缀", "使用 VALID_FAIL 错误码导致框架自动追加\"参数校验不通过:\"前缀", "影响前端精确展示错误信息", "已修复验证通过（改用单参构造函数）"],
            ["BUG-002", "P2", "前端编译环境异常导致部分前端用例无法执行", "前端开发环境 Node.js 版本/依赖冲突", "影响前端界面用例执行，不影响后端功能", "待前端环境恢复后补充执行"],
            ["BUG-003", "P2", "POBM010304 接口环境不可连导致身份验证用例阻塞", "测试环境统一身份认证系统未就绪", "影响 POBM010304 身份验证用例执行，不影响其他功能", "待统一身份认证系统环境就绪后补充执行"],
        ],
        col_widths=[2, 1, 4, 3, 3, 3]
    )

    add_para(doc, "已知问题影响评估总结：", bold=True, indent=False)
    add_bullet(doc, "BUG-001（P1）已修复验证通过，不影响交付")
    add_bullet(doc, "BUG-002 和 BUG-003 均为 P2 级别，根因为环境问题，非代码缺陷")
    add_bullet(doc, "P2 缺陷不影响核心功能（数据库表、定时任务、建立机构校验、贴现业务校验、安全缺陷修复）的交付")
    add_bullet(doc, "二轮调试测试14项全部通过，验证了核心功能正确性")

    add_page_break(doc)

    # ===== 第九章 部署验证清单 =====
    add_heading(doc, "第九章 部署验证清单", level=1)
    add_para(doc, "上线部署后，按以下清单逐项验证：")

    add_table(doc,
        ["序号", "验证项", "验证方法", "预期结果"],
        [
            ["1", "PJGCS 表创建", "SELECT TABLE_NAME FROM USER_TABLES WHERE TABLE_NAME='PJGCS'", "返回1行记录"],
            ["2", "PJGGX 表创建", "SELECT TABLE_NAME FROM USER_TABLES WHERE TABLE_NAME='PJGGX'", "返回1行记录"],
            ["3", "PJGCS 字段数验证", "SELECT COUNT(*) FROM USER_TAB_COLUMNS WHERE TABLE_NAME='PJGCS'", "返回36"],
            ["4", "PJGGX 字段数验证", "SELECT COUNT(*) FROM USER_TAB_COLUMNS WHERE TABLE_NAME='PJGGX'", "返回16"],
            ["5", "定时任务注册", "SELECT TASK_NO, TASK_NAME FROM TT_TASK WHERE TASK_NO IN ('HNNXTK020103','HNNXTK020104')", "返回2行记录"],
            ["6", "CBS 文件配置", "检查 application.properties 中 cbsFile.cbsFilePath 配置", "路径配置正确且目录存在"],
            ["7", "日初任务执行", "手动触发 HNNXTK020103 和 HNNXTK020104，查看后端日志", "任务执行成功，PJGCS/PJGGX 表有数据"],
            ["8", "建立机构校验", "新建一个 PJGCS 中不存在的机构号", "提示\"机构[9999999999]在核心系统机构树中不存在，不能创建\""],
            ["9", "贴现业务校验", "使用非账务机构发起贴现业务提交", "提示\"业务发起机构非账务机构，无法发起贴现\""],
            ["10", "法人管理员校验", "使用机构管理员账号尝试批量新增机构管理员", "提示\"只有法人管理员可以批量新增机构管理员\""],
        ],
        col_widths=[1, 4, 7, 4]
    )

    # ===== 附录 =====
    add_page_break(doc)
    add_heading(doc, "附录", level=1)

    add_heading(doc, "附录 A：交付物清单", level=2)
    add_table(doc,
        ["序号", "交付物", "格式", "路径"],
        [
            ["1", "详细设计文档", ".docx", "本文件"],
            ["2", "测试用例文档", ".md", "aotutests-devtools/sync-org-tree-validation.md"],
            ["3", "测试报告", ".docx", "河南农商同步机构树数据并校验-测试报告-20260715.docx"],
            ["4", "需求文档(PRD)", ".md", "docs/prd/机构管理优化/河南农商银行关于新一代票据系统机构管理及机构管理员功能变更的需求 (20260707).md"],
            ["5", "Spec 文档", ".md", ".trae/specs/sync-org-tree-validation/spec.md"],
        ],
        col_widths=[1, 3, 2, 10]
    )

    add_heading(doc, "附录 B：参考资料", level=2)
    add_bullet(doc, "《河南农商银行关于新一代票据系统机构管理及机构管理员功能变更的需求》（2026年06月18日）")
    add_bullet(doc, "BEMP5.0 项目编码规范")
    add_bullet(doc, "BOPC010101MessageConverter（同步核心系统贷款余额信息）- 文件读取逻辑参考")
    add_bullet(doc, "POBM010304MessageConverter（同步机构柜员信息）- 身份验证接口参考")
    add_bullet(doc, "核心系统表字段说明.txt - PJGCS/PJGGX 表字段定义参考")

    add_heading(doc, "附录 C：代码文件清单", level=2)
    add_table(doc,
        ["序号", "文件类型", "文件路径", "说明"],
        [
            ["1", "Entity", "banks/ext-hnnxbank/hnnxbank-biz-as/src/main/java/com/hundsun/bemp/hnnxbank/biz/sm/dao/branch/entity/PjgcsBranchParam.java", "PJGCS 机构参数实体（36字段）"],
            ["2", "Entity", "banks/ext-hnnxbank/hnnxbank-biz-as/src/main/java/com/hundsun/bemp/hnnxbank/biz/sm/dao/branch/entity/PjgxBranchRelation.java", "PJGGX 机构关系实体（16字段）"],
            ["3", "DAO", "banks/ext-hnnxbank/hnnxbank-biz-as/src/main/java/com/hundsun/bemp/hnnxbank/biz/sm/dao/branch/PjgcsBranchParamDao.java", "PJGCS 数据访问接口"],
            ["4", "DAO", "banks/ext-hnnxbank/hnnxbank-biz-as/src/main/java/com/hundsun/bemp/hnnxbank/biz/sm/dao/branch/PjgxBranchRelationDao.java", "PJGGX 数据访问接口"],
            ["5", "Mapper XML", "banks/ext-hnnxbank/hnnxbank-biz-as/src/main/resources/com/hundsun/bemp/hnnxbank/biz/sm/dao/branch/PjgcsBranchParamDao.xml", "PJGCS Mapper XML"],
            ["6", "Mapper XML", "banks/ext-hnnxbank/hnnxbank-biz-as/src/main/resources/com/hundsun/bemp/hnnxbank/biz/sm/dao/branch/PjgxBranchRelationDao.xml", "PJGGX Mapper XML"],
            ["7", "Service", "banks/ext-hnnxbank/hnnxbank-biz-as/src/main/java/com/hundsun/bemp/hnnxbank/biz/sm/service/branch/PjgcsBranchParamService.java", "PJGCS 服务接口"],
            ["8", "Service", "banks/ext-hnnxbank/hnnxbank-biz-as/src/main/java/com/hundsun/bemp/hnnxbank/biz/sm/service/branch/PjgxBranchRelationService.java", "PJGGX 服务接口"],
            ["9", "ServiceImpl", "banks/ext-hnnxbank/hnnxbank-biz-as/src/main/java/com/hundsun/bemp/hnnxbank/biz/sm/service/impl/branch/PjgcsBranchParamServiceImpl.java", "PJGCS 服务实现"],
            ["10", "ServiceImpl", "banks/ext-hnnxbank/hnnxbank-biz-as/src/main/java/com/hundsun/bemp/hnnxbank/biz/sm/service/impl/branch/PjgxBranchRelationServiceImpl.java", "PJGGX 服务实现"],
            ["11", "Util", "banks/ext-hnnxbank/hnnxbank-biz-as/src/main/java/com/hundsun/bemp/hnnxbank/biz/sm/utils/CbsFileReaderUtil.java", "CBS 文件读取工具类"],
            ["12", "Job Service", "banks/ext-hnnxbank/hnnxbank-biz-api/src/main/java/com/hundsun/bemp/hnnxbank/biz/tk/job/service/SyncPjgcsBranchParamJobService.java", "PJGCS 同步任务接口"],
            ["13", "Job Service", "banks/ext-hnnxbank/hnnxbank-biz-api/src/main/java/com/hundsun/bemp/hnnxbank/biz/tk/job/service/SyncPjgxBranchRelationJobService.java", "PJGGX 同步任务接口"],
            ["14", "Job Impl", "banks/ext-hnnxbank/hnnxbank-biz-as/src/main/java/com/hundsun/bemp/hnnxbank/biz/tk/job/service/impl/SyncPjgcsBranchParamJobServiceImpl.java", "PJGCS 同步任务实现"],
            ["15", "Job Impl", "banks/ext-hnnxbank/hnnxbank-biz-as/src/main/java/com/hundsun/bemp/hnnxbank/biz/tk/job/service/impl/SyncPjgxBranchRelationJobServiceImpl.java", "PJGGX 同步任务实现"],
            ["16", "Controller", "banks/ext-hnnxbank/hnnxbank-biz-as/src/main/java/com/hundsun/bemp/hnnxbank/biz/sm/controller/branch/HnnxBankBranchController.java", "机构 Controller（建立机构校验）"],
            ["17", "Controller", "banks/ext-hnnxbank/hnnxbank-biz-as/src/main/java/com/hundsun/bemp/hnnxbank/biz/sm/controller/branch/HnnxbankBranchAdminController.java", "机构管理员 Controller（安全缺陷修复）"],
            ["18", "ServiceImpl", "banks/ext-hnnxbank/hnnxbank-biz-as/src/main/java/com/hundsun/bemp/hnnxbank/biz/ce/disc/service/impl/HnnxBankDiscBillServiceImpl.java", "贴现业务实现（账务机构校验）"],
            ["19", "DDL", "deploy/bemp-script/src/main/resources/banks/河南农信/V202301.03.079_202607151500_T202607150001_机构参数及机构关系表建表.ddl.sql", "建表脚本"],
            ["20", "DML", "deploy/bemp-script/src/main/resources/banks/河南农信/V202301.03.080_202607151600_T202607150002_日初同步机构参数及机构关系数据.dml.sql", "定时任务调度脚本"],
        ],
        col_widths=[1, 2, 9, 4]
    )

    # 注入 updateFields，使打开文档时自动更新目录
    inject_update_fields(doc)

    # 保存
    os.makedirs(OUTPUT_DIR, exist_ok=True)
    doc.save(OUTPUT_FILE)
    print(f"[OK] 详细设计文档已生成：{OUTPUT_FILE}")


if __name__ == "__main__":
    build_design_doc()
