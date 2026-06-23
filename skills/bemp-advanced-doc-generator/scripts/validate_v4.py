"""验证v4文档完整性"""
from docx import Document

v4_path = r"d:\code\QJ\BEMP5.0DEV\.trae\skills\bemp-advanced-doc-generator\output\机构管理和管理员管理功能优化-详细设计文档-v4-20260617.docx"
doc = Document(v4_path)

print("=== v4文档验证 ===\n")

# 统计信息
headings = [p for p in doc.paragraphs if p.style.name.startswith('Heading')]
tables = doc.tables
paragraphs = [p for p in doc.paragraphs if p.text.strip()]

print(f"标题数量: {len(headings)}")
print(f"表格数量: {len(tables)}")
print(f"段落数量: {len(paragraphs)}")

# 检查关键内容
print("\n=== 关键内容检查 ===")
full_text = '\n'.join([p.text for p in doc.paragraphs])

checks = {
    "批量导入机构": "批量导入机构" in full_text,
    "批量导入管理员": "批量导入管理员" in full_text,
    "模板下载": "模板下载" in full_text or "模版下载" in full_text,
    "批量复制角色": "批量复制角色" in full_text,
    "BranchImportVo": "BranchImportVo" in full_text,
    "BranchAdminImportVo": "BranchAdminImportVo" in full_text,
    "HnnxBatchCopyRoleReq": "HnnxBatchCopyRoleReq" in full_text,
}

for key, value in checks.items():
    print(f"{key}: {'✓' if value else '✗'}")

# 检查API接口表格
print("\n=== API接口表格检查 ===")
api_table_found = False
for i, table in enumerate(tables):
    if len(table.rows) > 0 and len(table.rows[0].cells) >= 3:
        first_row = [cell.text.strip() for cell in table.rows[0].cells]
        if '接口路径' in first_row[0] or '接口' in str(first_row):
            api_table_found = True
            print(f"表格{i}: 找到API接口表格")
            print(f"  - 行数: {len(table.rows)}")
            if len(table.rows) >= 8:
                print(f"  - ✓ 包含7个接口定义")
            break

if not api_table_found:
    print("✗ 未找到API接口表格")

# 检查错误码表格
print("\n=== 错误码表格检查 ===")
error_table_found = False
for i, table in enumerate(tables):
    if len(table.rows) > 0 and len(table.rows[0].cells) >= 4:
        first_row = [cell.text.strip() for cell in table.rows[0].cells]
        if '错误码' in str(first_row) and ('错误信息' in str(first_row) or '说明' in str(first_row)):
            error_table_found = True
            print(f"表格{i}: 找到错误码表格")
            print(f"  - 行数: {len(table.rows)}")
            if len(table.rows) >= 7:
                print(f"  - ✓ 包含6个错误码定义")
            break

if not error_table_found:
    print("✗ 未找到错误码表格")

# 检查栏位描述表格
print("\n=== 栏位描述表格检查 ===")
field_table_found = False
for i, table in enumerate(tables):
    if len(table.rows) > 0 and len(table.rows[0].cells) == 2:
        first_row = [cell.text.strip() for cell in table.rows[0].cells]
        if '数据名称' in str(first_row[0]):
            field_table_found = True
            print(f"表格{i}: 找到栏位描述表格")
            print(f"  - 行数: {len(table.rows)}")
            if len(table.rows) >= 15:
                print(f"  - ✓ 包含15个栏位定义")
            break

if not field_table_found:
    print("✗ 未找到栏位描述表格")

# 检查字段映射表格
print("\n=== 字段映射表格检查 ===")
field_mapping_found = False
for i, table in enumerate(tables):
    if len(table.rows) >= 5 and len(table.rows[0].cells) == 6:
        first_row = [cell.text.strip() for cell in table.rows[0].cells]
        if '字段名称' in first_row[0]:
            field_mapping_found = True
            print(f"表格{i}: 找到字段映射表格")
            print(f"  - 行数: {len(table.rows)}")
            if len(table.rows) >= 5:
                print(f"  - ✓ 包含4个字段定义")
            break

if not field_mapping_found:
    print("✗ 未找到字段映射表格")

print("\n=== 文档结构 ===")
h1_headings = [p.text.strip() for p in headings if p.style.name == 'Heading 1']
for h1 in h1_headings:
    print(f"  {h1}")

print("\n✓ v4文档验证完成")
