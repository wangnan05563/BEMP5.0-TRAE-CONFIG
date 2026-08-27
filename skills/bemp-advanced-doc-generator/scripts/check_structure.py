from pathlib import Path
SKILL_ROOT = Path(__file__).resolve().parent.parent
PROJECT_ROOT = SKILL_ROOT.parent.parent.parent

from pathlib import Path
"""检查文档实际结构"""
from docx import Document

doc_path = str(SKILL_ROOT / "output" / "机构管理和管理员管理功能优化-详细设计文档-20260617.docx")
doc = Document(doc_path)

print("=" * 70)
print("文档结构检查")
print("=" * 70)

# 遍历所有段落，打印标题结构
for i, para in enumerate(doc.paragraphs):
    style_name = para.style.name if para.style else "None"
    text = para.text.strip()
    
    if text and ('Heading' in style_name or style_name in ['1', '2', '3']):
        indent = "  " * (int(style_name) if style_name.isdigit() else 0)
        print(f"{indent}[{style_name}] {text[:60]}")
    
    # 检查表格前的段落
    if i < len(doc.paragraphs) - 1:
        next_para = doc.paragraphs[i + 1]
        if hasattr(next_para, '_element') and next_para._element.tag.endswith('tbl'):
            print(f"  [表格] 在段落 {i} 之后")

print("\n" + "=" * 70)
print("表格统计")
print("=" * 70)
print(f"总表格数: {len(doc.tables)}")

for i, table in enumerate(doc.tables):
    rows = len(table.rows)
    cols = len(table.columns)
    first_cell = table.rows[0].cells[0].text[:30] if rows > 0 and cols > 0 else ""
    print(f"表格 {i+1}: {rows}行 x {cols}列, 首单元格: {first_cell}")
