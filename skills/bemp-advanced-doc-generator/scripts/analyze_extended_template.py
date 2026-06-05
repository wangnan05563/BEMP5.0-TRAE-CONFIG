"""分析扩展后模板的H1/H2结构"""
import paths
from docx import Document

path = str(paths.detail_design_template())
doc = Document(path)

print(f"段落数: {len(doc.paragraphs)}")
for i, p in enumerate(doc.paragraphs):
    if p.style and p.style.name.startswith('Heading'):
        text = p.text.strip()[:60]
        print(f"[{i:3d}] {p.style.name}: {text}")
