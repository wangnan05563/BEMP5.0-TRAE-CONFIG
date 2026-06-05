"""深入分析v4文档的完整结构"""
import paths
from docx import Document
from docx.oxml.ns import qn

v4_path = str(paths.BANK_REQUIREMENTS_DIR / '承兑行额度管理-详细设计说明书-v4.docx')
doc = Document(v4_path)

print('='*80)
print('v4文档完整章节结构')
print('='*80)

for i, p in enumerate(doc.paragraphs):
    if p.style and p.style.name.startswith('Heading'):
        indent = '  ' * (int(p.style.name.split()[-1]) - 1) if p.style.name.split()[-1].isdigit() else ''
        text = p.text.strip()
        if text:
            print(f'[{i:3d}] {p.style.name}: {indent}{text[:80]}')

print('\n' + '='*80)
print('v4文档完整表格列表')
print('='*80)
for i, tbl in enumerate(doc.tables):
    rows = len(tbl.rows)
    cols = len(tbl.columns)
    first_row = ' | '.join(c.text.strip()[:20] for c in tbl.rows[0].cells) if rows else ''
    print(f'表格[{i:2d}]: {rows}行 x {cols}列 | {first_row[:100]}')
