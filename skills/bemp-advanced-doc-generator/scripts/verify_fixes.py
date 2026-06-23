from docx import Document

doc_path = r"D:\code\QJ\BEMP5.0DEV\.trae\skills\bemp-advanced-doc-generator\output\机构管理和管理员管理功能优化-详细设计文档-20260617.docx"
doc = Document(doc_path)

print("=" * 80)
print("验证修复效果")
print("=" * 80)

# 问题1：检查"组件内部的模块列表及说明"章节
print("\n【问题1】组件内部的模块列表及说明")
found = False
for i, para in enumerate(doc.paragraphs):
    if '组件内部的模块列表及说明' in para.text:
        found = True
        print(f"找到标题：{para.text}")
        # 打印后续10个段落
        for j in range(i+1, min(i+15, len(doc.paragraphs))):
            next_para = doc.paragraphs[j]
            if next_para.style.name.startswith('Heading'):
                break
            print(f"  [{j}] {next_para.text[:100]}")
        break

if not found:
    print("未找到该章节")

# 问题2：检查"模块2设计说明"下的表格
print("\n【问题2】模块2设计说明下的表格结构")
in_module2 = False
table_count = 0
for i, para in enumerate(doc.paragraphs):
    if para.style.name == 'Heading 1' and '模块2' in para.text:
        in_module2 = True
        print(f"进入模块2设计说明")
        continue
    
    if in_module2:
        if para.style.name == 'Heading 1':
            print(f"离开模块2设计说明")
            break
        
        if para.style.name == 'Heading 2':
            print(f"  H2: {para.text}")
        
        # 检查表格
        if para._element.getparent() is not None:
            parent = para._element.getparent()
            for sibling in parent:
                if sibling.tag.endswith('tbl'):
                    table_count += 1
                    if table_count <= 3:  # 只显示前3个表格
                        print(f"    [表格{table_count}]")

print(f"\n共发现 {table_count} 个表格")

# 检查 businessSubmodules 数据
print("\n【数据检查】businessSubmodules 字段")
import json
import os

design_data_path = r"D:\code\QJ\BEMP5.0DEV\.trae\skills\bemp-advanced-doc-generator\output\_design-data-20260617.json"
if os.path.exists(design_data_path):
    with open(design_data_path, 'r', encoding='utf-8') as f:
        design_data = json.load(f)
    
    bsm = design_data.get('businessSubmodules', [])
    print(f"businessSubmodules 数量: {len(bsm)}")
    for i, sub in enumerate(bsm[:5], 1):
        if isinstance(sub, dict):
            print(f"  {i}. {sub.get('name', 'N/A')}: {sub.get('description', 'N/A')[:50]}")
else:
    print("design_data 文件不存在")

print("\n" + "=" * 80)
