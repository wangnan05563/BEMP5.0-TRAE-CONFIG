from pathlib import Path
SKILL_ROOT = Path(__file__).resolve().parent.parent
PROJECT_ROOT = SKILL_ROOT.parent.parent.parent

from pathlib import Path
"""追踪模块2在生成过程中的状态变化"""
import sys
import os
sys.path.insert(0, str(SKILL_ROOT / "scripts"))

from docx import Document

def check_module2_state(doc, stage):
    """检查模块2的状态"""
    print(f"\n{'='*70}")
    print(f"[{stage}] 模块2状态检查")
    print('='*70)
    
    # 找到模块2的位置
    module2_pos = None
    next_h1_pos = None
    
    for i, para in enumerate(doc.paragraphs):
        if para.style and para.style.name == 'Heading 1':
            text = para.text.strip()
            if '模块2' in text:
                module2_pos = i
                print(f"模块2位置: {i}")
            elif module2_pos is not None and next_h1_pos is None:
                next_h1_pos = i
                print(f"下一个H1位置: {i} ({text})")
                break
    
    if module2_pos is None:
        print("❌ 未找到模块2")
        return
    
    if next_h1_pos is None:
        next_h1_pos = len(doc.paragraphs)
        print(f"模块2是最后一个H1，到文档末尾")
    
    # 统计模块2下的内容
    h2_count = 0
    para_count = 0
    
    print(f"\n模块2下的内容 (位置 {module2_pos+1} 到 {next_h1_pos-1}):")
    for i in range(module2_pos + 1, next_h1_pos):
        para = doc.paragraphs[i]
        style_name = para.style.name if para.style else 'None'
        text = para.text.strip()
        
        if style_name == 'Heading 2':
            h2_count += 1
            print(f"  [H2] 位置{i}: {text[:60] if text else '(空)'}")
        elif text and para_count < 3:
            para_count += 1
            print(f"  [段落] 位置{i}: {text[:60]}")
    
    print(f"\n统计: H2={h2_count}, 段落={para_count}")
    return h2_count

# 加载生成的文档
doc_path = str(SKILL_ROOT / "output" / "机构管理和管理员管理功能优化-详细设计文档-20260617.docx")
doc = Document(doc_path)

# 检查最终状态
final_h2_count = check_module2_state(doc, "最终文档")

if final_h2_count == 0:
    print("\n" + "="*70)
    print("问题确认: 模块2的H2标题在生成过程中被删除")
    print("="*70)
    print("\n可能原因:")
    print("1. _clear_content_between_headings 被错误调用")
    print("2. 后续处理步骤清理了模块2的内容")
    print("3. H2标题创建后未被正确保留")
else:
    print(f"\n✅ 模块2有 {final_h2_count} 个H2标题")
