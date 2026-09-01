# -*- coding: utf-8 -*-
"""
BEMP 河南农信"反洗钱收票待签收/背书待签收阶段校验（六阶段收官⑤⑥）"交付文档生成器
change-id: add-hnnx-aml-sign-stage-check（HNNS-EB-STD-REQ-002，PRD v3.5.0 §3.2.3/§3.2.4 + Q-13）
基于实际代码实现与测试证据生成 .docx，结构遵循交付文档标准：
需求概述/方案设计/详细设计/测试设计/测试结果/已知问题/部署说明 + 附录。
先例来源：gen_aml4_design.py / gen_org_mgmt_design.py（全代码构建，样式工具函数同源，
design-pipeline.yaml precedent-script 模式，keep_as_precedent）。
数据来源（编写依据，按优先级）：
  1. PRD docs/prd/反洗钱校验功能/HNNS-EB-STD-REQ-002-反洗钱功能优化-v3.md（v3.5.0）
  2. spec/tasks/checklist：.trae/specs/add-hnnx-aml-sign-stage-check/（全部勾选）
  3. 实现：banks/ext-hnnxbank/hnnxbank-biz-as/.../HnnxBankEbank2005AtomImpl.java（260 行）
  4. 功能测试报告：aotutests-devtools/testcases/反洗钱名单校验/六阶段功能测试报告-AML-FT-20260831-01.md
  5. sonar 扫描报告 + 后端代码走查报告（2026-08-31）
"""
import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

OUTPUT_DIR = r"d:\code\QJ\BEMP5.0DEV\docs\delivery\反洗钱签收校验"
OUTPUT_FILE = os.path.join(OUTPUT_DIR, "HNNS-EB-STD-REQ-002-反洗钱签收校验-详细设计说明书-20260831.docx")

# ===================== 样式工具函数（与 gen_aml4_design.py 先例一致） =====================

def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = tcPr.find(qn('w:tcBorders'))
    if tcBorders is None:
        tcBorders = OxmlElement('w:tcBorders')
        tcPr.append(tcBorders)
    for edge in ('top', 'left', 'bottom', 'right'):
        tag = 'w:{}'.format(edge)
        element = tcBorders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            tcBorders.append(element)
        element.set(qn('w:val'), 'single')
        element.set(qn('w:sz'), '4')
        element.set(qn('w:color'), '000000')


def set_cell_background(cell, color="D9D9D9"):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color)
    tcPr.append(shd)


def set_run_font(run, name="宋体", size=10.5, bold=False, color=None):
    run.font.name = name
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    rFonts.set(qn('w:eastAsia'), name)
    rFonts.set(qn('w:ascii'), name)
    rFonts.set(qn('w:hAnsi'), name)


def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        set_run_font(run, name="黑体", size=max(12, 18 - level * 2), bold=True, color="000000")
    return p


def add_para(doc, text, bold=False, indent=True, size=10.5):
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.first_line_indent = Cm(0.74)
    p.paragraph_format.line_spacing = 1.5
    run = p.add_run(text)
    set_run_font(run, name="宋体", size=size, bold=bold)
    return p


def add_bullet(doc, text, size=10.5):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.line_spacing = 1.5
    run = p.add_run(text)
    set_run_font(run, name="宋体", size=size)
    return p


def add_code_para(doc, text):
    """流程链路/口径原文用等宽风格段落呈现，避免表格嵌套。"""
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.3
    p.paragraph_format.left_indent = Cm(0.74)
    run = p.add_run(text)
    set_run_font(run, name="宋体", size=9.5)
    return p


def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr_cells[i].text = ""
        p = hdr_cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(header)
        set_run_font(run, name="宋体", size=10, bold=True)
        set_cell_background(hdr_cells[i])
        set_cell_border(hdr_cells[i])
        hdr_cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    for r_idx, row in enumerate(rows, 1):
        row_cells = table.rows[r_idx].cells
        for c_idx, val in enumerate(row):
            row_cells[c_idx].text = ""
            p = row_cells[c_idx].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(str(val))
            set_run_font(run, name="宋体", size=10)
            set_cell_border(row_cells[c_idx])
            row_cells[c_idx].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    if col_widths:
        for i, width in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(width)
    doc.add_paragraph()
    return table


def add_toc_field(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run()
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = 'TOC \\o "1-3" \\h \\z \\u'
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')
    fldChar3 = OxmlElement('w:t')
    fldChar3.text = '右键此处选择"更新域"以生成目录'
    fldChar4 = OxmlElement('w:fldChar')
    fldChar4.set(qn('w:fldCharType'), 'end')
    r_element = run._r
    r_element.append(fldChar1)
    r_element.append(instrText)
    r_element.append(fldChar2)
    r_element.append(fldChar3)
    r_element.append(fldChar4)


def add_page_break(doc):
    p = doc.add_paragraph()
    run = p.add_run()
    run.add_break(WD_BREAK.PAGE)


def setup_page_and_header_footer(doc):
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(3.0)
        section.right_margin = Cm(2.5)
        header = section.header
        if not header.paragraphs[0].runs:
            hp = header.paragraphs[0]
            hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = hp.add_run("河南农商票据交易管理平台 - 详细设计说明书")
            set_run_font(run, name="宋体", size=9, color="808080")
        footer = section.footer
        fp = footer.paragraphs[0]
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = fp.add_run()
        set_run_font(run, name="宋体", size=9, color="808080")
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')
        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = 'PAGE'
        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'end')
        run._r.append(fldChar1)
        run._r.append(instrText)
        run._r.append(fldChar2)


def inject_update_fields(doc):
    settings = doc.settings.element
    update_fields = settings.find(qn('w:updateFields'))
    if update_fields is None:
        update_fields = OxmlElement('w:updateFields')
        update_fields.set(qn('w:val'), 'true')
        settings.append(update_fields)


# ===================== 文档生成主流程 =====================

def build_design_doc():
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = '宋体'
    style.font.size = Pt(10.5)
    rPr = style.element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    rFonts.set(qn('w:eastAsia'), '宋体')

    setup_page_and_header_footer(doc)

    # ===== 封面 =====
    for _ in range(3):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("河南农商银行")
    set_run_font(run, name="黑体", size=26, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("新一代票据系统个性化开发")
    set_run_font(run, name="黑体", size=22, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("反洗钱收票待签收/背书待签收校验（六阶段收官⑤⑥）")
    set_run_font(run, name="黑体", size=20, bold=True)

    for _ in range(2):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("详 细 设 计 说 明 书")
    set_run_font(run, name="黑体", size=20, bold=True)

    for _ in range(6):
        doc.add_paragraph()

    info_table = doc.add_table(rows=6, cols=2)
    info_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    info_data = [
        ("项目名称", "BEMP5.0 河南农商银行个性化开发"),
        ("需求名称", "反洗钱收票待签收/背书待签收校验（HNNS-EB-STD-REQ-002）"),
        ("需求依据", "PRD v3.5.0 §3.2.3（六阶段）/§3.2.4（取值逻辑）+ 待确认清单 Q-13 确认结论"),
        ("变更标识", "add-hnnx-aml-sign-stage-check（六阶段收官增量）"),
        ("文档版本", "V1.0"),
        ("编写日期", "2026-08-31"),
    ]
    for i, (k, v) in enumerate(info_data):
        cells = info_table.rows[i].cells
        cells[0].text = ""
        cells[1].text = ""
        p1 = cells[0].paragraphs[0]
        p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r1 = p1.add_run(k)
        set_run_font(r1, name="宋体", size=12, bold=True)
        p2 = cells[1].paragraphs[0]
        p2.alignment = WD_ALIGN_PARAGRAPH.LEFT
        r2 = p2.add_run(v)
        set_run_font(r2, name="宋体", size=12)
        for c in cells:
            set_cell_border(c)

    add_page_break(doc)

    # ===== 修订记录 =====
    add_heading(doc, "修订记录", level=1)
    add_table(doc,
        ["版本", "日期", "修订内容", "修订人"],
        [
            ["V1.0", "2026-08-31", "初始版本：覆盖反洗钱收票待签收/背书待签收校验（⑤⑥六阶段收官增量）需求概述、方案设计（bean 顶替/范围判定/取数链路/异常转应答）、逐文件详细设计、测试设计（16 用例+17 单测映射）、测试结果（单测 17 全绿/功能测试 16/16/数据就绪）、已知问题清单与部署说明", "BEMP文档交付工程师"],
        ],
        col_widths=[2, 3, 8, 3]
    )
    add_page_break(doc)

    # ===== 目录 =====
    add_heading(doc, "目录", level=1)
    add_toc_field(doc)
    add_page_break(doc)

    # ===== 第一章 需求概述 =====
    add_heading(doc, "第一章 需求概述", level=1)

    add_heading(doc, "1.1 业务背景与需求目标", level=2)
    add_para(doc, "20260831 版公司金融部需求 docx《关于新一代票据业务管理系统反洗钱功能优化的需求》将原文校验位置由四阶段扩展为六阶段，新增⑤收票待签收、⑥背书待签收两个校验阶段（网银【待签收业务】→【同意签收】场景）。系统现状：市场交易侧与网银端四阶段（①直贴②承兑③背书④客户维护）已实现并交付（详见《反洗钱四阶段校验-详细设计说明书-v1.docx》），签收环节为最后一处名单校验缺口——命中中互金名单的企业若在签收环节不拦截，票据权利将经签收落地完成转移。")
    add_para(doc, "本期目标：⑤收票待签收（校验出票人、收票人、承兑行）与⑥背书待签收（校验出票人、承兑行、背书人、被背书人）在网银统一签收链路落地，校验点位于业务落地前，命中即拦截并返回面客文案（角色前缀+不暴露名单来源+96288 热线），复用既有客户侧校验开关与两轮匹配口径，零新增参数/错误码/SQL。")

    add_heading(doc, "1.2 需求范围", level=2)
    add_para(doc, "触发链路：网银统一签收报文 PICE070105 → Ebank2005Service.batchComrclDrftSign（单笔直调/多笔线程池两路径共用）→ 签收 Atom（bean 顶替后为 HnnxBankEbank2005AtomImpl）。两分支校验范围：")
    add_table(doc,
        ["阶段", "触发条件（范围判定）", "校验角色", "交易编号（TRANS_NO）"],
        [
            ["⑤收票待签收", "signFlag=SU00（同意签收）", "出票人、收票人、承兑行", "NES.003.20.00P（ELEC_ACCEPTANCE_ACCEPT_ISSUE）"],
            ["⑥背书待签收", "signFlag=SU00（同意签收）", "出票人、承兑行、背书人、被背书人", "NES.006.20.00P（ELEC_ENDORSEMENT）+ NES.015.20.00P（不得转让撤销，CANCEL_BAN_ENDRSMT_FLAG）"],
        ],
        col_widths=[3, 4, 4.5, 4.5]
    )
    add_para(doc, "SU01（拒绝签收）与提示付款/提示承兑/清偿应答/回购式贴现赎回等其他签收类型不校验，直接走产品化逻辑。名单匹配口径与既有四阶段一致：统一社会信用代码（certType='22'）优先命中，为空/未命中回退客户名称兜底；名单数据口径 DATA_SRC='18' 且 VLD_ST='1'，失效时间空视为长期有效；黑名单拦截、灰名单一期提示中断（人工审批流为二期）。")

    add_heading(doc, "1.3 范围外（明确不做）", level=2)
    for t in [
        "提示付款签收（NES.011.*）、提示承兑（NES.002.20.00P）、清偿应答（NES.013.*）、回购式贴现赎回（NES.008.20.20P）等其他签收类型的名单校验；",
        "产品化 channel-as / btrc2-as 源码改动（仅经 bean 顶替覆写，Ebank2005AtomImpl 与 TransInfoSign 等产品类零改动）；",
        "新增开关/错误码/SQL 脚本（复用 hnnx.cust.antiMoneyList.block 与 HNNX0BE320011）；",
        "灰名单人工审批流（二期）；市场交易侧逻辑与文案；",
        "AntiMoneyBillDialog 弹窗改造（签收为单笔单一主体场景，无多票据明细弹窗需求）。",
    ]:
        add_bullet(doc, t)

    add_heading(doc, "1.4 需求确认闭环", level=2)
    add_para(doc, "需求梳理阶段待确认清单中与本次增量相关的 Q-13 已全部闭环（v3.5.0，2026-08-31 用户确认）：")
    add_table(doc,
        ["事项", "确认结论", "闭环状态"],
        [
            ["Q-13-A：收票待签收/背书待签收两新阶段是否纳入本期实现", "纳入本期开发范围，对应用例随实现编制", "已闭环（用户确认 2026-08-31）"],
            ["Q-13-B：背书待签收 trans_from/trans_to 角色映射方向", "待签收业务 from/to 方向就是反向的，口径无问题——trans_from=被背书人、trans_to=背书人，按此实现无需复核（该方向与既有 tb_endrsmt_info 语义相反，系 tb_trans_info_sign 表独立定义）", "已闭环（用户确认 2026-08-31）"],
        ],
        col_widths=[5.5, 7.5, 3]
    )
    add_para(doc, "测试假设验证结果：本次增量无遗留测试假设项。Q-11（面客文案按原文模板调整）与 Q-12（①~④网银三分支报文段待联调）为四阶段范围既有事项，其结论与口径已在四阶段交付文档中闭环登记，不阻塞本次增量。Q-13-B 反向口径已通过单元测试 testBuildEndorsementSignRoleItems_reverseFromTo（from/to 交叉命名防反向错位）在测试中验证。")

    add_page_break(doc)

    # ===== 第二章 方案设计 =====
    add_heading(doc, "第二章 方案设计（六阶段架构定位）", level=1)

    add_heading(doc, "2.1 六阶段校验总体架构与⑤⑥链路定位", level=2)
    add_para(doc, "反洗钱中互金名单校验共六个阶段，前四阶段已交付，⑤⑥为本次增量（校验位置均以电票系统收到申请后、业务落地前为时点，与 Q-02 确认的\"业务申请提交前校验并拦截\"口径一致）：")
    add_table(doc,
        ["阶段", "操作路径（原文）", "校验内容", "本期状态"],
        [
            ["①直贴业务发起", "网银【贴现】【贴现申请】【提交】", "贴现申请人、出票人", "已交付（四阶段）"],
            ["②承兑业务发起", "网银【出票】【提示承兑】【提交】", "出票人、收票人", "已交付（四阶段）"],
            ["③背书转让", "网银【背书】【提交】", "持票人、收票人、承兑行", "已交付（四阶段）"],
            ["④客户信息维护", "业务管理子系统-企业客户维护", "客户名", "已交付（四阶段）"],
            ["⑤收票待签收", "网银【待签收业务】【提示收票】【同意签收】（企业作为收票人）", "出票人、收票人、承兑行", "本期实现"],
            ["⑥背书待签收", "网银【待签收业务】【背书】/【不得转让撤销】【同意签收】（企业作为被背书人收票）", "背书人、被背书人、承兑行、出票人", "本期实现"],
        ],
        col_widths=[3, 6, 4, 3]
    )
    add_para(doc, "⑤⑥在网银统一签收链路中的位置如下（校验点插入于查得待签收记录之后、业务落地之前；若在 super 之后校验，ECDS/CPES 报文已发出无法撤回，票据权利将完成转移）：")
    add_code_para(doc, "PICE070105（网银签收应答报文）")
    add_code_para(doc, "  → Ebank2005ServiceImpl.batchComrclDrftSign（单笔直调 / 多笔 Ebank2005CallableImpl 线程池，两路径共用）")
    add_code_para(doc, "  → ebank2005Atom.comrclDrftSign（bean 顶替后运行时实例=HnnxBankEbank2005AtomImpl）")
    add_code_para(doc, "      → [覆写层] transId 守卫 → 查 TB_TRANS_INFO_SIGN（getTransInfoSignByTransId）")
    add_code_para(doc, "      → [新增校验] isAmlCheckScope 范围判定 → 取数组装 → validateCustBizList 名单匹配")
    add_code_para(doc, "      → （命中：抛 BempRuntimeException HNNX0BE320011，业务不落地）")
    add_code_para(doc, "      → super.comrclDrftSign（产品化逻辑：dealAcceptTransAndSendDraft* 内发 ECDS/CPES 报文，业务落地）")

    add_heading(doc, "2.2 bean 顶替机制", level=2)
    add_para(doc, "新建 HnnxBankEbank2005AtomImpl 标注 @CloudComponent + @CustomizedBean 并继承产品化 Ebank2005AtomImpl（与四阶段 HnnxBankEbank2004AtomImpl 同构模式）。BempBeanDefinitionRegistryProcessor（framework/fw-common L47-98）按\"@CustomizedBean 类的直接父类名→产品 bean 定义\"顶替：启动时删除产品 Ebank2005AtomImpl 的 ebank2005AtomImpl bean 定义，以该 bean 名注册本类定义，从而：")
    for t in [
        "Ebank2005ServiceImpl 中按 Ebank2005Atom 接口类型 @Autowired 注入（L46-47，无 @Qualifier）的即本类实例——单笔路径 L104 直调、多笔路径 L69 经构造器传 Ebank2005CallableImpl（其 L33-36 接收同一 bean 引用），两条路径均命中本类；",
        "产品 Ebank2005AtomImpl 仅标注 @Component（无 @CustomizedBean、无 Proxy 插件层），本类是唯一顶替者，不存在四阶段 Proxy 层那种同层竞争问题，故直接继承产品类（与四阶段的设计差异点）；",
        "启动日志佐证：Served_startup_20260901_002153.log L176\"Customized bean hnnxBankEbank2005AtomImpl, remove the bean ebank2005AtomImpl\"（顶替注册）、L3014 运行时实例 HnnxBankEbank2005AtomImpl$$EnhancerBySpringCGLIB（bean 顶替生效确认）。",
    ]:
        add_bullet(doc, t)

    add_heading(doc, "2.3 校验范围判定设计（isAmlCheckScope）", level=2)
    add_para(doc, "开发前置核实结论（spec\"开发前置核实项 1\"，留痕于 tasks.md SubTask 1.1 与主类注释）：TB_TRANS_INFO_SIGN.TRANS_TYPE 存的是大类标识（BtrcCommonConstant.SIGN_TYPE=\"2\" 签收/APPLY_TYPE=\"1\" 申请，所有待签收记录同值），不能用于区分业务类型；产品化 comrclDrftSign 内部分流也全部按 transInfoDto.getTransNo() 判断。故本次范围判定同样以 TRANS_NO（NES 报文编号）为准，码值清单（触发器落库代码核实）：")
    add_table(doc,
        ["分类", "TRANS_NO 码值", "常量来源", "落库依据（触发器）"],
        [
            ["⑤提示收票签收", "NES.003.20.00P", "TransCodeConstant.Acceptance.ELEC_ACCEPTANCE_ACCEPT_ISSUE", "Cim003OrgNes003TransTriggerImpl L116 / Ecds034Org003TriggerImpl L63（CPES/ECDS 落库同值）"],
            ["⑥背书签收", "NES.006.20.00P", "TransCodeConstant.Endorsement.ELEC_ENDORSEMENT", "Cim003OrgNes006TransTriggerImpl L106 / Ecds034Org010TriggerImpl L65"],
            ["⑥不得转让撤销", "NES.015.20.00P", "TransCodeConstant.Common.CANCEL_BAN_ENDRSMT_FLAG", "Cim003OrgNes015TransTriggerServiceImpl L116（落库 TRANS_TYPE=SIGN_TYPE）"],
            ["范围外（不校验）", "NES.011.*（提示付款）/ NES.002.20.00P（提示承兑）/ NES.013.*（清偿应答）/ NES.008.20.20P（回购式贴现赎回）等", "—", "直接透传产品化逻辑"],
        ],
        col_widths=[3.2, 4.8, 4.8, 4.2]
    )
    add_para(doc, "signFlag 取本次请求应答标志 reqInfo.getSignFlag()（非库内 SIGN_FLAG——查库所得为待签收登记初始值，不含本次应答意图），判定常量 EBankConsts.SignUpMarkEnum.SU00/SU01（与产品 busiValidate 同款，非魔法值）：仅 SU00 同意签收才校验；SU01 拒绝签收（拒绝动作不发生票据权利转移）与空值不校验。范围外交易零开销返回——不查票据主档、不触发承兑行机构反查、不进入名单匹配，保证\"其他签收类型行为与改造前完全一致\"（有单测断言零调用）。")

    add_heading(doc, "2.4 ⑤⑥取数链路设计（PRD §3.2.4 用户确认口径）", level=2)
    add_para(doc, "公共取数链路（两阶段一致）：")
    add_code_para(doc, "PICE070105MessageConverter 接口 → 获取网银发送的 transId（交易ID）")
    add_code_para(doc, "  → 以 transId 查 tb_trans_info_sign（tb_trans_info_sign.id = transId）")
    add_code_para(doc, "  → 取 tb_trans_info_sign.bill_id → 以 bill_id 查 tb_bill_info（tb_bill_info.id = bill_id）")
    add_para(doc, "⑤收票待签收阶段——校验角色与取值（PRD §3.2.4 二节）：")
    add_table(doc,
        ["校验角色", "取数来源表", "字段（名称 / 统一社会信用代码）"],
        [
            ["出票人", "tb_bill_info", "drwr_name / drwr_soc_code"],
            ["收票人", "tb_bill_info", "pyee_name / pyee_soc_code"],
            ["承兑行（名称）", "tb_bill_info", "acpt_bank_name"],
            ["承兑行（统一社会信用代码）", "tb_bill_info → tm_cpes_branch", "acpt_bank_no 关联 tm_cpes_branch.trans_brch_bank_no（且 brch_status='ST01'），取 tm_cpes_branch.soc_code（复用 HnnxAntiMoneyValidateUtil.queryAcptBankSocCode）"],
        ],
        col_widths=[4, 4.2, 8.8]
    )
    add_para(doc, "⑥背书待签收阶段——校验角色与取值（PRD §3.2.4 三节）：")
    add_table(doc,
        ["校验角色", "取数来源表", "字段（名称 / 统一社会信用代码）"],
        [
            ["出票人", "tb_bill_info", "drwr_name / drwr_soc_code（与⑤同源，经 billId）"],
            ["承兑行（名称）", "tb_bill_info", "acpt_bank_name"],
            ["承兑行（统一社会信用代码）", "tb_bill_info → tm_cpes_branch", "acpt_bank_no → trans_brch_bank_no（brch_status='ST01'）取 soc_code，与⑤同源"],
            ["被背书人", "tb_trans_info_sign", "trans_from_name / trans_from_soc_code（Q-13-B 反向口径：from=被背书人）"],
            ["背书人", "tb_trans_info_sign", "trans_to_name / trans_to_soc_code（Q-13-B 反向口径：to=背书人）"],
        ],
        col_widths=[4, 4.2, 8.8]
    )
    add_para(doc, "Q-13-B 角色映射确认说明：背书待签收阶段的 from/to 角色映射按业务方确认口径执行——trans_from=被背书人、trans_to=背书人。该口径与既有 tb_endrsmt_info 的语义（trans_from=转出方/背书人、trans_to=受让方/被背书人）方向相反，系 tb_trans_info_sign 表独立定义；业务方已确认\"待签收业务 from/to 方向就是反向的，口径无问题\"，按此实现无需复核。")

    add_heading(doc, "2.5 异常转应答机制（dealException 链）", level=2)
    add_para(doc, "与四阶段 Ebank2004 dealException 模式同源（开发前置核实项 2 已核实）：校验命中抛 BempRuntimeException(HNNX0BE320011, {0}角色前缀, {1}角色描述) 后，覆写类不做任何 try-catch 吞异常处理，沿 super.comrclDrftSign 调用链上抛至 Ebank2005ServiceImpl.batchComrclDrftSign——单笔路径 catch 后经 Ebank2005ServiceImpl.dealException（instanceof BaseException 取 errorCode/errorMessage），多笔路径 Ebank2005CallableImpl.call 内 catch 后走同一 dealException，填入应答 retCode/retMsg，由 PICE070105 应答报文将面客文案返回客户。面客文案模板（hnnx_mt_zh_CN.properties L11，零新增错误码）：")
    add_code_para(doc, "HNNX0BE320011={0}：{1}暂不符合相关政策要求，暂不支持提供该服务，详情请咨询开户机构或96288")
    add_para(doc, "文案对客户不暴露\"中互金关注名单/黑灰名单\"字样，仅以\"暂不符合相关政策要求\"话术替代并附 96288 客服热线；黑/灰名单共用同一模板（单测命中六法 listChrc 黑灰混合全绿，文案同构）。")

    add_heading(doc, "2.6 复用约束与关键设计决策", level=2)
    add_para(doc, "复用约束（SHALL NOT 违反项，代码走查 §22/§24 逐项核验通过）：")
    add_table(doc,
        ["复用项", "说明", "落点"],
        [
            ["客户侧校验开关", "hnnx.cust.antiMoneyList.block（TM_BUSINESS_PARAMETER，默认开启，按法人机构号隔离读取）：关闭跳过校验不查名单；异常按开关降级——开→拦截留痕（fail-closed）/关→放行留痕", "HnnxAntiMoneyValidateUtil.validateCustBizList 内部，覆写层不包开关判断（避免双重读参，与四阶段一致）"],
            ["面客文案", "HNNX0BE320011（角色前缀+不暴露名单来源+96288），五角色前缀：出票人/收票人/承兑行/被背书人/背书人", "hnnx_mt_zh_CN.properties L11，零新增错误码"],
            ["两轮匹配口径", "统一社会信用代码 certType='22' 优先命中 → 客户名称兜底；名单数据口径 DATA_SRC='18'、VLD_ST='1'、失效时间空=长期有效", "HnnxAntiMoneyValidateUtil 既有实现，与市场侧/四阶段同一 SQL 零漂移"],
            ["承兑行反查", "queryAcptBankSocCode（acpt_bank_no → tm_cpes_branch.trans_brch_bank_no，brch_status='ST01' 取 soc_code），失败返回 null 降级名称兜底不阻断", "口径收拢于 HnnxAntiMoneyValidateUtil，与市场交易侧同源零漂移"],
        ],
        col_widths=[3, 8, 6]
    )
    add_para(doc, "关键设计决策（每项说明\"为什么\"）：")
    add_table(doc,
        ["决策", "理由"],
        [
            ["校验点位于查得 TB_TRANS_INFO_SIGN 之后、业务落地（dealAcceptTransAndSendDraft*，其内发 ECDS/CPES 报文）之前", "若在 super 之后校验，报文已发出无法撤回，票据权利将完成转移；命中即中断，签收不落地"],
            ["覆写层与 super 各执行一次 TB_TRANS_INFO_SIGN 主键查询（接受重复查询）", "覆写层拿不到 super 内部中间态，与其在 super 后校验（报文已发出）不如提前查一次，保持覆写与产品逻辑解耦——四阶段同款取舍，主键级查询开销可接受（代码走查提示级 1 项，设计取舍已文档化）"],
            ["transId 为空/非 18 位时不做范围判定直接透传 super", "此类请求由 super 抛产品原生校验异常（EBANK_ERROR_CODE_0BE320807042/VALID_FAIL），覆写层不得抢先以不同语义报错，保持产品报错行为不变（双层 fail-safe 守卫）"],
            ["待签收记录不存在（transInfoDto=null）时跳过校验透传 super", "super 会抛产品原生 0BE320807050，null 场景零额外开销（无票据查询/反查/名单查询），与\"保持产品报错行为不变\"原则同款"],
            ["getBillInfoById 票据主档查无即抛产品\"票据不存在\"异常，不静默放行", "与四阶段同契约——数据异常不应静默放行签收"],
            ["直接继承产品类（不经 Proxy 层）", "产品 Ebank2005AtomImpl 仅 @Component 无 @CustomizedBean，本类是唯一顶替者，无四阶段 Proxy 层同层竞争问题"],
        ],
        col_widths=[6.5, 10.5]
    )

    add_page_break(doc)

    # ===== 第三章 详细设计 =====
    add_heading(doc, "第三章 详细设计（逐文件实现说明）", level=1)

    add_heading(doc, "3.1 新增/修改文件清单", level=2)
    add_table(doc,
        ["文件", "类型", "说明"],
        [
            ["banks/ext-hnnxbank/hnnxbank-biz-as/src/main/java/com/hundsun/bemp/hnnxbank/biz/channel/atom/ebank/impl/HnnxBankEbank2005AtomImpl.java", "新增（约 260 行）", "网银统一签收 Atom 个性化覆写：反洗钱⑤⑥阶段校验（本次唯一主实现文件）"],
            ["banks/ext-hnnxbank/hnnxbank-biz-as/src/test/java/com/hundsun/bemp/hnnxbank/biz/channel/atom/ebank/impl/HnnxBankEbank2005AtomAmlSignRoleTest.java", "新增（单元测试）", "17 用例：范围判定 3+⑤组装 2+⑥组装 2+端到端 10，JDK Proxy stub+反射静态字段注入模式"],
            ["产品化 served/cs/channel-as、served/pc/btrc2-as", "零改动", "Ebank2005AtomImpl/Ebank2005ServiceImpl/Ebank2005CallableImpl/TransInfoSign 等只读依赖（git status 核实）"],
            ["hnnx_mt_zh_CN.properties / HnnxCommonConst / HnnxAntiMoneyValidateUtil / pom.xml", "零改动", "复用既有文案/常量/校验工具；hnnxbank-biz-as 对 bemp-channel-as 依赖四阶段已加，无 pom 变更"],
        ],
        col_widths=[8, 2.8, 6.2]
    )

    add_heading(doc, "3.2 HnnxBankEbank2005AtomImpl.java（主实现）", level=2)
    add_para(doc, "类结构：@CloudComponent + @CustomizedBean 顶替产品 bean 名 ebank2005AtomImpl（L68-70）；五面客角色前缀常量（L72-81：出票人/收票人/承兑行/被背书人/背书人，去重恰为五）；核心方法均为 static 纯函数（仅范围判定/字段映射/校验转发），供单元测试直接覆盖。逐方法设计：")
    add_table(doc,
        ["方法（行号）", "职责", "关键实现要点"],
        [
            ["comrclDrftSign（L92-107，覆写）", "统一签收入口：查得待签收记录后、业务落地前插入校验", "五步流程：①取 reqInfo（L94）→②transId 双条件守卫（L95-96，null!=transId 且 max18NumericPattern 全串匹配，不满足跳过整个校验块直接 L106 透传 super）→③主键查询 getTransInfoSignByTransId（L97-99）→④非空守卫（L100-101，null 时 super 抛产品原生 0BE320807050）→⑤validateSignAmlList（L102-103，reqLegalNo 取 baseRequest.getReqLegalNo() 会话法人号，开关按法人隔离读取）→⑥无条件透传 super（L106）"],
            ["validateSignAmlList（L127-142，static）", "⑤⑥两分支校验统一入口", "先 isAmlCheckScope 范围判定（L128，范围外零开销 return）；getBillInfoById 票据主档主键查询（L132，查无抛产品异常）；queryAcptBankSocCode 承兑行反查（L133，失败 null 降级名称兜底）；按 transNo 分流：NES.003→⑤三角色（L134-136），NES.006/NES.015→⑥四角色（L137-141，else 分支）"],
            ["isAmlCheckScope（L154-163，static）", "校验范围判定", "signFlag 非 SU00 返回 false（L156-158，拒绝签收/空值不校验）；transNo 三码值显式 equals 无通配（L160-162：NES.003.20.00P/NES.006.20.00P/NES.015.20.00P）；范围外返回 false 后由调用方透传产品化逻辑，不产生任何额外查询"],
            ["buildCollectionSignRoleItems（L177-198，static）", "组装⑤三角色项", "组装顺序：出票人（L182-186）→收票人（L187-191，pyee_ 票据主档取数）→承兑行（L192-196，acpt_bank_name+反查代码）；billInfoDto=null 防御返回空集合（L178-180，理论不可达——getBillInfoById 查无已抛异常）；round1（信用代码轮）全量先于 round2（名称轮），多角色命中时按组装顺序首个命中前缀返回"],
            ["buildEndorsementSignRoleItems（L213-244，static）", "组装⑥四角色项（含 Q-13-B 反向口径）", "组装顺序：被背书人=trans_from_name/trans_from_soc_code（L219-224）→背书人=trans_to_name/trans_to_soc_code（L225-230）→承兑行（L231-236）→出票人（L237-242）；方法注释写明\"待签收业务 TB_TRANS_INFO_SIGN 的 from/to 方向与 tb_endrsmt_info 相反，业务方 2026-08-31 确认口径无问题（Q-13-B）\"（L201-206）"],
            ["validateCustBizList（L257-259，private static）", "客户侧名单校验转发入口（收拢调用点）", "单一转发至 HnnxAntiMoneyValidateUtil.validateCustBizList——开关前置判断（关闭跳过不查名单）、两轮匹配、命中抛 HNNX0BE320011、异常按开关降级均在工具内实现；本类无 PARAM_KEY 字面量、无双重读参（全文审查确认）"],
        ],
        col_widths=[4.5, 4, 8.5]
    )
    add_para(doc, "类注释同时留痕三项开发前置核实结论：TRANS_TYPE 实为大类标识按 TRANS_NO 判定（Task 1）、异常转应答 dealException 两路径机制（Task 1）、signFlag 取本次请求应答标志非库内值（L55-57）。")

    add_heading(doc, "3.3 复用组件（只读不改）", level=2)
    add_table(doc,
        ["组件", "复用内容", "边界约束"],
        [
            ["HnnxAntiMoneyValidateUtil（hnnxbank-biz-as antimoney/validate）", "validateCustBizList(reqLegalNo, roles)：开关前置/两轮匹配/命中抛 HNNX0BE320011/异常按开关降级；RoleCheckItem（roleLabel/custName/socCode）；queryAcptBankSocCode(acptBankNo)：ST01 反查", "只读复用零改动；覆写层单一转发无分叉，与市场侧/四阶段同源零漂移"],
            ["BaseEbankAtomImpl.getBillInfoById（静态方法，产品化）", "票据主档 billId 主键查询", "查无即抛产品\"票据不存在\"异常（与四阶段同契约，不静默放行）"],
            ["产品化 Ebank2005ServiceImpl / Ebank2005CallableImpl", "签收调度（单笔直调/多笔线程池）与异常转应答（dealException）", "只读依赖，bean 顶替后自动生效，零改动"],
        ],
        col_widths=[5, 7, 5]
    )

    add_heading(doc, "3.4 数据设计（涉及表与字段映射，全部只读零 DDL）", level=2)
    add_table(doc,
        ["表", "用途", "读取字段（实测列存在性经 USER_TAB_COLUMNS 复核）", "写入"],
        [
            ["TB_TRANS_INFO_SIGN", "待签收记录（主键 transId 查询；⑥背书双方要素来源）", "8 读取列：ID/BILL_ID/TRANS_NO/TRANS_TYPE/TRANS_FROM_NAME/TRANS_FROM_SOC_CODE/TRANS_TO_NAME/TRANS_TO_SOC_CODE 等", "无"],
            ["TB_BILL_INFO", "票据主档（billId 主键查询；⑤⑥出票人/收票人/承兑行名称来源）", "8 读取列：ID/DRWR_NAME/DRWR_SOC_CODE/PYEE_NAME/PYEE_SOC_CODE/ACPT_BANK_NAME/ACPT_BANK_NO 等（含 mapper 外列 PYEE_ACCT_NAME 同样存在，PYEE_SOC_CODE 命中——应用实际库 BEMP_HNNX 结构完整）", "无"],
            ["TM_CPES_BRANCH", "承兑行机构反查（acpt_bank_no → trans_brch_bank_no，brch_status='ST01'）", "3 读取列：TRANS_BRCH_BANK_NO/BRCH_STATUS/SOC_CODE", "无"],
            ["HNNX_M_CUST_SPECIAL_INFO", "中互金名单（实测位于 BEMP_HNNX 主库）", "经 HnnxAntiMoneyValidateUtil 既有 SQL（DATA_SRC='18'/VLD_ST='1' 过滤、certType='22' 信用代码轮+名称轮）", "无"],
            ["TM_BUSINESS_PARAMETER", "客户侧校验开关（PARAM_KEY=hnnx.cust.antiMoneyList.block，LEGAL_NO='000000'，PARAM_VALUE='1' 默认开启）", "3 读取列（零新增参数记录）", "无"],
        ],
        col_widths=[3.5, 4.5, 7.5, 1.5]
    )
    add_para(doc, "算法说明：名单匹配为两轮串行——round1 统一社会信用代码轮（certType='22'+certNo，命中即拦截）全量先于 round2 客户名称轮（兜底）；角色项按⑤⑤三角色/⑥四角色组装顺序参与匹配，多角色命中时返回组装顺序首个命中前缀。所有查询均为主键级或既有工具内分片批量查询，无循环内远程调用（代码走查 §12 性能 4/4 通过）。")

    add_page_break(doc)

    # ===== 第四章 测试设计 =====
    add_heading(doc, "第四章 测试设计", level=1)

    add_heading(doc, "4.1 测试策略与执行口径", level=2)
    add_para(doc, "⑤⑥为网银端 PICE070105 报文驱动链路，管理端 UI 无入口，Playwright 无法直接模拟网银报文。执行口径与四阶段 P0-023~025 同款锚点：\"代码审查段+单测段+DB 段本期执行，报文应答段⏸报文联调阶段执行\"。测试类型覆盖：单元测试（17 用例，方法级纯函数覆盖）+ 代码事实审查（16 用例逐条代码行号级证据）+ DB 数据就绪检查（六项 SQL，run-oracle-jdbc.ps1 BEMP_HNNX 通道）+ 启动日志佐证（bean 顶替链）。")

    add_heading(doc, "4.2 功能测试用例清单（16 条，用例基线 v3.5.0\"六阶段收官增量\"章节）", level=2)
    add_table(doc,
        ["用例编号", "用例名称", "优先级", "执行段", "结果"],
        [
            ["TC-AML-P0-027", "⑤收票待签收-出票人命中-拦截+面客文案", "P0", "代码审查+单测+⏸报文段", "审查段通过/⏸报文段"],
            ["TC-AML-P0-028", "⑤收票待签收-收票人命中-拦截+面客文案", "P0", "代码审查+单测+⏸报文段", "审查段通过/⏸报文段"],
            ["TC-AML-P0-029", "⑥背书待签收-背书人命中（trans_to 反向口径）", "P0", "代码审查+单测+⏸报文段", "审查段通过/⏸报文段"],
            ["TC-AML-P0-030", "⑥背书待签收-被背书人命中（trans_from）", "P0", "代码审查+单测+⏸报文段", "审查段通过/⏸报文段"],
            ["TC-AML-P0-031", "⑥不得转让撤销 NES.015 同走⑥分支", "P0", "代码审查+单测+⏸报文段", "审查段通过/⏸报文段"],
            ["TC-AML-P0-032", "复用开关两态-关闭跳过/开启校验", "P0", "代码审查+DB+单测+⏸报文对照", "通过（DB+审查+单测段）/⏸报文对照"],
            ["TC-AML-P0-033", "取数口径与 PRD §3.2.4 逐项一致-含承兑行 ST01 反查+名称兜底", "P0", "代码审查+DB+单测", "通过"],
            ["TC-AML-P0-034", "面客文案 HNNX0BE320011 五角色前缀+96288+不暴露名单来源", "P0", "代码审查+单测", "通过"],
            ["TC-AML-P1-033", "范围限定-其他签收类型零开销不校验", "P1", "代码审查+单测+⏸可选对照", "通过（审查+单测段）/⏸可选对照"],
            ["TC-AML-P1-034", "SU01 拒绝签收跳过校验", "P1", "代码审查+单测", "通过"],
            ["TC-AML-P1-035", "异常降级 fail-closed-开→拦截留痕/关→放行留痕", "P1", "单测佐证+代码事实审查", "通过（佐证段）"],
            ["TC-AML-P1-036", "transId 空/非 18 位透传 super 原生报错 0BE320807042", "P1", "代码审查+单测边界", "通过"],
            ["TC-AML-P1-037", "bean 顶替-单笔多笔两路径命中", "P1", "代码审查+启动日志佐证", "通过"],
            ["TC-AML-P1-038", "名单数据口径-数据来源=18+有效标志=1 过滤在⑤⑥生效", "P1", "DB 基线核对+代码审查口径引用+⏸场景对照", "通过（DB+审查段）/⏸场景对照"],
            ["TC-AML-P2-031", "多角色命中-组装顺序首个命中前缀返回", "P2", "代码审查+单测", "通过"],
            ["TC-AML-P2-032", "待签收记录不存在-透传 super 原生报错 0BE320807050", "P2", "代码审查+⏸可选对照", "通过（审查段）/⏸可选对照"],
        ],
        col_widths=[2.8, 6.2, 1.5, 4, 2.5]
    )

    add_heading(doc, "4.3 单元测试设计与用例映射（17 用例）", level=2)
    add_para(doc, "测试类 HnnxBankEbank2005AtomAmlSignRoleTest（与被测类同包，JDK Proxy stub+反射静态字段注入模式，参数服务桩复用市场侧 HnnxAntiMoneyValidateTest 闭包）；工具层内部口径（两轮匹配细节/降级开关两分支/读参异常 fail-closed）由既有 HnnxAntiMoneyCustValidateTest 覆盖不重复。17 方法与功能用例映射：")
    add_table(doc,
        ["#", "单测方法", "结果", "佐证用例"],
        [
            ["1", "testIsAmlCheckScope_inScope", "通过", "P0-031 步骤 3"],
            ["2", "testIsAmlCheckScope_rejectSignSkip", "通过", "P1-034 步骤 3"],
            ["3", "testIsAmlCheckScope_otherTypesSkip", "通过", "P1-033/P1-036（空 signFlag 边界）"],
            ["4", "testBuildCollectionSignRoleItems", "通过", "P0-033 步骤 1/6、P2-031 步骤 4"],
            ["5", "testBuildCollectionSignRoleItems_nullBill_emptyList", "通过", "P0-033（⑤ null 防御）"],
            ["6", "testBuildEndorsementSignRoleItems_reverseFromTo", "通过", "P0-029 步骤 2、P0-033、P2-031（from/to 交叉断言锁定反向口径）"],
            ["7", "testBuildEndorsementSignRoleItems_nullBill_emptyList", "通过", "P0-033（⑥ null 防御）"],
            ["8", "testValidateSignAmlList_collectionSign_drwrHit", "通过", "P0-027、P0-034"],
            ["9", "testValidateSignAmlList_collectionSign_pyeeHit", "通过", "P0-028、P0-034"],
            ["10", "testValidateSignAmlList_collectionSign_acptBankHitNameFallback", "通过", "P0-033（名称兜底降级）、P0-034"],
            ["11", "testValidateSignAmlList_endorsementSign_endrserHit", "通过", "P0-029、P0-034"],
            ["12", "testValidateSignAmlList_endorsementSign_endrseeHit", "通过", "P0-030、P0-034"],
            ["13", "testValidateSignAmlList_cancelBanEndrsmt_endrserHit", "通过", "P0-031、P0-034"],
            ["14", "testValidateSignAmlList_rejectSign_zeroQuery", "通过", "P1-034 步骤 4"],
            ["15", "testValidateSignAmlList_otherType_zeroQuery", "通过", "P1-033 步骤 4"],
            ["16", "testValidateSignAmlList_switchOff_skipQuery", "通过", "P0-032 步骤 3"],
            ["17", "testValidateSignAmlList_degrade_switchOn_failClosed", "通过", "P1-035 步骤 3"],
        ],
        col_widths=[1, 6.5, 1.5, 8]
    )

    add_page_break(doc)

    # ===== 第五章 测试结果 =====
    add_heading(doc, "第五章 测试结果", level=1)

    add_heading(doc, "5.1 单元测试结果", level=2)
    add_para(doc, "执行命令：mvn test -q -o -pl ext-hnnxbank/hnnxbank-biz-as -Dtest=\"HnnxBankEbank2005AtomAmlSignRoleTest\"（banks 目录下，离线模式）。结果：Tests run: 17, Failures: 0, Errors: 0, Skipped: 0（Time elapsed 4.008s，Maven Surefire 实测输出），17 用例全绿。日志中的 ERROR 行（HnnxAntiMoneyValidateUtil 命中/降级留痕）为降级用例与命中介入用例的预期日志输出，非失败。")
    add_para(doc, "回归验证：mvn test -o -Dtest=\"HnnxAntiMoneyValidateTest,HnnxAntiMoneyCustValidateTest,HnnxBankEbank2004AtomAmlRoleTest,HnnxBankDiscBillAmlRoleTest,HnnxBankEbank2005AtomAmlSignRoleTest\" → Tests run: 79, Failures: 0, Errors: 0, Skipped: 0，BUILD SUCCESS（21.599s）。明细：市场侧 37+客户侧工具 14+四阶段 5+6+本次新增 17。编译验证：served 聚合与 banks ext-hnnxbank reactor mvn compile test-compile -q -o 均 exit 0。")

    add_heading(doc, "5.2 功能测试执行汇总（一轮，AML-FT-20260831-01）", level=2)
    add_table(doc,
        ["维度", "本期执行段落实", "通过", "失败", "阻塞", "⏸ 延后"],
        [
            ["P0（8 条）", "8", "8", "0", "0", "P0-027~032 六条报文段（P0-032 为报文对照）"],
            ["P1（6 条）", "6", "6", "0", "0", "P1-033/P1-038 可选报文对照"],
            ["P2（2 条）", "2", "2", "0", "0", "P2-032 可选报文对照"],
            ["合计", "16 条用例执行段全部落实", "16", "0", "0", "⏸ 报文段 8 条次+P2-032 可选对照"],
        ],
        col_widths=[2.5, 4, 1.5, 1.5, 1.5, 6]
    )
    add_para(doc, "环境：五服务全 UP（Redis 6379/ZooKeeper 21811/Served 8010/Adapter 8090/Frontend 8091），新 bean HnnxBankEbank2005AtomImpl 已随启动加载（bean 顶替日志佐证）。门禁结论：通过——本期执行段 16/16 通过且无未关闭缺陷。")

    add_heading(doc, "5.3 缺陷清单", level=2)
    add_para(doc, "本轮未发现代码缺陷（无 P0/P1/P2/P3 缺陷登记）。16 条用例本期执行段全部通过，取数口径/文案模板/开关复用/bean 顶替机制与 spec、PRD §3.2.3/§3.2.4、Q-13 用户确认结论逐项一致。测试备注 NOTE-01~06 均为数据/环境/配置类事项，非代码缺陷（详见第六章已知问题与附录 B 证据路径）。")

    add_heading(doc, "5.4 数据就绪结论（六项就绪检查实测）", level=2)
    add_table(doc,
        ["检查项", "实测结果", "结论"],
        [
            ["[检查1] ⑤在途记录", "4 行（NES.003.20.00P），BILL_ID 全部 JOIN 成功；DRWR_SOC_CODE=91861230741970469R（=R22.CERT_NO，round1 命中前提）、ACPT_BANK_NO=307584007998", "就绪"],
            ["[检查2] ⑥在途记录", "0 行（NES.006.20.00P）", "缺失（仅影响 P0-029/030 ⏸ 报文段，经网银在途业务构造，不直插）"],
            ["[检查3] NES.015 计数", "0 行", "符合预期（P0-031 报文段专项构造项）"],
            ["[检查4] ST01 反查链", "4 行，ACPT_BANK_NO=307584007998→BRCH_STATUS='ST01'→SOC_CODE=91440300192185379H 非空", "就绪"],
            ["[检查5] 开关参数", "1 条，PARAM_VALUE='1'/LEGAL_NO='000000'（默认开启，零新增）", "就绪"],
            ["[检查6] 名单基线", "R1（黑/21/长期有效/18/1）✓ R2（灰/21/20250101~20281231/18/1）✓ R22（CERT_TYPE='22'/91861230741970469R/长期有效/18/1）✓ 已就绪（四阶段\"待业务方准备\"标注闭环）；全库 DATA_SRC=18/VLD_ST=1 共 10 条（黑 8/灰 2），无口径不符记录", "就绪"],
        ],
        col_widths=[3.5, 9.5, 4]
    )
    add_para(doc, "就绪结论：可执行测试（本期执行段=代码审查+单测+DB 段全部落实；缺失项均为 ⏸ 报文段构造前提，不阻塞）。零新增直插合规：⑥/NES.015 在途记录经网银在途业务构造（直插破坏触发器/状态机业务语义），R3/R4 口径对照如需场景验证按四阶段 §9.2 直插 SQL（口径对照专用+执行后删除）。本轮全部为只读检查（SELECT），无数据变更，开关终态='1'。")

    add_heading(doc, "5.5 覆盖度分析", level=2)
    add_table(doc,
        ["覆盖维度", "覆盖情况"],
        [
            ["需求覆盖度", "spec ADDED Requirements 全部 3 项（⑤收票待签收/⑥背书待签收/校验范围与复用约束）及全部 7 个 Scenario 均有对应用例（P0-027~034/P1-033~038/P2-031~032）；PRD §3.2.3 ⑤⑥操作路径与校验角色、§3.2.4 取数逻辑逐项映射至 P0-033 取数口径用例与单测 build 系列"],
            ["测试类型覆盖度", "范围判定（3 单测+P1-033/034）、取数组装（4 单测+P0-033/P2-031）、端到端命中拦截（6 单测+P0-027~031）、开关与降级（2 单测+P0-032/P1-035）、边界与产品报错保持（P1-036/P2-032）、部署生效（P1-037 启动日志）"],
            ["机制级覆盖", "bean 顶替链（顶替注册日志 L176/运行时实例 L3014/单笔多笔两路径/唯一顶替者/异常转应答两路径）经 P1-037 五环节证据闭环"],
        ],
        col_widths=[3.5, 13.5]
    )

    add_page_break(doc)

    # ===== 第六章 已知问题清单 =====
    add_heading(doc, "第六章 已知问题清单", level=1)
    add_para(doc, "以下为 P2/P3 级及环境/配置类已知事项（无 P0/P1 未关闭缺陷），均在交付放行口径内登记：")
    add_table(doc,
        ["编号", "级别", "问题描述", "影响评估", "后续计划"],
        [
            ["KI-01", "遗留项", "⏸ 报文段主断言 6 条（P0-027~032：PICE070105 应答 retCode/retMsg、签收不落地、ECDS/CPES 报文不发出、正反对照）+ 可选报文对照 2 条（P1-033/P1-038）+ P2-032 可选对照", "⑤数据前提已就绪（检查 1/4/5/6），⑥在途记录需经网银在途业务构造（0 行，NOTE-01）；代码审查段+单测 17 全绿已覆盖主断言逻辑", "回归条件=网银外围 PICE070105 联调环境/报文模拟器就绪后随报文联调批次执行；R22 round1 优先命中断言（P0-027/028/029 步骤 3）同批执行"],
            ["KI-02", "已知项", "sonar java:S112 ×10（测试方法签名 throws Exception，行 399~629），属 BEMP 已知豁免清单（scan_config.json known_false_positive_rules）", "仅测试代码，主代码 0 问题；不影响生产", "按豁免清单登记；如需清零可在测试方法体未抛 checked 异常时删除 throws 子句后重扫"],
            ["KI-03", "已知项", "sonar java:S100 ×16（测试方法名下划线 BDD 风格命名）", "项目测试命名惯例，无功能影响", "维持 ACCEPTED 决策保留；如需调整属质量配置决策（调整 S100 规则 format 正则）"],
            ["KI-04", "环境提示", "TB_BILL_INFO 测试库列差异：Oracle MCP 直连演示库（localhost:orcl/BEMP）缺 PYEE_SOC_CODE 且含 mapper 外列 PYEE_ACCT_NAME；应用实际库 BEMP_HNNX 结构完整（8 列全部命中，USER_TAB_COLUMNS 复核留档）", "仅影响经 MCP 直连执行的脚本（失真）；应用运行不受影响", "统一走 run-oracle-jdbc.ps1 -Schema BEMP_HNNX 通道执行数据库检查（NOTE-03）"],
            ["KI-05", "环境提醒", "SonarQube ES 数据盘使用率 95%，flood_stage 水位临时调高（97%，transient 配置重启后失效）——重启后只读块会再次触发", "影响 SonarQube 扫描可用性（L3 自愈触发源），不影响业务系统", "建议尽快人工清理 ES 数据盘（旧索引/日志），清理后恢复标准水位；另 application.properties 中 sqa_ Token 已失效建议同步更新"],
            ["KI-06", "设计取舍", "覆写层与 super 各执行一次 TB_TRANS_INFO_SIGN 主键查询（重复查询）", "主键级查询开销可接受；为保持覆写与产品逻辑解耦的有意取舍（与四阶段同模式），代码走查提示级 1 项不修复", "已文档化（主类注释/tasks.md/走查报告），无需后续动作"],
            ["KI-07", "环境提示", "名单表分库声明与实测不符：测试数据准备 §十 声明的 ecifdb 名单库（11.1.199.35:1521/ecifdb）在本轮 JDBC 通道不可达，HNNX_M_CUST_SPECIAL_INFO 实测位于 BEMP_HNNX 主库", "仅文档口径差异，检查已按实测库执行", "建议后续版本更新 §十 分库声明，统一按实测库维护（NOTE-02）"],
        ],
        col_widths=[1.5, 1.8, 5.5, 4.2, 4]
    )

    add_page_break(doc)

    # ===== 第七章 部署说明 =====
    add_heading(doc, "第七章 部署说明", level=1)

    add_heading(doc, "7.1 SQL 脚本与配置项", level=2)
    add_para(doc, "本次增量零新增 SQL 脚本、零新增参数、零新增错误码：客户侧校验开关 hnnx.cust.antiMoneyList.block（TM_BUSINESS_PARAMETER，LEGAL_NO='000000'，PARAM_VALUE='1' 默认开启）与面客文案 HNNX0BE320011 均为四阶段已登记项，复用零变更。涉及 5 张表全部只读，无 DDL。")

    add_heading(doc, "7.2 部署内容与部署验证", level=2)
    add_para(doc, "部署内容：hnnxbank-biz-as 新增主类与测试类随应用构建发布（served 聚合与 banks ext-hnnxbank reactor 编译 BUILD SUCCESS 已验证）。部署顺序：随 hnnxbank 常规版本发布，无前置依赖变更。部署验证要点：")
    add_table(doc,
        ["验证项", "预期", "佐证"],
        [
            ["bean 顶替注册", "启动日志出现 Customized bean hnnxBankEbank2005AtomImpl, remove the bean ebank2005AtomImpl", "Served_startup_20260901_002153.log L176"],
            ["运行时实例", "ebank2005AtomImpl 运行时实例=HnnxBankEbank2005AtomImpl$$EnhancerBySpringCGLIB", "同日志 L3014（JresBeanPostProcessor INFO）"],
            ["开关参数在位", "TM_BUSINESS_PARAMETER 中 hnnx.cust.antiMoneyList.block PARAM_VALUE='1'", "DB 检查 5（默认开启，零新增）"],
            ["签收回归", "其他签收类型（提示付款/承兑/清偿应答）行为与改造前完全一致", "单测 zeroQuery 系列+P1-033 审查段"],
        ],
        col_widths=[3.5, 7, 6.5]
    )

    add_page_break(doc)

    # ===== 附录 =====
    add_heading(doc, "附录A 需求确认闭环结论（Q-13 详细）", level=1)
    add_table(doc,
        ["事项", "门禁判断", "确认结论", "闭环记录"],
        [
            ["Q-13-A：收票待签收/背书待签收两新阶段是否纳入本期实现", "不阻塞市场交易侧功能测试；实现范围不假设，待需求方确认后编制用例", "纳入本期开发范围（用户确认回复问题 13-A，2026-08-31）", "PRD v3.5.0：§3.2.3 本期范围说明、§3.2.4 实现范围同步更新；对应用例随实现编制（P0-027~031 等）"],
            ["Q-13-B：背书待签收 trans_from/trans_to 角色映射方向（与既有 tb_endrsmt_info 语义相反）", "默认按用户提供口径执行、实现前复核", "待签收业务 from/to 方向就是反向的，口径无问题——trans_from=被背书人、trans_to=背书人，按此实现无需复核（用户确认回复问题 13-B，2026-08-31）", "PRD v3.5.0 §3.2.4 角色映射确认块更新；实现注释留痕（buildEndorsementSignRoleItems L201-206）；单测 testBuildEndorsementSignRoleItems_reverseFromTo 交叉断言验证"],
        ],
        col_widths=[4.5, 3.5, 5, 4]
    )
    add_para(doc, "Q-13 全部闭环后两新阶段无阻塞项，转入方案设计/代码开发环节（bemp-personalized-developer），经代码评审（0 阻塞/0 严重）、sonar 扫描（主代码 0 问题）、启动服务（五服务全 UP）、用例评审、功能测试（16/16）后进入本交付阶段。")

    add_heading(doc, "附录B 参考资料与证据路径索引", level=1)
    add_table(doc,
        ["类别", "资料/路径"],
        [
            ["需求依据", "docs/prd/反洗钱校验功能/HNNS-EB-STD-REQ-002-反洗钱功能优化-v3.md（v3.5.0 §3.2.3/§3.2.4/Q-13）；需求原文：HNNS-EB-STD-REQ-002-票据业务系统反洗钱工作对接-日常优化类(公司金融部)(20260831).docx"],
            ["实现规格", ".trae/specs/add-hnnx-aml-sign-stage-check/spec.md + tasks.md + checklist.md（全部勾选）"],
            ["实现代码", "banks/ext-hnnxbank/hnnxbank-biz-as/src/main/java/com/hundsun/bemp/hnnxbank/biz/channel/atom/ebank/impl/HnnxBankEbank2005AtomImpl.java（260 行）；测试类 HnnxBankEbank2005AtomAmlSignRoleTest.java（同包 test 目录）"],
            ["复用组件（只读）", "HnnxAntiMoneyValidateUtil.java（validateCustBizList/RoleCheckItem/queryAcptBankSocCode）；产品化 Ebank2005AtomImpl/Ebank2005ServiceImpl（L46-47/L69/L104/L124）/Ebank2005CallableImpl（L27-46）"],
            ["功能测试报告", "aotutests-devtools/testcases/反洗钱名单校验/六阶段功能测试报告-AML-FT-20260831-01.md（16/16 通过、零缺陷、NOTE-01~06、门禁通过）"],
            ["单测基线", "Tests run: 17, Failures: 0, Errors: 0（HnnxBankEbank2005AtomAmlSignRoleTest）；回归 79 全绿（含既有 62）"],
            ["质量报告", "sonar 扫描报告 .trae/skills/bemp-sonarqube-mcp/reports/hnnxbank-ebank2005-aml-scan-report-20260831.md（主代码 0 问题，S1192/S3011 全核销，S112×10 豁免清单已知项，S100×16 ACCEPTED）；后端代码走查报告 .trae/skills/bemp-backend-code-review/reports/hnnxbank_2026-08-31_211000_incremental_report.md（0 阻塞/0 严重）"],
            ["启动日志", ".trae/skills/bemp-automation-startserver/logs/Served_startup_20260901_002153.log（L176 顶替注册/L3014 运行时实例）"],
            ["面客文案模板", "banks/ext-hnnxbank/hnnxbank-conf/src/main/resources/adapter/prod/mt/i18n/hnnx_mt_zh_CN.properties（L11 HNNX0BE320011）"],
            ["同域交付物", "docs/delivery/反洗钱四阶段校验/反洗钱四阶段校验-详细设计说明书-v1.docx（四阶段交付，本篇为其六阶段收官增量）"],
        ],
        col_widths=[3, 14]
    )

    inject_update_fields(doc)
    os.makedirs(OUTPUT_DIR, exist_ok=True)
    doc.save(OUTPUT_FILE)
    print("OK: {}".format(OUTPUT_FILE))


if __name__ == "__main__":
    build_design_doc()
