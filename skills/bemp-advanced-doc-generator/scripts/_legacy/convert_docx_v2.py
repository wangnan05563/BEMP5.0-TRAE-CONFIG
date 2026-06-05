# -*- coding: utf-8 -*-
import docx
import sys
import re

def convert_docx_to_markdown(docx_path, md_path):
    doc = docx.Document(docx_path)
    
    md_content = []
    md_content.append('# 2025年票据系统优化项目')
    md_content.append('## 业务需求说明书')
    md_content.append('')
    
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        
        style_name = para.style.name if para.style else ''
        
        # 检测标题级别
        if 'Heading 1' in style_name or '标题 1' in style_name:
            md_content.append(f'# {text}')
        elif 'Heading 2' in style_name or '标题 2' in style_name:
            md_content.append(f'## {text}')
        elif 'Heading 3' in style_name or '标题 3' in style_name:
            md_content.append(f'### {text}')
        elif 'Heading 4' in style_name or '标题 4' in style_name:
            md_content.append(f'#### {text}')
        elif 'Heading 5' in style_name or '标题 5' in style_name:
            md_content.append(f'##### {text}')
        elif 'Heading 6' in style_name or '标题 6' in style_name:
            md_content.append(f'###### {text}')
        else:
            # 正文内容，保留换行
            md_content.append(text)
    
    # 处理表格
    table_idx = 0
    for table in doc.tables:
        table_idx += 1
        md_content.append('')
        rows = []
        for row in table.rows:
            cells = []
            for cell in row.cells:
                cell_text = cell.text.strip().replace('\n', ' ').replace('\r', '')
                cells.append(cell_text)
            
            if cells and any(c for c in cells):
                rows.append('| ' + ' | '.join(cells) + ' |')
        
        if len(rows) > 0:
            header_count = len(rows[0].split('|')) - 2
            if header_count > 0:
                md_content.append(rows[0])
                md_content.append('| ' + ' | '.join(['---'] * header_count) + ' |')
                for row in rows[1:]:
                    md_content.append(row)
        md_content.append('')
    
    with open(md_path, 'w', encoding='utf-8') as f:
        f.write('\n'.join(md_content))
    
    print(f"✓ 转换完成: {docx_path}")
    print(f"✓ 输出文件: {md_path}")
    print(f"✓ 总段落数: {len(md_content)}")

if __name__ == '__main__':
    docx_path = r"d:\code\QJ\BEMP5.0DEV\附件6：2025年票据系统优化项目-业务需求说明书v1.0-20250806.docx"
    md_path = r"d:\code\QJ\BEMP5.0DEV\.trae\skills\bemp-advanced-doc-generator\output\requirement_full.md"
    convert_docx_to_markdown(docx_path, md_path)
