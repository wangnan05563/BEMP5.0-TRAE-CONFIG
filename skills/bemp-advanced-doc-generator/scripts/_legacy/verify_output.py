import sys
from docx import Document
from collections import Counter

output_file = r'd:\code\QJ\BEMP5.0DEV\.trae\skills\bemp-advanced-doc-generator\output\verify_result.txt'

doc = Document(r'd:\code\QJ\BEMP5.0DEV\.trae\skills\bemp-advanced-doc-generator\output\BEMP5.0DEV-概要设计说明书-20260601.docx')

lines = []

lines.append('=== 1. 关键章节内容检查 ===')
check_sections = ['编写目的', '适用范围', '读者对象', '系统概述', '设计约束', '设计策略']
for i, p in enumerate(doc.paragraphs):
    style = p.style.name if p.style else 'None'
    text = p.text.strip()
    for kw in check_sections:
        if kw in text and (style.startswith('Heading') or len(text) > 30):
            display = text[:120] if text else '(空)'
            lines.append(f'{i:3d} | {style:12s} | {display}')
            break

lines.append('')
lines.append('=== 2. 重复内容检测 ===')
all_texts = [p.text.strip() for p in doc.paragraphs if p.text.strip() and len(p.text.strip()) > 30]
counter = Counter(all_texts)
duplicates = {k: v for k, v in counter.items() if v > 1}
if duplicates:
    for text, count in duplicates.items():
        lines.append(f'  重复{count}次: {text[:80]}...')
else:
    lines.append('  未发现重复内容')

lines.append('')
lines.append('=== 3. 数据库设计章节检查 ===')
found_er = False
for i, p in enumerate(doc.paragraphs):
    style = p.style.name if p.style else 'None'
    text = p.text.strip()
    if '数据库设计' in text and style.startswith('Heading'):
        found_er = True
        lines.append(f'  找到ER图章节: {i:3d} | {style} | {text}')
    elif found_er and style.startswith('Heading 2'):
        lines.append(f'  ER分组: {i:3d} | {text[:60]}')
    elif found_er and style.startswith('Heading 1') and '数据库设计' not in text:
        break
if not found_er:
    lines.append('  未找到数据库设计章节')

lines.append('')
lines.append('=== 4. 附录章节位置 ===')
for i, p in enumerate(doc.paragraphs):
    style = p.style.name if p.style else 'None'
    text = p.text.strip()
    if '附录' in text and style.startswith('Heading'):
        lines.append(f'  附录位置: {i:3d} | {style} | {text}')

lines.append('')
lines.append('=== 5. 所有Heading 1章节 ===')
for i, p in enumerate(doc.paragraphs):
    style = p.style.name if p.style else 'None'
    text = p.text.strip()
    if style == 'Heading 1':
        lines.append(f'  {i:3d} | {text}')

lines.append('')
lines.append('=== 6. 编写目的/适用范围内容详情 ===')
for i, p in enumerate(doc.paragraphs):
    style = p.style.name if p.style else 'None'
    text = p.text.strip()
    if i >= 73 and i <= 78:
        display = text[:150] if text else '(空)'
        lines.append(f'{i:3d} | {style:12s} | {display}')

with open(output_file, 'w', encoding='utf-8') as f:
    f.write('\n'.join(lines))

print('Done')
