# -*- coding: utf-8 -*-
import docx
import sys
import re

def convert_docx_to_markdown(docx_path, md_path):
    doc = docx.Document(docx_path)
    
    md_content = []
    
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        
        style_name = para.style.name if para.style else ''
        
        if 'Heading' in style_name or '标题' in style_name:
            level = 1
            if '2' in style_name or '二级' in style_name:
                level = 2
            elif '3' in style_name or '三级' in style_name:
                level = 3
            elif '4' in style_name or '四级' in style_name:
                level = 4
            elif '5' in style_name or '五级' in style_name:
                level = 5
            elif '6' in style_name or '六级' in style_name:
                level = 6
            
            md_content.append(f"{'#' * level} {text}")
        else:
            md_content.append(text)
    
    for table in doc.tables:
        md_content.append('')
        rows = []
        for row in table.rows:
            cells = [cell.text.strip().replace('\n', ' ') for cell in row.cells]
            rows.append('| ' + ' | '.join(cells) + ' |')
        
        if len(rows) > 0:
            header_count = len(rows[0].split('|')) - 2
            md_content.append(rows[0])
            md_content.append('| ' + ' | '.join(['---'] * header_count) + ' |')
            for row in rows[1:]:
                md_content.append(row)
        md_content.append('')
    
    with open(md_path, 'w', encoding='utf-8') as f:
        f.write('\n'.join(md_content))
    
    print(f"Converted: {docx_path} -> {md_path}")

if __name__ == '__main__':
    docx_path = r"d:\code\QJ\BEMP5.0DEV\附件6：2025年票据系统优化项目-业务需求说明书v1.0-20250806.docx"
    md_path = r"d:\code\QJ\BEMP5.0DEV\.trae\skills\bemp-advanced-doc-generator\output\requirement_temp.md"
    convert_docx_to_markdown(docx_path, md_path)
