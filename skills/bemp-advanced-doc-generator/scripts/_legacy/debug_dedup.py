import sys, json, shutil
from docx import Document
from docx.shared import RGBColor

BLACK = RGBColor(0, 0, 0)
BLUE = RGBColor(0, 0, 255)

def is_blue_paragraph(p):
    for r in p.runs:
        if r.font.color and r.font.color.rgb and r.font.color.rgb == BLUE:
            return True
    return False

def clear_paragraph(p):
    for r in p.runs:
        r.text = ''
    if not p.runs:
        p.add_run('')

def write_paragraph(p, text):
    clear_paragraph(p)
    p.runs[0].text = text
    set_black(p)

def set_black(p):
    for r in p.runs:
        r.font.color.rgb = BLACK

def get_heading_context(paragraphs, current_idx):
    h1 = h2 = h3 = ''
    for i in range(current_idx - 1, -1, -1):
        p = paragraphs[i]
        style = p.style.name if p.style else ''
        if style == 'Heading 1':
            if not h1:
                h1 = p.text.strip()
            break
        elif style == 'Heading 2' and not h2:
            h2 = p.text.strip()
        elif style == 'Heading 3' and not h3:
            h3 = p.text.strip()
    return h1, h2, h3

template_path = r'd:\code\QJ\BEMP5.0DEV\.trae\skills\bemp-advanced-doc-generator\assets\template-outline-design.docx'
doc = Document(template_path)
paragraphs = doc.paragraphs

filled_contexts = set()

for i, p in enumerate(paragraphs):
    if i < 67 or i > 92:
        continue
    text = p.text.strip()
    h1, h2, h3 = get_heading_context(paragraphs, i)
    context_key = f'{h1}|{h2}|{h3}'
    is_blue = is_blue_paragraph(p)
    style = p.style.name if p.style else 'None'

    print(f'{i:3d} | {style:12s} | blue={is_blue} | ctx={context_key!r} | in_filled={context_key in filled_contexts} | text={text[:60]}')

    if context_key in filled_contexts:
        if is_blue:
            clear_paragraph(p)
            print(f'     -> SKIP+CLEAR (already filled)')
        continue

    if context_key not in filled_contexts:
        pass

    if not is_blue:
        print(f'     -> SKIP (not blue)')
        continue

    if '注意' in text and '正文' in text:
        clear_paragraph(p)
        print(f'     -> CLEAR (注意+正文)')
        continue

    if '注意' in text or '建议' in text or '编号' in text:
        clear_paragraph(p)
        print(f'     -> CLEAR (注意/建议/编号)')
        continue

    if h1 == '概述':
        if h2 == '编写目的':
            if '编写目的' in text or '此处填写' in text or '例如' in text:
                new_text = 'TEST: 本文档的编写目的内容...'
                write_paragraph(p, new_text)
                print(f'     -> WRITE: {new_text[:50]}')
            else:
                print(f'     -> NO MATCH in 编写目的 conditions')
        elif h2 == '适用范围':
            if '此处填写' in text or '例如' in text:
                new_text = 'TEST: 本文档的适用范围内容...'
                write_paragraph(p, new_text)
                print(f'     -> WRITE: {new_text[:50]}')
            else:
                print(f'     -> NO MATCH in 适用范围 conditions')
        else:
            print(f'     -> NO MATCH for h2={h2!r}')

    if context_key not in filled_contexts:
        if is_blue_paragraph(p):
            clear_paragraph(p)
            print(f'     -> POST-CLEAR (still blue after processing)')
        filled_contexts.add(context_key)
        print(f'     -> ADD to filled_contexts: {context_key!r}')

print('\n=== 验证结果 ===')
for i, p in enumerate(paragraphs):
    if i < 73 or i > 78:
        continue
    text = p.text.strip()[:80]
    style = p.style.name if p.style else 'None'
    print(f'{i:3d} | {style:12s} | {text}')
