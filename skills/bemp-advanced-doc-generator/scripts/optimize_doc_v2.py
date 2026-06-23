"""优化详细设计文档v2版本"""
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import copy

# 读取v1文档
v1_path = r"d:\code\QJ\BEMP5.0DEV\.trae\skills\bemp-advanced-doc-generator\output\机构管理和管理员管理功能优化-详细设计文档-20260617.docx"
v2_path = r"d:\code\QJ\BEMP5.0DEV\.trae\skills\bemp-advanced-doc-generator\output\机构管理和管理员管理功能优化-详细设计文档-v2-20260617.docx"

doc = Document(v1_path)

# === 优化1: 更新业务背景 ===
for para in doc.paragraphs:
    if "当前系统存在以下业务需求" in para.text:
        para.clear()
        run = para.add_run("系统目前需手工新增机构及管理员，人工处理量大；后台批量添加机构科技暂无相关功能，增加批量导入功能。本优化为机构管理和机构管理员管理模块增加批量导入、模板下载、批量复制角色功能，提升运维效率。")
        run.font.size = Pt(10.5)
        break

# === 优化2: 更新设计目标 ===
for table in doc.tables:
    for row in table.rows:
        if row.cells[0].text.strip() == "功能目标":
            row.cells[1].text = "实现机构批量导入、机构管理员批量导入、模板下载、批量复制角色等功能，提升运维效率"
            break

# === 优化3: 更新范围说明 ===
for table in doc.tables:
    for row in table.rows:
        if row.cells[0].text.strip() == "纳入范围":
            row.cells[1].text = "机构批量导入、机构管理员批量导入、模板下载、批量复制角色（机构级/管理员级）"
            break

# === 优化4: 更新模块1设计说明-功能描述 ===
for para in doc.paragraphs:
    if "机构管理包含以下核心职责" in para.text:
        para.clear()
        run = para.add_run("机构管理模块包含以下核心功能：\n1. 批量导入机构：支持通过Excel文件批量导入机构信息，校验机构号、机构名称、核算机构号、组织机构代码唯一性\n2. 模板下载：提供机构信息导入模板下载功能\n3. 批量复制角色：将选中机构的角色权限批量复制到目标机构，支持多选目标机构")
        run.font.size = Pt(10.5)
        break

# === 优化5: 更新模块1设计说明-模块职责 ===
for para in doc.paragraphs:
    if "机构管理和管理员管理功能优化核心业务处理逻辑" in para.text:
        para.clear()
        run = para.add_run("机构管理核心业务处理逻辑：\n1. 批量导入校验：校验上级机构选择、机构号/名称/核算机构号/组织机构代码唯一性、机构层级限制（最多4级）\n2. 批量导入执行：校验通过后批量插入机构数据\n3. 批量复制角色：用新角色覆盖旧角色，做关联角色的新增和删除，校验角色是否被柜员使用")
        run.font.size = Pt(10.5)
        break

# === 优化6: 更新模块2设计说明 ===
# 找到"模块2设计说明"标题后的段落
found_module2 = False
for i, para in enumerate(doc.paragraphs):
    if para.text.strip() == "模块2设计说明":
        found_module2 = True
        # 找到下一个非标题段落
        for j in range(i+1, min(i+10, len(doc.paragraphs))):
            next_para = doc.paragraphs[j]
            if not next_para.style.name.startswith('Heading') and next_para.text.strip():
                next_para.clear()
                run = next_para.add_run("机构管理员管理模块包含以下核心功能：\n1. 批量导入管理员：支持通过Excel文件批量导入机构管理员，只有法人管理员可操作，新增管理员默认为无效状态\n2. 模板下载：提供机构管理员导入模板下载功能\n3. 批量复制角色：将选中管理员的角色权限批量复制到目标管理员，只能分配管理员所在机构所拥有的角色")
                run.font.size = Pt(10.5)
                break
        break

# === 优化7: 更新接口定义章节 ===
# 找到"接口定义"标题后的内容
found_api_section = False
for i, para in enumerate(doc.paragraphs):
    if para.text.strip() == "接口定义":
        found_api_section = True
        # 找到"API 接口清单"标题后的段落
        for j in range(i+1, min(i+20, len(doc.paragraphs))):
            next_para = doc.paragraphs[j]
            if "各接口的请求参数" in next_para.text:
                next_para.clear()
                run = next_para.add_run("本模块涉及7个核心接口，详细定义如下：")
                run.font.size = Pt(10.5)
                break
        break

# === 优化8: 更新接口清单表格 ===
api_table_data = [
    ["接口路径", "方法", "说明"],
    ["/hnnxbank/sm/auth/branch/branch/func_batchImportValidate", "POST", "机构批量导入校验"],
    ["/hnnxbank/sm/auth/branch/branch/func_batchImport", "POST", "机构批量导入执行"],
    ["/hnnxbank/sm/auth/branch/branch/func_downloadModel", "GET", "机构导入模板下载"],
    ["/hnnxbank/sm/auth/branch/branchAdmin/func_batchImportValidate", "POST", "管理员批量导入校验"],
    ["/hnnxbank/sm/auth/branch/branchAdmin/func_batchImportBranchAdmin", "POST", "管理员批量导入执行"],
    ["/hnnxbank/sm/auth/branch/branchAdmin/func_downloadModel", "GET", "管理员导入模板下载"],
    ["/hnnxbank/sm/auth/branch/branchAdmin/func_batchCopyRole", "POST", "管理员批量复制角色"]
]

# 找到API接口清单表格并更新
for table in doc.tables:
    if len(table.rows) > 0 and len(table.rows[0].cells) >= 3:
        first_row_text = " | ".join([cell.text.strip() for cell in table.rows[0].cells[:3]])
        if "接口路径" in first_row_text or "API" in first_row_text:
            # 清空表格并重新填充
            for i, row_data in enumerate(api_table_data):
                if i < len(table.rows):
                    for j, cell_text in enumerate(row_data):
                        if j < len(table.rows[i].cells):
                            table.rows[i].cells[j].text = cell_text
            break

# === 优化9: 删除冗余章节（技术实现细节、安全策略、异常处理机制） ===
# 标记需要删除的段落索引
paragraphs_to_delete = []
skip_sections = ["技术实现细节", "安全策略", "异常处理机制"]
current_skip = False

for i, para in enumerate(doc.paragraphs):
    if para.style.name.startswith('Heading'):
        # 检查是否是需要删除的章节标题
        heading_text = para.text.strip()
        if any(skip in heading_text for skip in skip_sections):
            current_skip = True
        elif para.style.name == 'Heading 1' and not any(skip in heading_text for skip in skip_sections):
            current_skip = False
    
    if current_skip:
        paragraphs_to_delete.append(i)

# 从后往前删除段落，避免索引变化
for idx in reversed(paragraphs_to_delete):
    para = doc.paragraphs[idx]
    p = para._element
    p.getparent().remove(p)

# === 优化10: 更新附录内容 ===
for para in doc.paragraphs:
    if "数据表清单详见数据库设计文档" in para.text:
        para.clear()
        run = para.add_run("本模块涉及的数据表包括：机构信息表、机构管理员表、角色关联表、系统参数表等，详细表结构参见数据库设计文档。")
        run.font.size = Pt(10.5)
        break

# 保存v2文档
doc.save(v2_path)
print(f"✓ v2文档已生成: {v2_path}")
print(f"优化内容：")
print(f"  1. 更新业务背景描述")
print(f"  2. 更新设计目标")
print(f"  3. 更新范围说明")
print(f"  4. 更新模块1设计说明（机构管理）")
print(f"  5. 更新模块2设计说明（机构管理员管理）")
print(f"  6. 更新接口定义（7个API接口）")
print(f"  7. 删除冗余章节（技术实现细节、安全策略、异常处理机制）")
print(f"  8. 更新附录内容")
