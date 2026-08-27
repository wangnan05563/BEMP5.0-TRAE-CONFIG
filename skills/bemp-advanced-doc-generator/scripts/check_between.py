from pathlib import Path
SKILL_ROOT = Path(__file__).resolve().parent.parent
PROJECT_ROOT = SKILL_ROOT.parent.parent.parent

from pathlib import Path
"""检查模块2和技术实现细节之间的所有内容"""
from docx import Document

doc_path = str(SKILL_ROOT / "output" / "机构管理和管理员管理功能优化-详细设计文档-20260617.docx")
doc = Document(doc_path)

print("=" * 70)
print("检查位置78-85的所有段落")
print("=" * 70)

for i in range(78, 86):
    if i >= len(doc.paragraphs):
        print(f"[{i}] 超出范围")
        continue
    
    para = doc.paragraphs[i]
    style_name = para.style.name if para.style else 'None'
    text = para.text.strip()
    
    print(f"[{i}] 样式={style_name}, 文本={text[:80] if text else '(空)'}")
