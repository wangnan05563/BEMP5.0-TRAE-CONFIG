from pathlib import Path
SKILL_ROOT = Path(__file__).resolve().parent.parent
PROJECT_ROOT = SKILL_ROOT.parent.parent.parent

from pathlib import Path
"""详细诊断模块2的完整结构"""
from docx import Document

doc_path = str(SKILL_ROOT / "output" / "机构管理和管理员管理功能优化-详细设计文档-20260617.docx")
doc = Document(doc_path)

print("=" * 70)
print("模块2完整结构诊断")
print("=" * 70)

# 找到所有H1标题
h1_list = []
for i, para in enumerate(doc.paragraphs):
    if para.style and para.style.name == 'Heading 1':
        h1_list.append((i, para.text.strip()))

print("\n所有H1标题:")
for idx, (i, text) in enumerate(h1_list):
    print(f"  [{idx}] 位置{i}: {text}")

# 找到模块2的位置
module2_idx = None
module2_pos = None
for idx, (i, text) in enumerate(h1_list):
    if '模块2' in text:
        module2_idx = idx
        module2_pos = i
        break

if module2_pos is None:
    print("\n❌ 未找到模块2!")
else:
    print(f"\n模块2位置: H1列表索引{module2_idx}, 段落位置{module2_pos}")
    
    # 找到下一个H1的位置
    next_h1_pos = None
    if module2_idx + 1 < len(h1_list):
        next_h1_pos = h1_list[module2_idx + 1][0]
        print(f"下一个H1位置: {next_h1_pos} ({h1_list[module2_idx + 1][1]})")
    
    # 统计模块2下的所有内容
    print("\n模块2下的所有内容:")
    h2_count = 0
    table_count = 0
    para_count = 0
    
    start_pos = module2_pos + 1
    end_pos = next_h1_pos if next_h1_pos else len(doc.paragraphs)
    
    for i in range(start_pos, end_pos):
        para = doc.paragraphs[i]
        style_name = para.style.name if para.style else 'None'
        text = para.text.strip()
        
        if style_name == 'Heading 2':
            h2_count += 1
            print(f"  [H2] 位置{i}: {text[:80] if text else '(空)'}")
        elif style_name == 'Heading 3':
            print(f"  [H3] 位置{i}: {text[:80] if text else '(空)'}")
        elif 'Normal' in style_name or style_name == 'None':
            if text:
                para_count += 1
                if para_count <= 5:  # 只显示前5个段落
                    print(f"  [段落] 位置{i}: {text[:80]}")
    
    # 统计表格
    for table in doc.tables:
        # 检查表格是否在模块2范围内
        # 简单方法:检查表格前的段落
        table_count += 1
    
    print(f"\n统计:")
    print(f"  H2数量: {h2_count}")
    print(f"  段落数量: {para_count}")
    print(f"  表格数量: {table_count}")
    
    if h2_count == 0:
        print("\n❌ 模块2下没有H2标题!")
    else:
        print(f"\n✅ 模块2下有 {h2_count} 个H2标题")
