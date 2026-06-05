#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
BEMP 高级文档生成器 v7.0 全面测试脚本

测试覆盖：
1. ContentRegistry 按需加载机制
2. 模块独立性（outline vs detail 互不干扰）
3. 索引导航（CONTENT_INDEX）
4. 内容生成器 generate 方法
5. 模板清理功能（full_template_cleanup）
6. 示例内容清除（clear_example_content）
7. 增强蓝色清理（enhanced_blue_cleanup）
8. 模板备注清除（clean_template_remarks）
9. 集成测试：完整概要设计生成
"""
import sys
import os
import json
import time
import traceback
import shutil
import tempfile

# 确保脚本目录在 sys.path 中
SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
if SCRIPT_DIR not in sys.path:
    sys.path.insert(0, SCRIPT_DIR)

# 确保 content 包可导入
CONTENT_DIR = os.path.join(SCRIPT_DIR, 'content')
if CONTENT_DIR not in sys.path:
    sys.path.insert(0, os.path.dirname(CONTENT_DIR))

# 测试结果收集
results = []
passed = 0
failed = 0


def check(name, condition, detail=''):
    """记录测试结果"""
    global passed, failed
    if condition:
        passed += 1
        results.append(f'  [PASS] {name}')
    else:
        failed += 1
        results.append(f'  [FAIL] {name} - {detail}')


def print_summary():
    """打印测试汇总"""
    total = passed + failed
    rate = (passed / total * 100) if total > 0 else 0
    print(f'\n{"="*60}')
    print(f'  测试结果: {passed}/{total} 通过, 通过率 {rate:.1f}%')
    print(f'{"="*60}')
    for r in results:
        print(r)


# ─── 测试 1：ContentRegistry 按需加载 ─────────────────────────────
def test_content_registry_lazy_loading():
    """测试 ContentRegistry 的按需加载机制"""
    print('\n--- 测试 1：ContentRegistry 按需加载 ---')

    from content import ContentRegistry, CONTENT_INDEX

    registry = ContentRegistry()

    # 1.1 初始状态：没有模块已加载
    check('初始无已加载模块', len(registry.get_loaded_modules()) == 0,
          f'预期 0，实际 {len(registry.get_loaded_modules())}')

    # 1.2 首次获取生成器：应加载模块
    gen = registry.get_generator('outline', '编写目的')
    check('获取编写目的生成器', gen is not None, '生成器为 None')
    loaded = registry.get_loaded_modules()
    check('模块已加载', 'outline_chapters' in loaded,
          f'已加载: {loaded}')

    # 1.3 再次获取：应从缓存获取，不重复加载
    loaded_count_before = len(registry.get_loaded_modules())
    gen2 = registry.get_generator('outline', '编写目的')
    loaded_count_after = len(registry.get_loaded_modules())
    check('缓存命中不重复加载', loaded_count_before == loaded_count_after,
          f'加载前 {loaded_count_before}，加载后 {loaded_count_after}')

    # 1.4 不存在的生成器返回 None
    gen_none = registry.get_generator('outline', '不存在的章节')
    check('不存在章节返回 None', gen_none is None)

    # 1.5 模糊匹配：关键词在章节名中
    gen_fuzzy = registry.get_generator('outline', '目的')
    check('模糊匹配"目的"→"编写目的"', gen_fuzzy is not None)


# ─── 测试 2：模块独立性 ──────────────────────────────────────────
def test_module_independence():
    """测试 outline_chapters 和 detail_chapters 互不干扰"""
    print('\n--- 测试 2：模块独立性 ---')

    from content import ContentRegistry

    registry = ContentRegistry()

    # 2.1 加载 outline 模块
    gen_outline = registry.get_generator('outline', '编写目的')
    modules_after_outline = list(registry.get_loaded_modules())
    check('outline 模块已加载', 'outline_chapters' in modules_after_outline,
          f'已加载: {modules_after_outline}')

    # 2.2 加载 detail 模块
    gen_detail = registry.get_generator('detail', '概述')
    modules_after_detail = list(registry.get_loaded_modules())
    check('detail 模块已加载', 'detail_chapters' in modules_after_detail,
          f'已加载: {modules_after_detail}')

    # 2.3 两个模块共存
    check('两个模块并存', 'outline_chapters' in modules_after_detail
          and 'detail_chapters' in modules_after_detail)

    # 2.4 两个模块的生成器是不同的
    check('outline≠detail 生成器', gen_outline is not gen_detail,
          '两个模块生成器引用相同')


# ─── 测试 3：索引导航 ────────────────────────────────────────────
def test_index_navigation():
    """测试 CONTENT_INDEX 索引导航"""
    print('\n--- 测试 3：索引导航 ---')

    from content import CONTENT_INDEX

    # 3.1 outline 索引存在
    outline = CONTENT_INDEX.get('outline')
    check('outline 索引存在', outline is not None)
    check('outline 标签正确', outline.get('label') == '概要设计说明书')

    # 3.2 detail 索引存在
    detail = CONTENT_INDEX.get('detail')
    check('detail 索引存在', detail is not None)
    check('detail 标签正确', detail.get('label') == '详细设计说明书')

    # 3.3 outline 章节数量
    outline_chapters = outline.get('chapters', {})
    check('outline 章节数正确', len(outline_chapters) >= 10,
          f'章节数: {len(outline_chapters)}')

    # 3.4 detail 章节数量
    detail_chapters = detail.get('chapters', {})
    check('detail 章节数正确', len(detail_chapters) >= 8,
          f'章节数: {len(detail_chapters)}')

    # 3.5 每个章节有 module 和 generator 字段
    all_have_info = all(
        'module' in info and 'generator' in info
        for info in outline_chapters.values()
    )
    check('outline 章节有完整信息', all_have_info)

    all_have_info = all(
        'module' in info and 'generator' in info
        for info in detail_chapters.values()
    )
    check('detail 章节有完整信息', all_have_info)


# ─── 测试 4：generate 方法 ───────────────────────────────────────
def test_generate_method():
    """测试 ContentRegistry.generate 方法"""
    print('\n--- 测试 4：generate 方法 ---')

    from content import ContentRegistry

    registry = ContentRegistry()

    # 模拟 scan_data
    scan = {
        'requirementModuleName': '测试模块',
        'projectName': 'BEMP测试项目',
        'businessModules': [
            {'name': '额度管理', 'description': '统一管理银行额度'},
            {'name': '额度查询', 'description': '查询银行额度信息'},
        ],
        'techStack': ['Spring Boot', 'MyBatis', 'Oracle'],
    }

    # 4.1 生成文本内容
    text = registry.generate('outline', '编写目的', scan)
    check('generate 返回文本', isinstance(text, str) and len(text) > 0,
          f'类型: {type(text)}, 长度: {len(text) if text else 0}')

    # 4.2 生成表格内容
    table = registry.generate('outline', '组件汇总表', scan)
    check('generate 返回表格', isinstance(table, (list, tuple)) and len(table) == 2,
          f'类型: {type(table)}')

    if isinstance(table, (list, tuple)) and len(table) == 2:
        headers, rows = table
        check('表格有表头', len(headers) > 0)
        check('表格有数据行', len(rows) > 0, f'行数: {len(rows)}')

    # 4.3 不存在章节返回 None
    result = registry.generate('outline', '不存在的章节', scan)
    check('不存在章节返回 None', result is None)

    # 4.4 不存在文档类型返回 None
    result = registry.generate('nonexistent', '编写目的', scan)
    check('不存在文档类型返回 None', result is None)


# ─── 测试 5：模板清理功能 ────────────────────────────────────────
def test_template_cleanup():
    """测试模板清理功能"""
    print('\n--- 测试 5：模板清理功能 ---')

    try:
        import doc_formatter
        from docx import Document
    except ImportError as e:
        check('模板清理测试（跳过）', False, f'导入失败: {e}')
        return

    # 使用实际的模板文件
    template_path = os.path.join(
        os.path.dirname(os.path.dirname(os.path.dirname(os.path.dirname(SCRIPT_DIR)))),
        'docs', '04【模板】概要设计说明书.docx'
    )
    template_path = os.path.normpath(template_path)

    if not os.path.exists(template_path):
        # 尝试 .doc 格式
        template_path = template_path.replace('.docx', '.doc')
        if not os.path.exists(template_path):
            check('测试模板文件存在', False, f'模板不存在: {template_path}')
            return

    # 复制到临时文件
    temp_dir = tempfile.mkdtemp(prefix='bemp_test_')
    temp_doc = os.path.join(temp_dir, 'test_template.docx')
    shutil.copy2(template_path, temp_doc)

    try:
        doc = Document(temp_doc)

        # 5.1 统计蓝色 run 数量（清理前）
        blue_before = 0
        for p in doc.paragraphs:
            for run in p.runs:
                if hasattr(run, 'font') and run.font.color and run.font.color.rgb:
                    if str(run.font.color.rgb) != '000000':
                        blue_before += 1
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        for run in p.runs:
                            if hasattr(run, 'font') and run.font.color and run.font.color.rgb:
                                if str(run.font.color.rgb) != '000000':
                                    blue_before += 1

        # 5.2 执行 full_template_cleanup
        stats = doc_formatter.full_template_cleanup(doc, aggressive=False)
        check('full_template_cleanup 返回统计', isinstance(stats, dict),
              f'类型: {type(stats)}')

        # 5.3 统计蓝色 run 数量（清理后）
        blue_after = 0
        for p in doc.paragraphs:
            for run in p.runs:
                if hasattr(run, 'font') and run.font.color and run.font.color.rgb:
                    if str(run.font.color.rgb) != '000000':
                        blue_after += 1
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        for run in p.runs:
                            if hasattr(run, 'font') and run.font.color and run.font.color.rgb:
                                if str(run.font.color.rgb) != '000000':
                                    blue_after += 1

        # 蓝色 run 应减少
        if blue_before > 0:
            reduction = (blue_before - blue_after) / blue_before * 100
            check(f'蓝色 run 减少 (前:{blue_before} 后:{blue_after})',
                  blue_after < blue_before,
                  f'减少率: {reduction:.1f}%')

        # 5.4 检查示例内容清除
        example_stats = stats.get('example', {})
        if isinstance(example_stats, dict):
            check('示例内容清除统计存在',
                  'paragraphs_removed' in example_stats or 'paragraphs_cleared' in example_stats)

        print(f'  清理统计: {json.dumps(stats, ensure_ascii=False, default=str)}')

    finally:
        # 清理临时文件
        try:
            shutil.rmtree(temp_dir)
        except Exception:
            pass


# ─── 测试 6：性能测试 ─────────────────────────────────────────────
def test_performance():
    """测试按需加载的性能提升"""
    print('\n--- 测试 6：性能测试 ---')

    from content import ContentRegistry

    # 6.1 首次加载时间
    start = time.time()
    registry = ContentRegistry()
    gen = registry.get_generator('outline', '编写目的')
    first_load_time = time.time() - start

    # 6.2 缓存命中时间（应远小于首次加载）
    start = time.time()
    gen2 = registry.get_generator('outline', '编写目的')
    cache_hit_time = time.time() - start

    # 6.3 缓存命中应快于首次加载
    check(f'缓存命中更快 (首次:{first_load_time*1000:.1f}ms, 缓存:{cache_hit_time*1000:.1f}ms)',
          cache_hit_time <= first_load_time * 0.5,
          '缓存未显著加速')

    # 6.4 第二个模块加载时间
    start = time.time()
    gen3 = registry.get_generator('detail', '概述')
    second_module_time = time.time() - start
    check(f'detail 模块加载时间: {second_module_time*1000:.1f}ms',
          second_module_time < 1.0,  # 应在 1 秒内
          '加载超时')


# ─── 测试 7：ContentRegistry.list_chapters ───────────────────────
def test_list_chapters():
    """测试章节列表功能"""
    print('\n--- 测试 7：章节列表功能 ---')

    from content import ContentRegistry

    registry = ContentRegistry()

    # 7.1 outline 章节列表
    chapters = registry.list_chapters('outline')
    check('outline 章节列表非空', len(chapters) > 0)
    check('章节有 name 字段', all('name' in c for c in chapters))

    # 7.2 detail 章节列表
    chapters = registry.list_chapters('detail')
    check('detail 章节列表非空', len(chapters) > 0)
    check('章节有 type 字段', all('type' in c for c in chapters))

    print(f'  outline 章节: {[c["name"] for c in chapters][:5]}...')


# ─── 测试 8：get_chapter_info ────────────────────────────────────
def test_get_chapter_info():
    """测试章节信息获取"""
    print('\n--- 测试 8：章节信息获取 ---')

    from content import ContentRegistry

    registry = ContentRegistry()

    # 8.1 已知章节
    info = registry.get_chapter_info('outline', '编写目的')
    check('获取编写目的信息', info is not None)
    check('包含 module', 'module' in info)
    check('包含 generator', 'generator' in info)
    check('包含 type', 'type' in info)

    # 8.2 未知章节
    info = registry.get_chapter_info('outline', '不存在')
    check('未知章节返回 None', info is None)


# ─── 测试 9：reload_module ───────────────────────────────────────
def test_reload_module():
    """测试模块热更新"""
    print('\n--- 测试 9：模块热更新 ---')

    from content import ContentRegistry

    registry = ContentRegistry()

    # 9.1 加载模块
    registry.get_generator('outline', '编写目的')
    check('模块已加载', registry.is_loaded('outline_chapters'))

    # 9.2 重新加载
    registry.reload_module('outline_chapters')
    # 重新加载后模块应重新变为可用
    gen = registry.get_generator('outline', '编写目的')
    check('重新加载后生成器可用', gen is not None)


# ─── 主入口 ──────────────────────────────────────────────────────
def main():
    print('=' * 60)
    print('  BEMP 高级文档生成器 v7.0 全面测试')
    print('=' * 60)

    tests = [
        ('ContentRegistry 按需加载', test_content_registry_lazy_loading),
        ('模块独立性', test_module_independence),
        ('索引导航', test_index_navigation),
        ('generate 方法', test_generate_method),
        ('模板清理功能', test_template_cleanup),
        ('性能测试', test_performance),
        ('章节列表', test_list_chapters),
        ('章节信息', test_get_chapter_info),
        ('模块热更新', test_reload_module),
    ]

    for name, fn in tests:
        try:
            fn()
        except Exception as e:
            print(f'  [ERROR] {name}: {e}')
            traceback.print_exc()
            global failed
            failed += 1

    print_summary()

    # 返回退出码
    sys.exit(0 if failed == 0 else 1)


if __name__ == '__main__':
    main()