"""验证v2文档质量"""
from docx import Document
import json

v2_path = r"d:\code\QJ\BEMP5.0DEV\.trae\skills\bemp-advanced-doc-generator\output\机构管理和管理员管理功能优化-详细设计文档-v2-20260617.docx"
doc = Document(v2_path)

result = {
    "headings": [],
    "key_paragraphs": [],
    "table_count": len(doc.tables),
    "paragraph_count": len(doc.paragraphs)
}

# 提取标题结构
for para in doc.paragraphs:
    if para.style.name.startswith('Heading'):
        level = int(para.style.name.replace('Heading ', ''))
        result["headings"].append({
            "level": level,
            "text": para.text.strip()
        })

# 提取关键段落内容
for para in doc.paragraphs:
    text = para.text.strip()
    if len(text) > 20 and not para.style.name.startswith('Heading'):
        result["key_paragraphs"].append(text[:150])

# 检查关键内容
checks = {
    "has_batch_import": False,
    "has_template_download": False,
    "has_batch_copy_role": False,
    "has_7_apis": False,
    "has_branch_module": False,
    "has_admin_module": False,
    "removed_tech_details": True,
    "removed_security": True,
    "removed_exception": True
}

heading_texts = [h["text"] for h in result["headings"]]
all_text = " ".join([p.text for p in doc.paragraphs])

checks["has_batch_import"] = "批量导入" in all_text
checks["has_template_download"] = "模板下载" in all_text or "模版下载" in all_text
checks["has_batch_copy_role"] = "批量复制角色" in all_text
checks["has_branch_module"] = "机构管理" in all_text
checks["has_admin_module"] = "机构管理员管理" in all_text
checks["removed_tech_details"] = "技术实现细节" not in " ".join(heading_texts)
checks["removed_security"] = "安全策略" not in " ".join(heading_texts)
checks["removed_exception"] = "异常处理机制" not in " ".join(heading_texts)

# 检查API表格
for table in doc.tables:
    if len(table.rows) >= 7:
        first_cell = table.rows[0].cells[0].text.strip()
        if "接口" in first_cell or "路径" in first_cell:
            checks["has_7_apis"] = len(table.rows) >= 8  # 表头+7个接口

result["checks"] = checks

# 保存结果
with open(r"d:\code\QJ\BEMP5.0DEV\.trae\skills\bemp-advanced-doc-generator\output\v2_validation.json", 'w', encoding='utf-8') as f:
    json.dump(result, f, ensure_ascii=False, indent=2)

print("=== v2文档质量验证 ===")
print(f"标题数: {len(result['headings'])}")
print(f"段落数: {result['paragraph_count']}")
print(f"表格数: {result['table_count']}")
print(f"\n=== 关键内容检查 ===")
for k, v in checks.items():
    status = "✓" if v else "✗"
    print(f"  {status} {k}: {v}")

print(f"\n=== 标题结构 ===")
for h in result['headings']:
    indent = "  " * (h['level'] - 1)
    print(f"{indent}H{h['level']}: {h['text']}")
