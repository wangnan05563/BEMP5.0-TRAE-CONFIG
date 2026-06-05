"""验证详细设计文档质量"""
import paths
from docx import Document
import sys

doc_path = sys.argv[1] if len(sys.argv) > 1 else str(paths.BANK_REQUIREMENTS_DIR / '承兑行额度管理-详细设计文档-v2.docx')
doc = Document(doc_path)

print(f'段落数: {len(doc.paragraphs)}')
print(f'表格数: {len(doc.tables)}')
print(f'节数: {len(doc.sections)}')

# 检查标题层级
headings = []
for p in doc.paragraphs:
    if p.style and p.style.name.startswith('Heading'):
        headings.append((p.style.name, p.text.strip()[:60]))

print(f'\n标题数: {len(headings)}')
for style, text in headings[:30]:
    print(f'  {style}: {text}')

# 检查蓝色文本残留
blue_count = 0
for p in doc.paragraphs:
    for run in p.runs:
        if run.font.color and run.font.color.rgb:
            color = str(run.font.color.rgb)
            if color.startswith('00') and color not in ['000000']:
                blue_count += 1
                if blue_count <= 5:
                    print(f'  蓝色文本: {run.text[:30]} (颜色: {color})')

print(f'\n蓝色文本残留数: {blue_count}')

# 检查封面
print('\n前10段落:')
for i, p in enumerate(doc.paragraphs[:10]):
    print(f'  [{i}] {p.style.name}: {p.text.strip()[:80]}')

# 检查表格内容
print(f'\n表格概况:')
for i, table in enumerate(doc.tables[:5]):
    rows = len(table.rows)
    cols = len(table.columns)
    first_cell = table.cell(0, 0).text[:30] if rows > 0 and cols > 0 else ''
    print(f'  表格{i+1}: {rows}行x{cols}列, 首格: {first_cell}')
