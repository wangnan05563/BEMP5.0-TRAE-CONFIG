"""对比111和112号段落的样式"""
import paths
from docx import Document

path = str(paths.detail_design_template())
doc = Document(path)

print(f"总段落数: {len(doc.paragraphs)}")
for i in [110, 111, 112, 113]:
    p = doc.paragraphs[i]
    print(f"[{i}] style.name='{p.style.name}' text='{p.text.strip()}'")
    if p.style:
        print(f"     style_id={p.style.style_id}")
