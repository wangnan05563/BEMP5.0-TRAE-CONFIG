"""测试修订记录表格更新"""
import sys, os
sys.path.insert(0, os.path.dirname(__file__))
import paths
from docx import Document

doc_path = paths.OUTPUT_DIR / '河南农商-概要设计说明书-20260604.docx'
doc = Document(str(doc_path))

print('--- 所有表格的表头 ---')
for i, table in enumerate(doc.tables):
    if not table.rows:
        continue
    header = [cell.text.strip() for cell in table.rows[0].cells]
    print(f'Table#{i}: {header}')
    if len(table.rows) > 1:
        row1 = [cell.text.strip() for cell in table.rows[1].cells]
        print(f'  Row1: {row1}')

# 测试关键词匹配
revision_keywords = ('版本', '版本号', '日期', '修改人', '修改内容')
for i, table in enumerate(doc.tables):
    header_cells = table.rows[0].cells if table.rows else []
    header_text = ' '.join(cell.text for cell in header_cells)
    match_count = sum(1 for kw in revision_keywords if kw in header_text)
    print(f'\nTable#{i} header_text: "{header_text[:100]}"')
    print(f'  match_count: {match_count}')
    if match_count >= 1:
        print(f'  → 匹配成功！')
        for ci, cell in enumerate(header_cells):
            col_header = cell.text.strip()
            print(f'  列{ci}: "{col_header}"')
