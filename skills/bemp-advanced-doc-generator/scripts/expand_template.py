"""直接修改模板.docx，添加4个业务模块的H1和H2标题"""
import paths
from docx import Document
from docx.oxml.ns import qn
from copy import deepcopy
import shutil

template_path = str(paths.detail_design_template())

doc = Document(template_path)

# 找到"模块1设计说明"H1，复用其样式作为业务模块H1的模板
module1_h1 = None
for p in doc.paragraphs:
    if p.style and p.style.name == 'Heading 1' and '模块1设计说明' in p.text:
        module1_h1 = p
        break

if module1_h1 is None:
    print("未找到模块1设计说明H1")
    exit(1)

print(f"找到模块1设计说明H1: {module1_h1.text.strip()}")

# 找到"模块2设计说明"H1和其后续H2
module2_h1 = None
appendix_h1 = None
for p in doc.paragraphs:
    if p.style and p.style.name == 'Heading 1':
        if '模块2设计说明' in p.text:
            module2_h1 = p
        elif p.text.strip() == '附录':
            appendix_h1 = p
            break

# 找到"模块1设计说明"下的所有H2，作为业务模块H2的样式模板
module1_h2_template = None
for p in doc.paragraphs:
    if p.style and p.style.name == 'Heading 2' and module1_h1 in [module1_h1]:
        module1_h2_template = p
        break

# 找到"附录"H1前一个段落（用于插入新H1的位置）
print(f"模块1设计说明: {module1_h1.text.strip()}")
print(f"模块2设计说明: {module2_h1.text.strip() if module2_h1 else 'N/A'}")
print(f"附录: {appendix_h1.text.strip() if appendix_h1 else 'N/A'}")

# 业务模块列表
business_modules = [
    ("额度申请模块设计说明", [
        "4.1 功能描述", "4.2 界面", "4.3 性能", "4.4 输入项", "4.5 输出项",
        "4.6 接口", "4.7 类图", "4.8 顺序图", "4.9 活动图", "4.10 备注"
    ]),
    ("批复明细模块设计说明", [
        "5.1 功能描述", "5.2 界面", "5.3 性能", "5.4 输入项", "5.5 输出项",
        "5.6 接口", "5.7 类图", "5.8 顺序图", "5.9 活动图", "5.10 备注"
    ]),
    ("额度复核模块设计说明", [
        "6.1 功能描述", "6.2 界面", "6.3 性能", "6.4 输入项", "6.5 输出项",
        "6.6 接口", "6.7 类图", "6.8 顺序图", "6.9 活动图", "6.10 备注"
    ]),
    ("额度占用/释放模块设计说明", [
        "7.1 功能描述", "7.2 界面", "7.3 性能", "7.4 输入项", "7.5 输出项",
        "7.6 接口", "7.7 类图", "7.8 顺序图", "7.9 活动图", "7.10 备注"
    ]),
]

# 找到"模块1设计说明"H1的XML元素
m1_elem = module1_h1._element
m1_parent = m1_elem.getparent()
m1_idx = list(m1_parent).index(m1_elem)

# 删除现有的"模块1设计说明"和"模块2设计说明"两个H1及其内容
# 包括：模块1的H1 + 模块1的所有H2 + 之间的内容
# 然后模块2的H1 + 模块2的所有H2 + 之间的内容
# 直到附录H1

# 先找到附录H1的XML元素索引
appendix_elem = appendix_h1._element
appendix_idx = list(appendix_parent if (appendix_parent := appendix_elem.getparent()) else m1_parent).index(appendix_elem)

# 计算要删除的XML元素范围
m1_parent = m1_elem.getparent()
# 从m1_elem开始到appendix_elem之前（不包含）
to_remove = []
elem = m1_elem
while elem is not None and elem is not appendix_elem:
    to_remove.append(elem)
    elem = elem.getnext()

print(f"\n要删除的元素数: {len(to_remove)}")
print("前5个元素:")
for e in to_remove[:5]:
    if e.tag.endswith('}p'):
        pPr = e.find('.//' + qn('w:pPr') + '/' + qn('w:pStyle'))
        text_elems = e.findall('.//' + qn('w:t'))
        text = ''.join(t.text or '' for t in text_elems)[:50]
        style = pPr.get(qn('w:val')) if pPr is not None else 'Normal'
        print(f"  p[{style}]: {text}")

# 删除
for e in to_remove:
    m1_parent.remove(e)

# 找到删除后appendix_elem在parent中的新索引
new_appendix_idx = list(m1_parent).index(appendix_elem)

# 构造4个业务模块的XML元素
# H1使用module1_h1的样式模板，H2使用module1的第一个H2的样式模板

# 找到module1的第一个H2（在删除前是"功能描述"）
# 实际上module1_h1的H2在to_remove里，但我们需要其样式模板
# 由于已删除，重新读取
# 重新加载文档获取H2样式模板
doc2 = Document(template_path)
h2_template = None
for p in doc2.paragraphs:
    if p.style and p.style.name == 'Heading 2' and p.text.strip() == '功能描述':
        h2_template = p
        break

if h2_template is None:
    print("未找到H2样式模板")
    exit(1)

# 构造4个业务模块的H1和H2元素
new_elements = []

for h1_title, h2_titles in business_modules:
    # 构造H1
    h1_elem = deepcopy(m1_elem)
    # 清除所有run
    for r in list(h1_elem.findall(qn('w:r'))):
        h1_elem.remove(r)
    # 添加新文本
    new_r = h1_elem.makeelement(qn('w:r'), {})
    new_t = new_r.makeelement(qn('w:t'), {})
    new_t.text = h1_title
    new_t.set(qn('xml:space'), 'preserve')
    new_r.append(new_t)
    h1_elem.append(new_r)
    new_elements.append(h1_elem)

    # 构造H2
    for h2_title in h2_titles:
        h2_elem = deepcopy(h2_template._element)
        # 清除所有run
        for r in list(h2_elem.findall(qn('w:r'))):
            h2_elem.remove(r)
        # 添加新文本
        new_r = h2_elem.makeelement(qn('w:r'), {})
        new_t = new_r.makeelement(qn('w:t'), {})
        new_t.text = h2_title
        new_t.set(qn('xml:space'), 'preserve')
        new_r.append(new_t)
        h2_elem.append(new_r)
        new_elements.append(h2_elem)

# 插入到附录H1之前
for elem in reversed(new_elements):
    m1_parent.insert(new_appendix_idx, elem)

print(f"\n新增元素数: {len(new_elements)}")

# 保存为新模板
new_template_path = str(paths.detail_design_template()).replace('.docx', '-v5扩展.docx')
doc.save(new_template_path)
print(f"已保存: {new_template_path}")

# 重新加载验证
doc3 = Document(new_template_path)
h1_count = sum(1 for p in doc3.paragraphs if p.style and p.style.name == 'Heading 1')
h2_count = sum(1 for p in doc3.paragraphs if p.style and p.style.name == 'Heading 2')
print(f"新模板 - H1: {h1_count}, H2: {h2_count}")
