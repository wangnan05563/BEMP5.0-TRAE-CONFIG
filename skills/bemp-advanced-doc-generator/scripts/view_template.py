"""查看模板的H1结构"""
import paths
from docx import Document
import sys

doc = Document(str(paths.ASSETS_DIR / 'template-outline-design.docx'))
print('模板段落数:', len(doc.paragraphs))
for i, p in enumerate(doc.paragraphs):
    style = p.style.name if p.style else ''
    if 'Heading' in style or i < 30:
        print(f'  [{i}] [{style}] {p.text[:80]}')
