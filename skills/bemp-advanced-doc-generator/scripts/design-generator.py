"""
详细设计文档生成器 - 基于 .docx 模板填充
复用 outline-design-generator.py 的模板处理基础设施（蓝色文本清理、封面替换等）
填充详细设计专用的章节内容（系统概述、功能模块、接口设计、数据库设计等）
"""
import sys
import json
import re
import os
import shutil
import zipfile
from copy import deepcopy
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime

# 2026-06-04 新增：调试日志开关（排查章节窜行问题时开启）
DEBUG_LAYOUT = os.environ.get('BEMP_DESIGN_DEBUG', '').lower() in ('1', 'true', 'yes')
import sys as _sys_dbg
print(f'[DBG-INIT] BEMP_DESIGN_DEBUG={os.environ.get("BEMP_DESIGN_DEBUG")!r} DEBUG_LAYOUT={DEBUG_LAYOUT}', file=_sys_dbg.stderr)

# 2026-06-07 新增：模板保留模式。当 design_data.chapters 为空时设为 True，
# _clear_content_between_headings / fill_empty_chapter 等函数自动跳过正文清理。
# 设计原则：模块级变量避免大面积缩进改动，所有内容破坏性函数统一在此检查。
_PRESERVE_MODE = False

# 2026-06-04 新增：doc_rules.yaml 规则加载（由其他子任务创建，存在时生效）
try:
    import yaml  # type: ignore
except Exception:  # pragma: no cover - 缺 PyYAML 时优雅降级
    yaml = None

_RULES_CACHE = None


def _load_doc_rules():
    """从 doc_rules.yaml 加载规则，文件不存在或解析失败时返回空 dict

    规则示例（非硬编码业务）：
    ```yaml
    overview_keywords: ['概述', '项目概述', '系统概述']
    non_functional_chapters: ['界面', '性能', '安全性', '可靠性']
    empty_check_extra: ['模块复用分析', '设计约束']
    toc_levels: '1-3'
    debug_layout: false
    ```
    """
    global _RULES_CACHE
    if _RULES_CACHE is not None:
        return _RULES_CACHE
    if yaml is None:
        _RULES_CACHE = {}
        return _RULES_CACHE
    # 规则文件与本脚本同目录
    rule_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'doc_rules.yaml')
    if not os.path.isfile(rule_path):
        _RULES_CACHE = {}
        return _RULES_CACHE
    try:
        with open(rule_path, 'r', encoding='utf-8') as f:
            _RULES_CACHE = yaml.safe_load(f) or {}
    except Exception as e:  # 解析失败时降级
        print(f'[WARN] 加载 doc_rules.yaml 失败，使用内置默认: {e}', file=sys.stderr)
        _RULES_CACHE = {}
    return _RULES_CACHE


from doc_utils import (
    BLACK, BLUE, TABLE_FONT_SIZE_PT, is_blue_color,
    is_blue_paragraph, is_blue_cell, is_hyperlink_run, has_non_hyperlink_blue,
    is_blue_placeholder_text,
    set_black, clear_paragraph, write_paragraph, write_as_normal,
    content_hash, write_with_dedup, write_as_normal_with_dedup,
    add_paragraph_text, fill_cell_text, add_table_row, clear_blue_cell,
    remove_paragraph, format_table_styled, get_heading_context
)

# 2026-06-04 新增：引入统一格式化工具（缩进、表格样式、空章节占位、TOC域）
# 2026-06-04 优化：导入 force_insert_toc 与 doc_formatter 兼容接口
from doc_formatter import (
    apply_paragraph_indent,
    apply_body_indent_to_doc,
    apply_table_style,
    fill_empty_chapter,
    fill_empty_chapter_compat,
    has_toc_field,
    insert_toc_field_after,
    force_insert_toc,
    inject_update_fields,
    clear_all_blue_runs,
    remove_design_constraint_irrelevant,
    format_ui_paragraph,
    # 2026-06-06 新增：图片后冗余文字、标题编号重构、表格内容映射、标题样式统一
    remove_redundant_text_after_images,
    renumber_headings,
    clean_all_function_module_tables,
    clean_function_module_table,
    unify_heading_styles,
    build_function_module_rows_from_subsystems,
    # 2026-06-07 新增：自动提升 Normal/加粗子标题为对应 Heading 样式
    auto_promote_bold_subheadings,
)
# 2026-06-07 新增：通用化图表自动生成模块（类图/顺序图/活动图/流程图/时序图）
try:
    from _diagram_utils import generate_and_insert_diagram, DIAGRAM_TYPES as _DIAGRAM_TYPE_MAP
except ImportError:
    generate_and_insert_diagram = None
    _DIAGRAM_TYPE_MAP = {}

# 2026-06-07 新增：附录F清理模块（移除不协调的功能模块表格）
try:
    from _appendix_cleanup import clean_appendix_f_table
except ImportError:
    clean_appendix_f_table = lambda doc: 0

TABLE_FONT_SIZE = Pt(TABLE_FONT_SIZE_PT)

# 2026-06-04 新增：需检查的空章节关键词（H2/H3标题）
EMPTY_CHECK_KEYWORDS = [
    ('设计目标', '设计目标描述'),
    ('输入项', '输入项说明'),
    ('输出项', '输出项说明'),
    ('代码示例', '代码示例'),
    ('性能优化', '性能优化措施'),
    ('附录', '附录说明'),
    ('错误码', '错误码定义'),
    ('接口设计', '接口设计'),
    ('开发规范', '开发规范'),
    # 2026-06-05 移除以下关键词：它们由特殊处理函数填充（可能插入表格），
    # fill_empty_chapter 只检查段落文本不检查表格，会误判为空并覆盖为"不涉及"
    # ('使用范围', '使用范围'),
    # ('术语定义', '术语定义'),
    # ('参考资料', '参考资料'),
    # ('开发环境', '开发环境'),
    # ('界面', '界面'),
    # ('性能', '性能'),
    # ('目的', '目的'),
    # ('适用范围', '适用范围'),
    # 2026-06-05 新增：模块设计说明下的H2子节关键词（解决空章节问题）
    ('功能描述', '功能描述说明'),
    ('类图', '类图描述'),
    ('顺序图', '顺序图描述'),
    ('活动图', '活动图描述'),
    ('备注', '备注说明'),
]

# 2026-06-04 新增：扩展空章节关键词（含系统概述、设计约束、模块复用分析）
# 由 doc_rules.yaml 中 empty_check_extra 覆盖（如有）
EMPTY_CHECK_EXTRA = [
    ('系统概述', '系统概述'),
    ('设计约束', '设计约束'),
    ('模块复用分析', '模块复用分析'),
    ('组件结构图', '组件结构图'),
    ('系统组件', '系统组件'),
]

# 2026-06-04 新增：非功能属性章节标准命名（用于 fill_empty_chapter 兜底）
NON_FUNCTIONAL_CHAPTERS = [
    '界面', '性能', '安全性', '可靠性', '易用性',
    '可调试性', '可移植性', '可维护性',
]

# 2026-06-04 新增：概述章节的关键词集合（用于章节内容匹配）
# 含 "概述"、"项目概述"、"系统概述" 等多种写法
CHAPTER_KEYWORDS_OVERVIEW = ('概述', '项目概述', '系统概述')

# 2026-06-04 新增：通用错误码（用于错误码章节为空时兜底）
DEFAULT_ERROR_CODES = [
    ('E0000', '成功'),
    ('E0001', '系统异常'),
    ('E0002', '参数错误'),
    ('E0003', '鉴权失败'),
    ('E0004', '数据不存在'),
    ('E0005', '数据已存在'),
]

# 2026-06-04 新增：组件关键字（用于检测"系统组件"标题下的子组件标题）
COMPONENT_SUBHEADING_KEYWORDS = ('组件',)

# 模板标题 → JSON章节标题的映射关系（支持H1+H2组合匹配）
# 键格式: "H1标题" 或 "H1标题|H2标题"
# 值: JSON中对应的章节标题列表（按优先级排序）
TEMPLATE_CHAPTER_MAP = {
    # 概述下的H2映射
    '概述|目的': ['1.1 业务背景', '业务背景', '编写目的'],
    '概述|编写目的': ['1.1 业务背景', '业务背景', '编写目的'],
    '概述|读者对象': ['读者对象', '目标读者'],
    '概述|适用范围': ['1.3 范围说明', '范围说明', '使用范围'],
    '概述|术语定义': [],  # 从appendix.glossary填充
    '概述|开发环境': [],  # 通用内容
    '概述|参考资料': [],  # 从appendix.references填充
    # H1级别映射
    '概述': ['概述'],
    '设计目标': ['设计目标'],
    '设计策略': ['设计策略'],
    '设计原则': ['设计原则'],
    '设计约束': ['设计约束'],
    '外部接口': ['外部接口', '接口设计'],
    '组件内部的模块列表及说明': ['组件汇总', '组件内部的模块列表及说明', '模块列表及说明'],
    '组件汇总': ['组件汇总', '组件内部的模块列表及说明', '模块列表及说明'],
    '技术实现': ['技术实现', '技术方案'],
    '关键技术': ['关键技术', '关键技术难点'],
    '非功能性': ['非功能性设计', '非功能性要求', '非功能性'],
    '模块复用': ['模块复用分析', '模块复用'],
    # 4个业务模块的H1+H2映射（按H2序号前缀精确匹配section）
    # 模板H2是"4.1 功能描述"/"5.1 功能描述"等带编号的，section也用对应编号
    # 这里返回空列表走特殊处理（在主循环中H1匹配后由_fill_chapter_h1_sections处理）
    # 业务模块H1标题由_find_matching_chapter中的模式匹配动态识别（含"模块设计说明"），
    # 无需在此硬编码具体模块名
    # 附录
    '附录': ['附录'],
}

# 模块级H2映射：模板H2标题（去除编号前缀）→ JSON section标题（用于模块章节下的子节匹配）
# 2026-06-05 新增：解决"模块1设计说明"下H2子节无法匹配JSON章节数据的问题
MODULE_H2_SECTION_MAP = {
    '功能描述': ['模块职责', '功能描述', '模块划分'],
    '输入项': ['接口边界', '接口定义', '参数说明'],
    '输出项': ['接口边界', '接口定义', '参数说明'],
    '接口': ['接口边界', '接口定义', '接口列表'],
    '类图': [],  # 由UML图表生成器填充
    '顺序图': [],  # 由UML图表生成器填充
    '活动图': [],  # 由UML图表生成器填充
    '备注': [],  # 一般不需要填充
}


# ==================== 2026-06-06 章节分类与 H1 填充函数 ====================

# 业务模块识别：业务模块 ch 通常含有"模块设计说明"/"功能模块"/"业务模块"等关键词
_BUSINESS_MODULE_TITLE_KEYWORDS = ('模块设计说明', '功能模块', '业务模块', '功能子模块')
# 业务模块的"标准 H2 子节"关键词集合（用于识别该 ch 是否为业务模块 ch）
_BUSINESS_MODULE_H2_KEYWORDS = (
    '功能描述', '界面', '性能', '输入项', '输出项', '接口',
    '类图', '顺序图', '活动图', '备注',
    '功能职责', '功能', '输入', '输出', '参数',
    '业务规则', '业务逻辑', '业务流程',
)
# 概述类 ch 标题关键词
_OVERVIEW_TITLE_KEYWORDS = ('概述', '项目概述', '系统概述', '引言', '简介')
# 设计约束类 ch 标题关键词
_CONSTRAINT_TITLE_KEYWORDS = ('设计约束', '约束', '限制条件', '前提条件')
# 组件列表类 ch 标题关键词
_COMPONENT_LIST_TITLE_KEYWORDS = ('组件列表', '模块列表', '组件汇总', '模块清单')
# 附录类 ch 标题关键词
_APPENDIX_TITLE_KEYWORDS = ('附录', '附表', '附件')


def _strip_chapter_number(title):
    """去除章节标题前缀编号（如"1.1 业务背景" → "业务背景"，"第X章 概述" → "概述"）"""
    if not title:
        return ''
    t = re.sub(r'^\d+(\.\d+)*\s*', '', title).strip()
    t = re.sub(r'^第[一二三四五六七八九十百零\d]+章\s*', '', t).strip()
    return t


def _classify_chapters(chapters):
    """根据 chapters 的 title/content 特征，将其分类为以下五类（避免硬编码模块名）

    返回: (overview_chs, constraint_ch, component_list_ch, business_chs, appendix_ch, other_chs)
      - overview_chs: list（一般仅一个"概述"ch）
      - constraint_ch: dict 或 None（"设计约束"ch）
      - component_list_ch: dict 或 None（"组件列表"ch）
      - business_chs: list（按 chapters 中顺序排列的业务模块 ch）
      - appendix_ch: dict 或 None（"附录"ch，可能含 sections）
      - other_chs: list（未归入以上五类的其他 ch，将动态追加到文档末尾）

    分类规则（按优先级）：
    1. 标题含"附录" → appendix
    2. 标题含"设计约束"/"约束"/"限制" → constraint
    3. 标题含"组件列表"/"模块列表" → component_list
    4. 标题含"概述"/"项目概述"/"系统概述" → overview
    5. 含 sections 且 sections 标题命中业务模块标准子节 → business
    6. 标题含"模块设计说明" → business（兜底）
    7. 标题含"功能模块"/"业务模块" → business（兜底）
    8. 含有 sections（≥1）且非上面前 4 类 → business（2026-06-06 扩展：
       详细设计文档中数据 chapters 多为带 sections 的独立章节，
       应优先归为 business 以便分配到模板"模块N设计说明"H1 锚点）
    9. 其余 → other

    返回: (overview_chs, constraint_ch, component_list_ch, business_chs, appendix_ch, other_chs)
      - overview_chs: list（一般仅一个"概述"ch）
      - constraint_ch: dict 或 None（"设计约束"ch）
      - component_list_ch: dict 或 None（"组件列表"ch）
      - business_chs: list（按 chapters 中顺序排列的业务模块 ch）
      - appendix_ch: dict 或 None（"附录"ch，可能含 sections）
      - other_chs: list（未归入以上五类的其他 ch，将动态追加到附录之前）
    """
    if not chapters:
        return [], None, None, [], None, []

    overview_chs = []
    constraint_ch = None
    component_list_ch = None
    business_chs = []
    appendix_ch = None
    other_chs = []

    def is_business_module(ch):
        """判断一个 ch 是否为业务模块 ch"""
        title_clean = _strip_chapter_number(ch.get('title', ''))
        # 规则 1: 标题命中业务模块关键词
        for kw in _BUSINESS_MODULE_TITLE_KEYWORDS:
            if kw in title_clean:
                return True
        # 规则 2: 含 sections 且 sections 标题命中业务模块 H2 标准子节
        sections = ch.get('sections', []) or []
        if sections:
            for sec in sections:
                sec_title_clean = _strip_chapter_number(sec.get('title', ''))
                for kw in _BUSINESS_MODULE_H2_KEYWORDS:
                    if sec_title_clean == kw or sec_title_clean.endswith(kw):
                        return True
        # 规则 3: 2026-06-06 扩展：含 sections（≥1）即视为业务模块
        # 详细设计文档中"接口定义"/"数据模型"等章节通常含 sections，
        # 应当被识别为业务模块以便分配到模板的"模块N设计说明"H1 锚点
        if sections:
            return True
        return False

    for ch in chapters:
        title = (ch.get('title') or '').strip()
        title_clean = _strip_chapter_number(title)
        # 1. 附录
        if any(kw in title_clean for kw in _APPENDIX_TITLE_KEYWORDS):
            appendix_ch = ch
        # 2. 设计约束
        elif any(kw in title_clean for kw in _CONSTRAINT_TITLE_KEYWORDS):
            if constraint_ch is None:
                constraint_ch = ch
            else:
                other_chs.append(ch)
        # 3. 组件列表
        elif any(kw in title_clean for kw in _COMPONENT_LIST_TITLE_KEYWORDS):
            if component_list_ch is None:
                component_list_ch = ch
            else:
                other_chs.append(ch)
        # 4. 概述
        elif any(kw in title_clean for kw in _OVERVIEW_TITLE_KEYWORDS):
            overview_chs.append(ch)
        # 5/6/7/8. 业务模块
        elif is_business_module(ch):
            business_chs.append(ch)
        else:
            other_chs.append(ch)

    return overview_chs, constraint_ch, component_list_ch, business_chs, appendix_ch, other_chs


def _insert_section_as_h2(after_para, sec):
    """在指定段落后插入一个 section（按"子节标题+内容"形式），返回最后的插入段落

    sec 格式: {id, title, content: {description, headers, rows}}

    2026-06-07 增强：空内容占位
    - 若 sec.content 为空 dict / None，且 sec 标题命中图表类关键词（类图/顺序图/时序图/
      活动图/状态图/界面/性能），插入"不涉及：本章节无图表内容"的灰色占位段落
    - 占位段落使用 Normal 样式 + 灰色字体（区别于正常内容）
    """
    sec_title = (sec.get('title') or '').strip()
    sec_content = sec.get('content', {}) or {}
    insert_after = after_para

    try:
        trace_f = open(r'D:\code\QJ\BEMP5.0DEV\.trae\skills\bemp-advanced-doc-generator\scripts\_insert_section.log', 'a', encoding='utf-8')
        trace_f.write(f'[INSERT-SEC] title={sec_title[:30]!r} anchor_text={(after_para.text or "")[:30]!r} has_table={bool(sec_content.get("headers") and sec_content.get("rows"))}\n')
        trace_f.close()
    except Exception:
        pass

    # 1. 插入 H2 标题
    if sec_title:
        # 2026-06-07 通用化：剥除 sec_title 开头的硬编码编号（"4.1 功能描述" → "功能描述"）
        sec_title = _strip_chapter_number(sec_title)
        new_h2 = _insert_paragraph_after(insert_after, sec_title, style='Heading 2')
        set_black(new_h2)
        insert_after = new_h2

    # 2. 插入描述
    desc = sec_content.get('description') if isinstance(sec_content, dict) else None
    if isinstance(desc, str) and desc:
        new_p = _insert_paragraph_after(insert_after, desc)
        set_black(new_p)
        insert_after = new_p

    # 3. 插入表格
    headers = sec_content.get('headers') if isinstance(sec_content, dict) else None
    rows = sec_content.get('rows') if isinstance(sec_content, dict) else None
    if headers and rows:
        inserted_tbl = _add_table_after_paragraph(insert_after, headers, rows)
        # 2026-06-06 修复关键 bug：表格插入后，下一节的 H2 如果用 addnext 插入到 desc 后面，
        # 会把刚插入的表格推到末尾。正确做法：在表格后再插入一个空段落作为新的锚点。
        if inserted_tbl is not None:
            # 在表格后插入一个空段落（用 OXML addnext 直接挂在表格后面）
            anchor_p = OxmlElement('w:p')
            inserted_tbl.addnext(anchor_p)
            from docx.text.paragraph import Paragraph
            table_para = Paragraph(anchor_p, insert_after._parent)
            set_black(table_para)
            insert_after = table_para

    # 4. 2026-06-07 新增：空内容占位
    # 若 sec 完全无 description/headers/rows（空 dict），且标题属于图表类关键词，
    # 插入"不涉及"灰色占位段落（避免文档中空标题悬空）
    if not desc and not (headers and rows):
        chart_keywords = ('类图', '顺序图', '时序图', '活动图', '状态图',
                          '流程图', '架构图', '网络图', '部署图', '界面', '性能')
        # 标题剥除编号后判断
        sec_title_clean = re.sub(r'^\s*\d+(\.\d+)*\s*', '', sec_title).strip()
        if any(kw in sec_title_clean for kw in chart_keywords):
            placeholder_p = _insert_paragraph_after(
                insert_after, '不涉及：本章节无图表内容或无需展示对应图示。')
            # 设置灰色字体（区别于正常内容）
            from docx.shared import RGBColor
            from docx.oxml.ns import qn as _qn
            for r in placeholder_p.runs:
                r.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
                r.italic = True
            insert_after = placeholder_p

    return insert_after


def _insert_h1_after(after_para, h1_text):
    """在指定段落后插入一个 H1 标题，返回该 H1 段落对象

    2026-06-06 新增：用于将未匹配的 chapters 追加为新 H1。
    2026-06-07 修订：通用化剥除 h1_text 开头硬编码的章节编号（"1 "、"2 " 等），
    避免与模板样式 numId 自动编号叠加。Word 会按 H1 出现顺序自动渲染 1/2/3。
    """
    h1_text = _strip_chapter_number(h1_text)
    new_h1 = _insert_paragraph_after(after_para, h1_text, style='Heading 1')
    set_black(new_h1)
    return new_h1


def _fill_overview_h1(doc, h1_para, overview_ch, design_data=None):
    """填充 H1 = "概述"：按模板 H2 顺序（目的/术语定义/开发环境/参考资料）匹配
    概述 ch 的 sections，剩余 sections 在尾部追加为 H2。

    工作流程：
    1. 模板 H2: 目的 / 术语定义 / 开发环境 / 参考资料（顺序固定）
    2. 概述 ch.sections 按上述顺序匹配：目的←业务背景/编写目的；术语定义←术语定义；...
    3. 剩余 sections（如"设计目标"/"适用范围"）在"参考资料"后追加为 H2
    4. 已分配 sections 不会被重复处理
    """
    if not overview_ch:
        return
    sections = overview_ch.get('sections', []) or []
    if not sections:
        # 仅有 content 字段，直接填充
        content = overview_ch.get('content', {})
        if isinstance(content, str) and content:
            new_p = _insert_paragraph_after(h1_para, content)
            set_black(new_p)
        elif isinstance(content, dict):
            desc = content.get('description')
            if desc:
                new_p = _insert_paragraph_after(h1_para, desc)
                set_black(new_p)
                insert_after = new_p
            else:
                insert_after = h1_para
            headers = content.get('headers')
            rows = content.get('rows')
            if headers and rows:
                _add_table_after_paragraph(insert_after, headers, rows)
        return

    # 模板 H2 顺序与"标准 H2 → section 标题别名"映射
    template_h2_order = [
        ('目的', ('1.1 业务背景', '业务背景', '编写目的', '目的')),
        ('术语定义', ('1.4 术语定义', '术语定义', '术语和缩写', '术语')),
        ('开发环境', ('1.5 开发环境', '开发环境')),
        ('参考资料', ('1.6 参考资料', '参考资料', '参考文档')),
    ]
    # 反向索引：从 h2_key 找到 aliases
    h2_key_aliases = dict(template_h2_order)

    used_section_ids = set()
    module_name = design_data.get('moduleName', '本项目') if isinstance(design_data, dict) else '本项目'

    # 2026-06-06 修复：原实现每个 section 都用 h1_para 作为锚点，导致插入顺序与
    # 模板 H2 顺序相反（H1, 1.6, 1.5, 1.4, 1.1）。改为维护 last_para 累积锚点，
    # 让所有 H2 按时间顺序依次追加在 H1 之后。
    last_para = h1_para

    # 2026-06-06 修复：再按章节编号（1.1, 1.2 ...）升序排序所有 sections，
    # 确保 "1.2 设计目标" 出现在 "1.1 业务背景" 之后，避免模板 H2 顺序与编号错位
    def _section_chapter_num(sec):
        import re as _re
        m = _re.match(r'^\s*(\d+)\.(\d+)', sec.get('title', '') or '')
        if m:
            return (int(m.group(1)), int(m.group(2)))
        return (99, 99)

    sections_sorted = sorted(sections, key=_section_chapter_num)

    # 2026-06-06 重构：改为按 sections 编号顺序遍历，对每个 section 判断其属于哪个
    # "标准 H2"（目的/术语定义/开发环境/参考资料），对应调用原特殊处理函数；
    # 其余 section 走通用 _insert_section_as_h2。所有插入均以 last_para 为锚点，
    # 保持自然顺序 1.1 → 1.2 → 1.3 → 1.4 → 1.5 → 1.6
    for sec in sections_sorted:
        sec_id = sec.get('id') or sec.get('title', '')
        if sec_id in used_section_ids:
            continue
        sec_title_clean = _strip_chapter_number(sec.get('title', ''))
        matched_h2_key = None
        for h2_key, aliases in template_h2_order:
            for alias in aliases:
                alias_clean = _strip_chapter_number(alias)
                if sec_title_clean == alias_clean or sec_title_clean.endswith(alias_clean):
                    matched_h2_key = h2_key
                    break
            if matched_h2_key:
                break
        if matched_h2_key in ('目的', '术语定义', '开发环境', '参考资料'):
            used_section_ids.add(sec_id)
            last_para = _insert_section_as_h2(last_para, sec) or last_para
            continue
        # 普通 section
        content = sec.get('content', {}) or {}
        has_desc = isinstance(content, dict) and content.get('description')
        has_table = isinstance(content, dict) and content.get('headers') and content.get('rows')
        if has_desc or has_table:
            last_para = _insert_section_as_h2(last_para, sec) or last_para


def _find_last_para_under_h1(doc, h1_para):
    """在 H1 标题下找到最后一个段落（用于在尾部追加 H2）"""
    h1_elem = h1_para._element
    parent = h1_elem.getparent()
    if parent is None:
        return h1_para

    found_h1 = False
    last_para = h1_para
    for elem in list(parent):
        if elem is h1_elem:
            found_h1 = True
            continue
        if not found_h1:
            continue
        if elem.tag.endswith('}p'):
            pPr = elem.find('.//' + qn('w:pPr') + '/' + qn('w:pStyle'))
            if pPr is not None:
                style_val = pPr.get(qn('w:val'), '')
                if style_val in ['1', 'Heading1', 'heading1', 'heading 1', 'Heading 1']:
                    break
            from docx.text.paragraph import Paragraph
            last_para = Paragraph(elem, h1_para._parent)
        elif elem.tag.endswith('}tbl'):
            # 表格后通常紧跟一个空段落
            continue
    return last_para


def _fill_appendix_h1(doc, h1_para, appendix_ch):
    """填充 H1 = "附录"：按 ABCDEF 分类组织 sections，统一排版风格

    2026-06-06 重构：解决原模板"附录"章节内容混乱问题
    1) sections 按 A-F 分类（术语/规则/数据字典/接口/错误码/参考）
    2) 每类附加统一的"分类标识 H2"（如"附录A 术语表"）
    3) 同类 sections 合并到同一 H2 下，按出现顺序渲染
    4) 排版统一：H2 黑体 14pt + 黑色，描述段落 12pt + 黑色（与正文一致）
    5) 若 appendix_ch 为空，插入标准 A-F 目录占位
    """
    if not appendix_ch:
        return
    sections = appendix_ch.get('sections', []) or []

    # 1. 无 sections 时填充 content（直接作为 H1 下的描述）
    if not sections:
        content = appendix_ch.get('content', {})
        if isinstance(content, dict):
            desc = content.get('description')
            if desc:
                new_p = _insert_paragraph_after(h1_para, desc)
                set_black(new_p)
        elif isinstance(content, str) and content:
            new_p = _insert_paragraph_after(h1_para, content)
            set_black(new_p)
        return

    # 2. 按 A-F 分类
    grouped = {'A': [], 'B': [], 'C': [], 'D': [], 'E': [], 'F': []}
    for sec in sections:
        sec_title_raw = (sec.get('title') or '').strip()
        sec_title_clean = re.sub(r'^\d+(\.\d+)*\s*', '', sec_title_raw).strip()
        cat = _classify_appendix_section(sec_title_clean)
        grouped[cat].append(sec)

    # 3. 按 A→F 顺序输出（保持原顺序：同类内按 chapters 出现顺序）
    insert_after = h1_para
    output_categories = [c for c in ['A', 'B', 'C', 'D', 'E', 'F'] if grouped.get(c)]
    for cat in output_categories:
        label = _APPENDIX_CATEGORY_LABELS.get(cat, cat)
        # 插入分类 H2 标题（如"附录A 术语表"）
        cat_h2 = _insert_paragraph_after(insert_after, label, style='Heading 2')
        set_black(cat_h2)
        insert_after = cat_h2
        # 该分类下每个 section 作为 H3（嵌套更清晰）
        for sec in grouped[cat]:
            sec_title_raw = (sec.get('title') or '').strip()
            # 去除硬编码编号（如"A.1"/"A1"），避免与分类标签重复
            sec_title_clean = re.sub(r'^[A-F][\.\s]\d*[\s\.]*', '', sec_title_raw).strip()
            sec_title_clean = re.sub(r'^\d+(\.\d+)*\s*', '', sec_title_clean).strip()
            if not sec_title_clean:
                sec_title_clean = '条目'
            h3 = _insert_paragraph_after(insert_after, sec_title_clean, style='Heading 3')
            set_black(h3)
            insert_after = h3
            # section 内容（description + table）
            sec_content = sec.get('content', {}) or {}
            if isinstance(sec_content, dict):
                desc = sec_content.get('description')
                if isinstance(desc, str) and desc:
                    desc_p = _insert_paragraph_after(insert_after, desc)
                    set_black(desc_p)
                    insert_after = desc_p
                headers = sec_content.get('headers')
                rows = sec_content.get('rows')
                if headers and rows:
                    tbl = _add_table_after_paragraph(insert_after, headers, rows)
                    if tbl is not None:
                        # 在表格后插入空段落作为下一节的锚点
                        anchor_p = OxmlElement('w:p')
                        tbl.addnext(anchor_p)
                        from docx.text.paragraph import Paragraph
                        anchor_para = Paragraph(anchor_p, insert_after._parent)
                        set_black(anchor_para)
                        insert_after = anchor_para
            elif isinstance(sec_content, str) and sec_content:
                content_p = _insert_paragraph_after(insert_after, sec_content)
                set_black(content_p)
                insert_after = content_p

    # 4. 若没有任何分类命中且 sections 为空，插入"参考资料"占位
    if not output_categories:
        placeholder = _insert_paragraph_after(insert_after, 'F. 参考资料', style='Heading 2')
        set_black(placeholder)
        empty_p = _insert_paragraph_after(placeholder, '暂无')
        set_black(empty_p)


def _classify_appendix_section(sec_title):
    """根据 appendix section 标题将其分类为 A-F 六类

    2026-06-06 新增：附录内容 ABCDEF 分类（解决原模板内容混乱问题）
    A - 术语表（Terminology / Glossary）
    B - 业务规则清单（Business Rules / 业务约束）
    C - 数据字典（Data Dictionary / 字段说明）
    D - 接口清单（API / 接口列表 / WebService）
    E - 错误码对照表（Error Code / 返回码）
    F - 参考资料（Reference / 参考文献）

    Args:
        sec_title: section 标题（已去除编号前缀）

    Returns:
        分类标识字符串（'A'/'B'/'C'/'D'/'E'/'F'），未匹配时返回 'F'（默认归到参考资料）
    """
    if not sec_title:
        return 'F'
    t = sec_title.lower()
    rules = [
        ('A', ['术语', 'glossary', 'terminology', '缩略语', '缩写', '名词解释']),
        ('B', ['业务规则', '业务约束', 'business rule', '规则清单', '约束清单']),
        ('C', ['数据字典', '字段说明', 'data dictionary', '数据项', '表结构', '字段定义', '表字段']),
        ('D', ['接口清单', '接口列表', 'api', 'webservice', 'wsdl', '接口定义', '外部接口', '服务清单']),
        ('E', ['错误码', '错误代码', 'error code', '返回码', '应答码', '异常码']),
    ]
    for tag, kws in rules:
        for kw in kws:
            if kw in t or kw in sec_title:
                return tag
    return 'F'


# 附录分类视觉标识（标签文本 + 简短描述，用于排版统一）
_APPENDIX_CATEGORY_LABELS = {
    'A': 'A. 术语表',
    'B': 'B. 业务规则清单',
    'C': 'C. 数据字典',
    'D': 'D. 接口清单',
    'E': 'E. 错误码对照表',
    'F': 'F. 参考资料',
}


def _insert_appendix_extension_chapters(doc, appendix_h1, extension_chapters):
    """在附录 H1 之前插入扩展章节（术语表/数据字典/参考等通用附录）

    2026-06-06 新增：用于在"附录"章节前添加必要的扩展章节（如术语表、
    文档信息、版本历史），保证新章节的标题格式/排版风格与文档其他
    章节完全一致（统一通过 _append_chapter_as_h1 实现）。

    Args:
        doc: docx Document 对象
        appendix_h1: 模板原"附录"H1 段落对象（扩展章节插入到其前）
        extension_chapters: list[dict] 扩展章节列表，每项含 title/sections/content
    """
    if not extension_chapters or appendix_h1 is None:
        return 0
    inserted = 0
    for ch in extension_chapters:
        title = (ch.get('title') or '').strip()
        if not title:
            continue
        # 复用 _append_chapter_as_h1（position='before'）实现统一 H1 + H2 排版
        _append_chapter_as_h1(doc, appendix_h1, ch, position='before')
        inserted += 1
    return inserted


def _build_default_appendix_extension_chapters(design_data=None):
    """根据 design_data 生成默认的扩展附录章节（术语表/版本历史/文档信息）

    2026-06-06 新增：通用化设计 — 仅当 design_data 中有对应数据时，
    才生成对应的扩展章节。优先取术语表、版本历史、文档说明等
    结构化字段，避免无意义的占位章节。

    Returns:
        list[dict] 扩展章节列表，每项格式与 chapters 兼容
    """
    if not isinstance(design_data, dict):
        return []
    extensions = []

    # 2026-06-07 增强：通用化 — 术语表/参考文献/错误码对照表/数据字典
    # 不仅读取顶层字段，也读取 design_data['appendix'] 内部字段（适配两种数据布局）
    appendix_payload = design_data.get('appendix') or {}
    if not isinstance(appendix_payload, dict):
        appendix_payload = {}

    def _get_glossary():
        # 优先顶层 glossary，回退到 appendix.glossary
        g = design_data.get('glossary') or appendix_payload.get('glossary') or []
        return g if isinstance(g, list) else []

    def _get_references():
        r = design_data.get('references') or appendix_payload.get('references') or []
        return r if isinstance(r, list) else []

    def _get_version_history():
        v = design_data.get('versionHistory') or appendix_payload.get('versionHistory') or []
        return v if isinstance(v, list) else []

    def _get_error_codes():
        e = design_data.get('errorCodes') or appendix_payload.get('errorCodes') or []
        return e if isinstance(e, list) else []

    def _get_data_dict():
        d = design_data.get('dataDictionary') or appendix_payload.get('dataDictionary') or []
        return d if isinstance(d, list) else []

    # 1. 术语表（若 design_data.glossary 或 appendix.glossary 存在）
    glossary = _get_glossary()
    if glossary:
        rows = []
        for item in glossary:
            if isinstance(item, dict):
                rows.append([
                    item.get('term', ''),
                    item.get('abbr', ''),
                    item.get('fullName', '') or item.get('description', ''),
                ])
            elif isinstance(item, str):
                rows.append([item, '', ''])
        if rows:
            extensions.append({
                'title': '术语与缩略语',
                'sections': [{
                    'id': 'glossary',
                    'title': '术语表',
                    'content': {
                        'description': '本节列出文档中使用的核心术语、缩略语及其含义。',
                        'headers': ['术语/缩略语', '英文', '解释'],
                        'rows': rows,
                    },
                }],
            })

    # 2. 数据字典（若 design_data.dataDictionary 或 appendix.dataDictionary 存在）
    dd = _get_data_dict()
    if dd:
        rows = []
        for item in dd:
            if isinstance(item, dict):
                rows.append([
                    item.get('table', ''),
                    item.get('field', ''),
                    item.get('type', ''),
                    item.get('description', ''),
                ])
        if rows:
            extensions.append({
                'title': '数据字典',
                'sections': [{
                    'id': 'data-dict',
                    'title': '数据字典',
                    'content': {
                        'description': '本节列出本模块涉及的核心数据表与字段说明。',
                        'headers': ['表名', '字段', '类型', '说明'],
                        'rows': rows,
                    },
                }],
            })

    # 3. 错误码（若 design_data.errorCodes 或 appendix.errorCodes 存在）
    ec = _get_error_codes()
    if ec:
        rows = []
        for item in ec:
            if isinstance(item, dict):
                rows.append([
                    item.get('code', ''),
                    item.get('message', ''),
                    item.get('cause', ''),
                    item.get('action', ''),
                ])
        if rows:
            extensions.append({
                'title': '错误码对照表',
                'sections': [{
                    'id': 'error-codes',
                    'title': '错误码',
                    'content': {
                        'description': '本节列出本模块涉及的错误码及处理建议。',
                        'headers': ['错误码', '错误信息', '可能原因', '处理建议'],
                        'rows': rows,
                    },
                }],
            })
    return extensions


def _render_template(text, design_data=None, module_name=None):
    """通用化：替换文本中的 ${fieldName} 占位符

    2026-06-07 新增：coverPage.title 等字段可能含 ${moduleName}/${currentDate} 占位符，
    需在写入 Word 前替换为 design_data 实际值。设计原则：
    - 通用化：支持任意 ${fieldName} 占位符，不硬编码具体字段
    - 安全：缺失字段时保留原占位符（不抛错）
    - 内置：自动识别 currentDate（YYYY年MM月）、moduleName、company、product
    - 兜底：design_data 为 None 时仅替换内置变量

    示例：
        _render_template('${moduleName}详细设计文档', {'moduleName': '客户号合并'})
        → '客户号合并详细设计文档'
    """
    if not text or '${' not in text:
        return text

    import datetime as _dt

    # 内置变量（兜底集）
    builtin = {
        'currentDate': _dt.datetime.now().strftime('%Y年%m月'),
        'currentYear': str(_dt.datetime.now().year),
        'currentMonth': f"{_dt.datetime.now().month:02d}",
    }
    if module_name:
        builtin['moduleName'] = module_name

    # 合并 design_data 顶层字段（优先级 > 内置）
    if isinstance(design_data, dict):
        # 直接合并顶层
        for k, v in design_data.items():
            if isinstance(v, (str, int, float)):
                builtin.setdefault(k, str(v))
        # coverPage 子字段（company/product/version/department）
        cp = design_data.get('coverPage', {})
        if isinstance(cp, dict):
            for k, v in cp.items():
                if isinstance(v, (str, int, float)):
                    builtin.setdefault(f'coverPage.{k}', str(v))
        # docProps.core
        dp = design_data.get('docProps', {})
        if isinstance(dp, dict):
            core = dp.get('core', {})
            if isinstance(core, dict):
                for k, v in core.items():
                    if isinstance(v, (str, int, float)):
                        builtin.setdefault(f'docProps.core.{k}', str(v))

    def _replace(m):
        key = m.group(1).strip()
        return str(builtin.get(key, m.group(0)))  # 缺失时保留原占位符

    return re.sub(r'\$\{([^{}]+)\}', _replace, str(text))


def _strip_chapter_number(text):
    """通用化：剥除文本开头的"编号+空白"（如"4.1 "、"12.4.1 "、"1 "）

    2026-06-07 新增：模板 H1/H2 样式已绑定 numId 自动编号（见 docs/07 styles.xml
    中 styleId=1/2 的 <w:numPr><w:numId w:val="1"/></w:numPr>），Word 会自动渲染章节
    编号。如果在 H1/H2 文本里再写一次"4.1 功能描述"，会与自动编号叠加产生
    "4.1 4.1 功能描述"的视觉重复。该函数用于在所有 H1/H2 文本写入前剥除
    任何"数字开头+空白"前缀。
    """
    if not text:
        return text
    # 兼容：1、4.1、12.4.1、4.1\t、4.1\u3000（半角/全角空白）
    return re.sub(r'^\s*\d+(?:\.\d+)*[\s\u3000]+', '', str(text).strip())


def _strip_all_heading_hardcoded_numbers(doc):
    """通用化兜底：扫描所有 H1~H8 段落，剥除文本开头的硬编码章节编号。

    2026-06-07 新增：模板 H1~H8 样式已绑定 numId 自动编号，Word 会按出现顺序
    自动渲染 1/2/3 与 1.1/1.2/2.1 与 1.1.1/1.1.2/... 等多级编号。如果在标题
    文本里出现"4.1 功能描述"等写死编号，会与自动编号叠加产生"4.1 4.1 功能描述"。

    该函数在所有标题写入完成后做一次全文清理，确保：
    - 不依赖调用方在每个插入点都正确剥除
    - 模板原生标题段落（如"1 概述"、"1.1 目的"）中的编号也会被剥除
    - JSON 数据中误带的编号（如"4.2 模块划分"、"8.1 核心算法"、"1.1 业务背景"）也会被剥除

    返回：被清理的段落数量
    """
    H_STYLES = (
        _H1_STYLE_ALIASES | _H2_STYLE_ALIASES |
        _H3_STYLE_ALIASES | _H4_STYLE_ALIASES |
        _H5_STYLE_ALIASES | _H6_STYLE_ALIASES |
        _H7_STYLE_ALIASES | _H8_STYLE_ALIASES
    ) if '_H8_STYLE_ALIASES' in globals() else {
        '1', '2', '3', '4', '5', '6', '7', '8',
        'Heading1', 'Heading2', 'Heading3', 'Heading4',
        'Heading5', 'Heading6', 'Heading7', 'Heading8',
        'heading1', 'heading2', 'heading3', 'heading4',
        'heading5', 'heading6', 'heading7', 'heading8',
        'Heading 1', 'Heading 2', 'Heading 3', 'Heading 4',
        'Heading 5', 'Heading 6', 'Heading 7', 'Heading 8',
    }
    stripped = 0
    for para in doc.paragraphs:
        pStyle = para._element.find('.//' + qn('w:pPr') + '/' + qn('w:pStyle'))
        if pStyle is None:
            continue
        sv = pStyle.get(qn('w:val'), '')
        if sv not in H_STYLES:
            continue
        old_text = para.text or ''
        new_text = _strip_chapter_number(old_text)
        if new_text and new_text != old_text:
            _set_paragraph_text(para._element, new_text)
            stripped += 1
    if DEBUG_LAYOUT and stripped:
        print(f'[DEBUG] _strip_all_heading_hardcoded_numbers: stripped={stripped}', file=sys.stderr)
    return stripped


# 兼容旧调用名（_strip_all_h1h2_hardcoded_numbers → 新名 _strip_all_heading_hardcoded_numbers）
def _strip_all_h1h2_hardcoded_numbers(doc):
    """兼容旧调用：转发到 _strip_all_heading_hardcoded_numbers（已扩展到 H1~H8）"""
    return _strip_all_heading_hardcoded_numbers(doc)


# H1~H8 样式别名集合（兼容 w:val 取值的多种写法）
_H1_STYLE_ALIASES = {'1', 'Heading1', 'heading1', 'heading 1', 'Heading 1'}
_H2_STYLE_ALIASES = {'2', 'Heading2', 'heading2', 'heading 2', 'Heading 2'}
_H3_STYLE_ALIASES = {'3', 'Heading3', 'heading3', 'heading 3', 'Heading 3'}
_H4_STYLE_ALIASES = {'4', 'Heading4', 'heading4', 'heading 4', 'Heading 4'}
_H5_STYLE_ALIASES = {'5', 'Heading5', 'heading5', 'heading 5', 'Heading 5'}
_H6_STYLE_ALIASES = {'6', 'Heading6', 'heading6', 'heading 6', 'Heading 6'}
_H7_STYLE_ALIASES = {'7', 'Heading7', 'heading7', 'heading 7', 'Heading 7'}
_H8_STYLE_ALIASES = {'8', 'Heading8', 'heading8', 'heading 8', 'Heading 8'}


def _renumber_h2_under_h1(h1_para, chapter_no):
    """重新编号 H1 下的所有 H2 子标题（4.1/4.2/... → chapter_no.1/chapter_no.2/...）

    2026-06-06 新增：模板中"模块1设计说明"等 H1 下的 H2 编号是硬编码的
    （如 4.1~4.10），不会随实际模块位置变化。本函数根据 chapter_no 动态
    修正为 chapter_no.1~chapter_no.N。

    2026-06-07 修订：模板 H1/H2 样式已绑定 numId 自动编号（见 docs/07
    styles.xml 中 styleId=1/2 的 <w:numPr><w:numId w:val="1"/></w:numPr>），
    Word 会按出现顺序自动渲染章节编号。本函数改为 **只剥除** 旧编号前缀
    （避免与自动编号叠加产生"4.1 4.1 功能描述"），不再写入新编号文本。
    chapter_no 参数保留以兼容调用方，但不再影响最终文本。
    """
    if h1_para is None:
        return 0
    h1_elem = h1_para._element
    parent = h1_elem.getparent()
    if parent is None:
        return 0
    h1_idx = list(parent).index(h1_elem)
    # 找下一个 H1
    next_h1_elem = None
    for j in range(h1_idx + 1, len(list(parent))):
        elem = list(parent)[j]
        if elem.tag.endswith('}p'):
            pPr = elem.find('.//' + qn('w:pPr') + '/' + qn('w:pStyle'))
            if pPr is not None:
                style_val = pPr.get(qn('w:val'), '')
                if style_val in ['1', 'Heading1', 'heading1', 'heading 1', 'Heading 1']:
                    next_h1_elem = elem
                    break
    # 收集所有 H2 段落元素
    h2_paras = []
    elem = h1_elem.getnext()
    while elem is not None and elem is not next_h1_elem:
        if elem.tag.endswith('}p'):
            pPr = elem.find('.//' + qn('w:pPr') + '/' + qn('w:pStyle'))
            if pPr is not None:
                style_val = pPr.get(qn('w:val'), '')
                if style_val in ['2', 'Heading2', 'heading2', 'heading 2', 'Heading 2']:
                    h2_paras.append(elem)
        elem = elem.getnext()
    # 通用化：只剥除旧编号前缀，不再写入新编号（让模板 numId 自动渲染）
    renamed = 0
    for h2_elem in h2_paras:
        t_elems = h2_elem.findall('.//' + qn('w:t'))
        old_text = ''.join(t.text or '' for t in t_elems).strip()
        new_text = _strip_chapter_number(old_text)
        if new_text != old_text:
            _set_paragraph_text(h2_elem, new_text)
            renamed += 1
    if DEBUG_LAYOUT:
        print(f'[DEBUG] _renumber_h2_under_h1: chapter_no={chapter_no}, stripped={renamed}', file=sys.stderr)
    return renamed


# 2026-06-07 新增：模块级"当前正在处理的 doc"引用。
# 用于在不修改 56 个 _insert_paragraph_xxx 调用点签名的前提下，
# 让 _resolve_style_for_insert 能拿到 doc 来解析模板 styleId。
# 调用约定：任何处理 doc 的入口函数（如 generate_design_document）应将
# 自己的 doc 赋给 _CURRENT_DOC，并在退出时清理。
_CURRENT_DOC = None


def _set_current_doc(doc):
    global _CURRENT_DOC
    _CURRENT_DOC = doc


def _resolve_style_for_insert(style, doc=None):
    """解析插入段落时使用的 Heading 样式 styleId（通用化）

    2026-06-07 新增。模板中 Heading 样式 styleId 可能是 'Heading 1' / 'Heading1' / '1'
    等多种形式。统一通过 _resolve_heading_style_id 解析为模板实际可识别的 styleId，
    避免设置 'Heading 1' 后 Word 回退到 Normal。

    优先级：调用方显式传入的 doc > 模块级 _CURRENT_DOC > 原样返回

    Args:
        style: 调用方传入的样式字符串，例如 'Heading 1' / 'Heading2'
        doc: Document 对象（python-docx）；若为 None，使用 _CURRENT_DOC

    Returns:
        str: 解析后的 styleId（若无法解析则原样返回）
    """
    if not style:
        return style
    s = (style or '').strip()
    if not s:
        return s
    # 解析 heading 级别
    import re as _re
    m = _re.match(r'^[Hh]eading\s*(\d)$', s)
    if not m:
        return style
    level = int(m.group(1))
    if level < 1 or level > 6:
        return style
    target_doc = doc if doc is not None else _CURRENT_DOC
    if target_doc is None:
        return style
    try:
        from doc_formatter import _resolve_heading_style_id
        return _resolve_heading_style_id(target_doc, level)
    except Exception:
        return style


# H1/H2 样式别名集合（兼容 w:val 取值的多种写法）
_H1_STYLE_ALIASES = {'1', 'Heading1', 'heading1', 'heading 1', 'Heading 1'}
_H2_STYLE_ALIASES = {'2', 'Heading2', 'heading2', 'heading 2', 'Heading 2'}


def _get_pstyle_val(para_or_elem):
    """读取段落 OXML 元素的 pStyle w:val 字符串（如 'Heading 1' / '2'）

    2026-06-06 新增：python-docx 的 p.style.name 对手动构建的 OxmlElement 段落
    无法解析样式表，p.style.name 总是返回 Normal。直接读 OXML 才是可靠方案。
    """
    elem = para_or_elem._element if hasattr(para_or_elem, '_element') else para_or_elem
    pPr = elem.find('.//' + qn('w:pPr'))
    if pPr is None:
        return ''
    pStyle = pPr.find(qn('w:pStyle'))
    if pStyle is None:
        return ''
    return pStyle.get(qn('w:val'), '') or ''


def _get_h1_chapter_no(doc, h1_para):
    """计算 h1_para 在文档所有 H1 中的 1-indexed 位置。

    2026-06-06 新增：替代基于 doc.paragraphs + p.style.name 的脆弱查找。
    """
    if h1_para is None:
        return None
    target_elem = h1_para._element
    target_parent = target_elem.getparent()
    if target_parent is None:
        return None
    # 找 target_elem 在 parent 中的位置
    try:
        target_idx = list(target_parent).index(target_elem)
    except ValueError:
        return None
    # 扫描 parent 中所有 H1 段落
    chapter_no = 0
    for j, child in enumerate(list(target_parent)):
        if j > target_idx:
            break
        if child.tag.endswith('}p'):
            sv = _get_pstyle_val(Paragraph(child, doc)) if False else _get_pstyle_val(child)
            if sv in _H1_STYLE_ALIASES:
                chapter_no += 1
    return chapter_no if chapter_no > 0 else None


def _fill_business_module_h1(doc, h1_para, biz_ch, design_data=None, chapters=None, chapter_no=None):
    """填充 H1 = "模块1设计说明" / "模块2设计说明" / 动态追加的业务模块 H1

    工作流程：
    1. 模板 H2 标准子节（功能描述/界面/性能/...）按 sections 匹配填充
    2. 剩余 sections 追加为 H2
    3. 若 biz_ch 自身没有 sections（结构异常），用其 content 兜底
    4. 若指定 chapter_no，则重编号 H2 子节（修复 4.1~4.10 模板硬编码问题）

    2026-06-06 修复：传入 chapter_no 后，会调用 _renumber_h2_under_h1
    将 H1 下的 H2 编号从模板硬编码值改为按 chapter_no 顺序编号。
    """
    if not biz_ch:
        return
    sections = biz_ch.get('sections', []) or []
    if not sections:
        content = biz_ch.get('content', {})
        if isinstance(content, dict):
            desc = content.get('description')
            insert_after = h1_para
            if desc:
                new_p = _insert_paragraph_after(insert_after, desc)
                set_black(new_p)
                insert_after = new_p
            headers = content.get('headers')
            rows = content.get('rows')
            if headers and rows:
                _add_table_after_paragraph(insert_after, headers, rows)
        return

    # 模板 H2 标准子节 → section 标题别名
    h2_alias_map = {
        '功能描述': ('功能描述', '功能职责', '模块职责', '功能', '模块划分'),
        '界面': ('界面', '界面设计', '界面布局'),
        '性能': ('性能', '性能指标', '性能要求'),
        '输入项': ('输入项', '输入', '参数说明', '接口边界'),
        '输出项': ('输出项', '输出', '参数说明', '接口边界'),
        '接口': ('接口', '接口设计', '接口定义', '接口列表', '外部接口'),
        '类图': ('类图',),
        '顺序图': ('顺序图', '时序图'),
        '活动图': ('活动图',),
        '备注': ('备注', '说明', '注释'),
    }

    used_sec_ids = set()

    # 找到模板 H1 下所有 H2 子标题
    h1_elem = h1_para._element
    parent = h1_elem.getparent()
    if parent is None:
        return

    h1_idx = list(parent).index(h1_elem)
    next_h1_elem = None
    for j in range(h1_idx + 1, len(list(parent))):
        elem = list(parent)[j]
        if elem.tag.endswith('}p'):
            pPr = elem.find('.//' + qn('w:pPr') + '/' + qn('w:pStyle'))
            if pPr is not None:
                style_val = pPr.get(qn('w:val'), '')
                if style_val in ['1', 'Heading1', 'heading1', 'heading 1', 'Heading 1']:
                    next_h1_elem = elem
                    break

    # 遍历 H1 下的所有 H2，按模板子节填充
    elem = h1_elem.getnext()
    h2_elems = []
    while elem is not None and elem is not next_h1_elem:
        if elem.tag.endswith('}p'):
            pPr = elem.find('.//' + qn('w:pPr') + '/' + qn('w:pStyle'))
            if pPr is not None:
                style_val = pPr.get(qn('w:val'), '')
                if style_val in ['2', 'Heading2', 'heading2', 'heading 2', 'Heading 2']:
                    text_elems = elem.findall('.//' + qn('w:t'))
                    h2_text = ''.join(t.text or '' for t in text_elems).strip()
                    h2_elems.append((elem, h2_text))
        elem = elem.getnext()

    for h2_elem, h2_text in h2_elems:
        h2_clean = _strip_chapter_number(h2_text)
        aliases = h2_alias_map.get(h2_clean, ())

        matched_sec = None
        for sec in sections:
            sec_id = sec.get('id') or sec.get('title', '')
            if sec_id in used_sec_ids:
                continue
            sec_title_clean = _strip_chapter_number(sec.get('title', ''))
            for alias in aliases:
                if sec_title_clean == alias or sec_title_clean.endswith(alias):
                    matched_sec = sec
                    break
            if matched_sec:
                break

        if not matched_sec:
            continue

        used_sec_ids.add(matched_sec.get('id') or matched_sec.get('title', ''))

        # 清理 H2 标题下的原占位内容，再插入 section 内容
        from docx.text.paragraph import Paragraph
        h2_para = Paragraph(h2_elem, h1_para._parent)
        _clear_content_between_h2_heading(h2_para)
        _insert_section_as_h2(h2_para, matched_sec)

    # 剩余 sections 追加为 H2（在 H1 末尾）
    for sec in sections:
        sec_id = sec.get('id') or sec.get('title', '')
        if sec_id in used_sec_ids:
            continue
        last = _find_last_para_under_h1(doc, h1_para)
        _insert_section_as_h2(last, sec)

    # 2026-06-06 新增：重编号 H1 下的 H2（4.1~4.10 → chapter_no.1~chapter_no.N）
    if chapter_no is not None:
        _renumber_h2_under_h1(h1_para, chapter_no)


def _append_chapter_as_h1(doc, anchor_para, ch, position='after'):
    """在指定段落附近追加一个 chapter 作为新 H1（标题+所有 sections/内容）

    2026-06-06 新增：position 参数控制插入位置
    - 'after' (默认): 追加到 anchor_para 之后（保持向后兼容）
    - 'before': 插入到 anchor_para 之前（用于在附录 H1 前插入章节）
    """
    ch_title = (ch.get('title') or '').strip() or '新增章节'
    if position == 'before':
        new_h1 = _insert_h1_before(anchor_para, ch_title)
    else:
        new_h1 = _insert_h1_after(anchor_para, ch_title)
    last = new_h1

    sections = ch.get('sections', []) or []
    if sections:
        for sec in sections:
            last = _insert_section_as_h2(last, sec)
    else:
        content = ch.get('content', {})
        if isinstance(content, dict):
            desc = content.get('description')
            if desc:
                new_p = _insert_paragraph_after(last, desc)
                set_black(new_p)
                last = new_p
            headers = content.get('headers')
            rows = content.get('rows')
            if headers and rows:
                _add_table_after_paragraph(last, headers, rows)
        elif isinstance(content, str) and content:
            new_p = _insert_paragraph_after(last, content)
            set_black(new_p)
            last = new_p

    # 2026-06-06 新增：追加的 H1 同样需要重编号 H2 子节（JSON 中 sections 的 title
    # 可能是 "4.1 功能描述" 这种硬编码编号，需替换为新 H1 在文档中的实际章节号）
    # 2026-06-06 修复：用 OXML 检查 pStyle 而非 p.style.name（手动构建的段落
    # 无法通过 python-docx 自动解析样式表，p.style.name 始终返回 Normal）
    chapter_no = _get_h1_chapter_no(doc, new_h1)
    if chapter_no is not None:
        _renumber_h2_under_h1(new_h1, chapter_no)
    return new_h1


def _insert_h1_before(before_para, h1_text):
    """在指定段落前插入一个 H1 标题，返回该 H1 段落对象

    2026-06-06 新增：与 _insert_h1_after 对称，配合 _append_chapter_as_h1(position='before') 使用
    """
    new_h1 = _insert_paragraph_before(before_para, h1_text, style='Heading 1')
    set_black(new_h1)
    return new_h1


# ==================== 章节分类与 H1 填充函数结束 ====================


def _find_matching_chapter(h1_title, h2_title, chapter_map):
    """根据模板H1+H2标题组合查找匹配的JSON章节

    优先级：
    1. H1+H2组合精确匹配映射表（空列表=有映射但无JSON对应，返回None交由特殊处理）
    2. H1精确匹配映射表
    3. H2单独匹配映射表
    4. 模糊匹配（去除编号前缀后比较）

    2026-06-04 优化：模板标题可能带编号前缀（如"1.1 编写目的"），
    去除编号后再查映射表，避免因编号导致精确匹配失败。
    """
    # 预处理：去除编号前缀（如"1.1 "、"2.3 "等），生成 clean 版本
    h1_clean = re.sub(r'^\d+(\.\d+)*\s*', '', h1_title) if h1_title else h1_title
    h2_clean = re.sub(r'^\d+(\.\d+)*\s*', '', h2_title) if h2_title else h2_title

    # 1. H1+H2组合匹配（同时尝试原始标题和去编号标题）
    if h2_title:
        for combo_key in [f'{h1_title}|{h2_title}', f'{h1_clean}|{h2_clean}',
                          f'{h1_title}|{h2_clean}', f'{h1_clean}|{h2_title}']:
            if combo_key in TEMPLATE_CHAPTER_MAP:
                aliases = TEMPLATE_CHAPTER_MAP[combo_key]
                # 空列表=有映射但无JSON对应，返回None交由特殊处理逻辑
                if not aliases:
                    return None
                for alias in aliases:
                    if alias in chapter_map:
                        return chapter_map[alias]

    # 2. H1精确匹配（同时尝试原始标题和去编号标题）
    for h1_key in [h1_title, h1_clean]:
        if h1_key:
            h1_aliases = TEMPLATE_CHAPTER_MAP.get(h1_key, [])
            for alias in h1_aliases:
                if alias in chapter_map:
                    return chapter_map[alias]

    # 2.5 业务模块H1标题含"模块设计说明"时，不走后续模糊匹配，交由H1级别逻辑处理
    if h1_title and '模块设计说明' in h1_title:
        return None

    # 3. H2单独匹配（当H1+H2组合无映射时，同时尝试原始标题和去编号标题）
    if h2_title:
        for h2_key in [h2_title, h2_clean]:
            if h2_key:
                h2_aliases = TEMPLATE_CHAPTER_MAP.get(h2_key, [])
                for alias in h2_aliases:
                    if alias in chapter_map:
                        return chapter_map[alias]

    # 4. 模糊匹配：去掉"第X章 "前缀和编号前缀后比较
    search_title = h2_title or h1_title
    # 对 search_title 也去除编号前缀（模板标题可能带编号如"4.7 类图"）
    search_clean = re.sub(r'^\d+(\.\d+)*\s*', '', search_title)
    for ch_title, ch_data in chapter_map.items():
        clean_title = re.sub(r'^第[一二三四五六七八九十]+章\s*', '', ch_title)
        # 去掉编号前缀（如"3.1 "、"4.7 "等）
        clean_title = re.sub(r'^\d+(\.\d+)*\s*', '', clean_title)
        if clean_title == search_clean or search_clean in clean_title or clean_title in search_clean:
            return ch_data

    return None


def generate_design_from_template(template_path, design_data_path, output_path, diagram_dir=None):
    """基于 .docx 模板生成详细设计文档"""
    with open(r'D:\code\QJ\BEMP5.0DEV\.trae\skills\bemp-advanced-doc-generator\scripts\_fn_trace.log', 'a', encoding='utf-8') as _ftr:
        _ftr.write('[DBG-FN] generate_design_from_template called\n')
        _ftr.flush()
    shutil.copy2(template_path, output_path)
    doc = Document(output_path)
    # 2026-06-07 新增：注册当前 doc，让 _insert_paragraph_xxx 自动解析模板 styleId
    _set_current_doc(doc)

    with open(design_data_path, 'r', encoding='utf-8') as f:
        design_data = json.load(f)

    module_name = design_data.get('moduleName', '本项目')
    chapters = design_data.get('chapters', [])
    cover_page = design_data.get('coverPage', {})
    revision_history = design_data.get('revisionHistory', {})

    # 2026-06-07 修复：保留模式。当 design_data 无章节数据时，或 _preserve 标志为 true 时，
    # 不清理模板正文内容，仅替换封面/修订记录等字段，保留模板原始占位文字供用户后续编辑。
    # 原因：_clear_content_between_headings 在 chapters 为空时仍会清除模板段落，
    # 导致大部分内容丢失（从 305 段 → 184 段，丢失 121 段 ~40%）
    global _PRESERVE_MODE
    _PRESERVE_MODE = design_data.get('_preserve', False) or (len(chapters) == 0)

    # 2026-06-04 新增：加载 doc_rules.yaml 配置（rules.debug_layout 控制调试日志）
    rules = _load_doc_rules()
    if rules.get('debug_layout'):
        global DEBUG_LAYOUT
        DEBUG_LAYOUT = True

    # 2026-06-04 新增：统一目录标题（去除全角/半角空格），防止模板使用"目  录"导致匹配失败
    _normalize_toc_heading(doc)

    filled_contexts = set()
    paragraphs = doc.paragraphs
    total = len(paragraphs)

    # 第一遍：封面替换 + 蓝色文本清理
    cover_replaced = False
    for i, p in enumerate(paragraphs):
        text = p.text.strip()

        # 封面标题替换（只替换第一个含XXX或详细设计的非标题段落）
        if not cover_replaced and ('XXX' in text or '/项目' in text):
            if not p.style or not p.style.name.startswith('Heading'):
                if cover_page.get('title'):
                    # 2026-06-07 修复：通用化占位符替换（${moduleName}/${currentDate}）
                    # coverPage.title 数据中可能含 ${fieldName} 占位符，
                    # 需在写入前替换为 design_data 实际值，否则字面 ${moduleName} 会
                    # 原样显示在封面上
                    rendered_title = _render_template(cover_page['title'], design_data, module_name)
                    write_paragraph(p, rendered_title)
                else:
                    write_paragraph(p, f'{module_name} 详细设计文档')
                cover_replaced = True
                continue

        # 封面副标题"详细设计说明书"保留不变
        # 部门替换（匹配模板中的部门占位文字）
        _DEPT_PLACEHOLDERS = ('信息科技部', '技术开发部', '研发部', '技术部')
        if any(ph in text for ph in _DEPT_PLACEHOLDERS):
            dept = cover_page.get('department', '业务部门')
            write_paragraph(p, dept)
            continue

        # 日期替换
        if re.match(r'\d{4}年\d+月', text):
            write_paragraph(p, datetime.now().strftime('%Y年%m月'))
            continue

        # 清理蓝色占位段落 (2026-06-07 修复：保留模式下跳过，避免删除模板正文)
        is_heading = p.style and p.style.name.startswith('Heading')
        if not _PRESERVE_MODE:
            if is_heading and is_blue_paragraph(p) and not text:
                remove_paragraph(p)
                continue

            # 清理蓝色占位文本（非标题）
            if is_blue_paragraph(p) and is_blue_placeholder_text(text) and not is_heading:
                remove_paragraph(p)
                continue

            # 2026-06-04 增强：彻底清除蓝色 run 内容（非超链接、非占位的蓝色文字变黑且清空）
            if not is_heading and is_blue_paragraph(p) and text:
                # 表格内 run 由后续 _clear_blue_table_cells 处理
                if p._element.getparent() is not None and p._element.getparent().tag.endswith('}tc'):
                    pass
                else:
                    clear_all_blue_runs(p)
                    # 清除后如段落为空，移除整段
                    if not p.text.strip():
                        remove_paragraph(p)
                        continue

        # 蓝色标题保留但变黑（保留模式也执行，因为只改颜色不删内容）
        if is_heading and is_blue_paragraph(p) and text and not is_blue_placeholder_text(text):
            set_black(p)

    # 第二遍：根据标题上下文填充内容
    # 构建章节内容映射：标题 → 内容
    chapter_map = {}
    for ch in chapters:
        ch_title = ch.get('title', '')
        chapter_map[ch_title] = ch
        if ch.get('sections'):
            for sec in ch['sections']:
                sec_title = sec.get('title', '')
                chapter_map[sec_title] = sec

    # 获取附录数据
    appendix = design_data.get('appendix', {})
    glossary = appendix.get('glossary', [])
    references = appendix.get('references', [])

    paragraphs = doc.paragraphs
    total = len(paragraphs)

    # 收集所有已匹配的JSON章节标题
    matched_json_chapters = set()

    # 2026-06-07 修复：保留模式下跳过所有内容填充/清理，直接跳到 TOC 域检查
    if _PRESERVE_MODE:
        # 仅保留必要操作：修订记录更新 + 标题编号重构 + 样式统一 + 段落缩进
        _update_revision_table(doc, module_name, revision_history)
        _ensure_toc_field(doc, levels=rules.get('toc_levels', '1-3'))
        try:
            renumber_headings(doc, rules=rules)
        except Exception as e:
            print(f'[WARN] 标题编号重构失败: {e}', file=sys.stderr)
        try:
            auto_promote_bold_subheadings(doc, rules=rules)
        except Exception as e:
            print(f'[WARN] 子标题样式提升失败: {e}', file=sys.stderr)
        try:
             u_stats = unify_heading_styles(doc, rules=rules)
             print(f'[INFO] 标题样式统一: {u_stats}', file=sys.stderr)
        except Exception as e:
            print(f'[WARN] 标题样式统一失败: {e}', file=sys.stderr)
        try:
            s_stats = _strip_all_h1h2_hardcoded_numbers(doc)
            print(f'[INFO] H1/H2 硬编码编号清理: 已剥除 {s_stats} 处编号前缀', file=sys.stderr)
        except Exception as e:
            print(f'[WARN] 编号前缀剥除失败: {e}', file=sys.stderr)
        apply_body_indent_to_doc(doc, chars=2, skip_headings=True, skip_tables=True)
        for table in doc.tables:
            apply_table_style(table, style_name='Table Grid', font_size=TABLE_FONT_SIZE, header_bold=True)
        doc.save(output_path)
        return {'success': True, 'outputPath': output_path, 'preserveMode': True}

    for i, p in enumerate(paragraphs):
        text = p.text.strip()
        h1, h2, h3 = get_heading_context(paragraphs, i)
        context_key = f'{h1}|{h2}|{h3}'

        if context_key in filled_contexts:
            continue

        is_heading = p.style and p.style.name.startswith('Heading')

        if is_heading and text:
            # H1匹配chapter：进入"按H2循环填充模式"
            if p.style.name == 'Heading 1':
                chapter_data = None
                # 1. 精确标题匹配
                for ch in chapters:
                    if ch.get('title') == h1:
                        chapter_data = ch
                        break
                # 2. 2026-06-05 新增：模块级标题关键词匹配（如"模块1设计说明"）
                #    模板的"模块1设计说明"不直接对应JSON的章节标题，需要匹配到包含模块数据的章节
                if not chapter_data and '模块' in h1 and '设计说明' in h1:
                    for ch in chapters:
                        ch_title = ch.get('title', '')
                        # 匹配"功能模块划分"、"模块划分"等含模块相关数据的章节
                        if '模块' in ch_title or '功能模块' in ch_title:
                            chapter_data = ch
                            break
                if chapter_data and chapter_data.get('sections'):
                    filled_contexts.add(context_key)
                    matched_json_chapters.add(chapter_data.get('title', ''))
                    # 2026-06-05 修复：先插入业务子模块（如承兑行额度管理），再填充H2子节
                    _fill_business_submodules(doc, p, design_data)
                    _fill_chapter_h1_sections(doc, p, chapter_data, design_data, chapters)
                    # 2026-06-05 修复：将H1下所有H2的context_key也加入filled_contexts，
                    # 防止后续主循环中重复处理这些H2并覆盖已填充的内容
                    for j in range(i + 1, total):
                        pp = paragraphs[j]
                        if pp.style and pp.style.name.startswith('Heading'):
                            if pp.style.name == 'Heading 1':
                                break
                            if pp.style.name == 'Heading 2':
                                h1c, h2c, h3c = get_heading_context(paragraphs, j)
                                ctx_key = f'{h1c}|{h2c}|{h3c}'
                                filled_contexts.add(ctx_key)
                    continue

            # 2026-06-05 修复：特殊处理必须在 _find_matching_chapter 之前，
            # 否则模糊匹配会抢占特殊处理的优先级（如"适用范围"被匹配到"1.3 范围说明"）
            text_clean = re.sub(r'^\d+(\.\d+)*\s*', '', text)

            # 2026-06-06 修复：概述/项目概述/系统概述的H1级填充由
            # "H1精确分配"段统一处理（_fill_overview_h1），此处跳过避免重复。
            # H2级别的"目的"/"适用范围"等仍走下方的特殊处理逻辑。
            if p.style.name == 'Heading 1' and (
                text in CHAPTER_KEYWORDS_OVERVIEW or
                any(kw in text for kw in CHAPTER_KEYWORDS_OVERVIEW if text != kw)
            ):
                # 标记为已处理（H1 已在 H1 精确分配中处理）
                filled_contexts.add(context_key)
                continue

            # 特殊处理：术语定义 → 从appendix.glossary填充（空时生成默认术语）
            if text_clean in ('术语定义', '术语和缩写', '术语'):
                filled_contexts.add(context_key)
                _clear_content_between_headings(p, paragraphs)
                _fill_glossary_table(doc, p, glossary)
                continue

            # 特殊处理：参考资料 → 从appendix.references填充（空时生成默认参考文档）
            if text_clean in ('参考资料', '参考文档'):
                filled_contexts.add(context_key)
                _clear_content_between_headings(p, paragraphs)
                _fill_references(doc, p, references)
                continue

            # 特殊处理：开发环境 → 通用内容
            if text_clean == '开发环境':
                filled_contexts.add(context_key)
                _clear_content_between_headings(p, paragraphs)
                _fill_dev_environment(doc, p, design_data)
                continue

            # 特殊处理：界面 → 从需求文档提取
            if text_clean == '界面':
                filled_contexts.add(context_key)
                _clear_content_between_headings(p, paragraphs)
                _fill_ui_description(doc, p, design_data)
                continue

            # 特殊处理：性能 → 通用内容
            if text_clean == '性能':
                filled_contexts.add(context_key)
                _clear_content_between_headings(p, paragraphs)
                _fill_performance(doc, p, design_data)
                continue

            # 特殊处理：目的/编写目的 → 从1.1业务背景提取
            if text_clean in ('目的', '编写目的'):
                filled_contexts.add(context_key)
                _clear_content_between_headings(p, paragraphs)
                _fill_purpose(doc, p, chapter_map, module_name)
                continue

            # 特殊处理：适用范围/使用范围 → 从1.3范围说明提取
            if text_clean in ('适用范围', '使用范围', '范围'):
                filled_contexts.add(context_key)
                _clear_content_between_headings(p, paragraphs)
                _fill_scope(doc, p, chapter_map, module_name)
                continue

            # 特殊处理：读者对象 → 通用内容
            if text_clean in ('读者对象', '目标读者'):
                filled_contexts.add(context_key)
                _clear_content_between_headings(p, paragraphs)
                _fill_reader_audience(doc, p, module_name)
                continue

            # 特殊处理：组件内部的模块列表及说明 → 从chapters提取组件汇总表
            if text == '组件内部的模块列表及说明' or '组件内部的模块列表' in text:
                filled_contexts.add(context_key)
                _clear_content_between_headings(p, paragraphs)
                _fill_component_module_list(doc, p, chapters, design_data)
                continue

            # 特殊处理：模块1设计说明 → 插入业务子模块 + 填充H2子节
            if text and '模块1设计说明' in text:
                filled_contexts.add(context_key)
                _fill_business_submodules(doc, p, design_data)
                # 2026-06-05 修复：同时填充H2子节（功能描述/输入项/输出项等）
                # 找到匹配的章节数据（如"第二章 功能模块划分"）
                module_chapter = None
                for ch in chapters:
                    ch_title = ch.get('title', '')
                    if '模块' in ch_title or '功能模块' in ch_title:
                        module_chapter = ch
                        break
                if module_chapter and module_chapter.get('sections'):
                    _fill_chapter_h1_sections(doc, p, module_chapter, design_data, chapters)
                # 标记H1下所有H2为已处理，防止后续主循环重复处理
                # 2026-06-05 修复：对"界面"/"性能"H2也显式处理，然后统一标记为已填充
                _refresh_paragraphs = doc.paragraphs
                for j in range(i + 1, len(_refresh_paragraphs)):
                    pp = _refresh_paragraphs[j]
                    if pp.style and pp.style.name.startswith('Heading'):
                        if pp.style.name == 'Heading 1':
                            break
                        if pp.style.name == 'Heading 2':
                            h1c, h2c, h3c = get_heading_context(_refresh_paragraphs, j)
                            ctx_key = f'{h1c}|{h2c}|{h3c}'
                            h2_text = h2c or ''
                            h2_clean = re.sub(r'^\d+(\.\d+)*\s*', '', h2_text)
                            # 显式处理"界面"和"性能"（它们有独立的特殊处理逻辑）
                            if h2_clean == '界面':
                                _clear_content_between_headings(pp, _refresh_paragraphs)
                                _fill_ui_description(doc, pp, design_data)
                            elif h2_clean == '性能':
                                _clear_content_between_headings(pp, _refresh_paragraphs)
                                _fill_performance(doc, pp, design_data)
                            filled_contexts.add(ctx_key)
                continue

            # 特殊处理：附录 → 从chapters汇总生成附录内容
            if text_clean == '附录' or '附录' in text_clean:
                filled_contexts.add(context_key)
                _clear_content_between_headings(p, paragraphs)
                _fill_appendix(doc, p, design_data, chapters)
                continue

            # H1+H2组合或H2匹配：使用_find_matching_chapter（兜底匹配）
            matched_chapter = _find_matching_chapter(h1, text, chapter_map)
            if matched_chapter:
                filled_contexts.add(context_key)
                # 记录已匹配的JSON章节标题
                if 'sections' in matched_chapter:
                    matched_json_chapters.add(matched_chapter.get('title', ''))
                # 2026-06-04 优化：仅在 content 存在时调用 _fill_chapter_content，
                # 否则可能误用其他章节的 data 导致窜行
                if _has_meaningful_content(matched_chapter):
                    _fill_chapter_content(doc, p, i, matched_chapter, paragraphs)
                else:
                    _clear_content_between_headings(p, paragraphs)
                    _insert_placeholder_after(p, '不涉及')
                continue

    # ==================== 2026-06-06 修复：H1 精确分配 + 末尾追加 ====================
    # 原"第三遍"逻辑：将所有未匹配 chapters 强制塞到"模块2设计说明"下，
    # 导致设计目标表格错位、模块2 编号错误、附录丢失等问题。
    # 新策略：
    # 1. 通过 _classify_chapters 识别 chapters 的类型
    # 2. 将每个分类 ch 精准分配到模板对应的 H1 标题下
    # 3. 未归入的 ch（other_chs）以"标题作为新 H1"形式追加到文档末尾

    overview_chs, constraint_ch, component_list_ch, business_chs, appendix_ch, other_chs = \
        _classify_chapters(chapters)

    # 2026-06-07 新增：通用化兜底 — 当 chapters 未含"附录"标题（_classify_chapters 返回 None），
    # 但 design_data 顶层有 appendix 数据（glossary/references/errorCodes/dataDictionary），
    # 合成一个虚拟 appendix_ch 注入流程，确保模板无"附录" H1 时也能渲染出 ABCDEF 分类
    if appendix_ch is None and isinstance(design_data, dict):
        top_appendix = design_data.get('appendix') or {}
        if isinstance(top_appendix, dict) and any(
                top_appendix.get(k) for k in
                ('glossary', 'references', 'errorCodes', 'dataDictionary', 'sections')):
            appendix_ch = {
                'id': 'appendix',
                'title': '附录',
                'sections': [],
                # 把顶层字段平铺进去，_fill_appendix_h1 / _fill_appendix 会按需读取
                'glossary': top_appendix.get('glossary', []),
                'references': top_appendix.get('references', []),
                'errorCodes': top_appendix.get('errorCodes', []),
                'dataDictionary': top_appendix.get('dataDictionary', []),
                '_synthesized_from_top_level': True,
            }

    # 找到模板中所有的 H1 标题及其位置
    template_h1_elems = []  # [(h1_para, h1_text_clean), ...]
    for p in doc.paragraphs:
        if p.style and p.style.name == 'Heading 1':
            t = _strip_chapter_number(p.text)
            if t:
                template_h1_elems.append((p, t))

    def _find_template_h1_by_predicate(predicate):
        """在模板 H1 中按 predicate(h1_text_clean) 查找第一个匹配的 H1 段落"""
        for hp, htext in template_h1_elems:
            if predicate(htext):
                return hp
        return None

    # 1. 概述 ch → 模板 H1 概述
    if overview_chs:
        overview_ch = overview_chs[0]  # 多个取第一个
        overview_h1 = _find_template_h1_by_predicate(lambda t: '概述' in t)
        # 2026-06-06 修复：模板可能没有"概述"H1（如"模块名"），此时用第一个 H1 作为兜底，
        # 否则"概述"内容会被整体跳过，导致 1.1~1.6 等 H2 出现但无 H1 归属
        if overview_h1 is None and template_h1_elems:
            overview_h1 = template_h1_elems[0][0]
            # 将 H1 标题改为"概述"（避免"模块名"残留）
            if overview_h1.text.strip() != '概述':
                _set_paragraph_text(overview_h1._element, '概述')
        if overview_h1 is not None:
            # 2026-06-06 修复：填充前先清空 H1 与下一个 H1 之间的旧内容（模板中残留的
            # "1.1 业务背景"等空 H2），避免与新插入的 sections 重复出现
            # stop_at_next_h1=False：嵌套 H2/H3 一并删除，确保完全重写
            _clear_content_between_headings(overview_h1, list(doc.paragraphs), stop_at_next_h1=False)
            ctx = f'Heading 1|{overview_h1.text.strip()}'
            if ctx not in filled_contexts:
                filled_contexts.add(ctx)
                matched_json_chapters.add(overview_ch.get('title', ''))
                _fill_overview_h1(doc, overview_h1, overview_ch, design_data)
                # 标记 H1 下所有 H2 为已处理，避免重复填充
                for j, pp in enumerate(doc.paragraphs):
                    if pp is overview_h1:
                        continue
                    if pp.style and pp.style.name == 'Heading 1':
                        if '概述' in _strip_chapter_number(pp.text):
                            continue
                        break
                    if pp.style and pp.style.name == 'Heading 2':
                        h1c, h2c, h3c = get_heading_context(doc.paragraphs, j)
                        filled_contexts.add(f'{h1c}|{h2c}|{h3c}')

    # 2. 组件列表 ch → 模板 H1 "组件内部的模块列表及说明"（数据从 business_chs 提取）
    if component_list_ch is not None or business_chs:
        comp_h1 = _find_template_h1_by_predicate(
            lambda t: '组件' in t and '列表' in t or '模块列表' in t or '组件汇总' in t)
        if comp_h1 is not None:
            ctx = f'Heading 1|{comp_h1.text.strip()}'
            if ctx not in filled_contexts:
                filled_contexts.add(ctx)
                # 2026-06-06 修复：数据源从 business_chs 提取（避免使用业务无关的 component_list_ch）
                _clear_content_between_headings(comp_h1, list(doc.paragraphs))
                _fill_component_module_list(doc, comp_h1, business_chs, design_data)

    # 3. 业务模块 ch 按顺序分配到模板 H1 "模块1设计说明"/"模块2设计说明"/...
    # 重新扫描模板 H1 列表
    template_h1_elems = [(p, _strip_chapter_number(p.text)) for p in doc.paragraphs
                         if p.style and p.style.name == 'Heading 1']
    biz_module_h1s = [(p, t) for p, t in template_h1_elems
                      if '模块' in t and '设计说明' in t]
    # 业务模块 ch 按需重新编号（如果模板 H1 数量 != business_chs 数量，
    # 多的追加为新 H1，少的丢弃）
    if business_chs:
        with open(r'D:\code\QJ\BEMP5.0DEV\.trae\skills\bemp-advanced-doc-generator\scripts\_fn_trace.log', 'a', encoding='utf-8') as _ftr:
            _ftr.write(f'[DBG-LOOP] business_chs count={len(business_chs)}, biz_module_h1s count={len(biz_module_h1s)}\n')
            _ftr.flush()
        # 2026-06-06 新增：计算每个业务模块 H1 的"章节号"（= 在文档 H1 列表中的 1-indexed 位置）
        # 用于 H2 子节按实际位置编号（避免模板硬编码 4.1~4.10 错位）
        # 注意：H1 样式可能为 "1"、"2"、"Heading 1"、"Heading1" 等多种命名，需兼容
        # 2026-06-06 修复：用 OXML 检测样式（手动创建的 H1 在 doc.paragraphs 中可访问，
        # 但 p.style.name 不会自动解析 OXML 的 pStyle，必须读 _element 的 pStyle）
        all_h1s_in_doc = [p for p in doc.paragraphs
                          if _get_pstyle_val(p) in _H1_STYLE_ALIASES]
        for idx, bch in enumerate(business_chs):
            if idx < len(biz_module_h1s):
                hp, _ = biz_module_h1s[idx]
                # 计算章节号 = 该 H1 在文档中所有 H1 的位置
                chapter_no = _get_h1_chapter_no(doc, hp) or (idx + 2)
                ctx = f'Heading 1|{hp.text.strip()}'
                if ctx not in filled_contexts:
                    filled_contexts.add(ctx)
                    matched_json_chapters.add(bch.get('title', ''))
                    _fill_business_module_h1(doc, hp, bch, design_data, chapters, chapter_no=chapter_no)
                    # 标记 H2 为已处理
                    for j, pp in enumerate(doc.paragraphs):
                        if pp is hp:
                            continue
                        if pp.style and pp.style.name == 'Heading 1':
                            break
                        if pp.style and pp.style.name == 'Heading 2':
                            h1c, h2c, h3c = get_heading_context(doc.paragraphs, j)
                            filled_contexts.add(f'{h1c}|{h2c}|{h3c}')
            else:
                # 模板业务模块 H1 不够，追加为新 H1（_append_chapter_as_h1 内部已重编号 H2）
                anchor = biz_module_h1s[-1][0] if biz_module_h1s else _find_template_h1_by_predicate(
                    lambda t: '附录' in t)
                if anchor is None:
                    anchor = doc.paragraphs[-1]
                _append_chapter_as_h1(doc, anchor, bch)

    # 4. 附录 ch → 模板 H1 "附录"
    if appendix_ch is not None:
        appendix_h1 = _find_template_h1_by_predicate(lambda t: '附录' in t)
        if appendix_h1 is not None:
            ctx = f'Heading 1|{appendix_h1.text.strip()}'
            if ctx not in filled_contexts:
                filled_contexts.add(ctx)
                matched_json_chapters.add(appendix_ch.get('title', ''))
                # 2026-06-06 新增：在附录 H1 之前插入扩展章节（术语表/数据字典/错误码对照表）
                # 通用化设计：仅当 design_data 中存在对应数据时才生成对应扩展章节
                extension_chapters = _build_default_appendix_extension_chapters(design_data)
                if extension_chapters:
                    _insert_appendix_extension_chapters(doc, appendix_h1, extension_chapters)
                _fill_appendix_h1(doc, appendix_h1, appendix_ch)
        else:
            # 2026-06-07 新增：通用化兜底 — 模板无"附录" H1 占位时，自动在文档末尾追加附录章节
            # 触发条件：JSON 含 appendix 数据（如 glossary/references/errorCodes 等）但模板未提供
            # "附录" H1；此时不应丢弃数据，而是按统一 H1 风格追加到文档末尾（避免信息丢失）
            has_appendix_payload = bool(
                appendix_ch.get('sections') or
                (appendix_ch.get('references')) or
                (appendix_ch.get('glossary')) or
                (appendix_ch.get('errorCodes')) or
                (appendix_ch.get('dataDictionary'))
            )
            if has_appendix_payload:
                appendix_h1 = _insert_paragraph_after(doc.paragraphs[-1], '附录', style='Heading 1')
                for r in appendix_h1.runs:
                    r.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
                filled_contexts.add(f'Heading 1|附录')
                matched_json_chapters.add(appendix_ch.get('title', ''))
                # 在新追加的附录前，插入扩展章节（术语表/数据字典/参考等）
                extension_chapters = _build_default_appendix_extension_chapters(design_data)
                if extension_chapters:
                    _insert_appendix_extension_chapters(doc, appendix_h1, extension_chapters)
                _fill_appendix_h1(doc, appendix_h1, appendix_ch)

    # 5. 其他未归入的 ch → 追加为新 H1
    # 2026-06-06 修复：追加位置由"文档末尾"改为"附录H1之前"，
    # 避免附录后出现错位的"7 第三章 核心业务流程"等冗余章节
    if other_chs:
        # 5.1 定位附录 H1 段落（如果存在），将 other_chs 插入其前
        appendix_anchor = _find_template_h1_by_predicate(lambda t: '附录' in t)
        # 5.2 兜底：使用文档最后一段
        for ch in other_chs:
            if appendix_anchor is not None:
                _append_chapter_as_h1(doc, appendix_anchor, ch, position='before')
            else:
                # 兜底：追加到文档末尾（与历史行为保持一致）
                anchor = doc.paragraphs[-1] if doc.paragraphs else None
                if anchor is not None:
                    _append_chapter_as_h1(doc, anchor, ch, position='after')
            matched_json_chapters.add(ch.get('title', ''))

    # ==================== 修复结束 ====================

    # 第四遍：处理表格中的蓝色占位文本 + 表格内蓝色 run 清除
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if is_blue_cell(cell):
                    clear_blue_cell(cell)
                # 2026-06-04 增强：彻底清除表格内蓝色 run 内容
                for para in cell.paragraphs:
                    clear_all_blue_runs(para)

    # 2026-06-04 新增：处理"设计约束"标题下与业务无关的内容
    business_keywords = _extract_business_keywords(design_data, module_name)
    if business_keywords:
        remove_design_constraint_irrelevant(list(doc.paragraphs), business_keywords)
        # 删除后若整个"设计约束"章节为空，补充"不涉及"占位
        _ensure_design_constraint_placeholder(doc, placeholder='不涉及')

    # 2026-06-04 新增：检查错误码章节，缺失则插入"本系统统一错误码定义"+ 示例表
    # 必须在 fill_empty_chapter 之前执行，避免被"不涉及"占位覆盖
    _ensure_error_code_section(doc, design_data)

    # 2026-06-04 新增：界面段落格式化（每行不超过40字符）
    # 必须在 fill_empty_chapter 之前执行，否则原文会被"不涉及"覆盖
    _format_ui_paragraphs(doc)

    # 2026-06-04 新增：删除指定章节范围内的空表格（组件结构图、系统组件、模块复用分析等）
    for kw, _label in EMPTY_CHECK_EXTRA:
        _remove_empty_table_between_headings(doc, start_keyword=kw, end_keyword=None,
                                             placeholder='不涉及')

    # 2026-06-04 新增：检查并填充空子标题（设计目标、输入项、输出项、代码示例、性能优化、附录等）
    fill_empty_chapter(doc, EMPTY_CHECK_KEYWORDS, placeholder='不涉及')

    # 2026-06-04 新增：扩展空章节检查（系统概述、设计约束、模块复用分析、组件结构图、系统组件）
    fill_empty_chapter(doc, EMPTY_CHECK_EXTRA, placeholder='不涉及')

    # 2026-06-04 新增：非功能属性章节兜底（界面、性能、安全性、可靠性、易用性、可调试性、可移植性、可维护性）
    fill_empty_chapter(doc, [(kw, kw) for kw in NON_FUNCTIONAL_CHAPTERS], placeholder='不涉及')

    # endif PRESERVE_TEMPLATE_BODY — 保留模式下跳过以上所有内容填充/清理逻辑

    # 2026-06-07 新增：清理附录F下的"功能模块"不协调表格
    # 根因：_fill_appendix_h1 或章节追加逻辑将模块功能清单表格误放到附录F术语章节下
    clean_appendix_f_table(doc)

    # 2026-06-04 新增：检查 TOC 域，无则插入（含 H1-H3 全层级）
    _ensure_toc_field(doc, levels=rules.get('toc_levels', '1-3'))

    # 2026-06-04 新增：系统组件章节下的子标题重新编号（组件1/组件2/...）
    _renumber_component_subheadings(doc)

    # 2026-06-04 新增：每个组件的"关键技术"章节按 component_name 动态生成（避免通用模板内容一致）
    _build_component_specific_tech(doc, design_data)

    # 2026-06-04 新增：正文段落首行缩进2字符（跳过标题、表格内）
    apply_body_indent_to_doc(doc, chars=2, skip_headings=True, skip_tables=True)

    # 2026-06-04 新增：所有表格统一样式（边框+字号+表头加粗+底色）
    for table in doc.tables:
        apply_table_style(table, style_name='Table Grid', font_size=TABLE_FONT_SIZE, header_bold=True)

    # 第五遍：插入架构图/网络图/部署图/UML图
    print(f'[INFO] diagram_dir 参数: {diagram_dir}', file=sys.stderr)
    if diagram_dir:
        print(f'[INFO] diagram_dir 是否为目录: {os.path.isdir(diagram_dir)}', file=sys.stderr)
        if os.path.isdir(diagram_dir):
            png_count = sum(1 for root, _, files in os.walk(diagram_dir) for f in files if f.lower().endswith('.png'))
            print(f'[INFO] diagram_dir 中 PNG 文件数: {png_count}', file=sys.stderr)
    if diagram_dir and os.path.isdir(diagram_dir):
        _insert_diagrams(doc, diagram_dir)
    else:
        print('[WARN] diagram_dir 为空或不是有效目录，跳过图表插入', file=sys.stderr)

    # 2026-06-06 新增：图片插入后，跨章节统一移除/剥除图片后冗余图说文字
    try:
        rt_stats = remove_redundant_text_after_images(doc, rules=rules)
        print(f'[INFO] 图片后冗余文字处理: {rt_stats}', file=sys.stderr)
    except Exception as e:
        print(f'[WARN] 图片后冗余文字处理失败: {e}', file=sys.stderr)

    # 2026-06-06 新增：功能模块表格清洗（去除技术实现列/行）
    try:
        t_stats = clean_all_function_module_tables(doc, rules=rules)
        if t_stats.get('tables_scanned', 0) > 0:
            print(f'[INFO] 功能模块表格清洗: {t_stats["tables_scanned"]} 张扫描, '
                  f'{t_stats["tables_cleaned"]} 张已清洗', file=sys.stderr)
    except Exception as e:
        print(f'[WARN] 功能模块表格清洗失败: {e}', file=sys.stderr)

    # 2026-06-06 新增：标题编号重构（按 OXML 层级动态重新编号，解决重复/错位）
    try:
        r_stats = renumber_headings(doc, rules=rules)
        print(f'[INFO] 标题编号重构: {r_stats}', file=sys.stderr)
    except Exception as e:
        print(f'[WARN] 标题编号重构失败: {e}', file=sys.stderr)

    # 2026-06-07 新增：自动提升 Normal/加粗的子标题为对应 Heading 样式
    # 解决"子标题未设置带样式标题"问题（模板子标题用 Normal + 加粗呈现，
    # 未带 Heading 样式，导致 TOC/大纲视图无法识别）
    try:
        p_stats = auto_promote_bold_subheadings(doc, rules=rules)
        if p_stats.get('promoted_h1', 0) > 0 or p_stats.get('promoted_h2', 0) > 0 or p_stats.get('promoted_h3', 0) > 0:
            print(f'[INFO] 子标题样式补齐: {p_stats}', file=sys.stderr)
    except Exception as e:
        print(f'[WARN] 子标题样式补齐失败: {e}', file=sys.stderr)

    # 2026-06-06 新增：标题样式统一（解决"4.1功能描述"和"4.2模块划分"字号不一致问题）
    try:
        u_stats = unify_heading_styles(doc, rules=rules)
        print(f'[INFO] 标题样式统一: {u_stats}', file=sys.stderr)
    except Exception as e:
        print(f'[WARN] 标题样式统一失败: {e}', file=sys.stderr)

    # 2026-06-07 新增：通用化兜底——剥除所有 H1/H2 段落文本中硬编码的章节编号。
    # 模板 H1/H2 样式已绑定 numId 自动编号，Word 会按出现顺序自动渲染。
    # 如果文本里写死了"4.1 功能描述"会与自动编号叠加产生"4.1 4.1 功能描述"。
    # 在所有 H1/H2 写入与样式统一完成后做一次全文清理。
    try:
        s_stats = _strip_all_h1h2_hardcoded_numbers(doc)
        if s_stats:
            print(f'[INFO] H1/H2 硬编码编号清理: 已剥除 {s_stats} 处编号前缀', file=sys.stderr)
    except Exception as e:
        print(f'[WARN] H1/H2 硬编码编号清理失败: {e}', file=sys.stderr)

    # 第六遍：更新修订记录表格
    _update_revision_table(doc, module_name, revision_history)

    # 保存文档
    doc.save(output_path)

    # 注入 updateFields=true 使 Word 打开时自动更新目录
    # 2026-06-04 优化：使用临时文件+重命名模式，避免 zipfile 追加条目
    inject_update_fields(output_path)

    return output_path


def _fill_chapter_h1_sections(doc, h1_para, chapter_data, design_data=None, chapters=None):
    """为H1下的每个H2子节，按section内容填充

    工作流程：
    1. 遍历H1下所有H2子标题
    2. 找到H1的XML元素
    3. 依次处理每个H2：清理H2下原占位内容，填充对应section的内容

    2026-06-05 新增：design_data 和 chapters 参数，用于在无匹配section时
    从需求数据中提取"输入项"/"输出项"/"接口"的具体内容，避免回退到"不涉及"占位。
    """
    h1_elem = h1_para._element
    h1_parent = h1_elem.getparent()
    if h1_parent is None:
        return

    # 构建section标题→内容映射（支持多种匹配方式）
    sections = chapter_data.get('sections', [])
    section_map = {}  # 按H2标题精确匹配
    for sec in sections:
        sec_title = sec.get('title', '')
        section_map[sec_title] = sec

    # 找到所有H1下的H2标题
    h1_idx_in_parent = list(h1_parent).index(h1_elem)

    # 找到下一个H1的位置作为边界
    next_h1_elem = None
    for j in range(h1_idx_in_parent + 1, len(list(h1_parent))):
        elem = list(h1_parent)[j]
        if elem.tag.endswith('}p'):
            pPr = elem.find('.//' + qn('w:pPr') + '/' + qn('w:pStyle'))
            if pPr is not None:
                style_val = pPr.get(qn('w:val'), '')
                if style_val in ['1', 'Heading1', 'heading1', 'heading 1', 'Heading 1']:
                    next_h1_elem = elem
                    break

    # 遍历H1下的所有段落
    elem = h1_elem.getnext()
    h2_elems = []  # 收集所有H2元素
    while elem is not None and elem is not next_h1_elem:
        if elem.tag.endswith('}p'):
            pPr = elem.find('.//' + qn('w:pPr') + '/' + qn('w:pStyle'))
            if pPr is not None:
                style_val = pPr.get(qn('w:val'), '')
                if style_val in ['2', 'Heading2', 'heading2', 'heading 2', 'Heading 2']:
                    # 提取文本
                    text_elems = elem.findall('.//' + qn('w:t'))
                    h2_text = ''.join(t.text or '' for t in text_elems).strip()
                    h2_elems.append((elem, h2_text))
        elem = elem.getnext()

    # 依次为每个H2填充内容
    for h2_elem, h2_text in h2_elems:
        # 1. 优先精确匹配（含编号的完整标题）
        matched_section = section_map.get(h2_text)

        # 2. 规范化匹配：去掉编号前缀和空白后比较
        #    模板 H2 可能带编号前缀如"4.7 类图"，JSON section 标题通常为"类图"
        if not matched_section:
            h2_normalized = re.sub(r'^\d+(\.\d+)*\s*', '', h2_text)  # 去掉编号前缀
            h2_normalized = re.sub(r'\s+', '', h2_normalized)
            for sec_title, sec in section_map.items():
                sec_normalized = re.sub(r'^\d+(\.\d+)*\s*', '', sec_title)
                sec_normalized = re.sub(r'\s+', '', sec_normalized)
                if sec_normalized == h2_normalized:
                    matched_section = sec
                    break

        # 3. 2026-06-05 新增：模块级H2映射（解决"模块1设计说明"下子节匹配问题）
        if not matched_section and h2_normalized in MODULE_H2_SECTION_MAP:
            aliases = MODULE_H2_SECTION_MAP[h2_normalized]
            for alias in aliases:
                if alias in section_map:
                    matched_section = section_map[alias]
                    break
                # 也尝试规范化匹配
                for sec_title, sec in section_map.items():
                    sec_normalized = re.sub(r'^\d+(\.\d+)*\s*', '', sec_title)
                    sec_normalized = re.sub(r'\s+', '', sec_normalized)
                    if sec_normalized == alias:
                        matched_section = sec
                        break
                if matched_section:
                    break

        if not matched_section:
            # 2026-06-05 新增：无匹配section时的H2兜底填充
            # 界面/性能/接口/类图/顺序图/活动图/备注 等H2在JSON中可能没有对应section
            # 需要插入合理的默认内容，避免文档中出现空章节
            h2_clean = re.sub(r'^\d+(\.\d+)*\s*', '', h2_text)
            from docx.text.paragraph import Paragraph
            h2_para = Paragraph(h2_elem, h1_para._parent)
            _clear_content_between_h2_heading(h2_para)

            if h2_clean == '界面':
                placeholder = '界面设计详见需求文档中的界面原型图；界面布局主要包括查询区、列表区、操作区三个部分，菜单位置与查询条件以原型图为准。'
                new_p_elem = _create_paragraph_xml(placeholder)
                h2_elem.addnext(new_p_elem)
            elif h2_clean == '性能':
                perf_items = [
                    '核心操作响应时间：< 500ms',
                    '列表查询响应时间：< 1s',
                    '并发用户数：支持50+并发',
                    '数据量：支持10万+记录查询',
                ]
                insert_anchor = h2_elem
                for item in perf_items:
                    new_p_elem = _create_paragraph_xml(item)
                    insert_anchor.addnext(new_p_elem)
                    insert_anchor = new_p_elem
            elif h2_clean == '输入项':
                # 2026-06-05 修复：从chapters中提取栏位描述表中的"输入"字段，
                # 避免回退到"不涉及"占位（fill_empty_chapter 兜底）
                input_items = _extract_field_items(chapters, '输入')
                if input_items:
                    insert_anchor = h2_elem
                    for item in input_items:
                        new_p_elem = _create_paragraph_xml(item)
                        insert_anchor.addnext(new_p_elem)
                        insert_anchor = new_p_elem
                else:
                    placeholder = '本模块输入项主要包括查询条件、表单输入字段等，具体定义详见需求文档中的栏位描述表。'
                    new_p_elem = _create_paragraph_xml(placeholder)
                    h2_elem.addnext(new_p_elem)
            elif h2_clean == '输出项':
                # 2026-06-05 修复：从chapters中提取栏位描述表中的"输出"字段
                output_items = _extract_field_items(chapters, '输出')
                if output_items:
                    insert_anchor = h2_elem
                    for item in output_items:
                        new_p_elem = _create_paragraph_xml(item)
                        insert_anchor.addnext(new_p_elem)
                        insert_anchor = new_p_elem
                else:
                    placeholder = '本模块输出项主要包括列表展示、详情展示等结果数据，具体定义详见需求文档中的栏位描述表。'
                    new_p_elem = _create_paragraph_xml(placeholder)
                    h2_elem.addnext(new_p_elem)
            elif h2_clean == '接口':
                # 2026-06-05 修复：从chapters中提取接口相关数据，而非仅填充占位文本
                interface_items = _extract_interface_items(chapters, design_data)
                if interface_items:
                    insert_anchor = h2_elem
                    for item in interface_items:
                        new_p_elem = _create_paragraph_xml(item)
                        insert_anchor.addnext(new_p_elem)
                        insert_anchor = new_p_elem
                else:
                    placeholder = '本模块接口设计详见接口文档，主要接口包括查询、新增、修改、删除等标准CRUD操作，具体接口定义参见各子模块设计说明。'
                    new_p_elem = _create_paragraph_xml(placeholder)
                    h2_elem.addnext(new_p_elem)
            elif h2_clean in ('类图', '顺序图', '活动图', '业务流程图', '时序图'):
                # 2026-06-07 修复：使用专业绘图工具（graphviz/matplotlib）自动生成图表
                # 替代原有的"详见UML设计文档"占位文字
                h2_para_obj = Paragraph(h2_elem, h1_para._parent) if isinstance(h2_elem, type(OxmlElement('w:p'))) else None
                if h2_para_obj is not None and generate_and_insert_diagram is not None:
                    # 获取 H1 父标题作为图表上下文
                    h1_text = h1_para.text.strip() if hasattr(h1_para, 'text') else ''
                    try:
                        inserted = generate_and_insert_diagram(doc, h2_para_obj, h2_clean, parent_context=h1_text)
                    except Exception:
                        inserted = False
                    if not inserted:
                        # 降级：图表生成失败时仍用占位文字
                        placeholder = f'（{h2_clean}详见UML设计文档，此处为占位说明）'
                        new_p_elem = _create_paragraph_xml(placeholder)
                        h2_elem.addnext(new_p_elem)
                else:
                    placeholder = f'（{h2_clean}详见UML设计文档，此处为占位说明）'
                    new_p_elem = _create_paragraph_xml(placeholder)
                    h2_elem.addnext(new_p_elem)
            elif h2_clean == '备注':
                new_p_elem = _create_paragraph_xml('暂无')
                h2_elem.addnext(new_p_elem)
            continue

        # 清理H2下的原占位内容（仅当有匹配内容时才清理）
        from docx.text.paragraph import Paragraph
        h2_para = Paragraph(h2_elem, h1_para._parent)
        _clear_content_between_h2_heading(h2_para)

        # 填充section内容
        content = matched_section.get('content', {})

        # 找到H2元素在parent中的位置，新内容插在H2后
        insert_anchor = h2_elem

        if isinstance(content, dict):
            description = content.get('description', '')
            if description:
                new_p_elem = _create_paragraph_xml(description)
                insert_anchor.addnext(new_p_elem)
                insert_anchor = new_p_elem

            # 表格
            headers = content.get('headers', [])
            rows = content.get('rows', [])
            if headers and rows:
                # 在H2后插入表格
                # 2026-06-05 修复：若只有表格无描述段落，先插入一个描述段落。
                # fill_empty_chapter 只检查段落文本不检查表格，无描述段落时会被误判为空并覆盖。
                if not description:
                    table_desc = f'以下表格列出了{matched_section.get("title", "")}的相关信息：'
                    new_p_elem = _create_paragraph_xml(table_desc)
                    insert_anchor.addnext(new_p_elem)
                    insert_anchor = new_p_elem
                _add_table_after_h2(insert_anchor, headers, rows)


def _extract_field_items(chapters, field_type):
    """从chapters中提取指定类型（'输入'/'输出'）的栏位描述字段

    遍历所有章节的section数据，查找包含field_type关键字的表格行，
    以"字段名 | 类型 | 说明"格式返回，用于填充"输入项"/"输出项"章节。
    如果chapters为None则返回空列表。
    """
    if not chapters:
        return []
    items = []
    for ch in chapters:
        for sec in ch.get('sections', []):
            content = sec.get('content', {})
            if isinstance(content, dict):
                headers = content.get('headers', [])
                rows = content.get('rows', [])
                if headers and rows:
                    # 查找标题行中包含field_type关键字的列索引
                    target_cols = []
                    for j, h in enumerate(headers):
                        if field_type in str(h):
                            target_cols.append(j)
                    if target_cols:
                        # 提取所有行中对应列的内容
                        for row in rows:
                            row_parts = []
                            for j, cell in enumerate(row):
                                cell_str = str(cell).strip()
                                if cell_str and cell_str != 'None':
                                    row_parts.append(f'{headers[j] if j < len(headers) else ""}: {cell_str}')
                            if row_parts:
                                items.append('；'.join(row_parts))
                    else:
                        # 无明确列标题匹配时，检查整行是否包含field_type关键字
                        for row in rows:
                            row_str = ' | '.join(str(c) for c in row if str(c).strip())
                            if field_type in row_str and row_str not in items:
                                items.append(row_str)
    return items


def _extract_interface_items(chapters, design_data):
    """从chapters中提取接口相关数据

    遍历所有章节的section数据，查找标题含"接口"/"API"/"方法"的表格，
    以"接口名 | 说明"格式返回。如果无数据则返回空列表。
    """
    if not chapters:
        return []
    items = []
    for ch in chapters:
        for sec in ch.get('sections', []):
            sec_title = sec.get('title', '')
            content = sec.get('content', {})
            if isinstance(content, dict):
                headers = content.get('headers', [])
                rows = content.get('rows', [])
                # 匹配section标题或表格标题含接口关键字
                if headers and any(
                    '接口' in str(h) or 'API' in str(h) or '方法' in str(h) or 'URL' in str(h)
                    for h in headers
                ):
                    if len(headers) >= 2:
                        # 多列表格：提取前两列作为接口名和说明
                        for row in rows:
                            name = str(row[0]).strip() if len(row) > 0 else ''
                            desc = str(row[1]).strip() if len(row) > 1 else ''
                            if name:
                                items.append(f'{name}：{desc}' if desc else name)
                    else:
                        # 单列表格
                        for row in rows:
                            items.append(' | '.join(str(c) for c in row))
    return items


def _clear_content_between_h2_heading(h2_para):
    """清理H2标题和下一个H2之间的所有非标题段落"""
    h2_elem = h2_para._element
    parent = h2_elem.getparent()
    if parent is None:
        return

    heading_style_vals = {
        '1', '2', '3', '4', '5', '6',
        'Heading1', 'Heading2', 'Heading3', 'Heading4', 'Heading5', 'Heading6',
        'heading1', 'heading2', 'heading3', 'heading4', 'heading5', 'heading6',
        'heading 1', 'heading 2', 'heading 3', 'heading 4', 'heading 5', 'heading 6',
        'Heading 1', 'Heading 2', 'Heading 3', 'Heading 4', 'Heading 5', 'Heading 6',
    }

    to_remove = []
    found = False
    for elem in list(parent):
        if elem is h2_elem:
            found = True
            continue
        if not found:
            continue
        if elem.tag.endswith('}p'):
            pPr = elem.find('.//' + qn('w:pPr') + '/' + qn('w:pStyle'))
            if pPr is not None:
                style_val = pPr.get(qn('w:val'), '')
                if style_val in heading_style_vals:
                    break
            to_remove.append(elem)

    for elem in to_remove:
        parent.remove(elem)


def _create_paragraph_xml(text):
    """创建普通段落XML元素"""
    new_p = OxmlElement('w:p')
    r = OxmlElement('w:r')
    # 设置黑色字体
    rPr = OxmlElement('w:rPr')
    color = OxmlElement('w:color')
    color.set(qn('w:val'), '000000')
    rPr.append(color)
    r.append(rPr)
    t = OxmlElement('w:t')
    t.text = text
    t.set(qn('xml:space'), 'preserve')
    r.append(t)
    new_p.append(r)
    return new_p


def _add_table_after_h2(h2_elem, headers, rows):
    """在H2元素后插入表格"""
    parent = h2_elem.getparent()
    if parent is None:
        return

    # 创建表格XML
    tbl = OxmlElement('w:tbl')

    # 表格属性
    tblPr = OxmlElement('w:tblPr')
    tblW = OxmlElement('w:tblW')
    tblW.set(qn('w:w'), '0')
    tblW.set(qn('w:type'), 'auto')
    tblPr.append(tblW)
    tblBorders = OxmlElement('w:tblBorders')
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement('w:' + border_name)
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:color'), '000000')
        tblBorders.append(border)
    tblPr.append(tblBorders)
    tbl.append(tblPr)

    # 表格网格
    num_cols = len(headers)
    tblGrid = OxmlElement('w:tblGrid')
    for _ in range(num_cols):
        gridCol = OxmlElement('w:gridCol')
        tblGrid.append(gridCol)
    tbl.append(tblGrid)

    # 表头行
    header_tr = OxmlElement('w:tr')
    for h in headers:
        tc = OxmlElement('w:tc')
        tcPr = OxmlElement('w:tcPr')
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), 'D9D9D9')
        tcPr.append(shd)
        tc.append(tcPr)
        p = OxmlElement('w:p')
        r = OxmlElement('w:r')
        rPr = OxmlElement('w:rPr')
        b = OxmlElement('w:b')
        rPr.append(b)
        sz = OxmlElement('w:sz')
        sz.set(qn('w:val'), '21')  # 10.5pt
        rPr.append(sz)
        r.append(rPr)
        t = OxmlElement('w:t')
        t.text = str(h)
        t.set(qn('xml:space'), 'preserve')
        r.append(t)
        p.append(r)
        tc.append(p)
        header_tr.append(tc)
    tbl.append(header_tr)

    # 数据行
    for row in rows:
        tr = OxmlElement('w:tr')
        for cell in row:
            tc = OxmlElement('w:tc')
            p = OxmlElement('w:p')
            r = OxmlElement('w:r')
            rPr = OxmlElement('w:rPr')
            sz = OxmlElement('w:sz')
            sz.set(qn('w:val'), '21')
            rPr.append(sz)
            r.append(rPr)
            t = OxmlElement('w:t')
            t.text = str(cell)
            t.set(qn('xml:space'), 'preserve')
            r.append(t)
            p.append(r)
            tc.append(p)
            tr.append(tc)
        tbl.append(tr)

    # 插入到H2元素之后
    h2_elem.addnext(tbl)
    # 在表格后加一个空段落（Word要求）
    empty_p = OxmlElement('w:p')
    tbl.addnext(empty_p)
    # 2026-06-04 优化：表格插入后应用统一样式（边框+字号+表头加粗+底色）
    # 注意：此处 tbl 是 OxmlElement，构建 Table 对象需要 parent；为简化，统一在主流程末尾 apply_table_style 兜底处理所有 doc.tables


def _insert_paragraph_after(paragraph, text, style=None, doc=None):
    new_p = OxmlElement('w:p')
    # 设置段落样式
    if style:
        pPr = OxmlElement('w:pPr')
        pStyle = OxmlElement('w:pStyle')
        # 2026-06-07 优化：通用化——若传入的是 'Heading N' / 'HeadingN'，
        # 通过 doc.styles 解析为模板实际 styleId（如 '1'/'2'/'3'）。
        actual_style = _resolve_style_for_insert(style, doc)
        pStyle.set(qn('w:val'), actual_style)
        pPr.append(pStyle)
        new_p.append(pPr)
    # 添加文本run
    r = OxmlElement('w:r')
    t = OxmlElement('w:t')
    t.text = text
    t.set(qn('xml:space'), 'preserve')
    r.append(t)
    new_p.append(r)
    # 插入到文档
    paragraph._element.addnext(new_p)
    # 返回python-docx段落对象
    from docx.text.paragraph import Paragraph
    return Paragraph(new_p, paragraph._parent)


def _insert_paragraph_before(paragraph, text, style=None, doc=None):
    """在指定段落前插入一个新段落（与 _insert_paragraph_after 对称）

    2026-06-06 新增：用于将未匹配的 chapters 追加为新 H1 到指定锚点之前
    （如"附录"H1 之前），避免在文档末尾错位。

    实现：OXML addprevious，将 new_p 插入到 paragraph 之前。
    """
    new_p = OxmlElement('w:p')
    if style:
        pPr = OxmlElement('w:pPr')
        pStyle = OxmlElement('w:pStyle')
        # 2026-06-07 优化：与 _insert_paragraph_after 对齐，解析模板 styleId
        actual_style = _resolve_style_for_insert(style, doc)
        pStyle.set(qn('w:val'), actual_style)
        pPr.append(pStyle)
        new_p.append(pPr)
    r = OxmlElement('w:r')
    t = OxmlElement('w:t')
    t.text = text
    t.set(qn('xml:space'), 'preserve')
    r.append(t)
    new_p.append(r)
    paragraph._element.addprevious(new_p)
    from docx.text.paragraph import Paragraph
    return Paragraph(new_p, paragraph._parent)


def _get_diagram_filename_keywords(keywords):
    """根据中文图表标题关键词，返回用于文件名模糊匹配的英文关键词列表

    2026-06-05 新增：解决实际图片文件名不符合预期命名规范时（如
    class_diagram.png 而非 class-diagram.png），精确匹配失败的问题。
    """
    keyword_map = {
        '系统架构': ['architecture', 'arch', 'system'],
        '组件结构图': ['component', 'comp', 'struct', 'architecture'],
        '架构图': ['architecture', 'arch'],
        '网络结构图': ['network', 'topology', 'net'],
        '网络拓扑': ['network', 'topology', 'net'],
        '网络图': ['network', 'net'],
        '部署图': ['deploy', 'deployment'],
        '部署架构': ['deploy', 'deployment'],
        '类图': ['class', 'uml'],
        '类图设计': ['class', 'uml'],
        '顺序图': ['sequence', 'seq', 'uml'],
        '时序图': ['sequence', 'seq', 'uml'],
        '序列图': ['sequence', 'seq', 'uml'],
        '活动图': ['activity', 'flow', 'uml'],
        '流程图': ['activity', 'flow', 'flowchart'],
        'ER图': ['er', 'entity', 'entity-relationship'],
        'E-R图': ['er', 'entity', 'entity-relationship'],
        '实体关系': ['er', 'entity', 'entity-relationship'],
        '数据流图': ['dataflow', 'data-flow', 'dfd'],
        '数据流': ['dataflow', 'data-flow', 'dfd'],
        '状态图': ['state', 'statechart', 'uml'],
        '状态机': ['state', 'statechart', 'uml'],
    }
    result = []
    for kw in keywords:
        if kw in keyword_map:
            result.extend(keyword_map[kw])
    return result


# ═══════════════════════════════════════════════════════════════════════
# 2026-06-07 重构：MMD→PNG 多后端降级渲染
# ═══════════════════════════════════════════════════════════════════════
# 背景：v12 中"类图/顺序图/活动图"3 张 UML 图消失。
# 根因：uml-generator.js 写出 .mmd 源文件但 mmdc 渲染需要 Puppeteer Chrome，
#       而 Chrome 在本环境未安装 → _insert_diagrams 找不到 PNG 只能插入
#       "不涉及"占位文字。
#
# 设计原则（通用化，无硬编码）：
# 1. 后端优先级：mmdc（Mermaid CLI）→ Graphviz via uml-renderer.py
# 2. 后端 1（mmdc）：Mermaid 原生支持，渲染效果最佳
# 3. 后端 2（Graphviz via uml-renderer.py）：Mermaid 不可用时的可靠降级
#    - 通过 subprocess 调用 scripts/uml-renderer.py
#    - 自动推断 UML 类型（class/sequence/activity）从 MMD 内容
#    - 自动探测系统常见 Graphviz 安装位置
# 4. 两个后端都失败时，记录日志并回退到 _insert_diagrams 的"不涉及"占位

_GRAPHVIZ_CANDIDATE_PATHS = [
    r'C:\Program Files\Graphviz\bin',
    r'C:\Program Files (x86)\Graphviz\bin',
    r'F:\Program Files\Graphviz\bin',
    r'F:\Program Files (x86)\Graphviz\bin',
    r'D:\Program Files\Graphviz\bin',
    r'C:\graphviz\bin',
    os.path.expanduser(r'~\scoop\apps\graphviz\current\bin'),
    '/usr/bin', '/usr/local/bin', '/opt/homebrew/bin',
    '/usr/local/Cellar/graphviz/*/bin',
]


def _ensure_graphviz_on_path():
    """通用化：将常见 Graphviz 安装位置加入 PATH。返回 dot.exe 路径或 None。

    2026-06-07 新增：自动探测而非硬编码具体路径，适配多平台（Windows/macOS/Linux）。
    已探测过的位置会被模块级缓存避免重复 IO。
    """
    cache_key = '_bemp_graphviz_dot_path'
    if hasattr(_ensure_graphviz_on_path, cache_key):
        return getattr(_ensure_graphviz_on_path, cache_key)
    from shutil import which
    existing = which('dot')
    if existing:
        setattr(_ensure_graphviz_on_path, cache_key, existing)
        return existing
    # 尝试将候选路径加入 PATH
    cur_path = os.environ.get('PATH', '')
    parts = cur_path.split(os.pathsep) if cur_path else []
    for cand in _GRAPHVIZ_CANDIDATE_PATHS:
        if '*' in cand:
            import glob
            cand_glob = sorted(glob.glob(cand), reverse=True)
            for g in cand_glob:
                if g not in parts:
                    parts.insert(0, g)
        else:
            if os.path.isdir(cand) and cand not in parts:
                parts.insert(0, cand)
    os.environ['PATH'] = os.pathsep.join(parts)
    result = which('dot')
    setattr(_ensure_graphviz_on_path, cache_key, result)
    return result


def _infer_uml_type_from_mmd(mmd_path):
    """通用化：从 .mmd 文件内容推断 UML 类型（class/sequence/activity）。

    2026-06-07 新增：用于 Graphviz 后端调用 uml-renderer.py 时选择正确的
    diagram_type 参数。匹配规则按 Mermaid 语法特性（首行关键字），
    无匹配时回退到"类图"（最常见）。

    Returns:
        str: 'class' / 'sequence' / 'activity'
    """
    try:
        with open(mmd_path, 'r', encoding='utf-8') as f:
            head = f.read(4096)
    except Exception:
        return 'class'
    head_l = head.lstrip().lower()
    if head_l.startswith('classdiagram') or 'classDiagram' in head[:200]:
        return 'class'
    if head_l.startswith('sequencediagram') or 'sequenceDiagram' in head[:200]:
        return 'sequence'
    if (head_l.startswith('flowchart') or head_l.startswith('graph')
            or 'flowchart' in head[:200] or '-->|' in head[:1000]):
        return 'activity'
    return 'class'


def _render_mmd_via_graphviz_fallback(mmd_path, png_path):
    """通用化：通过 Graphviz + scripts/uml-renderer.py 渲染 MMD → PNG。

    2026-06-07 新增：作为 mmdc 后端的降级方案。调用本目录下
    scripts/uml-renderer.py 的 CLI 入口（render_uml_auto）：
      python uml-renderer.py <type> <output.png> [business_module]

    通用化要点：
    - 不依赖 MMD 解析（uml-renderer.py 接受结构化数据生成）
    - 自动探测脚本路径（__file__ 同目录的 uml-renderer.py）
    - 自动设置 Graphviz PATH
    - 业务模块名取自设计数据，缺失时使用 MMD 文件名

    Returns:
        bool: 是否成功
    """
    dot_exe = _ensure_graphviz_on_path()
    if not dot_exe:
        return False
    renderer_py = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'uml-renderer.py')
    if not os.path.isfile(renderer_py):
        return False
    uml_type = _infer_uml_type_from_mmd(mmd_path)
    # 业务模块名 = MMD 文件名去后缀（避免硬编码具体业务）
    business_module = os.path.splitext(os.path.basename(mmd_path))[0]
    # 去掉后缀里常见的 -diagram / -mmd 前缀
    for prefix in ('-diagram', '_diagram', '-class', '-sequence', '-activity'):
        if business_module.endswith(prefix):
            business_module = business_module[:-len(prefix)]
            break
    try:
        import subprocess
        cmd = ['python', renderer_py, uml_type, png_path, business_module]
        r = subprocess.run(cmd, capture_output=True, timeout=30,
                           encoding='utf-8', errors='ignore')
        if r.returncode == 0 and os.path.exists(png_path) \
                and os.path.getsize(png_path) > 5 * 1024:
            return True
        # 记录 stderr 帮助诊断（截前 200 字符）
        err_msg = (r.stderr or r.stdout or '').strip()[:200]
        print(f'[WARN] _render_mmd_via_graphviz_fallback: {os.path.basename(mmd_path)} 失败: {err_msg}',
              file=sys.stderr)
        return False
    except Exception as e:
        print(f'[WARN] _render_mmd_via_graphviz_fallback: {os.path.basename(mmd_path)} 异常: {e}',
              file=sys.stderr)
        return False


def _try_install_puppeteer_chrome():
    """通用化降级：尝试用 npm 安装 Puppeteer Chrome 供 mmdc 使用。

    2026-06-07 新增：当 mmdc 因 Chrome 缺失而失败时，可触发
    `npx puppeteer browsers install chrome-headless-shell` 让 mmdc 复活。
    失败时静默返回（不影响主流程）。
    """
    import subprocess
    from shutil import which
    npx = which('npx') or which('npx.cmd')
    if not npx:
        return False
    try:
        r = subprocess.run(
            [npx, '-y', 'puppeteer', 'browsers', 'install', 'chrome-headless-shell'],
            capture_output=True, timeout=180, encoding='utf-8', errors='ignore',
        )
        return r.returncode == 0
    except Exception:
        return False


# 2026-06-07 新增：用于"无 .mmd 源文件时"直接合成 UML PNG 的关键字
# 通用化：使用文件名前缀匹配（不硬编码"承兑行"等具体业务）
_UML_PNG_SYNTH_KEYWORDS = [
    ('class', 'class-diagram'),
    ('sequence', 'sequence-diagram'),
    ('activity', 'activity-diagram'),
]


def _synthesize_missing_uml_pngs(diagram_dir):
    """通用化兜底：diagram_dir 中缺少 UML PNG（类图/顺序图/活动图）时，
    直接调用 uml-renderer.py 从业务模块名合成。

    2026-06-07 新增：解决 v12 中"顺序图/活动图"两张图消失的根因——
    EnhancedUmlService 流程失败时，.mmd 与 .png 都未生成，本函数
    作为最终兜底：扫描 diagram_dir 下若缺少 class/sequence/activity
    任一 PNG，调用 scripts/uml-renderer.py 重新合成。

    Returns:
        int: 合成的 PNG 数
    """
    if not diagram_dir or not os.path.isdir(diagram_dir):
        return 0
    if not _ensure_graphviz_on_path():
        return 0
    renderer_py = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'uml-renderer.py')
    if not os.path.isfile(renderer_py):
        return 0
    import subprocess
    synthesized = 0
    # 业务模块名 = diagram_dir 的父目录名（通用约定：./output/diagrams -> "output"）
    # 或 diagram_dir 同级目录名（更通用）；失败时回退到 "业务模块"
    parent_name = os.path.basename(os.path.dirname(diagram_dir.rstrip(os.sep)))
    business_module = parent_name if parent_name and parent_name not in ('.', 'output') else '业务模块'
    for uml_type, fname_base in _UML_PNG_SYNTH_KEYWORDS:
        target_png = os.path.join(diagram_dir, f'{fname_base}.png')
        if os.path.exists(target_png) and os.path.getsize(target_png) > 5 * 1024:
            continue  # 已有有效 PNG
        try:
            cmd = ['python', renderer_py, uml_type, target_png, business_module]
            r = subprocess.run(cmd, capture_output=True, timeout=30,
                               encoding='utf-8', errors='ignore')
            if r.returncode == 0 and os.path.exists(target_png) \
                    and os.path.getsize(target_png) > 5 * 1024:
                synthesized += 1
                print(f'[INFO] _synthesize_missing_uml_pngs: 已合成 {fname_base}.png '
                      f'(业务模块={business_module})', file=sys.stderr)
        except Exception as e:
            print(f'[WARN] _synthesize_missing_uml_pngs: {fname_base} 异常: {e}', file=sys.stderr)
    return synthesized


def _render_mmd_to_png_if_needed(diagram_dir):
    """通用化预处理：递归扫描 diagram_dir 及其子目录中所有 .mmd 源文件，
    对未对应 .png 的 .mmd 调用后端渲染（mmdc 优先，Graphviz 降级）。

    2026-06-07 重大更新：多后端降级链。
    - 背景：v12 中"类图/顺序图/活动图"3 张 UML 图消失。
    - 根因：mmdc 需要 Puppeteer Chrome，本环境未安装 → 渲染失败 →
      _insert_diagrams 找不到 PNG 只能插入"不涉及"占位文字。
    - 修复：先尝试 mmdc（最佳质量），失败/不可用时自动降级到
      scripts/uml-renderer.py（Graphviz dot，纯 Python + 通用 DOT 渲染）。
      仍失败则回退到"不涉及"占位（由 _insert_diagrams 处理）。

    设计原则：
    - 通用化：递归扫描所有子目录，匹配任意 .mmd 文件
    - 幂等：若 .png 已存在则跳过
    - 降级链：mmdc → Graphviz → 失败（不中断主流程）
    - 自动 PATH 探测：常见 Graphviz 安装位置（Windows/macOS/Linux）

    Args:
        diagram_dir: 图表目录根路径

    Returns:
        dict: {rendered: 成功渲染数, skipped: 跳过数, failed: 失败数, backend_used: 使用的后端}
    """
    if not diagram_dir or not os.path.isdir(diagram_dir):
        return {'rendered': 0, 'skipped': 0, 'failed': 0, 'backend_used': None}
    mmd_files = []
    for root, _dirs, files in os.walk(diagram_dir):
        for f in files:
            if f.lower().endswith('.mmd'):
                mmd_files.append(os.path.join(root, f))
    # 确保 Graphviz 在 PATH（为后端 2 准备）
    _ensure_graphviz_on_path()
    rendered = 0
    skipped = 0
    failed = 0
    backend_used = None
    import subprocess
    from shutil import which
    for mmd in mmd_files:
        png = os.path.splitext(mmd)[0] + '.png'
        if os.path.exists(png) and os.path.getsize(png) > 5 * 1024:
            skipped += 1
            continue
        # 后端 1：mmdc
        ok, err = False, ''
        mmdc_path = which('mmdc') or which('npx.cmd') or which('npx')
        if mmdc_path:
            try:
                cmd = [mmdc_path, '-i', mmd, '-o', png, '-w', '1400', '-H', '900', '-b', 'white', '-s', '2']
                if mmdc_path.lower().endswith('npx') or 'npx' in os.path.basename(mmdc_path):
                    cmd = [mmdc_path, '-y', '@mermaid-js/mermaid-cli', '-i', mmd, '-o', png,
                           '-w', '1400', '-H', '900', '-b', 'white', '-s', '2']
                r = subprocess.run(cmd, capture_output=True, timeout=60,
                                   encoding='utf-8', errors='ignore')
                if r.returncode == 0 and os.path.exists(png) \
                        and os.path.getsize(png) > 5 * 1024:
                    ok = True
                    backend_used = 'mmdc'
                else:
                    err = (r.stderr or r.stdout or '').strip()[:200]
            except Exception as e:
                err = str(e)
        # 后端 2：Graphviz via uml-renderer.py
        if not ok:
            if _render_mmd_via_graphviz_fallback(mmd, png):
                ok = True
                backend_used = 'graphviz'
        if ok:
            rendered += 1
            print(f'[INFO] _render_mmd_to_png_if_needed: {os.path.basename(mmd)} -> {os.path.basename(png)} OK '
                  f'(backend={backend_used})', file=sys.stderr)
        else:
            failed += 1
            warn = err or 'mmdc 与 graphviz 两条路径均失败'
            print(f'[WARN] _render_mmd_to_png_if_needed: {os.path.basename(mmd)} 渲染失败: {warn} '
                  f'（将回退到"不涉及"占位）', file=sys.stderr)
    # 2026-06-07 新增：兜底合成——diagram_dir 中无 .mmd 但缺少 UML PNG 时直接合成
    # 触发场景：EnhancedUmlService 完全失败时（既无 .mmd 也无 .png）
    # 注意：此分支独立于 mmd_files 列表，即使无任何 .mmd 也会执行
    synth_count = _synthesize_missing_uml_pngs(diagram_dir)
    if synth_count:
        rendered += synth_count
        backend_used = backend_used or 'graphviz-synth'
        print(f'[INFO] _render_mmd_to_png_if_needed: 兜底合成 UML PNG +{synth_count}', file=sys.stderr)
    if rendered or failed:
        print(f'[INFO] _render_mmd_to_png_if_needed: 渲染={rendered}, 跳过={skipped}, 失败={failed}, '
              f'后端={backend_used}', file=sys.stderr)
    return {'rendered': rendered, 'skipped': skipped, 'failed': failed, 'backend_used': backend_used}


def _insert_diagrams(doc, diagram_dir):
    """在文档中插入架构图、网络拓扑图、部署图、UML图（类图/顺序图/活动图）

    查找"系统架构"/"组件结构图"/"部署图"/"类图"/"顺序图"/"活动图"等标题，
    在标题下插入对应图片。支持在 diagram_dir 子目录中查找图片文件。
    支持多种图片格式：.png, .jpg, .jpeg。
    插入前先清除标题下的旧图片元素。

    2026-06-04 优化：
    - 支持 .jpg/.jpeg 格式（不仅限于 .png）
    - 查找图片时同时搜索多种扩展名，优先使用 .png
    - 标题匹配时也去除编号前缀（如"4.7 类图"）
    - 增加 ER 图、数据流图等常见图表类型

    2026-06-06 增强：
    - 对所有命中的图表标题（包括未找到图片的），统一在标题下插入"不涉及"占位
    - 保证"类图"/"顺序图"/"活动图"等 UML 标题下都有明确内容，
      避免出现"标题存在但下方无内容"的空白章节

    2026-06-07 增强：
    - 通用化：在扫描图片前先调用 _render_mmd_to_png_if_needed 把目录下
      所有 .mmd 源文件渲染成 .png，保证 UML 图（类图/顺序图/活动图）
      在文档中正确显示而非显示占位文字。
    """
    # 日志：确认图表目录和文件是否存在
    print(f'[INFO] _insert_diagrams: diagram_dir={diagram_dir}', file=sys.stderr)

    # 2026-06-07 新增：通用化预处理——把 diagram_dir 中所有 .mmd 渲染为 .png
    # 解决 v12 中"类图/顺序图/活动图"消失（仅有 mmd 源文件未渲染）的问题
    if os.path.isdir(diagram_dir):
        try:
            _render_mmd_to_png_if_needed(diagram_dir)
        except Exception as e:
            print(f'[WARN] _render_mmd_to_png_if_needed 失败: {e}', file=sys.stderr)
    if not os.path.isdir(diagram_dir):
        print(f'[WARN] _insert_diagrams: 目录不存在: {diagram_dir}', file=sys.stderr)
        # 2026-06-06 增强：目录不存在时仍扫描所有图表标题并插入"不涉及"占位
        _insert_placeholder_for_unmatched_diagrams(doc, all_keywords_set=[kw for kws, _ in [
            (['系统架构', '组件结构图', '架构图'], None),
            (['网络结构图', '网络拓扑', '网络图'], None),
            (['部署图', '部署架构'], None),
            (['类图', '类图设计'], None),
            (['顺序图', '时序图', '序列图'], None),
            (['活动图', '流程图'], None),
        ] for kw in kws])
        return

    # 递归收集 diagram_dir 及其子目录中的所有图片文件（支持 .png/.jpg/.jpeg）
    SUPPORTED_EXTS = ('.png', '.jpg', '.jpeg')
    all_images = {}  # filename → filepath，同文件名优先 .png
    for root, _dirs, files in os.walk(diagram_dir):
        for f in files:
            f_lower = f.lower()
            if f_lower.endswith(SUPPORTED_EXTS):
                # 同名文件优先保留 .png（后处理的覆盖先处理的）
                all_images[f] = os.path.join(root, f)
    print(f'[INFO] _insert_diagrams: 找到 {len(all_images)} 个图片文件: {list(all_images.keys())}', file=sys.stderr)

    # 标题关键词 → 对应图片文件名的映射（含 UML 图表）
    # 每个映射项支持多个候选文件名（按优先级排序，找到第一个即用）
    diagram_map = [
        (['系统架构', '组件结构图', '架构图'], ['architecture-diagram.png', 'architecture-diagram.jpg']),
        (['网络结构图', '网络拓扑', '网络图'], ['network-topology.png', 'network-topology.jpg']),
        (['部署图', '部署架构'], ['deployment-diagram.png', 'deployment-diagram.jpg']),
        (['类图', '类图设计'], ['class-diagram.png', 'class-diagram.jpg']),
        (['顺序图', '时序图', '序列图'], ['sequence-diagram.png', 'sequence-diagram.jpg']),
        (['活动图', '流程图'], ['activity-diagram.png', 'activity-diagram.jpg']),
        (['ER图', 'E-R图', '实体关系'], ['er-diagram.png', 'er-diagram.jpg']),
        (['数据流图', '数据流'], ['data-flow-diagram.png', 'data-flow-diagram.jpg']),
        (['状态图', '状态机'], ['state-diagram.png', 'state-diagram.jpg']),
    ]

    inserted_count = 0
    # 记录已使用的图片文件路径，避免同一张图重复插入
    used_images = set()
    # 记录已处理的标题（防止同标题下多张图重复插入）
    processed_titles = set()
    # 记录所有命中过的标题（用于"不涉及"占位的反查）
    matched_titles = []

    for keywords, filenames in diagram_map:
        # 按优先级查找图片文件（精确文件名匹配）
        img_path = None
        matched_filename = None
        for filename in filenames:
            if filename in all_images:
                img_path = all_images[filename]
                matched_filename = filename
                break

        # 2026-06-05 修复：精确文件名匹配失败时，回退到关键字模糊匹配
        # 实际图片目录中的文件名可能不符合预期命名规范（如 class_diagram.png 而非 class-diagram.png）
        if not img_path:
            # 从keywords中提取用于文件名匹配的关键词（英文）
            keyword_patterns = _get_diagram_filename_keywords(keywords)
            for fname, fpath in all_images.items():
                if fpath in used_images:
                    continue
                fname_lower = fname.lower()
                if any(kw in fname_lower for kw in keyword_patterns):
                    img_path = fpath
                    matched_filename = fname
                    print(f'[INFO] _insert_diagrams: 模糊匹配 {fname} → {keywords}', file=sys.stderr)
                    break

        if not img_path or not os.path.exists(img_path):
            # 仅对必要图表输出警告
            if any(fn in ('architecture-diagram.png', 'network-topology.png', 'deployment-diagram.png')
                   for fn in filenames):
                print(f'[WARN] _insert_diagrams: 必要图表缺失: {filenames}', file=sys.stderr)
            # 2026-06-06 增强：缺失图片时仍记录待插入"不涉及"占位的关键词
            # 后续 _insert_placeholder_for_unmatched_diagrams 会处理
            continue

        used_images.add(img_path)

        # 查找匹配关键词的标题段落（同时支持带编号前缀的标题）
        target_para = None
        for p in doc.paragraphs:
            if p.style and p.style.name.startswith('Heading'):
                p_text = p.text.strip()
                # 去除编号前缀后比较，支持"4.7 类图"等格式
                p_text_clean = re.sub(r'^\d+(\.\d+)*\s*', '', p_text)
                if any(kw in p_text or kw in p_text_clean for kw in keywords):
                    if p_text in processed_titles:
                        continue
                    target_para = p
                    processed_titles.add(p_text)
                    matched_titles.append((target_para, keywords))
                    break

        if target_para is None:
            print(f'[WARN] _insert_diagrams: 未找到匹配标题 {keywords}，跳过图片 {matched_filename}', file=sys.stderr)
            continue

        # 清除标题下的旧图片
        _clear_images_between_headings(target_para)

        # 插入新图片
        try:
            img_p = doc.add_paragraph()
            img_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = img_p.add_run()
            run.add_picture(img_path, width=Inches(6))
            target_para._element.addnext(img_p._element)
            inserted_count += 1
            print(f'[INFO] _insert_diagrams: 已插入 {matched_filename} 到标题 "{target_para.text.strip()}"', file=sys.stderr)
        except Exception as e:
            print(f'[WARN] 插入图表 {matched_filename} 失败: {e}', file=sys.stderr)

    print(f'[INFO] _insert_diagrams: 共插入 {inserted_count} 张图片', file=sys.stderr)

    # 2026-06-06 增强：扫描所有未匹配到图片的图表标题，插入"不涉及"占位
    # 这样保证"类图"/"顺序图"/"活动图"等标题下都有明确说明，避免空白章节
    all_diagram_keywords = [kw for kws, _ in diagram_map for kw in kws]
    _insert_placeholder_for_unmatched_diagrams(doc, all_keywords_set=all_diagram_keywords,
                                                processed_titles=processed_titles)


def _insert_placeholder_for_unmatched_diagrams(doc, all_keywords_set, processed_titles=None):
    """扫描文档中所有图表标题（类图/顺序图/活动图/部署图等），
    对未插入图片的标题在下方插入"不涉及"占位段落。

    2026-06-06 新增：解决"类图标题下空"或"顺序图无图"时无任何内容的问题。
    通过 doc_rules.yaml 中的 heading_style_unify.h2_keywords 也可启用。

    Args:
        doc: docx Document 对象
        all_keywords_set: 图表关键词集合（如 {'类图', '顺序图', '活动图'}）
        processed_titles: 已有图片的标题集合（这些跳过）
    """
    if processed_titles is None:
        processed_titles = set()
    if not all_keywords_set:
        return
    inserted = 0
    for p in doc.paragraphs:
        if not (p.style and p.style.name.startswith('Heading')):
            continue
        p_text = p.text.strip()
        if p_text in processed_titles:
            continue
        # 去除编号前缀
        p_text_clean = re.sub(r'^\d+(\.\d+)*\s*', '', p_text)
        if not any(kw in p_text or kw in p_text_clean for kw in all_keywords_set):
            continue
        # 标题与下一个标题之间是否已有图片（drawing/pict）
        if _has_image_between_headings(p):
            continue
        # 标题下是否已有非空内容（防止重复插入）
        if _has_content_after_heading(p, min_paragraphs=1):
            continue
        # 插入"不涉及"占位
        placeholder_p = doc.add_paragraph()
        placeholder_run = placeholder_p.add_run('不涉及')
        placeholder_run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
        p._element.addnext(placeholder_p._element)
        inserted += 1
        processed_titles.add(p_text)
    if inserted > 0:
        print(f'[INFO] _insert_placeholder_for_unmatched_diagrams: 插入"不涉及"占位 {inserted} 处', file=sys.stderr)
    return inserted


def _has_image_between_headings(heading_para):
    """判断标题与下一个标题之间是否已含图片元素（w:drawing / w:pict）"""
    heading_elem = heading_para._element
    parent = heading_elem.getparent()
    if parent is None:
        return False
    found = False
    for elem in list(parent):
        if elem is heading_elem:
            found = True
            continue
        if not found:
            continue
        # 遇到下一个标题则停止
        if elem.tag.endswith('}p'):
            pPr = elem.find('.//' + qn('w:pPr') + '/' + qn('w:pStyle'))
            if pPr is not None and pPr.get(qn('w:val'), '').startswith('Heading'):
                break
            if elem.findall('.//' + qn('w:drawing')) or elem.findall('.//' + qn('w:pict')):
                return True
    return False


def _has_content_after_heading(heading_para, min_paragraphs=1):
    """判断标题下是否已有非空内容（防止"不涉及"占位重复插入）"""
    heading_elem = heading_para._element
    parent = heading_elem.getparent()
    if parent is None:
        return False
    found = False
    non_empty_count = 0
    for elem in list(parent):
        if elem is heading_elem:
            found = True
            continue
        if not found:
            continue
        # 遇到下一个标题则停止
        if elem.tag.endswith('}p'):
            pPr = elem.find('.//' + qn('w:pPr') + '/' + qn('w:pStyle'))
            if pPr is not None and pPr.get(qn('w:val'), '').startswith('Heading'):
                break
            # 跳过当前 heading 本身
            t_elems = elem.findall('.//' + qn('w:t'))
            text = ''.join(t.text or '' for t in t_elems).strip()
            if text and text != '不涉及':
                non_empty_count += 1
                if non_empty_count >= min_paragraphs:
                    return True
    return False


def _clear_images_between_headings(heading_para):
    """清除标题段落与下一个标题之间的旧图片元素（w:drawing 和 w:pict）"""
    heading_elem = heading_para._element
    parent = heading_elem.getparent()
    if parent is None:
        return

    heading_style_vals = {
        '1', '2', '3', '4', '5', '6',
        'Heading1', 'Heading2', 'Heading3', 'Heading4', 'Heading5', 'Heading6',
        'heading1', 'heading2', 'heading3', 'heading4', 'heading5', 'heading6',
        'heading 1', 'heading 2', 'heading 3', 'heading 4', 'heading 5', 'heading 6',
        'Heading 1', 'Heading 2', 'Heading 3', 'Heading 4', 'Heading 5', 'Heading 6',
    }

    found_heading = False
    to_remove = []
    for elem in list(parent):
        if elem is heading_elem:
            found_heading = True
            continue
        if not found_heading:
            continue
        # 遇到下一个标题则停止
        if elem.tag.endswith('}p'):
            pPr = elem.find('.//' + qn('w:pPr') + '/' + qn('w:pStyle'))
            if pPr is not None:
                style_val = pPr.get(qn('w:val'), '')
                if style_val in heading_style_vals:
                    break
            # 检查段落中是否包含 w:drawing 或 w:pict（即图片）
            if elem.findall('.//' + qn('w:drawing')) or elem.findall('.//' + qn('w:pict')):
                to_remove.append(elem)

    for elem in to_remove:
        parent.remove(elem)


def _update_revision_table(doc, module_name, revision_history):
    """更新修订记录表格的第一行

    如果文档中存在修订记录表格，更新第一行数据行的版本号、日期、作者和描述。
    """
    for table in doc.tables:
        # 通过表头判断是否为修订记录表
        first_row_text = ''
        if table.rows:
            for cell in table.rows[0].cells:
                first_row_text += cell.text.strip()
        if '版本' not in first_row_text and '修订' not in first_row_text:
            continue

        # 找到数据行（跳过表头）
        if len(table.rows) < 2:
            continue

        data_row = table.rows[1]
        today = datetime.now().strftime('%Y-%m-%d')
        author = revision_history.get('author', '') if isinstance(revision_history, dict) else ''
        description = revision_history.get('description', '') if isinstance(revision_history, dict) else ''

        # 根据列数填充
        cells = data_row.cells
        if len(cells) >= 1:
            _set_cell_text(cells[0], 'V1.0')
        if len(cells) >= 2:
            _set_cell_text(cells[1], today)
        if len(cells) >= 3:
            _set_cell_text(cells[2], author or '开发组')
        if len(cells) >= 4:
            _set_cell_text(cells[3], description or f'{module_name}详细设计文档初稿')
        break


def _set_cell_text(cell, text):
    """设置单元格文本（保留格式）"""
    for para in cell.paragraphs:
        for r in para.runs:
            r.text = ''
    if cell.paragraphs and cell.paragraphs[0].runs:
        cell.paragraphs[0].runs[0].text = text
    elif cell.paragraphs:
        cell.paragraphs[0].add_run(text)


# 注意：_inject_update_fields 委托给 doc_formatter.inject_update_fields（见文件末尾）


def _fill_glossary_table(doc, heading_para, glossary):
    """填充术语定义表格（glossary 为空时生成通用默认术语）"""
    if not glossary:
        glossary = [
            {'term': 'API', 'definition': 'Application Programming Interface，应用程序编程接口'},
            {'term': 'CRUD', 'definition': 'Create/Read/Update/Delete，增删改查操作'},
            {'term': 'DAO', 'definition': 'Data Access Object，数据访问对象'},
            {'term': 'DTO', 'definition': 'Data Transfer Object，数据传输对象'},
            {'term': 'RPC', 'definition': 'Remote Procedure Call，远程过程调用'},
        ]
    headers = ['术语/缩写', '全称', '说明']
    rows = []
    for item in glossary:
        term = item.get('term', '')
        definition = item.get('definition', '')
        rows.append([term, '', definition])
    if rows:
        _add_table_after_paragraph(heading_para, headers, rows)


def _fill_references(doc, heading_para, references):
    """填充参考资料（references 为空时生成通用默认参考文档）"""
    if not references:
        references = [
            '《需求规格说明书》',
            '《概要设计说明书》',
            '《系统接口规范》',
            '《数据库设计文档》',
        ]
    for i, ref in enumerate(references, 1):
        new_p = _insert_paragraph_after(heading_para, f'[{i}] {ref}')
        set_black(new_p)
        heading_para = new_p


def _fill_dev_environment(doc, heading_para, design_data):
    """填充开发环境信息"""
    # 优先从JSON配置读取，未配置时使用通用默认值
    default_env = [
        '开发语言：Java',
        '框架：Spring Boot',
        '前端框架：Vue.js',
        '数据库：Oracle / MySQL',
        '构建工具：Maven',
        '版本管理：Git',
    ]
    env_items = design_data.get('devEnvironment', default_env)
    insert_after = heading_para
    for item in env_items:
        new_p = _insert_paragraph_after(insert_after, item)
        set_black(new_p)
        insert_after = new_p


def _fill_ui_description(doc, heading_para, design_data):
    """填充界面描述"""
    # 从chapters中提取界面相关信息
    chapters = design_data.get('chapters', [])
    ui_info = ''
    for ch in chapters:
        sections = ch.get('sections', [])
        for sec in sections:
            content = sec.get('content', {})
            if isinstance(content, dict):
                desc = content.get('description', '')
                if '界面' in desc or '菜单位置' in desc or '查询条件' in desc:
                    ui_info += desc + '\n'

    if ui_info:
        new_p = _insert_paragraph_after(heading_para, ui_info.strip())
        set_black(new_p)
    else:
        # 2026-06-04 优化：占位文本保持一定长度（>40字符）以触发 _format_ui_paragraphs 软换行
        placeholder = '界面设计详见需求文档中的界面原型图；界面布局主要包括查询区、列表区、操作区三个部分，菜单位置与查询条件以原型图为准。'
        new_p = _insert_paragraph_after(heading_para, placeholder)
        set_black(new_p)


def _fill_performance(doc, heading_para, design_data):
    """填充性能要求"""
    # 优先从JSON配置读取，未配置时使用通用占位
    default_perf = [
        '核心操作响应时间：< 500ms',
        '列表查询响应时间：< 1s',
        '并发用户数：支持50+并发',
        '数据量：支持10万+记录查询',
    ]
    perf_items = design_data.get('performance', default_perf)
    insert_after = heading_para
    for item in perf_items:
        new_p = _insert_paragraph_after(insert_after, item)
        set_black(new_p)
        insert_after = new_p


def _fill_purpose(doc, heading_para, chapter_map, module_name='本项目'):
    """填充目的章节，从1.1业务背景提取。检测模板占位文本并替换为有意义的内容。"""
    TEMPLATE_PLACEHOLDER_PREFIX = '描述当前业务痛点或需求来源'
    for key in ['1.1 业务背景', '业务背景']:
        if key in chapter_map:
            ch = chapter_map[key]
            content = ch.get('content', {})
            if isinstance(content, dict) and content.get('description'):
                desc = content['description']
                # 检测模板占位文本，替换为基于模块名的有意义的默认内容
                if desc.startswith(TEMPLATE_PLACEHOLDER_PREFIX):
                    desc = (f'{module_name}旨在实现相关业务功能，'
                            f'满足业务管理子系统的管控需求，'
                            f'确保数据一致性和操作完整性。')
                new_p = _insert_paragraph_after(heading_para, desc)
                set_black(new_p)
                return
    # 默认内容
    new_p = _insert_paragraph_after(heading_para, f'本文档详细描述{module_name}的设计方案，为开发、测试和维护提供技术依据。')
    set_black(new_p)


def _fill_scope(doc, heading_para, chapter_map, module_name='本项目'):
    """填充适用范围章节，从1.3范围说明提取"""
    for key in ['1.3 范围说明', '范围说明']:
        if key in chapter_map:
            ch = chapter_map[key]
            content = ch.get('content', {})
            if isinstance(content, dict):
                if content.get('headers') and content.get('rows'):
                    _add_table_after_paragraph(heading_para, content['headers'], content['rows'])
                    return
                if content.get('description'):
                    new_p = _insert_paragraph_after(heading_para, content['description'])
                    set_black(new_p)
                    return
    # 默认内容
    new_p = _insert_paragraph_after(heading_para, f'本文档适用于{module_name}的设计与开发。')
    set_black(new_p)


def _fill_reader_audience(doc, heading_para, module_name='本项目'):
    """填充读者对象章节（详细设计模板专用）

    2026-06-04 新增：详细设计模板中"读者对象"章节的默认内容。
    """
    default_readers = [
        f'{module_name}的开发人员',
        f'{module_name}的测试人员',
        f'{module_name}的项目经理',
        '系统架构师',
    ]
    for reader in default_readers:
        new_p = _insert_paragraph_after(heading_para, reader)
        set_black(new_p)
        heading_para = new_p


def _fill_component_module_list(doc, heading_para, chapters, design_data):
    """填充"组件内部的模块列表及说明"章节

    从 design_data 的 chapters 和 businessSubmodules 中提取模块列表，
    生成模块汇总表格。若无可提取数据则插入"不涉及"占位。
    """
    module_name = design_data.get('moduleName', '')
    business_submodules = design_data.get('businessSubmodules') or []

    # 收集模块条目：优先从 businessSubmodules 提取，其次从 chapters 提取
    rows = []
    if business_submodules:
        for idx, sub in enumerate(business_submodules, 1):
            if isinstance(sub, dict):
                name = sub.get('name') or sub.get('title') or f'子模块{idx}'
                desc = sub.get('description') or ''
                if isinstance(desc, dict):
                    desc = desc.get('description', '')
            elif isinstance(sub, str):
                name = sub
                desc = ''
            else:
                continue
            rows.append([str(idx), name, desc or f'{name}功能模块'])
    elif chapters:
        for idx, ch in enumerate(chapters, 1):
            title = ch.get('title', f'模块{idx}')
            desc = ''
            content = ch.get('content', {})
            if isinstance(content, dict):
                desc = content.get('description', '')
            rows.append([str(idx), title, desc or f'{title}功能模块'])

    if rows:
        headers = ['序号', '模块名称', '说明']
        _add_table_after_paragraph(heading_para, headers, rows)
    else:
        new_p = _insert_paragraph_after(heading_para, f'{module_name}组件内部的模块列表详见各模块设计说明。')
        set_black(new_p)


def _fill_business_submodules(doc, heading_para, design_data):
    """在"模块1设计说明"下插入业务子模块（H2级别标题+描述）

    从 design_data 的 businessSubmodules 中提取业务子模块，
    在标题后依次插入 H2 级别的子模块标题和简要描述。
    不清除模板中已有的 H2 子标题（功能描述/界面/性能等），子模块插入在它们之前。
    """
    business_submodules = design_data.get('businessSubmodules') or []
    if not business_submodules:
        return

    insert_after = heading_para
    for sub in business_submodules:
        if isinstance(sub, dict):
            sub_name = sub.get('name') or sub.get('title') or ''
            sub_desc = sub.get('description') or ''
            if isinstance(sub_desc, dict):
                sub_desc = sub_desc.get('description', '')
        elif isinstance(sub, str):
            sub_name = sub
            sub_desc = ''
        else:
            continue

        if not sub_name:
            continue

        # 插入 H2 级别子模块标题
        new_heading = _insert_paragraph_after(insert_after, sub_name, style='Heading2')
        set_black(new_heading)
        insert_after = new_heading

        # 插入子模块描述
        if sub_desc:
            new_p = _insert_paragraph_after(insert_after, sub_desc)
            set_black(new_p)
            insert_after = new_p


def _clear_content_between_headings(heading_para, paragraphs, stop_at_next_h1=True):
    """清理目标标题与下一标题之间的所有非标题段落（模板原始占位内容）

    2026-06-04 优化：
    - 仅清除"目标标题 → 下一标题"区间内的内容，避免越界
    - 表格（w:tbl）也一并删除，因为模板中的占位表格会导致内容窜行
    - DEBUG_LAYOUT=True 时输出每一步清理的标题/段落/表格，便于排查窜行

    2026-06-06 扩展：stop_at_next_h1=False 时，清除 H1 与下一 H1 之间的所有内容
    （包括嵌套的 H2/H3），用于"概述"等需要完全重写的章节（避免模板占位 H2 残留）

    2026-06-07 修复：_PRESERVE_MODE=True 时立即返回，保留模板正文原样不清理。
    """
    if _PRESERVE_MODE:
        return
    heading_element = heading_para._element
    parent = heading_element.getparent()
    if parent is None:
        return

    # Word XML中标题样式的pStyle值：1=Heading1, 2=Heading2, 3=Heading3等
    # 包含多种写法变体，确保正确识别所有级别的标题边界
    heading_style_vals = {
        '1', '2', '3', '4', '5', '6',
        'Heading1', 'Heading2', 'Heading3', 'Heading4', 'Heading5', 'Heading6',
        'heading1', 'heading2', 'heading3', 'heading4', 'heading5', 'heading6',
        'heading 1', 'heading 2', 'heading 3', 'heading 4', 'heading 5', 'heading 6',
        'Heading 1', 'Heading 2', 'Heading 3', 'Heading 4', 'Heading 5', 'Heading 6',
    }
    # 下一 H1 样式（仅 stop_at_next_h1=True 时用于早停）
    h1_only_vals = {
        '1', 'Heading1', 'heading1', 'heading 1', 'Heading 1',
    }

    heading_text = (heading_para.text or '').strip()
    if DEBUG_LAYOUT:
        print(f'[DEBUG] _clear_content_between_headings start: "{heading_text}" (stop_at_next_h1={stop_at_next_h1})', file=sys.stderr)

    # 从标题下一个元素开始，删除直到遇到下一个标题
    to_remove = []
    found_heading = False
    next_heading_text = None
    tbl_count = 0
    for elem in list(parent):
        if elem is heading_element:
            found_heading = True
            continue
        if not found_heading:
            continue
        # 检查是否是标题段落
        if elem.tag.endswith('}p'):
            pPr = elem.find('.//' + qn('w:pPr') + '/' + qn('w:pStyle'))
            if pPr is not None:
                style_val = pPr.get(qn('w:val'), '')
                if style_val in heading_style_vals:
                    if stop_at_next_h1:
                        # 取出下一标题文本以便调试
                        t_elems = elem.findall('.//' + qn('w:t'))
                        next_heading_text = ''.join(t.text or '' for t in t_elems).strip()
                        break  # 遇到下一个标题，停止
                    else:
                        # stop_at_next_h1=False 时，碰到 H1 才停止；嵌套 H2/H3 一并删除
                        if style_val in h1_only_vals:
                            t_elems = elem.findall('.//' + qn('w:t'))
                            next_heading_text = ''.join(t.text or '' for t in t_elems).strip()
                            break
            # 非标题段落（或嵌套标题），标记删除
            to_remove.append(elem)
        elif elem.tag.endswith('}tbl'):
            # 模板中的占位表格也需要删除，否则会导致内容窜行
            # 后续 _fill_chapter_content 会重新插入正确的表格
            to_remove.append(elem)
            tbl_count += 1

    if DEBUG_LAYOUT:
        print(f'[DEBUG]   next heading: "{next_heading_text}", will remove {len(to_remove)} elements (incl. {tbl_count} tables)', file=sys.stderr)

    for elem in to_remove:
        parent.remove(elem)


def _fill_chapter_content(doc, heading_para, heading_idx, chapter_data, paragraphs):
    """在标题后插入章节内容，支持chapter和section两种数据格式"""
    # 先清理模板原始占位内容
    _clear_content_between_headings(heading_para, paragraphs)
    
    sections = chapter_data.get('sections', [])
    body_texts = chapter_data.get('bodyTexts', [])
    tables = chapter_data.get('tables', [])
    content = chapter_data.get('content', None)

    # 当前插入位置参考段落
    insert_after = heading_para

    # 处理section级别的数据（只有content字段，没有sections字段）
    if not sections and content:
        if isinstance(content, str) and content:
            new_p = _insert_paragraph_after(insert_after, content)
            set_black(new_p)
            insert_after = new_p
        elif isinstance(content, dict):
            if content.get('description'):
                new_p = _insert_paragraph_after(insert_after, content['description'])
                set_black(new_p)
                insert_after = new_p
            if content.get('headers') and content.get('rows'):
                _add_table_after_paragraph(insert_after, content['headers'], content['rows'])
        return

    # 处理chapter级别的数据（有sections字段）
    # 插入正文文本
    for bt in body_texts:
        if bt:
            new_p = _insert_paragraph_after(insert_after, bt)
            set_black(new_p)
            insert_after = new_p

    # 插入表格
    for tbl_data in tables:
        headers = tbl_data.get('headers', [])
        rows = tbl_data.get('rows', [])
        if headers and rows:
            _add_table_after_paragraph(insert_after, headers, rows)

    # 插入子节
    for sec in sections:
        sec_title = sec.get('title', '')
        sec_content = sec.get('content', '')

        # 添加子节标题（Heading 2级别）
        new_heading = _insert_paragraph_after(insert_after, sec_title, style='Heading2')
        insert_after = new_heading

        # 添加子节内容
        if isinstance(sec_content, str) and sec_content:
            new_p = _insert_paragraph_after(insert_after, sec_content)
            set_black(new_p)
            insert_after = new_p
        elif isinstance(sec_content, dict):
            if sec_content.get('description'):
                new_p = _insert_paragraph_after(insert_after, sec_content['description'])
                set_black(new_p)
                insert_after = new_p
            if sec_content.get('headers') and sec_content.get('rows'):
                _add_table_after_paragraph(
                    insert_after,
                    sec_content['headers'],
                    sec_content['rows']
                )


def _add_table_after_paragraph(paragraph, headers, rows):
    """在段落后插入表格，返回插入的表格 OXML 元素（供调用方在表格后插入锚点段落）"""
    try:
        parent = paragraph._element.getparent()
        if parent is None:
            return None
        idx = list(parent).index(paragraph._element)

        tbl = OxmlElement('w:tbl')

        # 表格属性
        tblPr = OxmlElement('w:tblPr')
        tblStyle = OxmlElement('w:tblStyle')
        tblStyle.set(qn('w:val'), 'TableGrid')
        tblPr.append(tblStyle)
        tblW = OxmlElement('w:tblW')
        tblW.set(qn('w:w'), '5000')
        tblW.set(qn('w:type'), 'pct')
        tblPr.append(tblW)
        tbl.append(tblPr)

        # 表格网格
        tblGrid = OxmlElement('w:tblGrid')
        for _ in headers:
            gridCol = OxmlElement('w:gridCol')
            gridCol.set(qn('w:w'), str(int(9000 / len(headers))))
            tblGrid.append(gridCol)
        tbl.append(tblGrid)

        # 表头行
        tr_header = OxmlElement('w:tr')
        for h in headers:
            tc = OxmlElement('w:tc')
            p = OxmlElement('w:p')
            r = OxmlElement('w:r')
            rPr = OxmlElement('w:rPr')
            b = OxmlElement('w:b')
            rPr.append(b)
            r.append(rPr)
            t = OxmlElement('w:t')
            t.text = str(h)
            r.append(t)
            p.append(r)
            tc.append(p)
            tr_header.append(tc)
        tbl.append(tr_header)

        # 数据行
        for row in rows:
            tr = OxmlElement('w:tr')
            for cell_val in row:
                tc = OxmlElement('w:tc')
                p = OxmlElement('w:p')
                r = OxmlElement('w:r')
                t = OxmlElement('w:t')
                t.text = str(cell_val) if cell_val is not None else ''
                r.append(t)
                p.append(r)
                tc.append(p)
                tr.append(tc)
            tbl.append(tr)

        parent.insert(idx + 1, tbl)
        return tbl
    except Exception as e:
        print(f'[WARN] 表格插入失败: {e}', file=sys.stderr)
        return None


# ==================== 2026-06-04 新增辅助函数 ====================

def _extract_business_keywords(design_data, module_name):
    """从 design_data 提取业务关键词集合（用于"设计约束"业务相关性判定）

    来源（避免硬编码，按优先级扩展）：
    1. moduleName
    2. 显式字段 businessKeywords
    3. chapter.title / section.title 中的中文名词短语
    4. chapter.content.description / rows 中的中文业务名词（>=2字，<=10字）

    2026-06-06 修复：扩展数据源至 content.rows，避免关键词集合为空导致
    "设计约束"整章被误删（如承兑行额度管理场景下 design_data 无 moduleName，
    仅靠 chapter.title/section.title 提取的关键词集与约束内容无交集）。
    """
    if not isinstance(design_data, dict):
        return set()
    keywords = set()
    if module_name and isinstance(module_name, str):
        keywords.add(module_name)
    explicit = design_data.get('businessKeywords') or []
    for kw in explicit:
        if isinstance(kw, str) and kw.strip():
            keywords.add(kw.strip())
    # 从章节/节标题中提取
    for ch in design_data.get('chapters', []) or []:
        title = ch.get('title', '')
        if title:
            keywords.add(title)
        for sec in ch.get('sections', []) or []:
            t = sec.get('title', '')
            if t:
                keywords.add(t)
            # 2026-06-06 新增：section.content.rows 中的中文短语
            c = sec.get('content', {}) or {}
            if isinstance(c, dict):
                for row in c.get('rows', []) or []:
                    for cell in (row if isinstance(row, (list, tuple)) else [row]):
                        if not isinstance(cell, str):
                            continue
                        cell = cell.strip()
                        if 2 <= len(cell) <= 12 and re.search(r'[\u4e00-\u9fa5]', cell):
                            keywords.add(cell)
        # 2026-06-06 新增：chapter.content.description / rows
        c = ch.get('content', {}) or {}
        if isinstance(c, dict):
            desc = c.get('description')
            if isinstance(desc, str):
                # 仅保留描述中的中文名词短语
                for phrase in re.findall(r'[\u4e00-\u9fa5]{2,10}', desc):
                    keywords.add(phrase)
            for row in c.get('rows', []) or []:
                for cell in (row if isinstance(row, (list, tuple)) else [row]):
                    if not isinstance(cell, str):
                        continue
                    cell = cell.strip()
                    if 2 <= len(cell) <= 12 and re.search(r'[\u4e00-\u9fa5]', cell):
                        keywords.add(cell)
    # 过滤空值
    return {k for k in keywords if k}


def _ensure_design_constraint_placeholder(doc, placeholder='不涉及'):
    """在"设计约束"标题下若内容被清空，插入占位段落

    与 remove_design_constraint_irrelevant 配套使用：
    若业务相关性判定后整段被删除导致章节为空，则补一段"不涉及"。
    """
    if doc is None:
        return False
    paragraphs = list(doc.paragraphs)
    target_idx = None
    for i, p in enumerate(paragraphs):
        if not p.style or not p.style.name.startswith('Heading'):
            continue
        text = (p.text or '').strip()
        if text == '设计约束' or (text and '设计约束' in text and len(text) <= 10):
            target_idx = i
            break
    if target_idx is None:
        return False
    # 找到下一个标题
    next_idx = None
    for j in range(target_idx + 1, len(paragraphs)):
        if paragraphs[j].style and paragraphs[j].style.name.startswith('Heading'):
            next_idx = j
            break
    end_idx = next_idx if next_idx is not None else len(paragraphs)
    # 检查区间是否已有非空内容
    for j in range(target_idx + 1, end_idx):
        t = (paragraphs[j].text or '').strip()
        if t:
            return False
    # 区间为空 → 插入占位段落
    target_p = paragraphs[target_idx]
    parent = target_p._element.getparent()
    if parent is None:
        return False
    new_p = OxmlElement('w:p')
    r = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    color = OxmlElement('w:color')
    color.set(qn('w:val'), '000000')
    rPr.append(color)
    r.append(rPr)
    t_elem = OxmlElement('w:t')
    t_elem.set(qn('xml:space'), 'preserve')
    t_elem.text = placeholder
    r.append(t_elem)
    new_p.append(r)
    target_p._element.addnext(new_p)
    return True


def _ensure_error_code_section(doc, design_data):
    """确保错误码章节有内容：无则插入"不涉及"+ 示例错误码表"""
    if doc is None:
        return
    paragraphs = list(doc.paragraphs)
    target_idx = None
    for i, p in enumerate(paragraphs):
        if not p.style or not p.style.name.startswith('Heading'):
            continue
        text = (p.text or '').strip()
        # 兼容 "1.5 错误码"、"错误码定义"、"错误码" 等多种写法
        if text == '错误码' or '错误码定义' in text or text.endswith('错误码') or text.endswith('错误码定义'):
            target_idx = i
            break
    if target_idx is None:
        return
    # 找到下一个标题索引
    next_idx = None
    for j in range(target_idx + 1, len(paragraphs)):
        if paragraphs[j].style and paragraphs[j].style.name.startswith('Heading'):
            next_idx = j
            break
    end_idx = next_idx if next_idx is not None else len(paragraphs)
    # 检查区间内是否已有内容
    has_content = False
    for j in range(target_idx + 1, end_idx):
        t = (paragraphs[j].text or '').strip()
        if t:
            has_content = True
            break
    if has_content:
        return
    # 取 JSON 错误码定义
    error_codes = design_data.get('errorCodes') if isinstance(design_data, dict) else None
    if not error_codes:
        error_codes = [{'code': c, 'desc': d} for c, d in DEFAULT_ERROR_CODES]
    # 插入占位说明 + 表格
    target_p = paragraphs[target_idx]
    parent = target_p._element.getparent()
    if parent is None:
        return
    # 占位段落
    note = OxmlElement('w:p')
    r = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    color = OxmlElement('w:color')
    color.set(qn('w:val'), '000000')
    rPr.append(color)
    r.append(rPr)
    t = OxmlElement('w:t')
    t.set(qn('xml:space'), 'preserve')
    t.text = '本系统统一错误码定义如下：'
    r.append(t)
    note.append(r)
    target_p._element.addnext(note)
    # 错误码表
    headers = ['错误码', '说明']
    rows = [[ec.get('code', ''), ec.get('desc', '')] for ec in error_codes]
    from docx.text.paragraph import Paragraph as _P
    _add_table_after_paragraph(_P(note, doc), headers, rows)


def _ensure_toc_field(doc, levels='1-3'):
    """确保文档包含 TOC 域（含 H1-H3 全部层级）

    1. 检查是否已有 TOC 域（fldChar + instrText）
    2. 如有"目录"标题但无 TOC 域，在标题下插入动态 TOC 域
    3. 如无"目录"标题但也无 TOC 域，则在正文开头插入动态 TOC 域
       （2026-06-04 优化：使用 force_insert_toc 主动补全，不依赖模板"目录"标题）
    4. 2026-06-04 新增：levels 参数控制目录层级范围，默认 '1-3' 表示包含 H1/H2/H3
    """
    if doc is None:
        return False
    if has_toc_field(doc):
        return False
    # 优先查找"目录"标题
    target_para = None
    for p in doc.paragraphs:
        if not p.style or not p.style.name.startswith('Heading'):
            continue
        text = (p.text or '').strip()
        # 兼容"目录"/"目  录"/"目 录"等多种写法（空格由 _normalize_toc_heading 统一去除）
        if text == '目录' or text == '目  录' or text == '目 录':
            target_para = p
            break
    if target_para is not None:
        insert_toc_field_after(target_para)
        return True
    # 无"目录"标题时强制在正文开头插入 TOC 域
    return force_insert_toc(doc, after_paragraph=None, levels=levels)


def _format_ui_paragraphs(doc, max_chars=40):
    """遍历所有段落，对含"界面设计"标识的段落做软换行（每行不超过 max_chars 字符）"""
    if doc is None:
        return 0
    count = 0
    for p in doc.paragraphs:
        text = (p.text or '').strip()
        if not text:
            continue
        # 启发式匹配：段落含"界面"、"菜单"、"查询条件"、"按钮"等关键词时格式化
        if any(kw in text for kw in ['界面设计', '界面布局', '菜单位置', '查询条件', '按钮', '界面原型']):
            if len(text) > max_chars:
                format_ui_paragraph(p, max_chars=max_chars)
                count += 1
    return count
# 兼容旧引用：保留 _inject_update_fields 作为占位（实际已迁移到 doc_formatter）
def _inject_update_fields(output_path):
    """兼容旧引用：委托给 doc_formatter.inject_update_fields"""
    inject_update_fields(output_path)


# ═══════════════════════════════════════════════════════════════
# 2026-06-04 新增：详细设计文档质量优化函数
# ═══════════════════════════════════════════════════════════════

def _normalize_toc_heading(doc):
    """统一目录标题：去除全角/半角空格，保留"目录"作为标准标题

    模板中可能存在"目  录"/"目 录"/"目　录"等含空格的写法，
    去除后便于 _ensure_toc_field 命中匹配。
    """
    if doc is None:
        return False
    normalized = False
    for p in doc.paragraphs:
        if not p.style or not p.style.name.startswith('Heading'):
            continue
        text = (p.text or '').strip()
        # 去除内部所有空白（含全角空格 U+3000、半角空格 U+0020）
        compact = re.sub(r'[\s\u3000]+', '', text)
        if compact == '目录' and text != '目录':
            write_paragraph(p, '目录')
            normalized = True
    return normalized


def _has_meaningful_content(chapter_data):
    """判断章节数据是否包含可填充的实质内容"""
    if not isinstance(chapter_data, dict):
        return False
    sections = chapter_data.get('sections') or []
    if sections:
        return True
    body_texts = chapter_data.get('bodyTexts') or []
    if any(t for t in body_texts if t):
        return True
    tables = chapter_data.get('tables') or []
    if tables:
        return True
    content = chapter_data.get('content')
    if isinstance(content, str) and content.strip():
        return True
    if isinstance(content, dict):
        if content.get('description') or (content.get('headers') and content.get('rows')):
            return True
    return False


def _insert_placeholder_after(paragraph, placeholder='不涉及'):
    """在指定段落后插入占位段落（用于空章节兜底）"""
    if paragraph is None:
        return None
    new_p = _insert_paragraph_after(paragraph, placeholder)
    set_black(new_p)
    return new_p


def _find_overview_chapter_data(chapters, chapter_map):
    """从 JSON 章节列表中查找"概述"相关章节数据

    优先级：
    1. chapter.title 命中 CHAPTER_KEYWORDS_OVERVIEW
    2. chapter.sections 中含概述关键词
    """
    for ch in chapters or []:
        title = (ch.get('title') or '').strip()
        if title in CHAPTER_KEYWORDS_OVERVIEW:
            return ch
    # 次优：section 标题命中
    for ch in chapters or []:
        for sec in ch.get('sections', []) or []:
            sec_title = (sec.get('title') or '').strip()
            if sec_title in CHAPTER_KEYWORDS_OVERVIEW:
                return ch
    return None


def _remove_empty_table_between_headings(doc, start_keyword, end_keyword=None, placeholder='不涉及'):
    """在指定章节范围内删除空表格，并插入占位段落

    2026-06-04 新增：用于解决"组件结构图"/"系统组件"/"模块复用分析"等章节下空表格问题

    参数：
    - doc: docx Document 对象
    - start_keyword: 起始标题关键词（如"组件结构图"）
    - end_keyword: 结束标题关键词（None 表示到下一标题/文档末尾）
    - placeholder: 空表格删除后的占位文本

    返回：删除的表格数量
    """
    if doc is None or not start_keyword:
        return 0
    body = doc.element.body
    # 找到 start 与 end 之间的所有 w:tbl
    start_elem = None
    end_elem = None
    for child in list(body):
        if not child.tag.endswith('}p'):
            continue
        text_elems = child.findall('.//' + qn('w:t'))
        text = ''.join(t.text or '' for t in text_elems).strip()
        text_compact = re.sub(r'\s+', '', text)
        if start_elem is None and start_keyword in text_compact:
            start_elem = child
            continue
        if start_elem is not None and end_keyword and end_keyword in text_compact:
            end_elem = child
            break
    if start_elem is None:
        return 0
    # 收集区间内的所有 w:tbl
    in_range = False
    tables_to_check = []
    for child in list(body):
        if child is start_elem:
            in_range = True
            continue
        if end_elem is not None and child is end_elem:
            break
        if not in_range:
            continue
        if child.tag.endswith('}tbl'):
            tables_to_check.append(child)
    if not tables_to_check:
        return 0
    removed = 0
    inserted_placeholder = False
    for tbl in tables_to_check:
        # 判断表格是否为空（只有表头行，或所有单元格都为空）
        rows = tbl.findall('.//' + qn('w:tr'))
        if not rows or _is_table_empty(rows):
            # 找到一个安全位置插入占位段落：表格前的兄弟位置
            if not inserted_placeholder:
                # 在表格前插入一个占位段落
                new_p = OxmlElement('w:p')
                r = OxmlElement('w:r')
                rPr = OxmlElement('w:rPr')
                color = OxmlElement('w:color')
                color.set(qn('w:val'), '000000')
                rPr.append(color)
                r.append(rPr)
                t_elem = OxmlElement('w:t')
                t_elem.set(qn('xml:space'), 'preserve')
                t_elem.text = placeholder
                r.append(t_elem)
                new_p.append(r)
                tbl.addprevious(new_p)
                inserted_placeholder = True
            # 删除空表格
            tbl.getparent().remove(tbl)
            removed += 1
    if DEBUG_LAYOUT:
        print(f'[DEBUG] _remove_empty_table_between_headings "{start_keyword}": removed {removed} tables', file=sys.stderr)
    return removed


def _is_table_empty(rows):
    """判断表格行集合是否内容为空（只有 1 行或所有单元格都空白）"""
    if not rows or len(rows) <= 1:
        return True
    for row in rows:
        cells = row.findall('.//' + qn('w:tc'))
        for cell in cells:
            t_elems = cell.findall('.//' + qn('w:t'))
            cell_text = ''.join(t.text or '' for t in t_elems).strip()
            if cell_text:
                return False
    return True


def _renumber_component_subheadings(doc):
    """对"系统组件"标题下以"组件"开头的子标题按 H1/H2 顺序重新编号

    2026-06-04 新增：将"组件"/"组件1"/"组件N"等不同前缀的子标题统一为"组件1/组件2/..."

    返回：被重命名的子标题数量
    """
    if doc is None:
        return 0
    body = doc.element.body
    # 找到"系统组件"标题
    target_start = None
    target_style = None
    children = list(body)
    for child in children:
        if not child.tag.endswith('}p'):
            continue
        text_elems = child.findall('.//' + qn('w:t'))
        text = ''.join(t.text or '' for t in text_elems).strip()
        text_compact = re.sub(r'\s+', '', text)
        if text_compact == '系统组件':
            pPr = child.find('.//' + qn('w:pPr') + '/' + qn('w:pStyle'))
            target_style = pPr.get(qn('w:val'), '') if pPr is not None else ''
            target_start = child
            break
    if target_start is None:
        return 0
    # 确定标题的 H 级别（1/2/3）
    try:
        h_level = int(target_style)
    except (TypeError, ValueError):
        h_level = 1
    # 找到同一级别的下一个标题，作为章节边界
    sub_level_vals = {
        str(h_level + 1),
        f'Heading{h_level + 1}',
        f'heading{h_level + 1}',
        f'heading {h_level + 1}',
    }  # 子组件的样式应低一级
    boundary = None
    found = False
    for child in children:
        if child is target_start:
            found = True
            continue
        if not found:
            continue
        if not child.tag.endswith('}p'):
            continue
        pPr = child.find('.//' + qn('w:pPr') + '/' + qn('w:pStyle'))
        style_val = pPr.get(qn('w:val'), '') if pPr is not None else ''
        # 遇到同级或更高级标题 → 章节结束
        try:
            if pPr is not None and int(pPr.get(qn('w:val'), '0')) <= h_level:
                boundary = child
                break
        except (TypeError, ValueError):
            pass
    # 收集所有"组件"开头的子标题
    in_range = False
    found_start = False
    component_paras = []
    for child in children:
        if child is target_start:
            in_range = True
            continue
        if boundary is not None and child is boundary:
            break
        if not in_range:
            continue
        if not child.tag.endswith('}p'):
            continue
        pPr = child.find('.//' + qn('w:pPr') + '/' + qn('w:pStyle'))
        style_val = pPr.get(qn('w:val'), '') if pPr is not None else ''
        if style_val in sub_level_vals:
            t_elems = child.findall('.//' + qn('w:t'))
            text = ''.join(t.text or '' for t in t_elems).strip()
            # 匹配"组件"开头的标题（包括"组件1"、"组件:"等）
            if text.startswith('组件'):
                component_paras.append(child)
    # 按出现顺序重新编号
    renamed = 0
    for idx, para in enumerate(component_paras, start=1):
        _set_paragraph_text(para, f'组件{idx}')
        renamed += 1
    if DEBUG_LAYOUT:
        print(f'[DEBUG] _renumber_component_subheadings: renamed {renamed} component headings', file=sys.stderr)
    return renamed


def _set_paragraph_text(paragraph_elem, new_text):
    """将段落的文本内容替换为 new_text（保留首个 run 的样式，删除多余 run）"""
    if paragraph_elem is None:
        return
    runs = paragraph_elem.findall(qn('w:r'))
    if not runs:
        # 段落无 run，新建
        r = OxmlElement('w:r')
        t = OxmlElement('w:t')
        t.set(qn('xml:space'), 'preserve')
        t.text = new_text
        r.append(t)
        paragraph_elem.append(r)
        return
    # 清空除第一个 run 之外的所有 run 内容
    for r in runs[1:]:
        paragraph_elem.remove(r)
    # 设置第一个 run 文本
    first = runs[0]
    t_elems = first.findall(qn('w:t'))
    for t in t_elems[1:]:
        first.remove(t)
    if t_elems:
        t_elems[0].text = new_text
        t_elems[0].set(qn('xml:space'), 'preserve')
    else:
        t = OxmlElement('w:t')
        t.set(qn('xml:space'), 'preserve')
        t.text = new_text
        first.append(t)


def _build_component_specific_tech(doc, design_data):
    """为每个组件的"关键技术"章节按 component_name 动态生成内容

    2026-06-04 新增：避免使用通用模板导致所有组件"关键技术"内容一致

    数据来源：
    - design_data.requirementModuleName（业务子模块名）
    - design_data.componentTechMap（组件名→技术要点列表，优先）
    - 业务子模块的 subsections / sections 列表（fallback）

    若以上均无，使用业务子模块名构造"实现 X 子模块所需的关键技术"作为兜底。
    """
    if doc is None:
        return 0
    module_name = design_data.get('moduleName', '')
    req_module = design_data.get('requirementModuleName', module_name)
    tech_map = design_data.get('componentTechMap') or {}
    business_submodules = design_data.get('businessSubmodules') or []
    body = doc.element.body
    children = list(body)
    # 找到所有"组件N"开头的标题及其紧接的"关键技术"标题
    component_paras = []
    for child in children:
        if not child.tag.endswith('}p'):
            continue
        t_elems = child.findall('.//' + qn('w:t'))
        text = ''.join(t.text or '' for t in t_elems).strip()
        if re.match(r'^组件\d+$', text):
            component_paras.append((child, text))
    if not component_paras:
        return 0
    rewritten = 0
    for comp_para, comp_name in component_paras:
        # 找到该组件范围内的"关键技术"标题
        tech_para = _find_next_heading_in_range(body, comp_para, '关键技术',
                                                stop_headings={'组件', '关键', None})
        if tech_para is None:
            continue
        # 找到"关键技术"标题下的第一个非标题段落，作为待替换目标
        target_p = _find_first_para_after(tech_para)
        if target_p is None:
            continue
        # 按 component_name 查找对应的技术要点
        tech_items = _resolve_component_tech(comp_name, tech_map, business_submodules, req_module)
        if not tech_items:
            continue
        new_text = '；'.join(tech_items)
        _set_paragraph_text(target_p._element, new_text)
        rewritten += 1
    if DEBUG_LAYOUT:
        print(f'[DEBUG] _build_component_specific_tech: rewrote {rewritten} tech sections', file=sys.stderr)
    return rewritten


def _find_next_heading_in_range(body, start_para, heading_keyword, stop_headings=None):
    """在 start_para 之后查找第一个标题文本包含 heading_keyword 的段落

    stop_headings: 遇到标题文本以这些前缀开头则停止搜索（边界标题）
    返回 python-docx Paragraph 对象或 None
    """
    from docx.text.paragraph import Paragraph
    if body is None or start_para is None:
        return None
    found = False
    for child in list(body):
        if child is start_para:
            found = True
            continue
        if not found:
            continue
        if not child.tag.endswith('}p'):
            continue
        pPr = child.find('.//' + qn('w:pPr') + '/' + qn('w:pStyle'))
        style_val = pPr.get(qn('w:val'), '') if pPr is not None else ''
        # 仅匹配 H 样式
        is_heading_style = style_val in {
            '1', '2', '3', '4', '5', '6',
            'Heading1', 'Heading2', 'Heading3', 'Heading4', 'Heading5', 'Heading6',
        }
        if not is_heading_style:
            continue
        t_elems = child.findall('.//' + qn('w:t'))
        text = ''.join(t.text or '' for t in t_elems).strip()
        if heading_keyword in text:
            return Paragraph(child, start_para.getparent())
        # 边界标题（如新组件开头"组件N"）
        if stop_headings:
            for stop in stop_headings:
                if stop and text.startswith(stop):
                    return None
    return None


def _find_first_para_after(heading_para):
    """查找 heading_para 之后的第一个非标题段落"""
    from docx.text.paragraph import Paragraph
    if heading_para is None:
        return None
    parent = heading_para._element.getparent()
    if parent is None:
        return None
    found = False
    for elem in list(parent):
        if elem is heading_para._element:
            found = True
            continue
        if not found:
            continue
        if not elem.tag.endswith('}p'):
            continue
        pPr = elem.find('.//' + qn('w:pPr') + '/' + qn('w:pStyle'))
        style_val = pPr.get(qn('w:val'), '') if pPr is not None else ''
        if style_val.startswith('Heading') or style_val in {
            '1', '2', '3', '4', '5', '6',
            'Heading1', 'Heading2', 'Heading3', 'Heading4', 'Heading5', 'Heading6',
        }:
            return None
        return Paragraph(elem, parent)
    return None


def _resolve_component_tech(comp_name, tech_map, business_submodules, req_module):
    """根据组件名解析对应的关键技术要点列表

    解析优先级：
    1. tech_map[comp_name]（最精确）
    2. tech_map 中含 comp_name 数字序号的键（如"组件1"对应的 keys）
    3. business_submodules 中按索引匹配（comp_name='组件N' 取第 N-1 个）
    4. 使用 req_module + comp_name 构造通用兜底文案
    """
    # 1. 直接匹配
    if isinstance(tech_map, dict) and comp_name in tech_map:
        items = tech_map[comp_name]
        if isinstance(items, list) and items:
            return [str(x) for x in items if x]
    # 2. 序号推断（"组件1" → 第 1 个）
    m = re.match(r'^组件(\d+)$', comp_name)
    if m:
        idx = int(m.group(1)) - 1
        # 尝试在 business_submodules 中按索引取
        if isinstance(business_submodules, list) and 0 <= idx < len(business_submodules):
            sub = business_submodules[idx]
            if isinstance(sub, dict):
                techs = sub.get('keyTechnologies') or sub.get('techList') or []
                if isinstance(techs, list) and techs:
                    return [str(x) for x in techs if x]
                sub_name = sub.get('name') or sub.get('title') or ''
                if sub_name:
                    return _default_tech_for_submodule(sub_name, req_module)
            elif isinstance(sub, str) and sub:
                return _default_tech_for_submodule(sub, req_module)
    # 3. 兜底
    return _default_tech_for_submodule(comp_name, req_module)


def _default_tech_for_submodule(sub_name, req_module):
    """为指定子模块名生成通用关键技术描述（不硬编码具体业务）"""
    suffix = f'{req_module}的{sub_name}子模块' if sub_name else req_module
    return [
        f'{suffix}的业务流程建模与状态机设计',
        f'{suffix}涉及的数据表结构与索引设计',
        f'{suffix}的接口契约（入参/出参/错误码）定义',
        f'{suffix}的事务边界与并发控制策略',
    ]


def _fill_appendix(doc, heading_para, design_data, chapters):
    """填充附录章节：汇总各章节的关键信息作为附录内容

    2026-06-05 新增：解决详细设计文档"附录"章节为空的问题。
    从 design_data 和 chapters 中提取关键接口、数据表、错误码等汇总信息。

    2026-06-05 修复：附录内容过短（仅61字），增加栏位描述汇总、模块功能清单、
    业务规则等附录子节，从 chapters 中提取更丰富的数据。
    """
    module_name = design_data.get('moduleName', '本项目')
    insert_after = heading_para

    # 1. 附录A：接口清单
    sub_heading = _insert_paragraph_after(insert_after, '附录A：接口清单', style='Heading2')
    set_black(sub_heading)
    insert_after = sub_heading

    # 收集所有接口数据（复用 _extract_interface_items 逻辑，多关键字匹配）
    interface_rows = []
    interface_keywords = ('接口', 'API', '方法', 'URL', '请求', '服务')
    for ch in chapters:
        for sec in ch.get('sections', []):
            content = sec.get('content', {})
            if isinstance(content, dict):
                headers = content.get('headers', [])
                rows = content.get('rows', [])
                # 检测是否为接口相关表格（使用更宽松的关键字匹配）
                if headers and any(
                    any(kw in str(h) for kw in interface_keywords)
                    for h in headers
                ):
                    for row in rows:
                        interface_rows.append([str(c) for c in row])

    if interface_rows:
        headers = interface_rows[0] if len(interface_rows[0]) >= 2 else ['接口名称', '说明']
        data_rows = interface_rows[1:] if len(interface_rows) > 1 else []
        if len(headers) < 2:
            headers = ['接口名称', '说明']
        _add_table_after_paragraph(insert_after, headers, data_rows if data_rows else [['待补充', '']])
    else:
        new_p = _insert_paragraph_after(insert_after, f'{module_name}的接口清单详见各模块设计说明中的接口章节。')
        set_black(new_p)
        insert_after = new_p

    # 2. 附录B：数据表清单
    sub_heading2 = _insert_paragraph_after(insert_after, '附录B：数据表清单', style='Heading2')
    set_black(sub_heading2)
    insert_after = sub_heading2

    db_tables = design_data.get('dbTables') or []
    if db_tables:
        table_headers = ['表名', '说明']
        table_rows = [[t.get('name', ''), t.get('description', '')] for t in db_tables if isinstance(t, dict)]
        if table_rows:
            _add_table_after_paragraph(insert_after, table_headers, table_rows)
        else:
            new_p = _insert_paragraph_after(insert_after, '数据表清单详见数据库设计文档。')
            set_black(new_p)
            insert_after = new_p
    else:
        new_p = _insert_paragraph_after(insert_after, '数据表清单详见数据库设计文档。')
        set_black(new_p)
        insert_after = new_p

    # 3. 附录C：错误码清单
    sub_heading3 = _insert_paragraph_after(insert_after, '附录C：错误码清单', style='Heading2')
    set_black(sub_heading3)
    insert_after = sub_heading3

    error_codes = design_data.get('errorCodes') or []
    if error_codes:
        err_headers = ['错误码', '说明']
        err_rows = [[ec.get('code', ''), ec.get('desc', '')] for ec in error_codes]
        _add_table_after_paragraph(insert_after, err_headers, err_rows)
    else:
        new_p = _insert_paragraph_after(insert_after, '本系统统一错误码定义详见各模块设计说明。')
        set_black(new_p)
        insert_after = new_p

    # 4. 附录D：栏位描述汇总（2026-06-05 新增）
    sub_heading4 = _insert_paragraph_after(insert_after, '附录D：栏位描述汇总', style='Heading2')
    set_black(sub_heading4)
    insert_after = sub_heading4

    # 从 chapters 中收集所有栏位描述表格（含"字段"/"栏位"/"输入"/"输出"关键字的表格）
    field_tables = []
    field_keywords = ('字段', '栏位', '输入', '输出', '参数', '属性')
    for ch in chapters:
        for sec in ch.get('sections', []):
            content = sec.get('content', {})
            if isinstance(content, dict):
                headers = content.get('headers', [])
                rows = content.get('rows', [])
                if headers and any(
                    any(kw in str(h) for kw in field_keywords)
                    for h in headers
                ):
                    field_tables.append((sec.get('title', '未命名'), headers, rows))

    if field_tables:
        for tbl_title, tbl_headers, tbl_rows in field_tables:
            # 为每个栏位描述表添加子标题
            tbl_sub = _insert_paragraph_after(insert_after, tbl_title, style='Heading3')
            set_black(tbl_sub)
            insert_after = tbl_sub
            _add_table_after_paragraph(insert_after, tbl_headers, tbl_rows)
    else:
        new_p = _insert_paragraph_after(insert_after, '栏位描述详见各模块设计说明中的输入项/输出项章节。')
        set_black(new_p)
        insert_after = new_p

    # 5. 附录E：模块功能清单（2026-06-05 新增）
    sub_heading5 = _insert_paragraph_after(insert_after, '附录E：模块功能清单', style='Heading2')
    set_black(sub_heading5)
    insert_after = sub_heading5

    # 收集所有章节的标题和功能描述
    module_items = []
    for ch in chapters:
        ch_title = ch.get('title', '')
        ch_desc = ch.get('description', '')
        if ch_title:
            module_items.append([ch_title, ch_desc if ch_desc else '--'])
        for sec in ch.get('sections', []):
            sec_title = sec.get('title', '')
            sec_content = sec.get('content', '')
            if sec_title:
                sec_desc = ''
                if isinstance(sec_content, dict):
                    sec_desc = sec_content.get('description', '')
                elif isinstance(sec_content, str):
                    sec_desc = sec_content[:100]  # 截取前100字符作为摘要
                module_items.append([sec_title, sec_desc if sec_desc else '--'])

    if module_items:
        _add_table_after_paragraph(insert_after, ['功能模块', '说明'], module_items)
    else:
        new_p = _insert_paragraph_after(insert_after, '功能模块清单详见各模块设计说明。')
        set_black(new_p)
        insert_after = new_p

    # 6. 附录F：术语与缩略语（2026-06-05 新增）
    sub_heading6 = _insert_paragraph_after(insert_after, '附录F：术语与缩略语', style='Heading2')
    set_black(sub_heading6)
    insert_after = sub_heading6

    glossary = design_data.get('glossary') or []
    if glossary:
        gloss_headers = ['术语', '说明']
        gloss_rows = [[g.get('term', ''), g.get('definition', '')] for g in glossary if isinstance(g, dict)]
        if gloss_rows:
            _add_table_after_paragraph(insert_after, gloss_headers, gloss_rows)
        else:
            new_p = _insert_paragraph_after(insert_after, '术语与缩略语详见术语定义章节。')
            set_black(new_p)
    else:
        new_p = _insert_paragraph_after(insert_after, '术语与缩略语详见术语定义章节。')
        set_black(new_p)


if __name__ == '__main__':
    # Windows 控制台默认 GBK，强制 UTF-8 输出避免编码错误
    sys.stdout.reconfigure(encoding='utf-8', errors='replace')
    if len(sys.argv) < 4:
        print('Usage: python design-generator.py <template.docx> <design_data.json> <output.docx> [diagram_dir]', file=sys.stderr)
        sys.exit(1)
    diagram_dir = sys.argv[4] if len(sys.argv) > 4 and sys.argv[4] else None
    result = generate_design_from_template(sys.argv[1], sys.argv[2], sys.argv[3], diagram_dir)
    print(json.dumps({'success': True, 'outputPath': result}, ensure_ascii=False))
