"""基于模板样式，动态扩展4个业务模块H1和H2"""
import paths
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from copy import deepcopy
import json
import sys
import shutil
import re

template_path = str(paths.detail_design_template())
design_data_path = str(paths.OUTPUT_DIR / '_design-data-20260603.json')
output_path = str(paths.OUTPUT_DIR / '承兑行额度管理-详细设计文档-20260603.docx')

# 复制模板作为基础
shutil.copy(template_path, output_path)

with open(design_data_path, 'r', encoding='utf-8') as f:
    data = json.load(f)

doc = Document(output_path)

# 工具函数
def set_black(para):
    """设置段落为黑色字体"""
    for run in para.runs:
        run.font.color.rgb = None  # 重置为默认色
        if run.font.size is None:
            run.font.size = Pt(10.5)

def insert_paragraph_after(para, text, style=None):
    """在段落后插入新段落"""
    new_p = deepcopy(para._element)
    # 清除新段落的所有run
    for r in list(new_p.findall(qn('w:r'))):
        new_p.remove(r)
    # 清除原有pPr
    pPr = new_p.find(qn('w:pPr'))
    if pPr is not None:
        new_p.remove(pPr)
    # 设置样式
    if style:
        pPr_new = new_p.makeelement(qn('w:pPr'), {})
        pStyle = pPr_new.makeelement(qn('w:pStyle'), {qn('w:val'): style})
        pPr_new.append(pStyle)
        new_p.append(pPr_new)
    # 添加run
    new_r = new_p.makeelement(qn('w:r'), {})
    new_t = new_r.makeelement(qn('w:t'), {})
    new_t.text = text
    new_t.set(qn('xml:space'), 'preserve')
    new_r.append(new_t)
    new_p.append(new_r)
    para._element.addnext(new_p)
    # 返回Paragraph对象
    from docx.text.paragraph import Paragraph
    return Paragraph(new_p, para._parent)

# 找到模板中"模块1设计说明"和"模块2设计说明"对应的H1
# 改造为：保留一个作为参考样式，删除现有的两个，插入4个业务模块
existing_h1s = []
h1_module_pattern = re.compile(r'模块\d+设计说明')
for p in doc.paragraphs:
    if p.style and p.style.name == 'Heading 1':
        if h1_module_pattern.match(p.text.strip()):
            existing_h1s.append(p)

print(f"找到模块H1: {len(existing_h1s)}个")

# 策略：在附录H1前插入所有业务模块
# 1. 找到"附录"H1位置
appendix_h1 = None
for p in doc.paragraphs:
    if p.style and p.style.name == 'Heading 1' and p.text.strip() == '附录':
        appendix_h1 = p
        break

# 2. 找到4个业务模块H1数据
business_modules = [
    ch for ch in data['chapters']
    if '模块设计说明' in ch.get('title', '') and '附录' not in ch.get('title', '')
]
print(f"业务模块数: {len(business_modules)}")

# 3. 先删除现有的"模块1设计说明"和"模块2设计说明"H1及其内容
def remove_paragraph_completely(para):
    elem = para._element
    parent = elem.getparent()
    if parent is not None:
        # 删除该段落后到下一个Heading 1之间的所有内容
        next_elem = elem.getnext()
        to_remove = [elem]
        while next_elem is not None:
            if next_elem.tag.endswith('}p'):
                pPr = next_elem.find('.//' + qn('w:pPr') + '/' + qn('w:pStyle'))
                if pPr is not None:
                    style_val = pPr.get(qn('w:val'), '')
                    if style_val in ['1', 'Heading1', 'heading1', 'heading 1']:
                        break  # 遇到下一个H1停止
            to_remove.append(next_elem)
            next_elem = next_elem.getnext()
        for e in to_remove:
            parent.remove(e)

for h1 in existing_h1s:
    print(f"删除: {h1.text.strip()}")
    remove_paragraph_completely(h1)

# 4. 在附录H1前插入4个业务模块
if appendix_h1 is not None:
    # 找到附录H1的XML元素
    appendix_elem = appendix_h1._element
    appendix_parent = appendix_elem.getparent()
    appendix_idx = list(appendix_parent).index(appendix_elem)

    # 收集appendix前面要插入的H1和H2
    # 我们需要在appendix_elem前面插入新元素
    new_elements = []

    for module in business_modules:
        # 复制appendix_elem的样式
        h1_clone = deepcopy(appendix_elem)
        # 清除所有run
        for r in list(h1_clone.findall(qn('w:r'))):
            h1_clone.remove(r)
        # 添加新文本
        new_r = h1_clone.makeelement(qn('w:r'), {})
        new_t = new_r.makeelement(qn('w:t'), {})
        new_t.text = module['title']
        new_t.set(qn('xml:space'), 'preserve')
        new_r.append(new_t)
        h1_clone.append(new_r)
        new_elements.append(h1_clone)

        # 插入H2子节
        for section in module.get('sections', []):
            sec_title = section.get('title', '')
            sec_content = section.get('content', {})

            # H2标题
            h2_clone = deepcopy(appendix_elem)
            # 修改样式为Heading 2
            pPr = h2_clone.find(qn('w:pPr'))
            if pPr is None:
                pPr = h2_clone.makeelement(qn('w:pPr'), {})
                h2_clone.insert(0, pPr)
            pStyle = pPr.find(qn('w:pStyle'))
            if pStyle is None:
                pStyle = pPr.makeelement(qn('w:pStyle'), {})
                pPr.append(pStyle)
            pStyle.set(qn('w:val'), '2')  # Heading 2

            # 清除所有run
            for r in list(h2_clone.findall(qn('w:r'))):
                h2_clone.remove(r)
            # 添加新文本
            new_r = h2_clone.makeelement(qn('w:r'), {})
            new_t = new_r.makeelement(qn('w:t'), {})
            new_t.text = sec_title
            new_t.set(qn('xml:space'), 'preserve')
            new_r.append(new_t)
            h2_clone.append(new_r)
            new_elements.append(h2_clone)

            # H2内容（描述或表格）
            description = sec_content.get('description', '')
            if description:
                # 描述作为Normal段落
                p_clone = deepcopy(appendix_elem)
                # 修改样式为Normal
                pPr = p_clone.find(qn('w:pPr'))
                if pPr is None:
                    pPr = p_clone.makeelement(qn('w:pPr'), {})
                    p_clone.insert(0, pPr)
                pStyle = pPr.find(qn('w:pStyle'))
                if pStyle is not None:
                    pPr.remove(pStyle)

                for r in list(p_clone.findall(qn('w:r'))):
                    p_clone.remove(r)
                new_r = p_clone.makeelement(qn('w:r'), {})
                new_rpr = new_r.makeelement(qn('w:rPr'), {})
                new_color = new_rpr.makeelement(qn('w:color'), {qn('w:val'): '000000'})
                new_rpr.append(new_color)
                new_r.append(new_rpr)
                new_t = new_r.makeelement(qn('w:t'), {})
                new_t.text = description
                new_t.set(qn('xml:space'), 'preserve')
                new_r.append(new_t)
                p_clone.append(new_r)
                new_elements.append(p_clone)

            # 表格
            headers = sec_content.get('headers', [])
            rows = sec_content.get('rows', [])
            if headers and rows:
                # 表格用deeper处理，需要在XML层级构造
                # 找到模板中已有的一个表格作为样式参考
                pass  # 表格处理复杂，先跳过，后续单独处理

    # 插入新元素到appendix之前
    for elem in reversed(new_elements):
        appendix_parent.insert(appendix_idx, elem)

doc.save(output_path)
print(f"已保存: {output_path}")
print(f"文档段落数: {len(doc.paragraphs)}")
