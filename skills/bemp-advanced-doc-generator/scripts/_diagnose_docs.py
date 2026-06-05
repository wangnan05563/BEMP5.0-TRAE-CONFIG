"""诊断生成文档的内容质量"""
import sys, os
sys.path.insert(0, os.path.dirname(__file__))
import paths
from docx import Document
from docx.shared import RGBColor

def is_blue(rgb):
    if rgb is None:
        return False
    r, g, b = int(rgb[0:2], 16), int(rgb[2:4], 16), int(rgb[4:6], 16)
    return b > 150 and b > r * 1.5 and b > g * 1.3

def check_doc(doc_path, label):
    if not doc_path.exists():
        print(f'  文件不存在: {doc_path}')
        return

    doc = Document(str(doc_path))
    print(f'\n{"="*70}')
    print(f'诊断: {label} - {doc_path.name}')
    print(f'{"="*70}')

    # 1. 检查修订记录表格
    print('\n--- 1. 修订记录表格 ---')
    for i, table in enumerate(doc.tables):
        header_text = ' '.join(cell.text.strip() for cell in table.rows[0].cells) if table.rows else ''
        if any(kw in header_text for kw in ['版本', '日期', '修改人', '修改内容']):
            print(f'  找到修订记录表 (table#{i})')
            if len(table.rows) > 1:
                row1 = [cell.text.strip() for cell in table.rows[1].cells]
                print(f'  第一行数据: {row1}')
            break
    else:
        print('  未找到修订记录表!')

    # 2. 检查蓝色文本
    print('\n--- 2. 蓝色文本 ---')
    blue_count = 0
    blue_samples = []
    for p in doc.paragraphs:
        for run in p.runs:
            if run.font.color and run.font.color.rgb:
                if is_blue(str(run.font.color.rgb)):
                    blue_count += 1
                    if len(blue_samples) < 5:
                        blue_samples.append(f'  [{p.style.name}] {run.text[:60]}')
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for run in p.runs:
                        if run.font.color and run.font.color.rgb:
                            if is_blue(str(run.font.color.rgb)):
                                blue_count += 1
    print(f'  蓝色 run 总数: {blue_count}')
    for s in blue_samples:
        print(s)

    # 3. 检查各章节内容
    print('\n--- 3. 章节内容检查 ---')
    check_keywords = [
        '概述', '编写目的', '读者对象', '使用范围', '适用范围',
        '术语', '参考资料', '设计策略', '设计目标', '设计原则',
        '外部接口', '组件汇总', '技术实现', '关键技术',
        '非功能性', '附录', '设计约束', '模块复用',
        '类图', '顺序图', '活动图', '状态图',
    ]
    headings = []
    for p in doc.paragraphs:
        if p.style and p.style.name and p.style.name.startswith('Heading'):
            headings.append(p)

    for kw in check_keywords:
        found = False
        for h in headings:
            if kw in h.text:
                # 检查标题和下一标题之间的内容
                content_parts = []
                h_elem = h._element
                next_elem = h_elem.getnext()
                while next_elem is not None:
                    from docx.oxml.ns import qn
                    if next_elem.tag == qn('w:p'):
                        # 直接从 XML 读取文本，避免 Paragraph 构造错误
                        pPr = next_elem.find(qn('w:pPr'))
                        style_id = ''
                        if pPr is not None:
                            pStyle = pPr.find(qn('w:pStyle'))
                            if pStyle is not None:
                                style_id = pStyle.get(qn('w:val'), '')
                        # Heading 样式通常以 Heading 开头或包含 heading
                        if style_id and ('Heading' in style_id or 'heading' in style_id.lower() or 'TOC' in style_id):
                            break
                        # 提取文本
                        texts = []
                        for r in next_elem.findall(qn('w:r')):
                            t = r.find(qn('w:t'))
                            if t is not None and t.text:
                                texts.append(t.text)
                        text = ''.join(texts).strip()
                        if text:
                            content_parts.append(text[:80])
                    elif next_elem.tag == qn('w:tbl'):
                        content_parts.append('[表格]')
                    next_elem = next_elem.getnext()

                if content_parts:
                    print(f'  [{kw}] ✓ 有内容 ({len(content_parts)}段): {content_parts[0][:60]}')
                else:
                    print(f'  [{kw}] ✗ 空章节!')
                found = True
                break
        if not found:
            print(f'  [{kw}] - 标题未找到')

    # 4. 检查图片
    print('\n--- 4. 图片检查 ---')
    from docx.oxml.ns import qn
    drawing_count = 0
    for p in doc.paragraphs:
        drawings = p._element.findall('.//' + qn('w:drawing'))
        drawing_count += len(drawings)
    print(f'  图片总数: {drawing_count}')

    # 5. 检查占位符残留
    print('\n--- 5. 占位符残留 ---')
    placeholder_count = 0
    placeholder_samples = []
    for p in doc.paragraphs:
        text = p.text.strip()
        if '<' in text and '>' in text:
            placeholder_count += 1
            if len(placeholder_samples) < 3:
                placeholder_samples.append(f'  {text[:80]}')
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                text = cell.text.strip()
                if '<' in text and '>' in text:
                    placeholder_count += 1
    print(f'  占位符数量: {placeholder_count}')
    for s in placeholder_samples:
        print(s)

docs = [
    ('概要设计', paths.OUTPUT_DIR / '河南农商-概要设计说明书-20260604.docx'),
    ('详细设计', paths.OUTPUT_DIR / '河南农商-详细设计文档-20260604.docx'),
]
for label, doc_path in docs:
    check_doc(doc_path, label)
