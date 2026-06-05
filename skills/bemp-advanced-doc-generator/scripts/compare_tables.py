"""深入对比v4与new的H2子节内容、表格内容差异"""
import paths
from docx import Document
import json

v4_path = str(paths.BANK_REQUIREMENTS_DIR / '承兑行额度管理-详细设计说明书-v4.docx')
new_path = str(paths.OUTPUT_DIR / '承兑行额度管理-详细设计文档-20260603.docx')

print("=" * 80)
print("v4文档H2列表及表格:")
print("=" * 80)
doc_v4 = Document(v4_path)
v4_h2_to_table = {}
current_h2 = None
for p in doc_v4.paragraphs:
    if p.style and p.style.name == 'Heading 1':
        current_h2 = p.text.strip()
    elif p.style and p.style.name == 'Heading 2':
        current_h2 = p.text.strip()
        v4_h2_to_table[current_h2] = []

for i, tbl in enumerate(doc_v4.tables):
    rows = len(tbl.rows)
    cols = len(tbl.columns) if tbl.rows else 0
    first_row = ' | '.join(c.text.strip()[:15] for c in tbl.rows[0].cells) if tbl.rows else ''
    print(f"  表格[{i:2d}]: {rows}行x{cols}列 | {first_row}")

print(f"\nv4总表格数: {len(doc_v4.tables)}")

print("\n" + "=" * 80)
print("new文档H2列表及表格:")
print("=" * 80)
doc_new = Document(new_path)
for i, tbl in enumerate(doc_new.tables):
    rows = len(tbl.rows)
    cols = len(tbl.columns) if tbl.rows else 0
    first_row = ' | '.join(c.text.strip()[:15] for c in tbl.rows[0].cells) if tbl.rows else ''
    print(f"  表格[{i:3d}]: {rows}行x{cols}列 | {first_row}")
print(f"\nnew总表格数: {len(doc_new.tables)}")
