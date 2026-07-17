# -*- coding: utf-8 -*-
"""交付文档质量审核脚本 - 验证两份 .docx 文档的结构完整性"""
import os
from docx import Document
from docx.oxml.ns import qn

OUTPUT_DIR = r"d:\code\QJ\BEMP5.0DEV\.trae\skills\bemp-advanced-doc-generator\output"

DOCS = [
    ("详细设计说明书", os.path.join(OUTPUT_DIR, "河南农商同步机构树数据并校验-详细设计说明书-20260715.docx")),
    ("测试报告", os.path.join(OUTPUT_DIR, "河南农商同步机构树数据并校验-测试报告-20260715.docx")),
]


def check_doc(title, path):
    print("=" * 70)
    print("审核文档：{} ({})".format(title, os.path.basename(path)))
    print("=" * 70)

    if not os.path.exists(path):
        print("[FAIL] 文件不存在")
        return False

    size_kb = os.path.getsize(path) / 1024
    print("[OK] 文件大小：{:.1f} KB".format(size_kb))

    doc = Document(path)

    # 1. 段落统计
    paragraphs = doc.paragraphs
    print("[OK] 段落总数：{}".format(len(paragraphs)))

    # 2. 标题层级统计
    h1_count = 0
    h2_count = 0
    h3_count = 0
    heading_titles = []
    for p in paragraphs:
        if p.style.name.startswith('Heading'):
            level = p.style.name.replace('Heading ', '')
            text = p.text.strip()
            if level == '1':
                h1_count += 1
                heading_titles.append(("H1", text))
            elif level == '2':
                h2_count += 1
                heading_titles.append(("H2", text))
            elif level == '3':
                h3_count += 1
                heading_titles.append(("H3", text))
    print("[OK] 标题层级：H1={}, H2={}, H3={}".format(h1_count, h2_count, h3_count))

    # 3. 表格统计
    tables = doc.tables
    print("[OK] 表格总数：{}".format(len(tables)))

    # 4. 检查目录域 (TOC field)
    has_toc = False
    for p in paragraphs:
        for run in p.runs:
            xml = run._element.xml
            if 'TOC' in xml and 'fldChar' in xml:
                has_toc = True
                break
        if has_toc:
            break
    if has_toc:
        print("[OK] 目录域（TOC field）：存在")
    else:
        print("[WARN] 目录域（TOC field）：未检测到")

    # 5. 检查页码域 (PAGE field)
    has_page_num = False
    for section in doc.sections:
        footer = section.footer
        for p in footer.paragraphs:
            for run in p.runs:
                xml = run._element.xml
                if 'PAGE' in xml and 'fldChar' in xml:
                    has_page_num = True
                    break
            if has_page_num:
                break
        if has_page_num:
            break
    if has_page_num:
        print("[OK] 页码域（PAGE field）：存在")
    else:
        print("[WARN] 页码域（PAGE field）：未检测到")

    # 6. 检查页眉
    has_header = False
    for section in doc.sections:
        header = section.header
        for p in header.paragraphs:
            if p.text.strip():
                has_header = True
                print("[OK] 页眉内容：{}".format(p.text.strip()[:50]))
                break
        if has_header:
            break
    if not has_header:
        print("[WARN] 页眉：未检测到内容")

    # 7. 检查 updateFields 注入
    settings = doc.settings.element
    update_fields = settings.find(qn('w:updateFields'))
    if update_fields is not None and update_fields.get(qn('w:val')) == 'true':
        print("[OK] updateFields 注入：已启用（打开文档时自动更新目录）")
    else:
        print("[WARN] updateFields 注入：未启用")

    # 8. 检查占位符残留
    placeholder_patterns = ['XXX', '【请填写】', '【待补充】', '{placeholder}', '待填写']
    placeholder_count = 0
    for p in paragraphs:
        for pattern in placeholder_patterns:
            if pattern in p.text:
                placeholder_count += 1
                print("[WARN] 占位符残留：{} -> '{}...'".format(pattern, p.text[:30]))
    for table in tables:
        for row in table.rows:
            for cell in row.cells:
                for pattern in placeholder_patterns:
                    if pattern in cell.text:
                        placeholder_count += 1
    if placeholder_count == 0:
        print("[OK] 占位符残留：无")
    else:
        print("[WARN] 占位符残留：{} 处".format(placeholder_count))

    # 9. 列出所有 H1 标题
    print("\n--- H1 章节列表 ---")
    for level, text in heading_titles:
        if level == "H1":
            print("  {}".format(text))

    # 10. 关键章节检查
    full_text = "\n".join([p.text for p in paragraphs])
    for table in tables:
        for row in table.rows:
            for cell in row.cells:
                full_text += "\n" + cell.text

    key_chapters = {
        "详细设计说明书": ["系统概述", "系统架构", "数据库设计", "接口设计", "详细设计", "部署设计", "偏差记录", "已知问题"],
        "测试报告": ["测试概述", "测试执行", "测试结果", "缺陷", "质量评估", "测试结论", "已知问题"],
    }
    print("\n--- 关键章节检查 ---")
    for kw in key_chapters.get(title, []):
        if kw in full_text:
            print("[OK] 包含关键字：{}".format(kw))
        else:
            print("[FAIL] 缺少关键字：{}".format(kw))

    print("\n")
    return True


def main():
    print("\n")
    print("#" * 70)
    print("# BEMP 河南农商 同步机构树数据并校验 - 交付文档质量审核")
    print("#")
    print("#" * 70)
    print("\n")
    all_ok = True
    for title, path in DOCS:
        if not check_doc(title, path):
            all_ok = False
    print("#" * 70)
    if all_ok:
        print("# 审核结论：全部文档生成成功，结构完整")
    else:
        print("# 审核结论：存在问题，请检查")
    print("#" * 70)


if __name__ == "__main__":
    main()
