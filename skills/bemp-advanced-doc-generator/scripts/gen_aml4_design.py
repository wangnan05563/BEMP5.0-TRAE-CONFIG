# -*- coding: utf-8 -*-
"""
BEMP 河南农信"反洗钱四阶段校验"交付文档生成器
基于实际代码实现与测试报告数据生成 .docx，结构遵循交付文档标准：
需求概述/方案设计/详细设计/测试报告/部署说明/已知问题 + 附录。
生成模式与 gen_org_mgmt_design.py 先例一致（全代码构建，无模板遗留空章节）。
"""
import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

OUTPUT_DIR = r"d:\code\QJ\BEMP5.0DEV\docs\delivery\反洗钱四阶段校验"
OUTPUT_FILE = os.path.join(OUTPUT_DIR, "反洗钱四阶段校验-详细设计说明书-v1.docx")

# ===================== 样式工具函数（与 gen_org_mgmt_design.py 先例一致） =====================

def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = tcPr.find(qn('w:tcBorders'))
    if tcBorders is None:
        tcBorders = OxmlElement('w:tcBorders')
        tcPr.append(tcBorders)
    for edge in ('top', 'left', 'bottom', 'right'):
        tag = 'w:{}'.format(edge)
        element = tcBorders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            tcBorders.append(element)
        element.set(qn('w:val'), 'single')
        element.set(qn('w:sz'), '4')
        element.set(qn('w:color'), '000000')


def set_cell_background(cell, color="D9D9D9"):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color)
    tcPr.append(shd)


def set_run_font(run, name="宋体", size=10.5, bold=False, color=None):
    run.font.name = name
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    rFonts.set(qn('w:eastAsia'), name)
    rFonts.set(qn('w:ascii'), name)
    rFonts.set(qn('w:hAnsi'), name)


def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        set_run_font(run, name="黑体", size=max(12, 18 - level * 2), bold=True, color="000000")
    return p


def add_para(doc, text, bold=False, indent=True, size=10.5):
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.first_line_indent = Cm(0.74)
    p.paragraph_format.line_spacing = 1.5
    run = p.add_run(text)
    set_run_font(run, name="宋体", size=size, bold=bold)
    return p


def add_bullet(doc, text, size=10.5):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.line_spacing = 1.5
    run = p.add_run(text)
    set_run_font(run, name="宋体", size=size)
    return p


def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr_cells[i].text = ""
        p = hdr_cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(header)
        set_run_font(run, name="宋体", size=10, bold=True)
        set_cell_background(hdr_cells[i])
        set_cell_border(hdr_cells[i])
        hdr_cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    for r_idx, row in enumerate(rows, 1):
        row_cells = table.rows[r_idx].cells
        for c_idx, val in enumerate(row):
            row_cells[c_idx].text = ""
            p = row_cells[c_idx].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(str(val))
            set_run_font(run, name="宋体", size=10)
            set_cell_border(row_cells[c_idx])
            row_cells[c_idx].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    if col_widths:
        for i, width in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(width)
    doc.add_paragraph()
    return table


def add_toc_field(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run()
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = 'TOC \\o "1-3" \\h \\z \\u'
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')
    fldChar3 = OxmlElement('w:t')
    fldChar3.text = '右键此处选择"更新域"以生成目录'
    fldChar4 = OxmlElement('w:fldChar')
    fldChar4.set(qn('w:fldCharType'), 'end')
    r_element = run._r
    r_element.append(fldChar1)
    r_element.append(instrText)
    r_element.append(fldChar2)
    r_element.append(fldChar3)
    r_element.append(fldChar4)


def add_page_break(doc):
    p = doc.add_paragraph()
    run = p.add_run()
    run.add_break(WD_BREAK.PAGE)


def setup_page_and_header_footer(doc):
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(3.0)
        section.right_margin = Cm(2.5)
        header = section.header
        if not header.paragraphs[0].runs:
            hp = header.paragraphs[0]
            hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = hp.add_run("河南农商票据交易管理平台 - 详细设计说明书")
            set_run_font(run, name="宋体", size=9, color="808080")
        footer = section.footer
        fp = footer.paragraphs[0]
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = fp.add_run()
        set_run_font(run, name="宋体", size=9, color="808080")
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')
        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = 'PAGE'
        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'end')
        run._r.append(fldChar1)
        run._r.append(instrText)
        run._r.append(fldChar2)


def inject_update_fields(doc):
    settings = doc.settings.element
    update_fields = settings.find(qn('w:updateFields'))
    if update_fields is None:
        update_fields = OxmlElement('w:updateFields')
        update_fields.set(qn('w:val'), 'true')
        settings.append(update_fields)


# ===================== 文档生成主流程 =====================

def build_design_doc():
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = '宋体'
    style.font.size = Pt(10.5)
    rPr = style.element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    rFonts.set(qn('w:eastAsia'), '宋体')

    setup_page_and_header_footer(doc)

    # ===== 封面 =====
    for _ in range(3):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("河南农商银行")
    set_run_font(run, name="黑体", size=26, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("新一代票据系统个性化开发")
    set_run_font(run, name="黑体", size=22, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("反洗钱四阶段校验")
    set_run_font(run, name="黑体", size=22, bold=True)

    for _ in range(2):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("详 细 设 计 说 明 书")
    set_run_font(run, name="黑体", size=20, bold=True)

    for _ in range(6):
        doc.add_paragraph()

    info_table = doc.add_table(rows=5, cols=2)
    info_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    info_data = [
        ("项目名称", "BEMP5.0 河南农商银行个性化开发"),
        ("需求名称", "反洗钱四阶段校验与面客文案（HNNS-EB-STD-REQ-002）"),
        ("需求依据", "PRD v3.4.8 §3.2.3 + 待确认清单 Q-11/Q-12 结论"),
        ("文档版本", "V1.0"),
        ("编写日期", "2026-08-29"),
    ]
    for i, (k, v) in enumerate(info_data):
        cells = info_table.rows[i].cells
        cells[0].text = ""
        cells[1].text = ""
        p1 = cells[0].paragraphs[0]
        p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r1 = p1.add_run(k)
        set_run_font(r1, name="宋体", size=12, bold=True)
        p2 = cells[1].paragraphs[0]
        p2.alignment = WD_ALIGN_PARAGRAPH.LEFT
        r2 = p2.add_run(v)
        set_run_font(r2, name="宋体", size=12)
        for c in cells:
            set_cell_border(c)

    add_page_break(doc)

    # ===== 修订记录 =====
    add_heading(doc, "修订记录", level=1)
    add_table(doc,
        ["版本", "日期", "修订内容", "修订人"],
        [
            ["V1.0", "2026-08-29", "初始版本：覆盖反洗钱四阶段校验需求概述、方案设计、逐文件详细设计、测试报告（单测/评审/功能测试/二轮回归）、部署说明与已知问题清单", "BEMP文档交付工程师"],
        ],
        col_widths=[2, 3, 8, 3]
    )
    add_page_break(doc)

    # ===== 目录 =====
    add_heading(doc, "目录", level=1)
    add_toc_field(doc)
    add_page_break(doc)

    # ===== 第一章 需求概述 =====
    add_heading(doc, "第一章 需求概述", level=1)

    add_heading(doc, "1.1 业务背景与需求目标", level=2)
    add_para(doc, "公司金融部《关于新一代票据业务管理系统反洗钱功能优化的需求》明确了四个校验阶段与面向客户的报错文案模板。系统现状：仅实现市场交易侧（转贴现/回购/再贴现）中互金名单校验，四阶段（直贴贴现申请/提示承兑/背书转让/企业客户维护）未覆盖；且既有对内文案 HNNX0BE320009/010 暴露\"中互金名单\"来源，不适合面客场景。")
    add_para(doc, "本期目标：四阶段校验在票据系统侧落地（申请落地前拦截），面客文案与对内文案区分（对内文案保持不变），新增独立校验开关。")
    add_table(doc,
        ["项", "内容"],
        [
            ["需求编号", "HNNS-EB-STD-REQ-002（PRD v3.4.8 §3.2.3）"],
            ["变更标识", "add-hnnx-aml-four-stage-check（spec/tasks 已评审通过）"],
            ["适用银行", "河南农信（hnnxbank）"],
            ["本期交付", "四阶段校验 + 面客文案 HNNX0BE320011 + 独立开关参数 + 初始化 SQL"],
        ],
        col_widths=[3, 12]
    )

    add_heading(doc, "1.2 需求范围", level=2)
    add_para(doc, "四个校验阶段均复用既有两轮匹配口径，校验时点均为业务落地前（super 之前/落库前）。")
    add_table(doc,
        ["阶段", "业务入口", "校验角色", "校验时点"],
        [
            ["①直贴贴现申请", "网银渠道贴现分支（NES.007 三类型共用）+ 行内 eDisc commit 兜底", "贴现申请人、出票人", "申请落地前（super 之前）"],
            ["②提示承兑申请", "网银渠道承兑分支（NES.002）", "出票人、收票人", "申请落地前"],
            ["③背书转让", "网银渠道背书分支（NES.006，市场交易侧背书已接入不重复）", "持票人、收票人、承兑行", "申请落地前"],
            ["④企业客户维护", "业务管理-客户管理-企业客户维护（insert/update）", "客户名（socCode+custName）", "落库前"],
            ["面客文案", "HNNX0BE320011（黑/灰共用，不暴露名单性质）", "—", "命中即返回"],
            ["独立开关", "hnnx.cust.antiMoneyList.block（法人级，默认开启）", "—", "前置判断 + 异常降级依据"],
        ],
        col_widths=[3, 5.5, 4, 3]
    )

    add_heading(doc, "1.3 范围外（明确不做）", level=2)
    add_table(doc,
        ["序号", "范围外项", "说明"],
        [
            ["1", "直贴/承兑/背书的行内页面入口", "原文操作路径均为网银端发起；行内承兑申请页面 eAcptApplyMain 不在原文四阶段路径中"],
            ["2", "AntiMoneyBillDialog 弹窗改造", "四阶段为报文/表单单一主体场景，无多票据明细弹窗需求"],
            ["3", "灰名单人工审批流", "二期规划，一期命中提示中断"],
            ["4", "市场交易侧逻辑与文案变更", "既有校验逻辑与 HNNX0BE320009/010 对内文案保持不变"],
        ],
        col_widths=[1.5, 5, 9]
    )

    add_heading(doc, "1.4 需求确认闭环", level=2)
    add_para(doc, "需求梳理阶段待确认事项已全部闭环，测试假设验证结论如下。")
    add_table(doc,
        ["事项", "类型", "最终结论", "验证结果"],
        [
            ["Q-11 拦截文案口径", "测试假设", "黑/灰共用面客文案 HNNX0BE320011，按原文模板含角色前缀与 96288 热线，不暴露名单性质", "P0-015/P0-018/P1-026 逐字验证通过"],
            ["Q-12 网银端四阶段是否本期范围", "需求方确认", "用户确认纳入本期（含直贴/承兑/背书 + 客户维护）", "四阶段用例已编制并执行（v3.4.10 基线）"],
            ["证件类型码值 22", "已实现", "v3.4.6 已实现，无开发量", "代码核查确认"],
            ["四阶段取数口径", "用户确认", "spec 逐项记录（custNo/billId/hldrId/acpt_bank_no 链路）", "代码锚点核对一致（见 2.4 节）"],
        ],
        col_widths=[4, 2.2, 5.5, 4]
    )
    add_page_break(doc)

    # ===== 第二章 方案设计 =====
    add_heading(doc, "第二章 方案设计", level=1)

    add_heading(doc, "2.1 四阶段校验总体架构", level=2)
    add_para(doc, "统一收口设计：四阶段入口各异（网银渠道/eDisc/客户维护），但校验逻辑统一收口于 HnnxAntiMoneyValidateUtil.validateCustBizList——前置开关判断、round1 信用代码匹配、round2 名称兜底、命中抛 HNNX0BE320011、异常按开关降级，全部在工具类内实现，各入口不重复判断（避免双重读参）。")
    add_table(doc,
        ["层次", "组件", "职责"],
        [
            ["网银渠道层", "HnnxBankEbank2004AtomImpl（@CustomizedBean 顶替产品 bean eBank2004Atom）", "贴现/承兑/背书三分支覆写，校验于 super 前；命中异常经产品 dealException 转 retCode/retMsg 应答返回网银"],
            ["行内业务层", "HnnxBankDiscBillServiceImpl", "submitElecFlow commit 兜底校验（明细锁定后、业务校验前），与前端预检形成双保险"],
            ["客户维护层", "HnnxbankCustCorpServiceImpl", "覆写 insertCustCorp/updateCustCorp 前置校验，要素取 DTO 提交值（socCode+custName）"],
            ["公共校验层", "HnnxAntiMoneyValidateUtil.validateCustBizList", "开关前置、两轮匹配、面客文案、异常降级（fail-closed/fail-open）"],
            ["数据层", "HNNX_M_CUST_SPECIAL_INFO（名单）/ tm_cust_corp / tb_bill_info / tb_corp_bill_hldr / tm_cpes_branch", "名单存储与各角色取数"],
            ["配置层", "TM_BUSINESS_PARAMETER（开关参数）+ hnnx_mt_zh_CN.properties（面客文案）", "参数启停与文案模板"],
        ],
        col_widths=[2.5, 5.5, 7.5]
    )

    add_heading(doc, "2.2 开关机制", level=2)
    add_para(doc, "法人级参数 hnnx.cust.antiMoneyList.block（常量定义于 HnnxCommonConst.CUST_ANTI_MONEY_LIST_PARAM，L75），默认开启（PARAM_VALUE='1'，代码回退默认值同为 \"1\"）。与市场交易开关 hnnx.market.antiMoneyList.block 完全隔离：两开关读取路径无交集（L175 vs L202），独立开关可避免市场交易开关降级方向对四阶段的连带影响。开关实时生效：UPDATE 后无需重启或清缓存（二轮 V2c 实测确认）。")
    add_table(doc,
        ["开关值", "正常校验行为", "名单服务异常降级行为"],
        [
            ["'1'（开启）", "执行两轮匹配，命中拦截", "拦截留痕（fail-closed）"],
            ["'0'（关闭）", "跳过校验，日志 INFO 留痕后正常流转", "放行留痕（fail-open）"],
            ["读取异常", "按 fail-closed 处理：LOGGER.error 留痕后抛 320011 降级文案", "同左（防御读取故障产生校验失效窗口）"],
        ],
        col_widths=[3, 6, 6.5]
    )

    add_heading(doc, "2.3 匹配口径（两轮匹配）", level=2)
    add_para(doc, "四阶段复用既有两轮匹配口径，与市场交易侧同一方法（queryMatchBySocCodes/queryMatchByNames）零漂移。")
    add_table(doc,
        ["轮次", "匹配要素", "实现", "说明"],
        [
            ["round1（优先）", "统一社会信用代码（certType='22'+certNo）", "queryMatchBySocCodes，in 查询按 ≤1000 分片（partition）", "有信用代码的角色参与，命中即拦截不进入 round2"],
            ["round2（兜底）", "客户名称", "queryMatchByNames；交易侧名称 trim 后参与匹配（L217-222）", "已由信用代码命中的角色排除（buildCustNameIndex），无信用代码的角色仅参与名称兜底"],
            ["数据条件", "—", "DATA_SRC='18'、VLD_ST='1'、名单类型黑(1)/灰(2)", "与市场侧同源"],
            ["时效判定", "—", "失效时间空=长期有效；有失效按 生效≤当前≤失效 闭区间", "与市场侧同一实现（两分支 criteria）"],
            ["多角色命中", "—", "throwIfCustHit 取 roles.get(0)，LinkedHashMap 保序遍历", "首个命中角色前缀返回；round1 全量先于 round2"],
        ],
        col_widths=[2.8, 3.8, 4.8, 4.5]
    )

    add_heading(doc, "2.4 四阶段取数链路", level=2)
    add_para(doc, "取数口径均经用户确认（spec 记录），代码实现与口径逐项核对一致。")
    add_table(doc,
        ["阶段", "角色", "取数要素", "数据来源链路"],
        [
            ["①网银直贴", "贴现申请人", "soc_code + cust_name", "报文 custNo → tm_cust_corp（cust_no 相等且 active_flag='1'，经产品化 getCustCorpByCustNo）；查库缺失回退报文 applSocCode/custName，不阻断"],
            ["①网银直贴", "出票人", "drwr_name + drwr_soc_code", "报文 billId → tb_bill_info（id 相等）"],
            ["①eDisc", "出票人/贴现申请人", "drwr_name（逐票）/ discCustName", "tb_bill_info；贴现申请人为空时 fillDiscountCustInfoFromDb（ET05 三条件）补全"],
            ["②提示承兑", "出票人", "drwr_name + drwr_soc_code", "报文 billId → tb_bill_info"],
            ["②提示承兑", "收票人", "pyee_name + pyee_soc_code", "报文 billId → tb_bill_info"],
            ["③背书转让", "收票人", "pyee_name（仅名称兜底）", "报文 billId → tb_bill_info"],
            ["③背书转让", "持票人", "hldr_soc_code + hldr_name", "报文 hldrId → tb_corp_bill_hldr（id 相等）"],
            ["③背书转让", "承兑行", "acpt_bank_name + soc_code（ST01 反查）", "tb_bill_info.acpt_bank_no → tm_cpes_branch.trans_brch_bank_no（brch_status='ST01'）；反查失败返回 null 降级名称兜底"],
            ["④客户维护", "客户", "socCode + custName", "DTO 提交值（CustCorpDto.socCode L75 / custName L29），即表单提交值"],
        ],
        col_widths=[2.4, 2.4, 3.6, 7.5]
    )

    add_heading(doc, "2.5 面客文案机制", level=2)
    add_para(doc, "错误码 HNNX0BE320011 定义于 hnnx_mt_zh_CN.properties（L11）：{0}：{1}暂不符合相关政策要求，暂不支持提供该服务，详情请咨询开户机构或96288。黑/灰名单共用同一面客文案（Q-11 结论，不暴露名单性质）；对内文案 HNNX0BE320009/010 保留不变，市场交易侧调用点仍使用原错误码，两侧分流无交叉（P0-022 审查确认）。")
    add_table(doc,
        ["场景", "{0} 角色前缀", "{1} 角色描述"],
        [
            ["①直贴贴现", "贴现申请人 / 出票人", "含客户名的角色描述（工具类 buildCustRoleDesc 组装）"],
            ["②提示承兑", "出票人 / 收票人", "同上"],
            ["③背书转让", "持票人 / 收票人 / 承兑行", "同上"],
            ["④客户维护", "客户名", "该客户[XX]（RoleCheckItem 用法二，调用方组装）"],
            ["异常降级", "系统", "系统繁忙，请稍后重试（拼接通顺性见已知问题 KP-03）"],
        ],
        col_widths=[3, 4.5, 8]
    )

    add_heading(doc, "2.6 关键设计决策", level=2)
    add_table(doc,
        ["决策", "理由"],
        [
            ["新增独立开关，不共用市场交易开关", "避免市场交易开关降级方向对四阶段产生连带影响；两开关读取路径无交集（L175 vs L202）"],
            ["继承 Ebank2004ProxyAtomImpl + @CustomizedBean 顶替产品 bean", "本类处于产品 Proxy 下一层，bean 顶替必然胜出（若与产品同层竞争存活不确定）；保留产品插件机制（ebank2004PluginAtomList）；fxbank/huisbank/jinzbank 三家既有先例同构"],
            ["校验位于 super 之前", "命中即中断：申请不落库、不向票交所/ECDS 发送 NES 报文；若在 super 之后校验，报文已发出无法撤回"],
            ["开关读取 fail-closed", "开关读取异常按开启降级（拦截留痕），避免校验失效窗口"],
            ["in 查询 ≤1000 分片（partition）", "Oracle IN 列表上限防护"],
            ["贴现申请人查库缺失回退报文要素", "客户档案查询故障不放大为申请阻断（不臆造不阻断）；库内 cust_name 是名单登记对账基准，查库口径优先"],
            ["票据主档主键级重复查询", "接受一次主键级重复查询，换取校验逻辑与产品化流程完全解耦（不在 super 前修改其上下文）"],
            ["命中异常不做 try-catch 吞处理", "沿 insertApplyTransInfo 上抛至产品 catch 段，经 dealException 填入应答 retCode/retMsg，与既有网银报错机制一致"],
        ],
        col_widths=[6, 9.5]
    )
    add_page_break(doc)

    # ===== 第三章 详细设计（逐文件实现说明） =====
    add_heading(doc, "第三章 详细设计（逐文件实现说明）", level=1)

    add_heading(doc, "3.1 新增/修改文件清单", level=2)
    add_table(doc,
        ["序号", "文件", "类型", "关键改动", "锚点"],
        [
            ["1", "hnnxbank-biz-api/.../common/HnnxCommonConst.java", "修改", "新增开关参数码常量 CUST_ANTI_MONEY_LIST_PARAM = \"hnnx.cust.antiMoneyList.block\"", "L75"],
            ["2", "hnnxbank-conf/.../mt/i18n/hnnx_mt_zh_CN.properties", "修改", "新增面客文案错误码 HNNX0BE320011（320009/010 对内文案保持原样）", "L11"],
            ["3", "hnnxbank-biz-as/.../antimoney/validate/HnnxAntiMoneyValidateUtil.java", "修改", "客户侧校验入口 validateCustBizList/isCustBizListIntercept/validateCustBizListInternal、RoleCheckItem、queryAcptBankSocCode 承兑行反查、读参 fail-closed、in 查询 ≤1000 分片", "L201/L379/L413/L1094/L1915"],
            ["4", "hnnxbank-biz-as/.../corp/HnnxbankCustCorpServiceImpl.java", "修改", "覆写 insertCustCorp/updateCustCorp 前置校验，要素=DTO socCode+custName", "L90/L104/L136"],
            ["5", "hnnxbank-biz-as/.../disc/HnnxBankDiscBillServiceImpl.java", "修改", "submitElecFlow eDisc commit 兜底 + collectDiscCustRoleItems 拆分", "L275/L495"],
            ["6", "hnnxbank-biz-as/.../ebank/impl/HnnxBankEbank2004AtomImpl.java", "新增", "extends Ebank2004ProxyAtomImpl + @CustomizedBean 顶替产品 bean \"eBank2004Atom\"；覆写贴现/承兑/背书三分支，super 前校验", "L58-284"],
            ["7", "hnnxbank-biz-as/pom.xml", "修改", "新增 bemp-channel-as 依赖（对齐 ext-fxbank，支撑 Atom 继承）", "L130-131"],
            ["8", "deploy/bemp-script/.../河南农信/V202301.03.081_202608290001_T202608290001_客户渠道侧中互金名单校验开关参数初始化.dml.sql", "新增", "TM_BUSINESS_PARAMETER DELETE+INSERT 幂等（含 LEGAL_NO 条件），默认 PARAM_VALUE='1'", "详见第五章"],
        ],
        col_widths=[1.2, 5.2, 1.5, 6.2, 2.4]
    )

    add_heading(doc, "3.2 HnnxCommonConst.java（开关常量）", level=2)
    add_para(doc, "L75 新增 public static final String CUST_ANTI_MONEY_LIST_PARAM = \"hnnx.cust.antiMoneyList.block\"。参数码常量与初始化 SQL 脚本 PARAM_KEY 字面量一致（P0-016 审查确认），避免脚本与代码漂移。")
    add_table(doc,
        ["要素", "说明"],
        [
            ["参数码", "hnnx.cust.antiMoneyList.block"],
            ["参数级别", "法人级（LEGAL_NO='000000' 初始化）"],
            ["默认值", "'1'（开启）；代码读取回退默认 \"1\""],
        ],
        col_widths=[3, 12]
    )

    add_heading(doc, "3.3 hnnx_mt_zh_CN.properties（面客文案）", level=2)
    add_para(doc, "L11 新增 HNNX0BE320011={0}：{1}暂不符合相关政策要求，暂不支持提供该服务，详情请咨询开户机构或96288。{0}=角色前缀，{1}=含客户名的角色描述。黑名单/灰名单命中均抛 320011；降级文案（L395/L446）{0}=系统、{1}=系统繁忙，请稍后重试。对内文案 320009/010 定义原样保留，市场侧调用点（throwIfHit L573/L577）与四阶段入口（throwIfCustHit/320011）完全分流。")
    add_table(doc,
        ["错误码", "用途", "文案要点"],
        [
            ["HNNX0BE320011（新增）", "四阶段面客拦截（黑/灰共用）", "角色前缀 + 不暴露名单性质 + 96288 热线"],
            ["HNNX0BE320009（不变）", "市场交易侧黑名单对内拦截", "暴露名单来源，仅对内使用"],
            ["HNNX0BE320010（不变）", "市场交易侧灰名单对内拦截", "联系风控人员，仅对内使用"],
        ],
        col_widths=[4.2, 5, 6.3]
    )

    add_heading(doc, "3.4 HnnxAntiMoneyValidateUtil.java（客户侧校验核心）", level=2)
    add_para(doc, "四阶段校验统一收口。为满足 SonarQube S3776（认知复杂度）要求，核心方法拆分为单一职责私有方法。")
    add_table(doc,
        ["方法/结构", "锚点", "职责与要点"],
        [
            ["isCustBizListIntercept", "L201-202", "BusinessParameterUtil.getParamValue 读参，回退默认 \"1\"；仅引用客户侧开关，与市场侧开关隔离"],
            ["validateCustBizList", "L379-397", "对外入口：空集合短路（先于读开关返回，零 DB 往返）→ 前置开关判断（关闭 INFO 留痕跳过）→ 读参异常 fail-closed（L386-396）→ 委托 internal"],
            ["validateCustBizListInternal", "L413 起", "round1 buildSocCodeIndex→matchCustBySocCodes → round2 buildCustNameIndex（排除已命中、无信用代码不入索引）→matchCustByNames → throwIfCustHit；异常按 interceptOn 降级（L442-449）"],
            ["S3776 拆分方法", "L458/L473/L490/L511", "buildSocCodeIndex / matchCustBySocCodes / buildCustNameIndex / matchCustByNames，单一职责可测"],
            ["throwIfCustHit / buildCustRoleDesc", "L528/L551", "命中即抛 320011（roles.get(0) 保序）；工具类组装 {1} 角色描述"],
            ["RoleCheckItem", "L1094", "客户侧角色校验项：roleLabel（面客文案前缀）/ custName（客户名称）/ socCode（可选统一社会信用代码）；{1} 支持工具类组装与调用方组装两种用法"],
            ["queryAcptBankSocCode", "L1915", "承兑行信用代码反查：acpt_bank_no → tm_cpes_branch.trans_brch_bank_no（brch_status='ST01'）；仅认 ST01、失败返回 null 不抛异常，该角色自动降级名称兜底；口径与市场侧收拢同源"],
            ["名称 trim 口径", "L217-222", "交易侧名称 trim 后参与 round2 匹配（distinctNames.add(name.trim())），与前端序列化 trim 形成双层防护，无漏匹配风险（二轮 V2b 实测）"],
            ["in 查询分片", "L685 partition", "Oracle IN 列表 ≤1000 上限防护"],
        ],
        col_widths=[3.8, 2.6, 9.1]
    )

    add_heading(doc, "3.5 HnnxbankCustCorpServiceImpl.java（④企业客户维护）", level=2)
    add_para(doc, "覆写 insertCustCorp（L90）/updateCustCorp（L104）：未命中调 super，产品化 CustCorpServiceImpl 零改动。校验要素=DTO 提交值（CustCorpDto.socCode L75 / custName L29），legalNo=baseRequest.getReqLegalNo()。")
    add_table(doc,
        ["要点", "说明"],
        [
            ["RoleCheckItem 用法二", "roleLabel=客户名 + roleDesc=该客户[XX]，由调用方组装文案 {1}（L113 注释/L128 组装）"],
            ["校验位置", "落库前（L136 validateCustBizList），命中拦截返回面客文案，数据不落库（P0-018/P0-019 实测 total=0）"],
            ["修改场景防护", "命中时原记录未被误改（二轮 V2a 快照复核一致）"],
            ["不做双重开关判断", "开关/两轮匹配/降级均在工具类内，本类不重复读参（L83 注释约定）"],
        ],
        col_widths=[4, 11.5]
    )

    add_heading(doc, "3.6 HnnxBankDiscBillServiceImpl.java（①eDisc commit 兜底）", level=2)
    add_para(doc, "submitElecFlow（L234）在明细锁定后、业务校验前（L275）执行四阶段校验兜底——前端预检已接入，commit 层兜底形成双保险。collectDiscCustRoleItems（L495）负责角色组装。")
    add_table(doc,
        ["要点", "说明"],
        [
            ["角色组装", "出票人逐票收集（tb_bill_info.drwr_name）；贴现申请人 discCustName 为空时复用 fillDiscountCustInfoFromDb（ET05 三条件）补全"],
            ["校验时点", "校验位于明细锁定后、业务校验前（L272-276 口径锚点），命中抛 320011 面客文案，角色前缀=贴现申请人/出票人"],
            ["幂等与复用", "validateCustBizList 统一入口，开关/匹配/降级由工具类内实现，此处不重复判断（L274 注释）"],
        ],
        col_widths=[4, 11.5]
    )

    add_heading(doc, "3.7 HnnxBankEbank2004AtomImpl.java（网银三分支覆写，新增）", level=2)
    add_para(doc, "网银统一申请 Atom 河南农信个性化覆写（新增文件，L1-285）。@CloudComponent + @CustomizedBean（L58-59）extends Ebank2004ProxyAtomImpl：启动时 BempBeanDefinitionRegistryProcessor 删除产品 Ebank2004AtomImpl 的 \"eBank2004Atom\" bean 定义并以产品 bean 名注册本类——Ebank2004ServiceImpl 的 @Qualifier 注入、单笔直调与多笔线程池（Ebank2004CallableImpl）、NES 策略类回调均命中本类，三个覆写方法经虚分派生效。")
    add_table(doc,
        ["覆写分支", "锚点", "校验角色与组装顺序", "要点"],
        [
            ["composeDiscApplyTransInfo（NES.007 直贴/强制/回购式共用）", "L84-93", "①=[贴现申请人, 出票人]", "billId 主键查票据主档；贴现申请人经 queryDiscApplyCustCorp（L258-270，异常捕获回退报文要素不阻断）；校验于 super 前"],
            ["doAcceptancApplyTransInfo（NES.001/002/003/撤票共用）", "L102-110", "②=[出票人, 收票人]（仅 NES.002）", "非提示承兑返回空集合短路零开销（buildAcceptanceRoleItems L190-207）"],
            ["composeEndrsmtApplyTransInfo（NES.006）", "L121-135", "③=[收票人, 持票人, 承兑行]", "收票人仅名称兜底；持票人经 getCorpBillHldr（父类公开方法）；承兑行经 queryAcptBankSocCode（ST01 反查，失败降级名称兜底）"],
            ["buildDiscRoleItems（回退策略）", "L153-179", "—", "custCorp=null 回退 reqDto.applSocCode/custName（L162-168）；reqDto=null 空集合（L154-156）；billInfoDto=null 仅校验贴现申请人（L170-177）"],
            ["预约提示承兑自动覆盖", "类注释 L49-51", "—", "Ebank2049 预约到期执行调用 batchComrclDrftApply→eBank2004Atom 共用提交点自动覆盖；Ebank2050 预约撤销不落地申请不在范围"],
        ],
        col_widths=[4.6, 1.8, 3.4, 5.7]
    )

    add_heading(doc, "3.8 hnnxbank-biz-as/pom.xml（依赖）", level=2)
    add_para(doc, "L130-131 新增 bemp-channel-as 依赖（${bemp-channel-as.version}），支撑 HnnxBankEbank2004AtomImpl 继承产品化 Ebank2004ProxyAtomImpl，与 ext-fxbank 同构。")

    add_heading(doc, "3.9 初始化 SQL（开关参数登记）", level=2)
    add_para(doc, "详见第五章部署说明。脚本命名沿用既有 V202301.03.081 格式，新任务编号 T202608290001。")
    add_table(doc,
        ["要素", "值"],
        [
            ["落点表", "TM_BUSINESS_PARAMETER"],
            ["幂等结构", "DELETE+INSERT（含 LEGAL_NO 条件），可重复执行"],
            ["默认值", "PARAM_VALUE='1'（默认开启）"],
            ["头注释", "含关联需求、代码落点、参数语义（四阶段校验启停与降级规则）"],
        ],
        col_widths=[4, 11.5]
    )
    add_page_break(doc)

    # ===== 第四章 测试报告 =====
    add_heading(doc, "第四章 测试报告（单测+评审+功能测试+二轮回归汇总）", level=1)

    add_heading(doc, "4.1 测试概述", level=2)
    add_table(doc,
        ["项", "内容"],
        [
            ["测试目的", "验证反洗钱四阶段校验功能、面客文案、开关机制、降级行为符合 PRD v3.4.8 §3.2.3 与 spec 约定"],
            ["测试范围", "四阶段校验入口 + 独立开关 + 面客文案 + 两轮匹配口径 + 异常降级 + 市场侧隔离性"],
            ["测试环境", "hnnxbank：前端 8091 / 后端 8010 / Redis 6379 / ZooKeeper 21811 / Adapter 8090 全 UP；业务库 10.20.42.211:1521/orcl（应用实际 schema=BEMP_HNNX，JDBC 通道）；业务日期 2025-08-28"],
            ["测试方法", "单元测试（mvn）+ 代码审查 + 功能测试一轮（Playwright）+ 二轮回归（Chrome DevTools 真实浏览器双通道）+ DB 断言（JDBC）+ 后端日志留痕核对"],
            ["测试日期", "2026-08-29（一轮 + 二轮同日完成）"],
            ["用例基线", "v3.4.10（用例评审两轮闭环放行口径），23 条增量用例（P0=12/P1=7/P2=4），路径 aotutests-devtools/testcases/反洗钱名单校验/"],
            ["测试数据", "名单基线 R1 黑（21/长期）/R2 灰（21/20250101~20281231）/R5 已失效（21/END=20250731）/R9 未生效（21/START=20300101）；R22 缺失待业务方准备"],
        ],
        col_widths=[2.8, 12.7]
    )

    add_heading(doc, "4.2 单元测试结果", level=2)
    add_para(doc, "新增 19 用例全部绿色；全模块 82 用例全绿（Tests run: 82, Failures: 0, Errors: 0），既有用例回归通过。")
    add_table(doc,
        ["测试类", "新增用例数", "覆盖要点"],
        [
            ["HnnxAntiMoneyCustValidateTest", "8", "开关关闭跳过、黑/灰命中文案、异常降级两分支（fail-closed 完整拼接断言/fail-open 放行）、信用代码优先匹配"],
            ["HnnxBankCustCorpServiceImplTest", "3", "客户维护覆写：新增/修改前置校验、RoleCheckItem 用法二组装"],
            ["HnnxBankEbank2004AtomAmlRoleTest", "4", "网银角色组装：查库优先回退报文、reqDto/billInfoDto/custCorp 空值分支"],
            ["HnnxBankDiscBillAmlRoleTest", "4", "eDisc 角色组装：逐票收集、fillDiscountCustInfoFromDb 补全"],
            ["既有回归", "63（合计 82）", "市场侧反洗钱既有用例 28+23+6+6 等全绿"],
        ],
        col_widths=[5, 2.2, 8.3]
    )

    add_heading(doc, "4.3 代码评审与质量扫描结果", level=2)
    add_table(doc,
        ["环节", "结果", "处置"],
        [
            ["代码评审", "严重 0 / 主要 3 / 次要+建议若干", "主要 3 项已修复并经原评审通道复验通过"],
            ["SonarQube（项目 bemp-ext-hnnxbank-antimoney）", "本次需求引入 6 Critical（S3776 认知复杂度）", "全部修复复验清零（校验方法拆分为单一职责私有方法）；MCP 不可用时按降级方案走 sonar-scanner CLI 执行"],
            ["SonarQube 遗留", "存量方法 getRiskBillNo/submitElecFlow 认知复杂度 Critical 2 项；new_coverage=0%（JaCoCo 报告未导入）", "存量方法非本次引入，列入已知问题 KP-02（建议专项重构）；覆盖率待 CI 配置后消除"],
        ],
        col_widths=[4.5, 5.5, 5.5]
    )

    add_heading(doc, "4.4 功能测试执行汇总（一轮，AML4-FT-20260829-01）", level=2)
    add_para(doc, "按\"本期执行段\"口径：执行 21 / 通过 21 / 失败 0；延后项均为评审放行口径中明确的延后段，不阻塞门禁。")
    add_table(doc,
        ["优先级", "用例数", "执行", "通过", "失败", "阻塞/延后"],
        [
            ["P0", "12", "11", "11", "0", "1（P0-026 round1 断言段，R22 数据前置；报文段另计）"],
            ["P1", "7", "6", "6", "0", "1（P1-027 报文段）"],
            ["P2", "4", "4", "4", "0", "0"],
            ["合计", "23", "21", "21", "0", "报文联调段 5 条 + round1 段 + 演练段 + eDisc 页面段 2 条"],
        ],
        col_widths=[1.8, 1.8, 1.8, 1.8, 1.8, 6.5]
    )

    add_heading(doc, "4.5 缺陷清单", level=2)
    add_para(doc, "本轮测试全程仅 1 个缺陷且已关闭，无 P0/P1 缺陷。")
    add_table(doc,
        ["编号", "严重度", "描述", "根因", "修复状态"],
        [
            ["DEF-AML4-001", "P2（数据类）", "灰名单客户④新增提交未被拦截且成功落库", "名单库 HNNX_M_CUST_SPECIAL_INFO 中 R2（灰）/R5（已失效）记录缺失，应用名单查询返回 0 行后放行——代码行为正确（未命中放行），数据未按测试数据准备就绪", "已按 SQL 直插 R2/R5，删除误落库档案后重测，拦截+文案+不落库三重断言通过 → 已关闭"],
            ["—", "—", "无 P0/P1 缺陷", "—", "二轮回归零新缺陷（DEF-AML4-002 起未启用）"],
        ],
        col_widths=[2.4, 2, 3.6, 5.5, 4]
    )

    add_heading(doc, "4.6 二轮回归结果（AML4-RT-20260829-02）", level=2)
    add_para(doc, "Chrome DevTools 真实浏览器通道（区别于一轮 Playwright，防同通道偶发/共因误判）：6/6 可执行场景全过，零新缺陷，一轮结论稳定复现。")
    add_table(doc,
        ["场景", "结果", "关键结论"],
        [
            ["R1 新增黑名单拦截全链路复测", "通过", "文案逐字命中；响应 HTTP 200 + retCode=HNNX0BE320011 + retMsg 完整面客文案；DB 零落库；日志 ERROR 留痕"],
            ["V2a 修改场景变体", "通过", "改名黑名单客户被拦截，原记录未被误改（快照复核一致）"],
            ["V2b 首尾空格变体", "通过", "带空格黑名单名被拦截；前端表单序列化 trim + 后端交易侧 trim（L217-222）双层口径一致，无漏匹配风险"],
            ["V2c 开关关闭/恢复变体", "通过", "开关='0' 放行落库且日志 INFO 留痕；恢复='1' 同场景恢复拦截；开关 UPDATE 后实时生效（无需重启/清缓存）"],
            ["Q 浏览器端质量监控", "通过", "零 TypeError/ReferenceError/ChunkLoadError；零 4xx/5xx；拦截响应结构 3 次完全一致"],
            ["E eDisc 页面段评估", "维持 BLOCKED", "TB_BILL_INFO 无出票人=名单客户的在途票据（0/21），需完整票据流转构造，业务表不直插；代码审查段两轮均通过"],
        ],
        col_widths=[4.2, 2.2, 9.1]
    )

    add_heading(doc, "4.7 覆盖度分析", level=2)
    add_table(doc,
        ["维度", "覆盖情况"],
        [
            ["需求覆盖", "spec 7 项 ADDED Requirements 全部有用例覆盖：面客文案（P0-015/018/022/P1-026）、独立开关（P0-016/017/P1-028/032）、直贴贴现（P0-020/021/023/P2-029）、提示承兑（P0-024/P1-027）、背书转让（P0-025/P2-030）、企业客户维护（P0-018/019/P1-026）、匹配口径（P0-026/P2-027/028）"],
            ["测试类型覆盖", "单元测试（19 新增+82 全绿）/ 代码审查 / 功能测试一轮（Playwright）/ 二轮回归（Chrome DevTools 双通道交叉确认）/ DB 断言 / 日志留痕核对"],
            ["测试假设验证", "Q-11 面客文案假设经 P0-015/P0-018/P1-026 逐字验证通过；时效口径经 P2-027 四象限（长期/有效期内/未生效/已失效）实测验证"],
        ],
        col_widths=[3, 12.5]
    )

    add_heading(doc, "4.8 延后测试项清单", level=2)
    add_para(doc, "延后项不阻塞本轮门禁，随对应条件就绪后跟踪执行（详见第六章 KP-01）。")
    add_table(doc,
        ["延后项", "关联用例", "延后原因/所需条件"],
        [
            ["网银报文应答段（retCode/retMsg 断言）", "P0-023/024/025、P1-027、P1-030 报文段", "需网银外围联调环境/报文模拟器；代码审查段已全部通过"],
            ["round1 信用代码优先命中断言", "P0-026 步骤 4、P1-022", "R22（CERT_TYPE='22' 名单记录）待业务方准备；单测信用代码优先用例已通过"],
            ["运行时异常注入演练", "P1-028 步骤 5/6", "测试环境不具备停库/超时注入手段；降级两分支已由单测（入参注入）+代码审查覆盖"],
            ["eDisc 页面段", "P0-020/021 步骤 1-6", "需经完整票据流转（签发/承兑/背书）构造在途批次，两轮评估均不可构造；代码审查段已通过"],
        ],
        col_widths=[4.5, 4.5, 6.5]
    )
    add_page_break(doc)

    # ===== 第五章 部署说明 =====
    add_heading(doc, "第五章 部署说明", level=1)

    add_heading(doc, "5.1 SQL 脚本执行", level=2)
    add_table(doc,
        ["项", "内容"],
        [
            ["脚本路径", "deploy/bemp-script/src/main/resources/banks/河南农信/V202301.03.081_202608290001_T202608290001_客户渠道侧中互金名单校验开关参数初始化.dml.sql"],
            ["落点表", "TM_BUSINESS_PARAMETER（法人级参数，LEGAL_NO='000000'）"],
            ["幂等性", "DELETE+INSERT 结构（含 LEGAL_NO 条件），可重复执行（实测重复执行后仍各 1 条记录）"],
            ["登记值", "PARAM_KEY=hnnx.cust.antiMoneyList.block、PARAM_VALUE='1'（默认开启）、PARAM_GROUP_CODE=CustAntiMoneyList（独立分组）"],
            ["PARAM_TITLE/REMARK", "说明四阶段校验语义与降级规则（开启→拦截留痕/关闭→放行留痕）"],
            ["头注释", "含关联需求（HNNS-EB-STD-REQ-002）、代码落点、参数语义"],
            ["执行时机", "后端部署前执行（应用读取该参数做前置判断，缺失时代码回退默认 \"1\"，不阻塞启动）"],
        ],
        col_widths=[3.5, 12]
    )

    add_heading(doc, "5.2 配置项说明", level=2)
    add_table(doc,
        ["配置项", "默认值", "级别", "语义", "本次变更"],
        [
            ["hnnx.cust.antiMoneyList.block", "'1'", "法人级", "四阶段校验启停 + 异常降级方向（开→拦截留痕/关→放行留痕）", "新增"],
            ["hnnx.market.antiMoneyList.block", "'1'", "法人级", "市场交易侧校验开关（既有）", "不变（与四阶段开关互不影响）"],
            ["HNNX0BE320011", "—", "错误码（properties）", "四阶段面客拦截文案（黑/灰共用）", "新增"],
            ["HNNX0BE320009/010", "—", "错误码（properties）", "市场交易侧对内拦截文案", "不变"],
        ],
        col_widths=[4.6, 1.6, 2.6, 4.7, 2]
    )

    add_heading(doc, "5.3 部署顺序与部署验证", level=2)
    add_para(doc, "涉及模块：hnnxbank-biz-api（常量）、hnnxbank-biz-as（校验工具/服务覆写/Atom/pom）、hnnxbank-conf（文案）、deploy/bemp-script（SQL）。")
    add_table(doc,
        ["步骤", "操作", "验证方法"],
        [
            ["1", "执行初始化 SQL（目标 schema=BEMP_HNNX）", "SELECT PARAM_VALUE FROM TM_BUSINESS_PARAMETER WHERE PARAM_KEY='hnnx.cust.antiMoneyList.block' 返回 '1'"],
            ["2", "部署后端（含 hnnxbank-biz-api/biz-as/conf）", "服务启动正常；日志无 bean 注册冲突（@CustomizedBean 顶替 eBank2004Atom 生效）"],
            ["3", "功能验证（拦截正向）", "企业客户维护提交黑名单客户名 → 弹窗\"客户名：该客户[XX]暂不符合相关政策要求…96288\"；日志 ERROR 命中留痕"],
            ["4", "功能验证（开关负向）", "开关临时置 '0' 提交同客户 → 放行且日志 INFO 跳过留痕；验证后复位 '1'"],
            ["5", "回归市场侧", "市场交易侧命中场景仍返回 HNNX0BE320009/010 对内文案，两侧无交叉"],
        ],
        col_widths=[1.4, 5.3, 8.8]
    )
    add_page_break(doc)

    # ===== 第六章 已知问题清单 =====
    add_heading(doc, "第六章 已知问题清单", level=1)
    add_para(doc, "以下事项不阻塞本期交付，均已在测试报告中如实记录，并给出影响评估与后续计划。")
    add_table(doc,
        ["编号", "类别", "描述", "影响评估", "紧急度", "后续计划"],
        [
            ["KP-01", "延后测试项", "①网银报文段 5 条（P0-023/024/025、P1-027/030 报文应答段，需网银联调环境）；②eDisc 页面段 P0-020/021（需完整票据流转数据）；③P0-026/P1-022 round1 段（R22 名单数据待业务方准备）；④P1-028 异常演练段", "核心校验逻辑已由单测+代码审查+真实浏览器通道四重验证；报文应答链路经产品 dealException 机制保障且响应结构实测规范（func_insertCustCorp retCode/retMsg）", "中", "随联调环境就绪、R22 数据就绪、票据流转数据构造后逐项回归"],
            ["KP-02", "SonarQube 遗留", "存量方法 getRiskBillNo/submitElecFlow 认知复杂度 Critical 2 项；new_coverage=0%（JaCoCo 报告未导入）", "存量方法非本次需求引入，无功能影响；覆盖率盲区影响质量度量完整性", "低", "建议专项重构任务；CI 配置 JaCoCo 报告导入后消除覆盖率盲区"],
            ["KP-03", "实现质量观察项", "降级拼接文案\"系统：系统繁忙，请稍后重试暂不符合相关政策要求…\"通顺性存疑（{0}=系统、{1}=系统繁忙，请稍后重试 与 320011 模板拼接）", "仅在名单服务异常降级时出现，出现频率低；语义仍可理解但不够顺", "低", "建议评估独立降级错误码与专用降级文案"],
            ["KP-04", "存量代码遗留", "HnnxbankCustCorpServiceImpl @Resource 注入方式、getSelfAndChildBranchNos 空 catch", "非本次引入；空 catch 存在异常吞噬风险", "低", "建议单独任务清理（注入方式规范化 + 异常日志补全）"],
            ["KP-05", "P3 改进建议", "修改弹窗后再开新增弹窗，addForm 字段重置不彻底（custCorp.vue，createTime/updateTime/createBrchNo 等残留）", "真实用户正常操作不触发业务异常；二轮 V2c 已验证显式清空后落库无污染", "低", "产品化跟踪：addFormReset 对部分字段重置完整性增强"],
            ["KP-06", "环境/运维", "四磁盘使用率 ≥95%（SonarQube ES 水位 transient 临时调整）；后端代码尚未 git 提交", "磁盘满将影响 SonarQube ES 与构建稳定性；代码未提交存在交付资产风险", "中", "清理磁盘并持久化水位配置；代码提交需用户确认后按规范执行：【河南农商】个性化开发【反洗钱四阶段校验】"],
            ["KP-07", "测试数据准备", "R22（CERT_TYPE='22' 码值名单）记录待业务方准备", "round1 信用代码优先命中断言（P0-026 步骤 4/P1-022）依赖该数据", "中", "业务方准备 R22 数据后回归 round1 断言段"],
        ],
        col_widths=[1.4, 2.2, 4.6, 3.8, 1.3, 3.2]
    )
    add_page_break(doc)

    # ===== 附录 =====
    add_heading(doc, "附录A 需求确认闭环结论", level=1)
    add_para(doc, "需求梳理阶段《待确认清单》（v3.4.8）事项最终结论已全部闭环并写入正文：Q-11（拦截文案口径，测试假设）按原文案模板实现为 HNNX0BE320011，经 P0-015/P0-018/P1-026 逐字验证通过；Q-12（网银端四阶段是否本期范围）经用户确认纳入本期，四阶段用例已编制（v3.4.10）并执行。四阶段取数口径经用户逐项确认（spec 记录），代码实现与口径核对一致。无遗留阻塞项。")
    add_table(doc,
        ["事项", "闭环状态"],
        [
            ["Q-11 面客文案口径", "已闭环（假设验证通过）"],
            ["Q-12 四阶段范围", "已闭环（用户确认纳入本期）"],
            ["四阶段取数口径", "已闭环（用户确认 + 代码核对一致）"],
            ["R22 名单数据", "移交已知问题 KP-07（数据准备，非需求阻塞）"],
        ],
        col_widths=[6, 6]
    )

    add_heading(doc, "附录B 参考资料", level=1)
    add_table(doc,
        ["资料", "路径"],
        [
            ["需求 spec", ".trae/specs/add-hnnx-aml-four-stage-check/spec.md"],
            ["实现任务", ".trae/specs/add-hnnx-aml-four-stage-check/tasks.md"],
            ["PRD（v3.4.8）", "docs/prd/反洗钱校验功能/HNNS-EB-STD-REQ-002-反洗钱功能优化-v3.md"],
            ["待确认清单", "docs/prd/反洗钱校验功能/HNNS-EB-STD-REQ-002-反洗钱功能优化-待确认-v3.md"],
            ["测试用例（v3.4.10）", "aotutests-devtools/testcases/反洗钱名单校验/（P0用例.md/P1用例.md/功能地图与优先级矩阵.md/测试数据准备.md）"],
            ["一轮功能测试报告", "aotutests-devtools/testcases/反洗钱名单校验/四阶段功能测试报告.md（AML4-FT-20260829-01）"],
            ["二轮回归测试报告", "aotutests-devtools/testcases/反洗钱名单校验/四阶段二轮回归测试报告.md（AML4-RT-20260829-02）"],
            ["二轮截图/日志", "aotutests-devtools/screenshots/2026-08-29/AML4-round2/、aotutests-devtools/console-logs/2026-08-29/AML4-round2-console.json"],
            ["初始化 SQL", "deploy/bemp-script/src/main/resources/banks/河南农信/V202301.03.081_202608290001_T202608290001_客户渠道侧中互金名单校验开关参数初始化.dml.sql"],
        ],
        col_widths=[4.5, 11]
    )

    add_heading(doc, "附录C 术语表", level=1)
    add_table(doc,
        ["术语", "说明"],
        [
            ["中互金名单", "中国互联网金融协会关注名单（含黑名单/灰名单），存储于 HNNX_M_CUST_SPECIAL_INFO"],
            ["四阶段校验", "直贴贴现申请/提示承兑/背书转让/企业客户维护四个业务阶段的名单校验"],
            ["面客文案", "面向客户的拦截提示（不暴露名单来源），错误码 HNNX0BE320011"],
            ["对内文案", "内部风控拦截提示（HNNX0BE320009/010），暴露名单来源，仅市场交易侧使用"],
            ["两轮匹配", "round1 统一社会信用代码优先匹配 + round2 客户名称兜底匹配"],
            ["fail-closed", "名单服务异常时按开关开启方向拦截并留痕，防止校验失效窗口"],
            ["bean 顶替", "@CustomizedBean 个性化类以产品 bean 名注册、删除产品 bean 定义的覆写机制"],
            ["R22", "CERT_TYPE='22'（统一社会信用代码）类型的名单记录，用于 round1 优先匹配"],
        ],
        col_widths=[3.5, 12]
    )

    inject_update_fields(doc)
    os.makedirs(OUTPUT_DIR, exist_ok=True)
    doc.save(OUTPUT_FILE)
    print(f"OK: {OUTPUT_FILE}")
    return OUTPUT_FILE


if __name__ == "__main__":
    build_design_doc()
