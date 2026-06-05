"""验证详细设计文档质量 - 增强版，检查表格"""
from docx import Document
from lxml import etree
import sys

doc_path = sys.argv[1]
doc = Document(doc_path)

nsmap = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}

# 构建段落到表格的映射
body = doc.element.body
elem_to_table = {}
for elem in body:
    if elem.tag.endswith('}tbl'):
        # 找到表格前的最近一个段落元素
        prev = elem.getprevious()
        if prev is not None:
            elem_to_table[id(prev)] = True

print(f'段落数: {len(doc.paragraphs)}')
print(f'表格数: {len(doc.tables)}')

# 检查标题内容，包括表格
headings = []
for i, p in enumerate(doc.paragraphs):
    if p.style and p.style.name.startswith('Heading'):
        # 检查标题后是否有内容（段落或表格）
        has_text = False
        has_table = False
        content_preview = ''
        for j in range(i+1, min(i+5, len(doc.paragraphs))):
            next_p = doc.paragraphs[j]
            if next_p.style and next_p.style.name.startswith('Heading'):
                break
            if next_p.text.strip():
                has_text = True
                content_preview = next_p.text.strip()[:60]
                break
        # 检查标题元素后是否有表格
        heading_elem = p._element
        next_elem = heading_elem.getnext()
        if next_elem is not None and next_elem.tag.endswith('}tbl'):
            has_table = True
            if not content_preview:
                content_preview = '[表格]'

        status = '有内容' if (has_text or has_table) else '空'
        detail = ''
        if has_table and not has_text:
            detail = '[仅表格]'
        elif has_table and has_text:
            detail = '[文本+表格]'
        print(f'  {p.style.name}: {p.text.strip()[:50]} [{status}]{detail} {content_preview}')

# 检查蓝色文本
blue_count = 0
for p in doc.paragraphs:
    for run in p.runs:
        if run.font.color and run.font.color.rgb:
            color = str(run.font.color.rgb)
            if color.startswith('00') and color not in ['000000']:
                blue_count += 1
print(f'\n蓝色文本残留数: {blue_count}')
