from docx.shared import RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from copy import deepcopy

BLACK = RGBColor(0, 0, 0)
BLUE = RGBColor(0, 0, 255)
TABLE_FONT_SIZE_PT = 10.5

BLUE_VARIANTS = [
    RGBColor(0, 0, 255),
    RGBColor(0, 0, 205),
    RGBColor(0, 0, 139),
    RGBColor(0, 0, 128),
    RGBColor(65, 105, 225),
    RGBColor(31, 73, 125),
    RGBColor(0, 112, 192),
    RGBColor(0, 51, 153),
    RGBColor(68, 114, 196),
    RGBColor(91, 155, 213),
]


def is_blue_color(rgb):
    if rgb is None:
        return False
    if rgb in BLUE_VARIANTS:
        return True
    r, g, b = rgb[0], rgb[1], rgb[2]
    if b > 150 and b > r * 1.8 and b > g * 1.5:
        return True
    return False


def is_blue_paragraph(p):
    for r in p.runs:
        if r.font.color and r.font.color.rgb and is_blue_color(r.font.color.rgb):
            return True
    return False


def is_blue_cell(cell):
    for para in cell.paragraphs:
        for r in para.runs:
            if r.font.color and r.font.color.rgb and is_blue_color(r.font.color.rgb):
                return True
    return False


def is_hyperlink_run(run):
    parent = run._element.getparent()
    if parent is None:
        return False
    if parent.tag.endswith('}hyperlink'):
        return True
    for ancestor in parent.iterancestors():
        if ancestor.tag.endswith('}hyperlink'):
            return True
    return False


def _is_hyperlink_style_run(run):
    if is_hyperlink_run(run):
        return True
    if run.font.underline and run.font.underline != False:
        return True
    rPr = run._element.find(qn('w:rPr'))
    if rPr is not None:
        u_elem = rPr.find(qn('w:u'))
        if u_elem is not None and u_elem.get(qn('w:val'), 'none') != 'none':
            return True
    return False


def has_non_hyperlink_blue(paragraph):
    for r in paragraph.runs:
        if r.font.color and r.font.color.rgb and is_blue_color(r.font.color.rgb) and r.text.strip():
            if not _is_hyperlink_style_run(r):
                return True
    return False


def set_black(p):
    for r in p.runs:
        if r.font.color and r.font.color.rgb and is_blue_color(r.font.color.rgb):
            if _is_hyperlink_style_run(r):
                continue
            r.font.color.rgb = BLACK
        elif r.font.color and r.font.color.rgb and r.font.color.rgb != BLACK:
            r.font.color.rgb = BLACK


def clear_paragraph(p):
    for r in p.runs:
        r.text = ''
    if not p.runs:
        p.add_run('')


PLACEHOLDER_KEYWORDS = [
    '此处填写', '例如：', '例如:', '例如', '<注意', '注意：', '注意:',
    '<如', '如：', '如:', '<同', '同>', '同）', '同)', '同 ',
    '<这里', '<这部分', '<说明', '<填写', '<根据', '<包括', '<以列', '<这一',
    '<这', '<接', '描述中', '文档中编号',
]


def is_blue_placeholder_text(text):
    if not text:
        return False
    s = text.strip()
    if not s:
        return False
    if s.startswith('<'):
        return True
    if s.startswith('注：') or s.startswith('注:'):
        return True
    for kw in PLACEHOLDER_KEYWORDS:
        if kw in s:
            return True
    return False


def write_paragraph(p, text):
    clear_paragraph(p)
    if '\n' in text:
        _write_with_breaks(p, text)
    else:
        p.runs[0].text = text
    set_black(p)


def write_as_normal(p, text):
    clear_paragraph(p)
    if '\n' in text:
        _write_with_breaks(p, text)
    else:
        p.runs[0].text = text
    set_black(p)
    try:
        normal_style = p.part.document.styles['Normal']
        p.style = normal_style
    except Exception:
        pass


def _write_with_breaks(p, text):
    raw_lines = text.split('\n')
    segments = []
    current = []
    for line in raw_lines:
        if line == '':
            if current:
                segments.append('\n'.join(current))
                current = []
        else:
            current.append(line)
    if current:
        segments.append('\n'.join(current))

    first_segment = segments[0]
    sub_lines = first_segment.split('\n')
    p.runs[0].text = sub_lines[0]
    first_run_elem = p.runs[0]._element
    insert_after = first_run_elem

    for line in sub_lines[1:]:
        br = OxmlElement('w:br')
        insert_after.addnext(br)
        new_run = OxmlElement('w:r')
        rPr = first_run_elem.find(qn('w:rPr'))
        if rPr is not None:
            new_run.append(deepcopy(rPr))
        new_t = OxmlElement('w:t')
        new_t.set(qn('xml:space'), 'preserve')
        new_t.text = line
        new_run.append(new_t)
        br.addnext(new_run)
        insert_after = new_run

    if len(segments) <= 1:
        return

    parent = p._element.getparent()
    insert_pos = list(parent).index(p._element) + 1
    pPr = p._element.find(qn('w:pPr'))
    pPr_template = deepcopy(pPr) if pPr is not None else None

    for seg in segments[1:]:
        new_p = OxmlElement('w:p')
        if pPr_template is not None:
            new_p.append(deepcopy(pPr_template))
        sub_lines = seg.split('\n')
        first_run = OxmlElement('w:r')
        rPr_orig = first_run_elem.find(qn('w:rPr'))
        if rPr_orig is not None:
            first_run.append(deepcopy(rPr_orig))
        t_elem = OxmlElement('w:t')
        t_elem.set(qn('xml:space'), 'preserve')
        t_elem.text = sub_lines[0]
        first_run.append(t_elem)
        new_p.append(first_run)
        for line in sub_lines[1:]:
            br = OxmlElement('w:br')
            new_p.append(br)
            r = OxmlElement('w:r')
            if rPr_orig is not None:
                r.append(deepcopy(rPr_orig))
            t = OxmlElement('w:t')
            t.set(qn('xml:space'), 'preserve')
            t.text = line
            r.append(t)
            new_p.append(r)
        parent.insert(insert_pos, new_p)
        insert_pos += 1


def content_hash(text):
    import hashlib
    return hashlib.md5(text.encode('utf-8')).hexdigest()


def write_with_dedup(p, text, content_hash_cache):
    h = content_hash(text)
    if h in content_hash_cache:
        clear_paragraph(p)
        return False
    content_hash_cache.add(h)
    write_paragraph(p, text)
    return True


def write_as_normal_with_dedup(p, text, content_hash_cache):
    h = content_hash(text)
    if h in content_hash_cache:
        clear_paragraph(p)
        return False
    content_hash_cache.add(h)
    write_as_normal(p, text)
    return True


def add_paragraph_text(p, text):
    if not p.runs:
        run = p.add_run(text)
        run.font.color.rgb = BLACK
    else:
        for r in p.runs[1:]:
            r.text = ''
        p.runs[0].text = text
    set_black(p)


def fill_cell_text(cell, text):
    from docx.shared import Pt
    for r in cell.paragraphs[0].runs:
        r.text = ''
    if cell.paragraphs[0].runs:
        cell.paragraphs[0].runs[0].text = text
        cell.paragraphs[0].runs[0].font.color.rgb = BLACK
        cell.paragraphs[0].runs[0].font.size = Pt(TABLE_FONT_SIZE_PT)
    elif cell.paragraphs:
        run = cell.paragraphs[0].add_run(text)
        run.font.color.rgb = BLACK
        run.font.size = Pt(TABLE_FONT_SIZE_PT)


def add_table_row(table, cells_data):
    row = table.add_row()
    for i, text in enumerate(cells_data):
        if i < len(row.cells):
            fill_cell_text(row.cells[i], str(text))
    return row


def clear_blue_cell(cell):
    for para in cell.paragraphs:
        for r in para.runs:
            r.text = ''
    if cell.paragraphs and cell.paragraphs[0].runs:
        cell.paragraphs[0].runs[0].font.color.rgb = BLACK


def remove_paragraph(p):
    p._element.getparent().remove(p._element)


def get_heading_context(paragraphs, current_idx):
    h1 = h2 = h3 = ''
    current_p = paragraphs[current_idx]
    current_style = current_p.style.name if current_p.style else ''

    if current_style == 'Heading 3':
        h3 = current_p.text.strip()
    elif current_style == 'Heading 2':
        h2 = current_p.text.strip()
    elif current_style == 'Heading 1':
        h1 = current_p.text.strip()

    found_h2 = bool(h2)
    for i in range(current_idx - 1, -1, -1):
        p = paragraphs[i]
        style = p.style.name if p.style else ''
        if style == 'Heading 1':
            if not h1:
                h1 = p.text.strip()
            break
        elif style == 'Heading 2':
            if not h2:
                h2 = p.text.strip()
            found_h2 = True
        elif style == 'Heading 3' and not h3 and not found_h2:
            h3 = p.text.strip()
    return h1, h2, h3


def format_table_styled(table, header_bg='D9E2F3', alt_row_bg='F2F2F3', font_size_pt=10.5):
    from docx.shared import Pt
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    tbl = table._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)

    existing_borders = tblPr.find(qn('w:tblBorders'))
    if existing_borders is not None:
        tblPr.remove(existing_borders)
    tblBorders = OxmlElement('w:tblBorders')
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), '999999')
        tblBorders.append(border)
    tblPr.append(tblBorders)

    for ri, row in enumerate(table.rows):
        for ci, cell in enumerate(row.cells):
            tc = cell._tc
            tcPr = tc.find(qn('w:tcPr'))
            if tcPr is None:
                tcPr = OxmlElement('w:tcPr')
                tc.insert(0, tcPr)

            existing_shd = tcPr.find(qn('w:shd'))
            if existing_shd is not None:
                tcPr.remove(existing_shd)
            shd = OxmlElement('w:shd')
            shd.set(qn('w:val'), 'clear')
            shd.set(qn('w:color'), 'auto')
            if ri == 0:
                shd.set(qn('w:fill'), header_bg)
            elif ri % 2 == 0:
                shd.set(qn('w:fill'), alt_row_bg)
            else:
                shd.set(qn('w:fill'), 'FFFFFF')
            tcPr.append(shd)

            for para in cell.paragraphs:
                for run in para.runs:
                    if run.font.size is None or run.font.size != Pt(font_size_pt):
                        run.font.size = Pt(font_size_pt)
                    if ri == 0:
                        run.font.bold = True

            pPr = cell.paragraphs[0]._element.find(qn('w:pPr'))
            if pPr is None:
                pPr = OxmlElement('w:pPr')
                cell.paragraphs[0]._element.insert(0, pPr)
            existing_spacing = pPr.find(qn('w:spacing'))
            if existing_spacing is not None:
                pPr.remove(existing_spacing)
            spacing = OxmlElement('w:spacing')
            spacing.set(qn('w:before'), '40')
            spacing.set(qn('w:after'), '40')
            pPr.append(spacing)
