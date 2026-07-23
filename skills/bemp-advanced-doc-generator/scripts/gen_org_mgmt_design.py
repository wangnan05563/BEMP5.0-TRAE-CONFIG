# -*- coding: utf-8 -*-
"""
BEMP 河南农商银行"机构管理优化"详细设计文档生成器
基于实际代码实现内容生成 .docx 文档，包含目录、页码、页眉页脚。
"""
import os
from datetime import datetime
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

OUTPUT_DIR = r"d:\code\QJ\BEMP5.0DEV\aotutests-devtools\delivery\机构管理优化"
OUTPUT_FILE = os.path.join(OUTPUT_DIR, "机构管理优化-详细设计说明书-v1.docx")

# ===================== 样式工具函数 =====================

def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = tcPr.find(qn('w:tcBorders'))
    if tcBorders is None:
        tcBorders = OxmlElement('w:tcBorders')
        tcPr.append(tcBorders)
    for edge in ('top', 'left', 'bottom', 'right'):
        if edge in kwargs:
            tag = 'w:{}'.format(edge)
            element = tcBorders.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                tcBorders.append(element)
            element.set(qn('w:val'), 'single')
            element.set(qn('w:sz'), '4')
            element.set(qn('w:color'), '000000')


def set_cell_background(cell, color="D9D9D9"):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color)
    tcPr.append(shd)


def set_run_font(run, name="宋体", size=10.5, bold=False, color=None):
    run.font.name = name
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    rFonts.set(qn('w:eastAsia'), name)
    rFonts.set(qn('w:ascii'), name)
    rFonts.set(qn('w:hAnsi'), name)


def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        set_run_font(run, name="黑体", size=max(12, 18 - level * 2), bold=True, color="000000")
    return p


def add_para(doc, text, bold=False, indent=True, size=10.5):
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.first_line_indent = Cm(0.74)
    p.paragraph_format.line_spacing = 1.5
    run = p.add_run(text)
    set_run_font(run, name="宋体", size=size, bold=bold)
    return p


def add_bullet(doc, text, size=10.5):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.line_spacing = 1.5
    run = p.add_run(text)
    set_run_font(run, name="宋体", size=size)
    return p


def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr_cells[i].text = ""
        p = hdr_cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(header)
        set_run_font(run, name="宋体", size=10, bold=True)
        set_cell_background(hdr_cells[i])
        set_cell_border(hdr_cells[i])
        hdr_cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    for r_idx, row in enumerate(rows, 1):
        row_cells = table.rows[r_idx].cells
        for c_idx, val in enumerate(row):
            row_cells[c_idx].text = ""
            p = row_cells[c_idx].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(str(val))
            set_run_font(run, name="宋体", size=10)
            set_cell_border(row_cells[c_idx])
            row_cells[c_idx].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    if col_widths:
        for i, width in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(width)
    doc.add_paragraph()
    return table


def add_toc_field(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run()
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = 'TOC \\o "1-3" \\h \\z \\u'
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')
    fldChar3 = OxmlElement('w:t')
    fldChar3.text = '右键此处选择"更新域"以生成目录'
    fldChar4 = OxmlElement('w:fldChar')
    fldChar4.set(qn('w:fldCharType'), 'end')
    r_element = run._r
    r_element.append(fldChar1)
    r_element.append(instrText)
    r_element.append(fldChar2)
    r_element.append(fldChar3)
    r_element.append(fldChar4)


def add_page_break(doc):
    p = doc.add_paragraph()
    run = p.add_run()
    run.add_break(WD_BREAK.PAGE)


def setup_page_and_header_footer(doc):
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(3.0)
        section.right_margin = Cm(2.5)
        header = section.header
        if not header.paragraphs[0].runs:
            hp = header.paragraphs[0]
            hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = hp.add_run("河南农商票据交易管理平台 - 详细设计说明书")
            set_run_font(run, name="宋体", size=9, color="808080")
        footer = section.footer
        fp = footer.paragraphs[0]
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = fp.add_run()
        set_run_font(run, name="宋体", size=9, color="808080")
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')
        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = 'PAGE'
        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'end')
        run._r.append(fldChar1)
        run._r.append(instrText)
        run._r.append(fldChar2)


def inject_update_fields(doc):
    settings = doc.settings.element
    update_fields = settings.find(qn('w:updateFields'))
    if update_fields is None:
        update_fields = OxmlElement('w:updateFields')
        update_fields.set(qn('w:val'), 'true')
        settings.append(update_fields)


# ===================== 文档生成主流程 =====================

def build_design_doc():
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = '宋体'
    style.font.size = Pt(10.5)
    rPr = style.element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    rFonts.set(qn('w:eastAsia'), '宋体')

    setup_page_and_header_footer(doc)

    # ===== 封面 =====
    for _ in range(3):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("河南农商银行")
    set_run_font(run, name="黑体", size=26, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("新一代票据系统个性化开发")
    set_run_font(run, name="黑体", size=22, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("机构管理优化")
    set_run_font(run, name="黑体", size=22, bold=True)

    for _ in range(2):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("详 细 设 计 说 明 书")
    set_run_font(run, name="黑体", size=20, bold=True)

    for _ in range(6):
        doc.add_paragraph()

    info_table = doc.add_table(rows=4, cols=2)
    info_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    info_data = [
        ("项目名称", "BEMP5.0 河南农商银行个性化开发"),
        ("需求名称", "机构管理优化（机构管理/机构管理员管理/机构数据同步）"),
        ("文档版本", "V1.0"),
        ("编写日期", "2026-07-22"),
    ]
    for i, (k, v) in enumerate(info_data):
        cells = info_table.rows[i].cells
        cells[0].text = ""
        cells[1].text = ""
        p1 = cells[0].paragraphs[0]
        p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r1 = p1.add_run(k)
        set_run_font(r1, name="宋体", size=12, bold=True)
        p2 = cells[1].paragraphs[0]
        p2.alignment = WD_ALIGN_PARAGRAPH.LEFT
        r2 = p2.add_run(v)
        set_run_font(r2, name="宋体", size=12)
        for c in cells:
            set_cell_border(c)

    add_page_break(doc)

    # ===== 修订记录 =====
    add_heading(doc, "修订记录", level=1)
    add_table(doc,
        ["版本", "日期", "修订内容", "修订人"],
        [
            ["V1.0", "2026-07-22", "初始版本，覆盖机构管理优化全部11个功能点（F-01~F-11）的详细设计", "BEMP文档交付工程师"],
        ],
        col_widths=[2, 3, 8, 3]
    )
    add_page_break(doc)

    # ===== 目录 =====
    add_heading(doc, "目录", level=1)
    add_toc_field(doc)
    add_page_break(doc)

    # ===== 第一章 系统概述 =====
    add_heading(doc, "第一章 系统概述", level=1)

    add_heading(doc, "1.1 业务背景", level=2)
    add_para(doc, "当前新一代票据系统机构管理菜单为省行法人管理员权限，机构及机构管理员均仅支持逐条手工新增，无法批量操作；机构树弹出框无查询条件，无法快速定位目标机构；机构数据未与核心系统自动同步，可能存在与核心系统不一致；贴现业务缺乏账务机构校验，非账务机构可发起贴现业务，存在管控风险。")
    add_para(doc, "行方公司金融部于2026年6月18日提交《关于新一代票据系统机构管理及机构管理员功能变更的需求》（需求编号20260707），优先级P2，要求实现批量导入、机构树查询、机构数据同步与贴现校验四类能力。本次开发覆盖11个功能点（F-01至F-11），涉及系统管理子系统的机构管理、机构管理员管理模块，以及业务管理子系统的贴现业务模块。")

    add_heading(doc, "1.2 设计目标", level=2)
    add_table(doc,
        ["目标类型", "目标描述"],
        [
            ["业务目标", "减少机构及管理员手工维护工作量；增强机构树检索能力；贴现业务由合规账务机构发起"],
            ["技术目标", "实现Excel批量导入；机构数据与核心系统ODS自动同步；机构管理员与统一身份认证系统对接"],
            ["合规目标", "机构信息符合票交所机构代码管理规范（一码绑定多机构）；机构层级不超过4级；账务机构校验满足管控要求"],
            ["兼容目标", "原有机构管理、贴现业务、机构管理员管理功能不受影响"],
        ],
        col_widths=[3, 13]
    )

    add_heading(doc, "1.3 需求范围", level=2)
    add_table(doc,
        ["范围维度", "内容说明"],
        [
            ["功能范围", "11个业务功能点（F-01至F-11）+ 9个菜单页面机构树改造"],
            ["系统范围", "新一代票据管理系统 - 系统管理子系统、业务管理子系统"],
            ["数据范围", "新增PJGCS（机构参数表）、PJGGX（机构关系表）；机构管理员表状态扩展"],
            ["接口范围", "统一身份认证接口（POBM010304MessageConverter，交易码1004）；核心机构数据同步定时任务（参考BOPC010101MessageConverter）"],
            ["不在范围", "票交所直接报文交互（仅间接影响机构代码映射）；除贴现外的其他业务账务机构校验"],
        ],
        col_widths=[3, 13]
    )

    add_heading(doc, "1.4 相关方", level=2)
    add_table(doc,
        ["相关方", "职责或分工", "联系人"],
        [
            ["总行科技部", "管理和把控项目进度", "连亚伦"],
            ["总行公司金融部", "需求梳理", "宋绍凯"],
            ["公司方", "负责项目实施、测试、上线", "史文豪"],
            ["行方测试", "负责系统测试", "靳向韶、宋绍凯"],
        ],
        col_widths=[4, 8, 4]
    )

    add_heading(doc, "1.5 术语与缩略语", level=2)
    add_table(doc,
        ["术语/缩略语", "全称", "说明"],
        [
            ["BEMP", "Bill Exchange Management Platform", "票据交易管理平台"],
            ["CBS", "Core Banking System", "核心银行系统"],
            ["TUA", "统一身份认证系统", "行内统一身份认证平台"],
            ["ODS", "Operational Data Store", "操作数据存储"],
            ["PJGCS", "机构参数表", "存储核心系统同步的机构基础信息"],
            ["PJGGX", "机构关系表", "存储机构间业务关系信息"],
            ["ESB", "Enterprise Service Bus", "企业服务总线"],
            ["SHCHPE", "上海票据交易所", "票交所"],
        ],
        col_widths=[3, 5, 8]
    )

    add_page_break(doc)

    # ===== 第二章 系统架构设计 =====
    add_heading(doc, "第二章 系统架构设计", level=1)

    add_heading(doc, "2.1 整体架构", level=2)
    add_para(doc, "本系统采用前后端分离的B/S架构，后端基于SpringBoot微服务框架，前端基于Vue.js框架。机构管理优化需求在现有系统架构基础上进行增量开发，遵循BEMP个性化开发规范，通过银行级扩展包（ext-hnnxbank）实现功能个性化，不影响产品化基线代码。")
    add_para(doc, "整体架构分为五层：前端展示层（Vue.js组件）、后端控制层（个性化Controller）、业务逻辑层（Service实现）、数据访问层（DAO+MyBatis）、数据存储层（Oracle数据库+文件系统）。外部系统通过适配器层（MessageConverter）进行报文转换与集成。")

    add_heading(doc, "2.2 技术选型", level=2)
    add_table(doc,
        ["技术层级", "技术选型", "版本/说明"],
        [
            ["前端框架", "Vue.js 2.x + iView UI", "单页面应用，组件化开发"],
            ["后端框架", "SpringBoot + Spring MVC", "微服务架构，RESTful API"],
            ["ORM框架", "MyBatis", "XML映射，支持批量操作"],
            ["数据库", "Oracle", "PJGCS/PJGGX新增表"],
            ["缓存", "Redis", "会话管理、字典缓存"],
            ["分布式协调", "ZooKeeper", "服务注册与发现"],
            ["报文转换", "MessageConverter", "ESB2协议，XML报文"],
            ["文件交互", "Local File System", "CBS同步文件读取"],
        ],
        col_widths=[3, 5, 8]
    )

    add_heading(doc, "2.3 部署拓扑", level=2)
    add_para(doc, "部署拓扑说明：")
    add_bullet(doc, "前端服务：Node.js webpack-dev-server，端口8091，编译前端Vue组件")
    add_bullet(doc, "后端服务：SpringBoot应用，端口8010，提供RESTful API")
    add_bullet(doc, "Redis缓存：端口6379，提供会话管理和字典缓存")
    add_bullet(doc, "ZooKeeper：端口2181，提供服务注册与发现")
    add_bullet(doc, "Oracle数据库：端口1521，存储业务数据（PJGCS/PJGGX等表）")
    add_bullet(doc, "外部系统：核心系统CBS（通过ODS文件交互）、统一身份认证TUA（通过ESB2报文交互）")

    add_heading(doc, "2.4 个性化开发架构", level=2)
    add_para(doc, "BEMP系统采用产品化基线+银行级个性化的架构模式。河南农商银行的个性化代码位于banks/ext-hnnxbank扩展包中，通过类级别@RequestMapping前缀/hnnxbank/实现路由覆盖。当请求路径包含/hnnxbank/前缀时，优先匹配个性化Controller；缺失前缀时回退到产品化Controller，导致个性化校验逻辑无法触发。")
    add_para(doc, "关键设计决策：F-10主系统校验集成于F-01批量导入流程中，通过个性化Controller路径保证PjgcsBranchParamService.checkBranchStatus()被正确调用。BUG-001根因即为新增机构路径缺失/hnnxbank/前缀导致PJGCS校验未触发，已在前端代码中修复（branch.vue L1455-1458代码注释确认）。")

    add_page_break(doc)

    # ===== 第三章 功能设计 =====
    add_heading(doc, "第三章 功能设计", level=1)

    add_heading(doc, "3.1 功能点总览", level=2)
    add_para(doc, "本次机构管理优化共实现11个功能点（F-01至F-11），覆盖机构管理、机构管理员管理、机构数据同步和贴现业务校验四大业务模块。全部功能点均已实现并通过代码审查验证。")
    add_table(doc,
        ["功能编号", "功能点名称", "优先级", "实现类", "实现状态"],
        [
            ["F-01", "机构批量导入", "P0", "HnnxBankBranchController", "已实现"],
            ["F-02", "机构导入模板下载", "P1", "HnnxBankBranchController", "已实现"],
            ["F-03", "机构批量复制角色", "P0", "HnnxBankBranchController", "已实现"],
            ["F-04", "机构树查询增强", "P0", "HnnxBankBranchQueryController", "已实现"],
            ["F-05", "管理员批量新增", "P0", "HnnxbankBranchAdminController", "已实现"],
            ["F-06", "管理员批量删除", "P0", "HnnxbankBranchAdminController", "已实现"],
            ["F-07", "管理员模板下载", "P1", "HnnxbankBranchAdminController", "已实现"],
            ["F-08", "管理员批量复制角色", "P0", "HnnxbankBranchAdminController", "已实现"],
            ["F-09", "机构数据同步定时任务", "P0", "SyncPjgcsBranchParamJobService + SyncPjgxBranchRelationJobService", "已实现"],
            ["F-10", "新增机构主系统校验", "P0", "PjgcsBranchParamServiceImpl（集成于F-01）", "已实现"],
            ["F-11", "贴现业务账务机构校验", "P0", "HnnxBankDiscBillServiceImpl（集成于submitElecFlow）", "已实现"],
        ],
        col_widths=[2, 4, 1.5, 6, 2.5]
    )

    # F-01 机构批量导入
    add_heading(doc, "3.2 F-01 机构批量导入", level=2)
    add_heading(doc, "3.2.1 功能描述", level=3)
    add_para(doc, "菜单位置：【系统管理子系统】-【系统管理】-【机构管理】。在删除按钮后面、角色分配前面增加批量导入按钮。通过Excel文件批量新增机构信息，含完整校验规则。操作流程为：点击按钮-弹出导入窗口-选择Excel文件-服务端校验-写入数据库-返回导入结果。")
    add_heading(doc, "3.2.2 接口设计", level=3)
    add_table(doc,
        ["属性", "值"],
        [
            ["接口路径", "POST /hnnxbank/sm/auth/branch/branch/func_batchImportValidate"],
            ["批量上传路径", "POST /hnnxbank/sm/auth/branch/branch/func_batchImport"],
            ["新增机构路径", "POST /hnnxbank/sm/auth/branch/branch/func_addBranch（集成F-10校验）"],
            ["请求参数", "branchDtoList（JSON数组，12列字段）"],
            ["实现类", "HnnxBankBranchController"],
            ["文件格式", "xls/xlsx"],
            ["文件大小", "不超过2M（前端校验）"],
        ],
        col_widths=[4, 12]
    )
    add_heading(doc, "3.2.3 校验规则", level=3)
    add_para(doc, "F-01机构批量导入实现以下校验规则：")
    add_bullet(doc, "新增机构必须选择上级机构（BR-ORG-01）")
    add_bullet(doc, "机构号、机构名称、核算机构号、组织机构代码均不可重复（文件内去重+系统已有数据校验，BR-ORG-02）")
    add_bullet(doc, "票交所机构代码非必输项，可维护行内机构非票交所参与机构（BR-ORG-03）")
    add_bullet(doc, "一个票交所机构代码可绑定多个行内机构（BR-ORG-04）")
    add_bullet(doc, "机构层级包括总行在内最多为4级（calculateBranchLevel方法，level>=4时阻止新增，BR-ORG-05）")
    add_bullet(doc, "主系统机构树清单校验：每个导入机构都需校验在PJGCS表中存在且状态正常（集成F-10，BR-ORG-06/07）")
    add_bullet(doc, "父级为分行时，机构类型不能为空且只能为县域或城区（BR-ORG-12）")
    add_bullet(doc, "失败策略：整批回滚，错误信息逐行收集后统一返回（StringJoiner收集错误后统一抛出，BR-ORG-08）")
    add_heading(doc, "3.2.4 导入预览列定义", level=3)
    add_table(doc,
        ["序号", "列名", "必填", "说明"],
        [
            ["1", "序号", "N", "行号"],
            ["2", "机构名称(*)", "Y", "机构中文名称"],
            ["3", "系统内机构号(*)", "Y", "机构号，不超过10位数字或字母"],
            ["4", "上级机构号(*)", "Y", "上级机构号"],
            ["5", "核算机构号", "N", "不超过10位数字或字母"],
            ["6", "大额行号(*)", "Y", "12位数字"],
            ["7", "组织机构代码(*)", "Y", "8位+横线+1位格式"],
            ["8", "票交所机构代码", "N", "非必填"],
            ["9", "机构地址(*)", "Y", "机构地址"],
            ["10", "虚拟柜员号(*)", "Y", "不超过20位数字或字母"],
            ["11", "机构类型", "N", "JG01=县域, JG02=城区（父级为分行时必填）"],
            ["12", "是否简单机构", "N", "1=是, 0=否"],
        ],
        col_widths=[1.5, 4, 1.5, 9]
    )

    # F-02
    add_heading(doc, "3.3 F-02 机构导入模板下载", level=2)
    add_para(doc, "菜单位置：【系统管理-机构管理】，在批量导入按钮后面增加模板下载按钮。通过form提交方式下载机构信息导入模板，模板含12列字段定义，与导入预览列完全一致。")
    add_table(doc,
        ["属性", "值"],
        [
            ["接口路径", "POST /hnnxbank/sm/auth/branch/branch/func_downloadModel"],
            ["实现类", "HnnxBankBranchController"],
            ["输出文件", "机构信息导入模板"],
            ["提交方式", "form.submit()（非Ajax，触发浏览器下载）"],
        ],
        col_widths=[4, 12]
    )

    # F-03
    add_heading(doc, "3.4 F-03 机构批量复制角色", level=2)
    add_para(doc, "菜单位置：【系统管理-机构管理】，在角色分配后面增加批量复制角色按钮。将选中机构角色全量复制到目标机构。界面交互流程为：点击按钮-弹出批量复制角色界面-选择目标机构号-自动回显机构名称-点击确定-全量覆盖目标机构角色。")
    add_table(doc,
        ["属性", "值"],
        [
            ["接口路径", "POST /hnnxbank/sm/auth/branch/branch/func_batchCopyRole"],
            ["实现类", "HnnxBankBranchController"],
            ["请求参数", "brchNo（源机构号）+ targetBrchNos（目标机构号数组）"],
            ["目标机构选择", "支持多选（branchSearchTree配置multiple+checkStrictly）"],
            ["复制方式", "全量覆盖：清除目标机构原有角色，替换为源机构角色"],
            ["失败策略", "单条失败不影响其他机构，返回结构化结果（successCount、failCount、failReasons）"],
        ],
        col_widths=[4, 12]
    )

    # F-04
    add_heading(doc, "3.5 F-04 机构树查询增强", level=2)
    add_para(doc, "在9个菜单页面的机构树弹出框中增加机构号、机构级别、机构名称三个查询条件，支持AND组合查询、自动展开匹配节点、高亮显示和无结果提示。")
    add_table(doc,
        ["属性", "值"],
        [
            ["树查询接口", "POST /hnnxbank/sm/auth/branch/queryBranch"],
            ["分页查询接口", "POST /hnnxbank/sm/auth/branch/queryBranchPage"],
            ["实现类", "HnnxBankBranchQueryController"],
            ["查询条件", "机构号（精确匹配）+ 机构级别（下拉框，字典BranchLevel）+ 机构名称（模糊匹配）"],
            ["条件组合", "三条件任意组合查询（AND关系）"],
            ["查询结果", "自动展开机构树至匹配节点+高亮显示+无结果提示"],
        ],
        col_widths=[4, 12]
    )
    add_heading(doc, "3.5.1 涉及页面清单", level=3)
    add_table(doc,
        ["序号", "所属子系统", "菜单路径", "改造内容"],
        [
            ["1", "系统管理子系统", "系统管理-机构管理", "机构树增加搜索框"],
            ["2", "系统管理子系统", "系统管理-机构管理员管理", "机构选择树增加搜索框"],
            ["3", "系统管理子系统", "系统管理-权属机构维护", "机构树增加搜索框"],
            ["4", "业务管理子系统", "票据准入管理", "机构名称选择树增加搜索框"],
            ["5", "业务管理子系统", "机构内部账户管理-机构内部账户维护", "机构名称选择树增加搜索框"],
            ["6", "业务管理子系统", "审批管理-审批-定义审批路线", "机构树选择增加搜索框"],
            ["7", "业务管理子系统", "审批管理-审批-审批机构产品管理", "机构树选择增加搜索框"],
            ["8", "业务管理子系统", "公共查询-余额管理", "机构名称支持搜索"],
            ["9", "业务管理子系统", "公共查询-业务发生查询", "机构名称支持搜索"],
        ],
        col_widths=[1.5, 4, 6, 4.5]
    )

    # F-05
    add_heading(doc, "3.6 F-05 管理员批量新增", level=2)
    add_para(doc, "菜单位置：【系统管理-机构管理员管理】，在删除按钮后面增加批量新增按钮。通过Excel文件批量新增机构管理员，含身份认证校验。新增管理员默认状态为无效(0)，首次登录修改密码后变为有效(1)。")
    add_table(doc,
        ["属性", "值"],
        [
            ["导入校验路径", "POST /hnnxbank/sm/auth/branch/branchAdmin/func_batchImportValidate"],
            ["导入上传路径", "POST /hnnxbank/sm/auth/branch/branchAdmin/func_batchImportBranchAdmin"],
            ["实现类", "HnnxbankBranchAdminController"],
            ["请求参数", "userDtoList（JSON数组，9列字段）"],
            ["权限控制", "仅法人管理员（userType=4）可操作"],
            ["初始密码", "从系统参数branch_admin_init_pwd读取"],
            ["身份认证", "调用TUA系统验证（POBM010304MessageConverter，交易码1004）"],
            ["失败策略", "单条失败不影响其他用户新增，返回结构化结果"],
        ],
        col_widths=[4, 12]
    )
    add_heading(doc, "3.6.1 管理员新增模板字段（9列）", level=3)
    add_table(doc,
        ["序号", "列名", "必填", "说明"],
        [
            ["1", "序号", "N", "行号"],
            ["2", "用户号(*)", "Y", "不超过20位数字或字母"],
            ["3", "姓名(*)", "Y", "不超过16位"],
            ["4", "机构号(*)", "Y", "管理员所属机构号"],
            ["5", "机构名称", "N", "自动回显"],
            ["6", "电话", "N", "数字或'-'组成"],
            ["7", "证件类型", "N", "DC01=居民身份证"],
            ["8", "证件号码", "N", "证件号码"],
            ["9", "备注", "N", "备注信息"],
        ],
        col_widths=[1.5, 4, 1.5, 9]
    )

    # F-06
    add_heading(doc, "3.7 F-06 管理员批量删除", level=2)
    add_para(doc, "菜单位置：【系统管理-机构管理员管理】，在批量新增按钮后面增加批量删除按钮。通过Excel文件批量删除机构管理员，采用逻辑删除方式，含二次确认弹窗和跨法人校验。")
    add_table(doc,
        ["属性", "值"],
        [
            ["接口路径", "POST /hnnxbank/sm/auth/branch/branchAdmin/func_batchDeleteBranchAdmin"],
            ["实现类", "HnnxbankBranchAdminController"],
            ["上传方式", "XMLHttpRequest（非axios，避免拦截器对FormData的处理）"],
            ["删除方式", "逻辑删除（置为无效状态），非物理删除"],
            ["权限控制", "仅法人管理员（userType=4）可操作"],
            ["跨法人校验", "用户必须属于当前法人，防止跨法人误删"],
            ["自删除保护", "法人管理员不可删除自身"],
            ["二次确认", "已实现（batchDeleteConfirm: 确定要批量删除机构管理员吗？）"],
            ["失败策略", "单条失败不影响其他用户删除，返回结构化结果"],
        ],
        col_widths=[4, 12]
    )

    # F-07
    add_heading(doc, "3.8 F-07 管理员模板下载", level=2)
    add_para(doc, "菜单位置：【系统管理-机构管理员管理】，下拉菜单形式含新增模板和删除模板两个选项。通过operateType参数区分新增/删除模板。")
    add_table(doc,
        ["属性", "值"],
        [
            ["接口路径", "POST /hnnxbank/sm/auth/branch/branchAdmin/func_downloadModel"],
            ["实现类", "HnnxbankBranchAdminController"],
            ["新增模板", "机构管理员新增模版.xls（9列字段）"],
            ["删除模板", "机构管理员删除模板.xls"],
            ["参数区分", "operateType: new（新增模板）/ delete（删除模板）"],
        ],
        col_widths=[4, 12]
    )

    # F-08
    add_heading(doc, "3.9 F-08 管理员批量复制角色", level=2)
    add_para(doc, "菜单位置：【系统管理-机构管理员管理】，在角色分配后面增加批量复制角色按钮。将源用户角色全量复制到目标管理员，含二次确认弹窗。")
    add_table(doc,
        ["属性", "值"],
        [
            ["接口路径", "POST /hnnxbank/sm/auth/branch/branchAdmin/func_batchCopyRole"],
            ["实现类", "HnnxbankBranchAdminController"],
            ["请求参数", "sourceUserNo（源用户号）+ targetUserNos（目标用户号数组）"],
            ["目标用户类型", "必须为机构管理员（userType=2）"],
            ["复制方式", "全量覆盖：用新角色覆盖旧角色"],
            ["二次确认", "已实现（confirmBatchCopyRole: 确定要批量复制角色吗？这将覆盖目标用户的原有角色。）"],
            ["失败策略", "单条失败不影响其他目标用户，返回结构化结果"],
        ],
        col_widths=[4, 12]
    )

    # F-09
    add_heading(doc, "3.10 F-09 机构数据同步定时任务", level=2)
    add_para(doc, "后台定时任务，每日6:00读取核心系统ODS提供的CBS_PJGCS/CBS_PJGGX文件，分别写入PJGCS和PJGGX表。采用全量替换策略（deleteAll+insertBatch），使用@Transactional保证原子性。")
    add_table(doc,
        ["任务编号", "任务名称", "CloudFunction", "触发时间", "功能", "实现类"],
        [
            ["JOB-01a", "PJGCS机构参数同步", "HNNXTK020111", "每日6:00", "读取CBS_PJGCS文件，写入PJGCS表", "SyncPjgcsBranchParamJobService"],
            ["JOB-01b", "PJGGX机构关系同步", "HNNXTK020112", "每日6:00", "读取CBS_PJGGX文件，写入PJGGX表", "SyncPjgxBranchRelationJobService"],
        ],
        col_widths=[2, 3, 3, 2, 4, 4]
    )
    add_heading(doc, "3.10.1 同步规则", level=3)
    add_bullet(doc, "文件来源：核心系统ODS日初提供CBS_PJGCS_YYYYMMDD.txt和CBS_PJGGX_YYYYMMDD.txt")
    add_bullet(doc, "文件路径：{cbsFilePath}{date}/CBS_PJGCS_{date}.txt，日期取票据系统营业日期前一天")
    add_bullet(doc, "文件格式：Tab分隔文本（实际分隔符0x03 ETX），无表头")
    add_bullet(doc, "数据处理策略：全量替换（先deleteAll再insertBatch），@Transactional保证原子性")
    add_bullet(doc, "记录状态过滤：仅同步JILUZT='0'的有效记录")
    add_bullet(doc, "失败处理：记录失败日志，不影响票据系统正常运行，保留前一次数据")
    add_bullet(doc, "读取逻辑参考：BOPC010101MessageConverter（CommonFileTransfer.readFileFormLoaclPath()）")

    # F-10
    add_heading(doc, "3.11 F-10 新增机构主系统校验", level=2)
    add_para(doc, "校验时机为新增机构时（含手工新增和批量导入），集成于F-01批量导入校验流程中，非独立接口暴露。实现类为PjgcsBranchParamServiceImpl。")
    add_para(doc, "核心规则：新建机构必须在PJGCS（主系统机构树清单）表中存在记录且状态正常（JILUZT='0'）。校验失败时阻止新增并返回详细错误信息：")
    add_bullet(doc, "机构不存在：机构[XXX]在核心系统机构树中不存在，不能创建")
    add_bullet(doc, "机构状态异常：机构[XXX]在核心系统状态异常，不能创建")
    add_para(doc, "实现方法：isBranchExistsInMainSystem(String brchNo)返回boolean；checkBranchStatus(String brchNo)返回详细错误信息（null表示正常）。")
    add_para(doc, "设计决策说明：F-10集成于F-01校验流程而非独立接口，是因为主系统校验是机构新增的前置必要条件，与批量导入校验逻辑紧密耦合。通过个性化Controller路径/hnnxbank/确保PJGCS校验被正确触发（BUG-001根因已修复）。")

    # F-11
    add_heading(doc, "3.12 F-11 贴现业务账务机构校验", level=2)
    add_para(doc, "校验时机为客户经理发起贴现业务申请、点击提交时（仅此一个业务场景校验），集成于submitElecFlow方法中，非独立接口暴露。实现类为HnnxBankDiscBillServiceImpl。")
    add_para(doc, "核心规则：电票系统校验业务发起机构是否为账务机构。账务机构判定通过PJGGX表中YWGXZL='ZNGWSJ'且JILUZT='0'判断。校验失败时阻止提交，抛出异常：业务发起机构非账务机构，无法发起贴现。")
    add_para(doc, "扩展性设计：通过PjgxBranchRelationService接口抽象，便于后续扩展至其他业务场景。当前校验范围仅限贴现业务，其他业务暂不校验。")

    add_page_break(doc)

    # ===== 第四章 接口设计 =====
    add_heading(doc, "第四章 接口设计", level=1)

    add_heading(doc, "4.1 接口总览", level=2)
    add_para(doc, "本次机构管理优化共涉及8个独立RESTful接口、2个集成实现和2个定时任务。所有个性化接口均包含/hnnxbank/前缀，确保路由优先匹配个性化Controller。")
    add_table(doc,
        ["接口编号", "接口名称", "接口路径", "请求方式", "所属功能点", "实现类"],
        [
            ["API-01", "机构批量导入校验", "/hnnxbank/sm/auth/branch/branch/func_batchImportValidate", "POST", "F-01", "HnnxBankBranchController"],
            ["API-02", "机构导入模板下载", "/hnnxbank/sm/auth/branch/branch/func_downloadModel", "POST", "F-02", "HnnxBankBranchController"],
            ["API-03", "机构批量复制角色", "/hnnxbank/sm/auth/branch/branch/func_batchCopyRole", "POST", "F-03", "HnnxBankBranchController"],
            ["API-04", "机构树查询", "/hnnxbank/sm/auth/branch/queryBranch + queryBranchPage", "POST", "F-04", "HnnxBankBranchQueryController"],
            ["API-05", "管理员批量新增", "/hnnxbank/sm/auth/branch/branchAdmin/func_batchImportBranchAdmin", "POST", "F-05", "HnnxbankBranchAdminController"],
            ["API-06", "管理员批量删除", "/hnnxbank/sm/auth/branch/branchAdmin/func_batchDeleteBranchAdmin", "POST", "F-06", "HnnxbankBranchAdminController"],
            ["API-07", "管理员模板下载", "/hnnxbank/sm/auth/branch/branchAdmin/func_downloadModel", "POST", "F-07", "HnnxbankBranchAdminController"],
            ["API-08", "管理员批量复制角色", "/hnnxbank/sm/auth/branch/branchAdmin/func_batchCopyRole", "POST", "F-08", "HnnxbankBranchAdminController"],
            ["API-09", "主系统机构校验", "集成于func_batchImportValidate校验流程", "-", "F-10", "PjgcsBranchParamServiceImpl"],
            ["API-10", "贴现账务机构校验", "集成于submitElecFlow方法", "-", "F-11", "HnnxBankDiscBillServiceImpl"],
        ],
        col_widths=[1.5, 3, 6, 1.5, 1.5, 4]
    )

    add_heading(doc, "4.2 批量操作返回结构", level=2)
    add_para(doc, "F-03/F-05/F-06/F-08四个批量操作接口采用统一的返回结构，支持单条失败不影响其他，返回结构化结果：")
    add_table(doc,
        ["字段名", "数据类型", "说明"],
        [
            ["retCode", "String", "返回码，000000表示整体成功"],
            ["retMsg", "String", "返回信息描述"],
            ["successCount", "Integer", "成功处理条数"],
            ["failCount", "Integer", "失败条数"],
            ["failReasons", "List<String>", "失败原因列表，每条对应一个失败记录"],
        ],
        col_widths=[3, 3, 10]
    )

    add_heading(doc, "4.3 外部系统接口", level=2)
    add_heading(doc, "4.3.1 统一身份认证接口（TUA）", level=3)
    add_table(doc,
        ["属性", "值"],
        [
            ["实现类", "POBM010304MessageConverter.java"],
            ["调用方式", "同步调用（基于ESB2）"],
            ["交易码", "trans_code = 1004"],
            ["报文协议", "XML（transaction -> body -> request）"],
            ["请求字段", "userNo、opUserNo、opBranchNo、txCd、chan"],
            ["响应字段", "legalNo、userNo、userName、brchNo、userType、isEnable"],
            ["用户状态映射", "userStat=0 -> isEnable=1（正常）；userStat=2 -> isEnable=2（无效）；其他 -> isEnable=0（锁定）"],
            ["失败处理", "retCode != 000000表示失败，retMsg返回失败原因"],
            ["性能要求", "单条验证响应时间不超过3秒"],
        ],
        col_widths=[4, 12]
    )

    add_heading(doc, "4.3.2 核心系统ODS文件交互", level=3)
    add_table(doc,
        ["属性", "值"],
        [
            ["实现类", "BOPC010101MessageConverter.java（参考）"],
            ["调用方式", "定时任务读取本地文件"],
            ["配置项", "cbsFile.cbsFilePath（CBS文件根目录）"],
            ["日期获取", "commomQueryService.beforeDate()（票据系统营业日期前一天）"],
            ["文件路径", "{cbsFilePath}{date}/CBS_PJGCS_{date}.txt、CBS_PJGGX_{date}.txt"],
            ["文件分隔符", "0x03（ETX），样本文件为Tab分隔"],
            ["文件读取", "CommonFileTransfer.readFileFormLoaclPath()"],
            ["触发时间", "每日6:00"],
        ],
        col_widths=[4, 12]
    )

    add_page_break(doc)

    # ===== 第五章 数据库设计 =====
    add_heading(doc, "第五章 数据库设计", level=1)

    add_heading(doc, "5.1 新增表：PJGCS（机构参数表）", level=2)
    add_para(doc, "PJGCS表存储核心系统同步的机构基础信息，由定时任务HNNXTK020111每日6:00从CBS_PJGCS文件同步。主键为YNGYJG（营业机构号）。字段定义与代码实体类PjgcsBranchParam.java及DAO XML映射完全一致。")
    add_table(doc,
        ["序号", "字段代码", "字段名称", "数据类型", "长度", "必填", "说明"],
        [
            ["1", "FAREDM", "法人代码", "VARCHAR2", "4", "Y", ""],
            ["2", "YNGYJG", "营业机构号", "VARCHAR2", "10", "Y", "主键"],
            ["3", "FENHDM", "分行代码", "VARCHAR2", "3", "N", ""],
            ["4", "JIGOLX", "机构类型", "VARCHAR2", "1", "N", "样本取值2/3/4/5（含义待确认Q-13）"],
            ["5", "JIGOMC", "机构中文名称", "VARCHAR2", "60", "Y", ""],
            ["6", "JIGOJC", "机构简称", "VARCHAR2", "20", "Y", ""],
            ["7", "JGYWMC", "机构英文名称", "VARCHAR2", "60", "N", ""],
            ["8", "JIGOJP", "机构简拼", "VARCHAR2", "10", "N", ""],
            ["9", "QSZXDH", "地区代号", "VARCHAR2", "4", "N", ""],
            ["10", "SHQUDH", "省区代号", "VARCHAR2", "10", "N", ""],
            ["11", "DIZHII", "地址", "VARCHAR2", "200", "N", ""],
            ["12", "DIANHH", "电话号码", "VARCHAR2", "20", "N", ""],
            ["13", "LNXIRM", "联系人", "VARCHAR2", "200", "N", ""],
            ["14", "LNXRDH", "联系人电话", "VARCHAR2", "25", "N", ""],
            ["15", "DZYJDZ", "E-mail", "VARCHAR2", "42", "N", ""],
            ["16", "WANGZH", "网址", "VARCHAR2", "62", "N", ""],
            ["17", "SFCZYH", "是否村镇银行", "VARCHAR2", "1", "N", ""],
            ["18", "JIBIEE", "行政级别", "VARCHAR2", "1", "N", ""],
            ["19", "QYNGRQ", "启用日期", "VARCHAR2", "8", "N", ""],
            ["20", "BYZD01", "备用字段", "VARCHAR2", "200", "N", ""],
            ["21", "WEIHRQ", "维护日期", "VARCHAR2", "8", "N", ""],
            ["22", "WEIHGY", "维护柜员", "VARCHAR2", "8", "N", ""],
            ["23", "WEIHJG", "维护机构", "VARCHAR2", "10", "N", ""],
            ["24", "JILUZT", "记录状态", "VARCHAR2", "1", "Y", "0-正常（已验证H-06）"],
        ],
        col_widths=[1, 2, 2.5, 2, 1, 1, 6.5]
    )

    add_heading(doc, "5.2 新增表：PJGGX（机构关系表）", level=2)
    add_para(doc, "PJGGX表存储机构间业务关系信息，由定时任务HNNXTK020112每日6:00从CBS_PJGGX文件同步。主键为YNGYJG + YWGXZL + FAREDM三字段联合主键。字段定义与代码实体类PjgxBranchRelation.java及DAO XML映射完全一致。")
    add_table(doc,
        ["序号", "字段代码", "字段名称", "数据类型", "长度", "必填", "说明"],
        [
            ["1", "FAREDM", "法人代码", "VARCHAR2", "4", "Y", "主键之一"],
            ["2", "YNGYJG", "营业机构号", "VARCHAR2", "10", "Y", "主键之一"],
            ["3", "YWGXZL", "业务关系种类", "VARCHAR2", "10", "Y", "主键之一；ZNGWSJ-账务机构关系（已验证H-07）"],
            ["4", "BIZHON", "币种", "VARCHAR2", "2", "N", ""],
            ["5", "YWGXJG", "业务关系机构", "VARCHAR2", "10", "N", ""],
            ["6", "YWGXJB", "业务关系级别", "VARCHAR2", "1", "N", ""],
            ["7", "GXQXJG", "关系权限机构", "VARCHAR2", "10", "N", ""],
            ["8", "SHMING", "说明信息", "VARCHAR2", "128", "N", ""],
            ["9", "BYZD01", "备用字段", "VARCHAR2", "200", "N", ""],
            ["10", "WEIHRQ", "维护日期", "VARCHAR2", "8", "N", ""],
            ["11", "WEIHGY", "维护柜员", "VARCHAR2", "8", "N", ""],
            ["12", "WEIHJG", "维护机构", "VARCHAR2", "10", "N", ""],
            ["13", "JILUZT", "记录状态", "VARCHAR2", "1", "Y", "0-正常（已验证H-06）"],
        ],
        col_widths=[1, 2, 2.5, 2, 1, 1, 6.5]
    )

    add_heading(doc, "5.3 数据访问层设计", level=2)
    add_table(doc,
        ["DAO名称", "提供方法", "说明"],
        [
            ["PjgcsBranchParamDao", "insertBatch、deleteAll、countByBrchNoAndStatus、selectByBrchNo", "PJGCS表数据访问"],
            ["PjgxBranchRelationDao", "insertBatch、deleteAll、countByBrchNoAndGxzLAndStatus", "PJGGX表数据访问"],
        ],
        col_widths=[4, 7, 5]
    )

    add_page_break(doc)

    # ===== 第六章 业务规则与约束 =====
    add_heading(doc, "第六章 业务规则与约束", level=1)

    add_heading(doc, "6.1 机构管理规则", level=2)
    add_table(doc,
        ["规则编号", "规则名称", "规则描述", "来源"],
        [
            ["BR-ORG-01", "上级机构必填", "新增机构必须选择上级机构", "行方原文"],
            ["BR-ORG-02", "字段唯一性", "机构号、机构名称、核算机构号、组织机构代码均不可重复", "行方原文"],
            ["BR-ORG-03", "票交所机构代码非必填", "可维护行内机构非票交所参与机构", "行方原文"],
            ["BR-ORG-04", "票交所机构代码一对多", "一个票交所机构代码可绑定多个行内机构", "行方原文+票交所规范"],
            ["BR-ORG-05", "机构层级上限", "机构层级包括总行在内最多为4级", "行方原文"],
            ["BR-ORG-06", "主系统存在性", "新建机构必须在PJGCS表中存在记录", "行方原文"],
            ["BR-ORG-07", "机构状态校验", "新建机构时校验JILUZT='0'为正常", "行方原文+已验证H-06"],
            ["BR-ORG-08", "失败回滚", "F-01导入校验阶段整批回滚（StringJoiner收集错误）", "代码核查确认"],
            ["BR-ORG-09", "角色复制覆盖", "全量覆盖目标机构原有角色", "行方原文"],
            ["BR-ORG-10", "角色使用校验", "角色被柜员使用时报错不可删除", "行方原文"],
            ["BR-ORG-11", "角色复制失败策略", "单条失败不影响其他机构，返回结构化结果", "v3代码核查修订"],
            ["BR-ORG-12", "机构类型校验", "父级为分行时机构类型必填且只能为县域或城区", "v3代码核查新增"],
        ],
        col_widths=[2, 3, 7, 4]
    )

    add_heading(doc, "6.2 机构管理员规则", level=2)
    add_table(doc,
        ["规则编号", "规则名称", "规则描述", "来源"],
        [
            ["BR-ADM-01", "批量新增权限", "仅法人管理员（userType=4）可操作", "行方原文"],
            ["BR-ADM-02", "批量删除权限", "仅法人管理员（userType=4）可操作", "行方原文"],
            ["BR-ADM-03", "角色分配权限", "法人管理员可为机构管理员分配角色", "行方原文"],
            ["BR-ADM-04", "初始密码来源", "从系统参数branch_admin_init_pwd读取", "行方原文"],
            ["BR-ADM-05", "密码有效期", "在系统参数设置中统一维护", "行方原文"],
            ["BR-ADM-06", "密码错误次数", "在系统参数设置中统一维护", "行方原文"],
            ["BR-ADM-07", "默认状态", "新增管理员默认为无效状态(0)", "行方原文"],
            ["BR-ADM-08", "首次登录激活", "初次登录修改密码后变为有效(1)", "行方原文"],
            ["BR-ADM-09", "身份认证", "新增管理员时向TUA请求验证", "行方原文"],
            ["BR-ADM-10", "身份认证失败策略", "单条失败不影响其他，返回结构化结果", "v3代码核查修订"],
            ["BR-ADM-11", "删除方式", "逻辑删除（置为无效），非物理删除", "代码核查确认"],
            ["BR-ADM-12", "跨法人校验", "用户必须属于当前法人，防止跨法人误删", "v3代码核查新增"],
            ["BR-ADM-13", "自删除保护", "法人管理员不可删除自身", "代码核查确认"],
            ["BR-ADM-14", "管理员号唯一性", "管理员号不可与系统已有管理员重复", "代码核查确认"],
            ["BR-ADM-15", "所属机构校验", "管理员所属机构必须在系统中已存在", "代码核查确认"],
            ["BR-ADM-16", "角色归属校验", "只能分配管理员所在机构所拥有的角色", "行方原文"],
            ["BR-ADM-17", "机构不具备角色报错", "机构不具备的角色报错", "行方原文"],
            ["BR-ADM-18", "目标用户类型校验", "目标用户必须为机构管理员（userType=2）", "v3代码核查新增"],
            ["BR-ADM-19", "角色复制失败策略", "单条失败不影响其他，返回结构化结果", "v3代码核查修订"],
        ],
        col_widths=[2, 3, 7, 4]
    )

    add_heading(doc, "6.3 机构数据同步规则", level=2)
    add_table(doc,
        ["规则编号", "规则名称", "规则描述", "来源"],
        [
            ["BR-SYNC-01", "文件来源", "核心系统ODS日初提供CBS_PJGCS/CBS_PJGGX文件", "行方原文"],
            ["BR-SYNC-02", "文件路径", "{cbsFilePath}{date}/CBS_PJGCS_{date}.txt", "行方原文"],
            ["BR-SYNC-03", "同步方式", "2个独立定时任务分别读取文件写入PJGCS/PJGGX表", "v3代码核查修订"],
            ["BR-SYNC-06", "文件分隔符", "0x03（ETX），样本为Tab分隔文本无表头", "代码扫描验证"],
            ["BR-SYNC-07", "同步时间", "每日6:00执行", "代码注释确认"],
            ["BR-SYNC-08", "数据处理策略", "全量替换（deleteAll+insertBatch，@Transactional）", "代码核查确认"],
            ["BR-SYNC-09", "失败处理", "记录日志，保留前一次数据", "代码核查确认"],
            ["BR-SYNC-10", "记录状态过滤", "仅同步JILUZT='0'的有效记录", "已验证H-06"],
            ["BR-SYNC-11", "CloudFunction编号", "PJGCS=HNNXTK020111, PJGGX=HNNXTK020112", "v3代码核查新增"],
        ],
        col_widths=[2, 3, 7, 4]
    )

    add_heading(doc, "6.4 贴现业务校验规则", level=2)
    add_table(doc,
        ["规则编号", "规则名称", "规则描述", "来源"],
        [
            ["BR-DISC-01", "校验时机", "客户经理发起贴现申请点击提交时（submitElecFlow）", "行方原文"],
            ["BR-DISC-02", "校验范围", "仅贴现业务场景校验，其他业务暂不校验", "行方原文"],
            ["BR-DISC-03", "账务机构判定", "PJGGX表中YWGXZL='ZNGWSJ'且JILUZT='0'", "行方原文+已验证H-07"],
            ["BR-DISC-04", "机构状态校验", "PJGCS表中JILUZT='0'", "已验证H-06"],
            ["BR-DISC-05", "校验失败处理", "阻止提交，抛出异常", "行方原文"],
            ["BR-DISC-06", "扩展性", "PjgxBranchRelationService接口抽象，便于扩展", "已实现H-11"],
        ],
        col_widths=[2, 3, 7, 4]
    )

    add_heading(doc, "6.5 非功能需求规则", level=2)
    add_table(doc,
        ["规则编号", "规则名称", "规则描述"],
        [
            ["BR-NFR-01", "批量导入响应时间", "单次导入100条以内，处理时间不超过30秒"],
            ["BR-NFR-02", "机构树查询响应时间", "查询条件输入后，结果展示不超过2秒"],
            ["BR-NFR-03", "批量复制角色响应时间", "单个目标机构操作不超过5秒"],
            ["BR-NFR-04", "数据同步时间", "日初同步任务执行时间不超过10分钟"],
            ["BR-NFR-05", "统一身份认证响应时间", "单条验证响应时间不超过3秒"],
            ["BR-NFR-06", "操作审计", "所有批量操作需记录操作日志"],
            ["BR-NFR-07", "文件安全", "导入文件需校验格式和内容"],
            ["BR-NFR-08", "并发控制", "同一机构/管理员被并发修改时加锁保护"],
        ],
        col_widths=[3, 4, 9]
    )

    add_page_break(doc)

    # ===== 第七章 与现有系统关联 =====
    add_heading(doc, "第七章 与现有系统关联", level=1)

    add_heading(doc, "7.1 与核心系统（CBS）的关联", level=2)
    add_table(doc,
        ["关联点", "关联说明", "影响范围"],
        [
            ["机构数据同步", "通过ODS日初提供的CBS_PJGCS/CBS_PJGGX文件同步机构数据", "F-09、F-10"],
            ["文件传输", "文件路径由cbsFile.cbsFilePath配置，日期取营业日期前一天", "F-09"],
            ["文件格式", "Tab分隔文本（实际分隔符0x03），无表头", "F-09"],
            ["机构状态同步", "同步JILUZT='0'的有效记录", "F-09、F-10"],
            ["账务机构关系", "通过YWGXZL='ZNGWSJ'标识账务机构", "F-11"],
        ],
        col_widths=[3, 8, 5]
    )

    add_heading(doc, "7.2 与统一身份认证系统（TUA）的关联", level=2)
    add_table(doc,
        ["关联点", "关联说明", "影响范围"],
        [
            ["用户身份验证", "新增机构管理员时调用TUA接口验证身份信息", "F-05"],
            ["接口实现", "参考POBM010304MessageConverter（交易码1004）", "F-05"],
            ["调用方式", "同步调用，循环逐条验证", "F-05"],
            ["性能要求", "单条验证响应时间不超过3秒", "F-05"],
        ],
        col_widths=[3, 8, 5]
    )

    add_heading(doc, "7.3 与票交所（SHCHPE）的关联", level=2)
    add_table(doc,
        ["关联点", "关联说明", "影响范围"],
        [
            ["票交所机构代码映射", "一个票交所机构代码可绑定多个行内机构", "F-01、F-04"],
            ["机构准入合规", "新增机构符合票交所机构代码管理规范", "F-01、F-10"],
            ["机构层级管理", "机构层级不超过4级", "F-01"],
        ],
        col_widths=[3, 8, 5]
    )

    add_heading(doc, "7.4 与现有票据系统模块的关联", level=2)
    add_table(doc,
        ["模块", "关联说明", "影响范围"],
        [
            ["系统管理-机构管理", "新增批量导入/复制角色/模板下载按钮；机构树查询增强；新增机构主系统校验", "F-01~F-04、F-10"],
            ["系统管理-机构管理员管理", "新增批量新增/删除/复制角色/模板下载按钮", "F-05~F-08"],
            ["业务管理-贴现业务", "贴现申请提交时增加账务机构校验", "F-11"],
            ["后台定时任务", "新增2个机构数据同步定时任务", "F-09"],
        ],
        col_widths=[4, 7, 5]
    )

    add_page_break(doc)

    # ===== 第八章 部署与配置 =====
    add_heading(doc, "第八章 部署与配置", level=1)

    add_heading(doc, "8.1 服务清单", level=2)
    add_table(doc,
        ["服务名称", "端口", "启动依赖", "说明"],
        [
            ["后端 SpringBoot", "8010", "Redis + ZooKeeper", "提供RESTful API"],
            ["前端 Node.js", "8091", "无", "webpack-dev-server编译前端Vue组件"],
            ["Redis", "6379", "无", "会话管理、字典缓存"],
            ["ZooKeeper", "2181", "无", "服务注册与发现"],
            ["Oracle数据库", "1521", "无", "存储业务数据"],
        ],
        col_widths=[4, 2, 4, 6]
    )

    add_heading(doc, "8.2 配置项", level=2)
    add_table(doc,
        ["配置项", "配置说明", "所属功能点"],
        [
            ["cbsFile.cbsFilePath", "CBS文件根目录，定时任务读取同步文件的根路径", "F-09"],
            ["branch_admin_init_pwd", "机构管理员初始密码，从系统参数读取", "F-05"],
            ["HNNXTK020111", "PJGCS机构参数同步定时任务CloudFunction编号", "F-09"],
            ["HNNXTK020112", "PJGGX机构关系同步定时任务CloudFunction编号", "F-09"],
            ["密码有效期", "在系统参数设置中统一维护", "F-05"],
            ["密码错误次数", "在系统参数设置中统一维护", "F-05"],
        ],
        col_widths=[5, 7, 4]
    )

    add_heading(doc, "8.3 访问地址", level=2)
    add_table(doc,
        ["访问项", "地址", "说明"],
        [
            ["前端地址", "http://127.0.0.1:8091", "Vue.js前端应用"],
            ["后端地址", "http://127.0.0.1:8010", "SpringBoot后端API"],
            ["个性化路径前缀", "/hnnxbank/", "所有个性化接口均包含此前缀"],
        ],
        col_widths=[4, 6, 6]
    )

    add_page_break(doc)

    # ===== 第九章 偏差记录 =====
    add_heading(doc, "第九章 偏差记录", level=1)
    add_para(doc, "本章记录需求文档（v2需求梳理）与实际代码实现（v3代码核查）之间的偏差，包括API路径修订、失败策略修订、定时任务修订和集成实现修订。")

    add_heading(doc, "9.1 API路径偏差", level=2)
    add_table(doc,
        ["接口编号", "v2文档路径（已修订）", "v3实际代码路径", "偏差原因"],
        [
            ["API-01", "func_batchImportBranch", "func_batchImportValidate", "先校验后导入，两步合一"],
            ["API-02", "downloadImportTemplate", "func_downloadModel", "方法命名规范统一"],
            ["API-04", "queryBranchTreeByCondition", "queryBranch + queryBranchPage", "拆分为两个独立接口"],
            ["API-05", "func_batchAddBranchUser", "func_batchImportBranchAdmin", "个性化路径+方法命名规范"],
            ["API-06", "func_batchDeleteBranchUser", "func_batchDeleteBranchAdmin", "个性化路径+方法命名规范"],
            ["API-07", "downloadUserTemplate", "func_downloadModel", "方法命名规范统一"],
            ["API-08", "func_batchCopyUserRole", "func_batchCopyRole", "方法命名规范统一"],
        ],
        col_widths=[2, 4, 4, 6]
    )

    add_heading(doc, "9.2 失败策略偏差", level=2)
    add_table(doc,
        ["功能点", "v2文档假设", "v3代码核查确认", "偏差原因"],
        [
            ["F-01", "整批回滚", "整批回滚（仅导入校验阶段，StringJoiner收集错误）", "一致，补充实现细节"],
            ["F-03", "原子性回滚", "单条失败不影响其他，返回结构化结果", "实际实现为逐个调用"],
            ["F-05", "身份认证失败整批回滚", "单条失败不影响其他，返回结构化结果", "实际实现为逐条处理"],
            ["F-06", "逻辑删除+前置校验", "逻辑删除+跨法人校验+自删除保护+单条策略", "代码核查发现跨法人校验"],
            ["F-08", "原子性回滚", "单条失败不影响其他，返回结构化结果", "实际实现为逐个调用"],
        ],
        col_widths=[2, 4, 5, 5]
    )

    add_heading(doc, "9.3 定时任务偏差", level=2)
    add_table(doc,
        ["项目", "v2文档定义", "v3代码核查确认", "偏差说明"],
        [
            ["定时任务数量", "1个定时任务（JOB-01）", "2个独立定时任务（JOB-01a + JOB-01b）", "PJGCS和PJGGX分别同步"],
            ["API-09（F-10）", "独立API", "集成于func_batchImportValidate", "集成实现，非独立接口"],
            ["API-10（F-11）", "独立API", "集成于submitElecFlow", "集成实现，非独立接口"],
        ],
        col_widths=[3, 4, 4, 5]
    )

    add_page_break(doc)

    # ===== 第十章 需求确认闭环 =====
    add_heading(doc, "第十章 需求确认闭环", level=1)

    add_heading(doc, "10.1 门禁结论", level=2)
    add_para(doc, "需求确认v2门禁结论：通过门禁（无暂缓）。v1的4项核心阻塞项（Q-01、Q-03、Q-04、Q-15）已全部通过v3代码核查解除。v1暂缓的8项用例（TC-DB-01至TC-DB-08）全部解除暂缓。14项待确认事项已确认，2项非核心流程待确认（Q-13、Q-14），1项前端待验证假设（H-10）。")

    add_heading(doc, "10.2 待确认事项闭环", level=2)
    add_table(doc,
        ["编号", "待确认事项", "v1状态", "v2状态", "确认依据"],
        [
            ["Q-01", "批量导入失败策略", "阻塞项（可假设）", "已确认", "代码核查：F-01整批回滚+其他单条策略"],
            ["Q-03", "身份认证失败导入策略", "阻塞项（可假设）", "已确认", "代码核查：单条失败不影响其他"],
            ["Q-04", "批量删除前置条件和方式", "阻塞项（可假设）", "已确认", "代码核查：逻辑删除+跨法人+自删除保护"],
            ["Q-15", "数据库字段扩展方案", "阻塞项（不可假设）", "已确认", "代码核查：模板字段已提取，代码已实现"],
            ["Q-13", "JIGOLX字段取值含义", "待确认", "待确认（非阻塞）", "不阻塞核心流程，需核心系统确认"],
            ["Q-14", "YWGXZL字段其他取值含义", "待确认", "待确认（非阻塞）", "不阻塞核心流程，需核心系统确认"],
        ],
        col_widths=[1.5, 4, 3, 3, 5]
    )

    add_heading(doc, "10.3 测试假设验证结论", level=2)
    add_table(doc,
        ["假设编号", "假设内容", "验证结论", "验证方式"],
        [
            ["H-01", "F-01批量导入失败策略：整批回滚", "已确认", "代码核查：StringJoiner收集错误后统一抛出"],
            ["H-02", "机构树查询结果自动展开+高亮+无结果提示", "已确认", "前端代码已实现"],
            ["H-03", "F-05身份认证失败策略：单条失败不影响其他", "已确认（修订）", "v3代码核查修订为单条策略"],
            ["H-04", "F-06批量删除策略：逻辑删除+跨法人+自删除保护", "已确认（修订）", "v3代码核查修订增加跨法人校验"],
            ["H-05", "F-09数据同步策略：全量替换+日初执行", "已确认", "代码核查：deleteAll+insertBatch+@Transactional"],
            ["H-06", "PJGCS表JILUZT='0'为正常状态", "已确认", "样本数据验证"],
            ["H-07", "PJGGX表YWGXZL='ZNGWSJ'为账务机构关系", "已确认", "样本数据验证"],
            ["H-08", "Excel模板字段定义", "已确认", "代码提取F-02模板12列、F-07模板9列"],
            ["H-09", "F-03/F-08角色复制失败策略：单条失败不影响其他", "已确认（修订）", "v3代码核查修订为单条策略"],
            ["H-10", "F-06批量删除二次确认弹窗", "成立（已验证）", "代码审查：batchDeleteConfirm文本确认"],
            ["H-11", "F-11贴现校验扩展性：接口抽象", "已确认", "PjgxBranchRelationService接口已实现"],
            ["H-12", "机构树查询交互：三条件AND+重置+无结果提示", "已确认", "前端代码已实现"],
            ["H-13", "JIGOLX字段取值含义", "部分成立", "字典加载机制确认，取值含义待数据库验证"],
            ["H-14", "YWGXZL字段其他取值含义", "待验证", "后端逻辑，前端不可达，待核心系统确认"],
            ["H-15", "数据库字段扩展方案", "已解除", "模板字段已提取，代码已实现"],
            ["H-16", "接口路径复用现有路径模式", "已确认", "v3已修订为实际代码路径"],
        ],
        col_widths=[1.5, 5, 3, 6]
    )

    # ===== 保存文档 =====
    inject_update_fields(doc)
    os.makedirs(OUTPUT_DIR, exist_ok=True)
    doc.save(OUTPUT_FILE)
    print(f"详细设计文档已生成：{OUTPUT_FILE}")
    file_size = os.path.getsize(OUTPUT_FILE)
    print(f"文件大小：{file_size} 字节 ({file_size/1024:.1f} KB)")


if __name__ == "__main__":
    build_design_doc()
