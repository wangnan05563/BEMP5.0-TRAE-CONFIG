"""追踪文档生成过程中模块2的变化"""
from docx import Document
import re

def count_h2_under_h1(doc, h1_text):
    """统计H1下的H2数量"""
    h1_found = False
    h2_count = 0
    
    for para in doc.paragraphs:
        if para.style and para.style.name == 'Heading 1':
            if h1_text in para.text:
                h1_found = True
                continue
            elif h1_found:
                # 遇到下一个H1，停止
                break
        
        if h1_found and para.style and para.style.name == 'Heading 2':
            h2_count += 1
    
    return h2_count

# 检查生成的文档
doc_path = r"d:\code\QJ\BEMP5.0DEV\.trae\skills\bemp-advanced-doc-generator\output\机构管理和管理员管理功能优化-详细设计文档-20260617.docx"
doc = Document(doc_path)

print("=" * 70)
print("模块2内容追踪")
print("=" * 70)

# 统计模块1和模块2的H2数量
module1_h2 = count_h2_under_h1(doc, "模块1设计说明")
module2_h2 = count_h2_under_h1(doc, "模块2设计说明")

print(f"\n模块1设计说明下的H2数量: {module1_h2}")
print(f"模块2设计说明下的H2数量: {module2_h2}")

# 列出模块2下的所有H2
print("\n模块2下的H2标题列表:")
in_module2 = False
for i, para in enumerate(doc.paragraphs):
    if para.style and para.style.name == 'Heading 1':
        if '模块2设计说明' in para.text:
            in_module2 = True
            print(f"  [{i}] H1: {para.text}")
            continue
        elif in_module2:
            # 遇到下一个H1，停止
            print(f"  [{i}] H1: {para.text} (下一个H1，停止)")
            break
    
    if in_module2 and para.style and para.style.name == 'Heading 2':
        print(f"  [{i}] H2: {para.text}")

# 检查模块2和下一个H1之间的所有段落
print("\n模块2后的前10个段落:")
in_module2 = False
count = 0
for i, para in enumerate(doc.paragraphs):
    if para.style and para.style.name == 'Heading 1':
        if '模块2设计说明' in para.text:
            in_module2 = True
            continue
        elif in_module2:
            # 遇到下一个H1
            print(f"  [{i}] H1: {para.text}")
            break
    
    if in_module2:
        count += 1
        if count <= 10:
            style_name = para.style.name if para.style else 'None'
            text = para.text.strip()[:80] if para.text.strip() else '(空)'
            print(f"  [{i}] {style_name}: {text}")
