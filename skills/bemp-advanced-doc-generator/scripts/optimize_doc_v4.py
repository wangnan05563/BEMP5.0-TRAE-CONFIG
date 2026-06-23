"""修复v4: 补充模块2内容 + 修复栏位描述表格"""
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

v3_path = r"d:\code\QJ\BEMP5.0DEV\.trae\skills\bemp-advanced-doc-generator\output\机构管理和管理员管理功能优化-详细设计文档-v3-20260617.docx"
v4_path = r"d:\code\QJ\BEMP5.0DEV\.trae\skills\bemp-advanced-doc-generator\output\机构管理和管理员管理功能优化-详细设计文档-v4-20260617.docx"

doc = Document(v3_path)

# === 修复1: 模块2设计说明后插入功能描述段落 ===
module2_idx = None
for i, para in enumerate(doc.paragraphs):
    if para.style.name == 'Heading 1' and '模块2' in para.text:
        module2_idx = i
        break

if module2_idx is not None:
    heading_para = doc.paragraphs[module2_idx]
    insert_after = heading_para._element
    
    module2_text = (
        "机构管理员管理模块包含以下核心功能：\n"
        "1. 批量导入管理员：支持通过Excel文件批量导入机构管理员，只有法人管理员可操作，"
        "新增管理员默认为无效状态，初次登录时需修改密码后变为有效；\n"
        "2. 模板下载：提供机构管理员导入模板下载功能；\n"
        "3. 批量复制角色：将选中管理员的角色权限批量复制到目标管理员，"
        "只能分配管理员所在机构所拥有的角色（包括机构角色和机构临时角色），"
        "用新角色覆盖旧角色，做关联角色的新增和删除。"
    )
    
    p = OxmlElement('w:p')
    r = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    sz = OxmlElement('w:sz')
    sz.set(qn('w:val'), '21')
    rPr.append(sz)
    r.append(rPr)
    r.append(t := OxmlElement('w:t'))
    t.text = module2_text
    t.set(qn('xml:space'), 'preserve')
    p.append(r)
    insert_after.addnext(p)
    print("✓ 模块2功能描述已补充")

# === 修复2: 更新表格13（附录D栏位描述汇总，32行x2列） ===
# 这个表格是2列结构，我们更新前几行为栏位描述
field_data_2col = [
    ["目标机构号", "输入-弹出框，支持多选，必输"],
    ["目标机构名称", "输入-文本框，灰显不可修改，选择目标机构号后回显，必输"],
    ["目标用户号", "输入-弹出框，支持多选，必输"],
    ["目标用户姓名", "输入-文本框，灰显不可修改，选择目标用户号后回显，必输"],
    ["机构号", "机构唯一标识，不可重复"],
    ["机构名称", "机构中文名称，不可重复"],
    ["核算机构号", "核算用机构标识，不可重复"],
    ["组织机构代码", "组织机构标识，不可重复"],
    ["票交所机构代码", "非必输项，可维护行内机构非票交所参与机构，一个代码可绑定多个行内机构"],
    ["上级机构号", "必输项，新增机构必须选择上级机构"],
    ["用户号", "管理员唯一标识"],
    ["用户姓名", "管理员姓名"],
    ["所属机构号", "管理员所属机构"],
    ["手机号", "管理员联系电话"],
    ["邮箱", "管理员邮箱地址"]
]

for table in doc.tables:
    if len(table.rows) > 20 and len(table.rows[0].cells) == 2:
        first_row = [cell.text.strip() for cell in table.rows[0].cells]
        if '数据名称' in first_row[0]:
            # 更新表头
            table.rows[0].cells[0].text = "数据名称"
            table.rows[0].cells[1].text = "输入/输出 | 表现形式 | 是否必输 | 数据约束 | 备注"
            
            # 填充数据
            for i, (name, desc) in enumerate(field_data_2col):
                if i + 1 < len(table.rows):
                    table.rows[i + 1].cells[0].text = name
                    table.rows[i + 1].cells[1].text = desc
            
            # 清空多余行
            for i in range(len(field_data_2col) + 1, len(table.rows)):
                table.rows[i].cells[0].text = ""
                table.rows[i].cells[1].text = ""
            
            print(f"✓ 栏位描述表格已更新（{len(field_data_2col)}个栏位）")
            break

# === 修复3: 更新表格14（5行x6列的字段映射表） ===
field_data_6col = [
    ["目标机构号", "targetBranchNo", "S(32)", "32", "是", "支持多选"],
    ["目标机构名称", "targetBranchName", "S(100)", "100", "是", "灰显，不可修改"],
    ["目标用户号", "targetUserNo", "S(32)", "32", "是", "支持多选"],
    ["目标用户姓名", "targetUserName", "S(100)", "100", "是", "灰显，不可修改"]
]

for table in doc.tables:
    if len(table.rows) >= 5 and len(table.rows[0].cells) == 6:
        first_row = [cell.text.strip() for cell in table.rows[0].cells]
        if '字段名称' in first_row[0]:
            for i, row_data in enumerate(field_data_6col):
                if i + 1 < len(table.rows):
                    for j, cell_text in enumerate(row_data):
                        table.rows[i + 1].cells[j].text = cell_text
            print("✓ 字段映射关系表格已更新（4个字段）")
            break

doc.save(v4_path)
print(f"\n✓ v4文档已生成: {v4_path}")
