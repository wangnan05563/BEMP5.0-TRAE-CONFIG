from pathlib import Path
SKILL_ROOT = Path(__file__).resolve().parent.parent
PROJECT_ROOT = SKILL_ROOT.parent.parent.parent

from pathlib import Path
"""追踪模块2内容消失的完整流程"""
from docx import Document
from docx.oxml.ns import qn

doc_path = str(SKILL_ROOT / "output" / "机构管理和管理员管理功能优化-详细设计文档-20260618.docx")
doc = Document(doc_path)

print("=" * 70)
print("模块2内容消失完整追踪")
print("=" * 70)

# 找到所有H1标题
h1_list = []
for i, para in enumerate(doc.paragraphs):
    style_name = para.style.name if para.style else 'None'
    if 'Heading 1' in style_name or style_name == '1':
        h1_list.append((i, para.text.strip()))

print("\n所有H1标题:")
for idx, (i, text) in enumerate(h1_list):
    print(f"  [{idx}] 位置{i}: {text}")

# 找到模块1和模块2的位置
module1_pos = None
module1_idx = None
module2_pos = None
module2_idx = None

for idx, (i, text) in enumerate(h1_list):
    if '模块1' in text:
        module1_pos = i
        module1_idx = idx
    if '模块2' in text:
        module2_pos = i
        module2_idx = idx

print(f"\n模块1位置: {module1_pos}")
print(f"模块2位置: {module2_pos}")

if module1_pos and module2_pos:
    # 统计模块1的内容
    next_h1_pos_m1 = h1_list[module1_idx + 1][0] if module1_idx + 1 < len(h1_list) else len(doc.paragraphs)
    h2_count_m1 = 0
    para_count_m1 = 0
    for i in range(module1_pos + 1, next_h1_pos_m1):
        para = doc.paragraphs[i]
        style_name = para.style.name if para.style else 'None'
        text = para.text.strip()
        if 'Heading 2' in style_name or style_name == '2':
            h2_count_m1 += 1
        elif text:
            para_count_m1 += 1
    
    print(f"\n模块1内容:")
    print(f"  H2数量: {h2_count_m1}")
    print(f"  段落数量: {para_count_m1}")
    
    # 统计模块2的内容
    next_h1_pos_m2 = h1_list[module2_idx + 1][0] if module2_idx + 1 < len(h1_list) else len(doc.paragraphs)
    h2_count_m2 = 0
    para_count_m2 = 0
    for i in range(module2_pos + 1, next_h1_pos_m2):
        para = doc.paragraphs[i]
        style_name = para.style.name if para.style else 'None'
        text = para.text.strip()
        if 'Heading 2' in style_name or style_name == '2':
            h2_count_m2 += 1
        elif text:
            para_count_m2 += 1
    
    print(f"\n模块2内容:")
    print(f"  H2数量: {h2_count_m2}")
    print(f"  段落数量: {para_count_m2}")
    
    # 检查模块2和下一个H1之间的所有元素
    print(f"\n模块2到下一个H1之间的元素（位置{module2_pos}到{next_h1_pos_m2}）:")
    for i in range(module2_pos, min(next_h1_pos_m2, module2_pos + 10)):
        para = doc.paragraphs[i]
        style_name = para.style.name if para.style else 'None'
        text = para.text.strip()
        print(f"  位置{i}: style={style_name}, text={text[:50] if text else '(空)'}")

# 检查XML结构
print("\n" + "=" * 70)
print("XML结构检查")
print("=" * 70)

body = doc.element.body
module2_elem = None
for elem in body:
    if elem.tag.endswith('}p'):
        pPr = elem.find(qn('w:pPr'))
        if pPr is not None:
            pStyle = pPr.find(qn('w:pStyle'))
            if pStyle is not None:
                style_val = pStyle.get(qn('w:val'), '')
                if style_val in ['1', 'Heading1', 'heading1', 'heading 1', 'Heading 1']:
                    text = ''.join(t.text or '' for t in elem.findall('.//' + qn('w:t'))).strip()
                    if '模块2' in text:
                        module2_elem = elem
                        print(f"找到模块2的XML元素: {elem.tag}")
                        break

if module2_elem is not None:
    # 检查模块2后面的元素
    print("\n模块2后面的5个元素:")
    count = 0
    next_elem = module2_elem.getnext()
    while next_elem is not None and count < 5:
        tag = next_elem.tag.split('}')[-1] if '}' in next_elem.tag else next_elem.tag
        if tag == 'p':
            pPr = next_elem.find(qn('w:pPr'))
            style_val = ''
            if pPr is not None:
                pStyle = pPr.find(qn('w:pStyle'))
                if pStyle is not None:
                    style_val = pStyle.get(qn('w:val'), '')
            text = ''.join(t.text or '' for t in next_elem.findall('.//' + qn('w:t'))).strip()
            print(f"  元素{count}: tag={tag}, style={style_val}, text={text[:50] if text else '(空)'}")
        elif tag == 'tbl':
            print(f"  元素{count}: tag={tag} (表格)")
        else:
            print(f"  元素{count}: tag={tag}")
        
        next_elem = next_elem.getnext()
        count += 1
