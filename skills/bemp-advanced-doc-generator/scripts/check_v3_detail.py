"""检查v3文档中模块2和附录D的实际内容"""
from docx import Document

v3_path = r"d:\code\QJ\BEMP5.0DEV\.trae\skills\bemp-advanced-doc-generator\output\机构管理和管理员管理功能优化-详细设计文档-v3-20260617.docx"
doc = Document(v3_path)

print("=== 模块2设计说明附近内容 ===")
for i, para in enumerate(doc.paragraphs):
    if '模块2' in para.text or '管理员' in para.text:
        print(f"  [{i}] ({para.style.name}) {para.text[:120]}")

print("\n=== 附录D附近内容 ===")
for i, para in enumerate(doc.paragraphs):
    if '附录D' in para.text or '栏位' in para.text:
        print(f"  [{i}] ({para.style.name}) {para.text[:120]}")

print("\n=== 所有表格预览 ===")
for i, table in enumerate(doc.tables):
    rows = len(table.rows)
    cols = len(table.rows[0].cells) if table.rows else 0
    first_row = [cell.text.strip()[:30] for cell in table.rows[0].cells] if table.rows else []
    print(f"  表格{i}: {rows}行x{cols}列 | 首行: {first_row}")
