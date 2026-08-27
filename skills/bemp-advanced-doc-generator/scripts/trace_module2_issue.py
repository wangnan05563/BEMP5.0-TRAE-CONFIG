from pathlib import Path
SKILL_ROOT = Path(__file__).resolve().parent.parent
PROJECT_ROOT = SKILL_ROOT.parent.parent.parent

from pathlib import Path
"""精确追踪模块2内容消失的原因"""
from docx import Document
from docx.oxml.ns import qn

doc_path = str(SKILL_ROOT / "output" / "机构管理和管理员管理功能优化-详细设计文档-20260617.docx")
doc = Document(doc_path)

print("=" * 70)
print("模块2内容消失原因追踪")
print("=" * 70)

# 找到所有H1标题
h1_list = []
for i, para in enumerate(doc.paragraphs):
    style_name = para.style.name if para.style else 'None'
    if 'Heading 1' in style_name or style_name == '1':
        h1_list.append((i, para.text.strip()))

print("\n所有H1标题:")
for idx, (i, text) in enumerate(h1_list):
    print(f"  [{idx}] 位置{i}: {text}")

# 找到模块2的位置
module2_pos = None
module2_idx = None
for idx, (i, text) in enumerate(h1_list):
    if '模块2' in text:
        module2_pos = i
        module2_idx = idx
        break

if module2_pos is None:
    print("\n未找到模块2!")
else:
    print(f"\n模块2位置: H1列表索引{module2_idx}, 段落位置{module2_pos}")
    
    # 找到下一个H1的位置
    next_h1_pos = None
    if module2_idx + 1 < len(h1_list):
        next_h1_pos = h1_list[module2_idx + 1][0]
        print(f"下一个H1位置: {next_h1_pos} ({h1_list[module2_idx + 1][1]})")
    
    # 详细统计模块2下的所有内容
    print(f"\n模块2下的所有内容（位置{module2_pos+1}到{next_h1_pos}）:")
    
    h2_count = 0
    h3_count = 0
    para_count = 0
    table_count = 0
    empty_para_count = 0
    
    start_pos = module2_pos + 1
    end_pos = next_h1_pos if next_h1_pos else len(doc.paragraphs)
    
    for i in range(start_pos, min(end_pos, start_pos + 50)):  # 只显示前50个元素
        para = doc.paragraphs[i]
        style_name = para.style.name if para.style else 'None'
        text = para.text.strip()
        
        if 'Heading 2' in style_name or style_name == '2':
            h2_count += 1
            print(f"  [H2] 位置{i}: {text[:80] if text else '(空)'}")
        elif 'Heading 3' in style_name or style_name == '3':
            h3_count += 1
            print(f"  [H3] 位置{i}: {text[:80] if text else '(空)'}")
        elif text:
            para_count += 1
            if para_count <= 10:  # 只显示前10个段落
                print(f"  [段落] 位置{i}: {text[:80]}")
        else:
            empty_para_count += 1
    
    # 统计表格
    for table in doc.tables:
        table_count += 1
    
    print(f"\n统计:")
    print(f"  H2数量: {h2_count}")
    print(f"  H3数量: {h3_count}")
    print(f"  段落数量: {para_count}")
    print(f"  空段落数量: {empty_para_count}")
    print(f"  表格数量: {table_count}")
    
    if h2_count == 0 and para_count == 0:
        print("\n结论: 模块2下完全没有内容!")
        print("可能原因:")
        print("  1. _fill_chapter_h1_sections 创建的内容被后续处理删除")
        print("  2. filled_contexts 检查导致模块2被跳过")
        print("  3. _clear_content_between_headings 删除了模块2的内容")
    elif h2_count > 0:
        print(f"\n模块2有 {h2_count} 个H2标题")
        if para_count == 0:
            print("但H2下没有段落内容!")

# 检查模块1的内容作为对比
print("\n" + "=" * 70)
print("模块1内容检查（作为对比）")
print("=" * 70)

module1_pos = None
module1_idx = None
for idx, (i, text) in enumerate(h1_list):
    if '模块1' in text:
        module1_pos = i
        module1_idx = idx
        break

if module1_pos:
    next_h1_pos_m1 = h1_list[module1_idx + 1][0] if module1_idx + 1 < len(h1_list) else len(doc.paragraphs)
    
    h2_count_m1 = 0
    para_count_m1 = 0
    
    for i in range(module1_pos + 1, next_h1_pos_m1):
        para = doc.paragraphs[i]
        style_name = para.style.name if para.style else 'None'
        text = para.text.strip()
        
        if 'Heading 2' in style_name or style_name == '2':
            h2_count_m1 += 1
        elif text:
            para_count_m1 += 1
    
    print(f"模块1 H2数量: {h2_count_m1}")
    print(f"模块1 段落数量: {para_count_m1}")
