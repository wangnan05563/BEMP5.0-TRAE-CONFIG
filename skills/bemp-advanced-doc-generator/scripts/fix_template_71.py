"""修复扩展模板中7.1功能描述丢失问题"""
import paths
from docx import Document
from docx.oxml.ns import qn
from copy import deepcopy

path = str(paths.detail_design_template())
doc = Document(path)

# 找到"额度占用/释放模块设计说明"H1
for i, p in enumerate(doc.paragraphs):
    if p.style and p.style.name == 'Heading 1' and '额度占用/释放' in p.text:
        # 找到下一个H2（7.2）
        # 在H1后插入"7.1 功能描述"
        # 先找"7.2 界面"H2作为模板
        for j in range(i+1, min(i+15, len(doc.paragraphs))):
            np = doc.paragraphs[j]
            if np.style and np.style.name == 'Heading 2' and '7.2' in np.text:
                # 复制np的XML元素
                elem = deepcopy(np._element)
                # 修改文本为"7.1 功能描述"
                # 清除所有run
                for r in list(elem.findall(qn('w:r'))):
                    elem.remove(r)
                # 添加新run
                new_r = elem.makeelement(qn('w:r'), {})
                new_t = new_r.makeelement(qn('w:t'), {})
                new_t.text = "7.1 功能描述"
                new_t.set(qn('xml:space'), 'preserve')
                new_r.append(new_t)
                elem.append(new_r)
                # 插入到np._element之前
                np._element.addprevious(elem)
                break
        break

doc.save(path)
print("已修复")

# 验证
doc2 = Document(path)
for p in doc2.paragraphs:
    if p.style and p.style.name.startswith('Heading') and ('7.' in p.text or '额度占用' in p.text):
        print(f"  {p.style.name}: {p.text.strip()}")
