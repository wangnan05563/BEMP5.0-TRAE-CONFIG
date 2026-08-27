from pathlib import Path
SKILL_ROOT = Path(__file__).resolve().parent.parent
PROJECT_ROOT = SKILL_ROOT.parent.parent.parent

from pathlib import Path
"""修复v3: 正确填充API接口表格 + 补充数据模型"""
from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

v2_path = str(SKILL_ROOT / "output" / "机构管理和管理员管理功能优化-详细设计文档-v2-20260617.docx")
v3_path = str(SKILL_ROOT / "output" / "机构管理和管理员管理功能优化-详细设计文档-v3-20260617.docx")

doc = Document(v2_path)

# === 修复1: 找到"API接口清单"标题后的第一个表格，重建为7接口表 ===
api_data = [
    ["/hnnxbank/sm/auth/branch/branch/func_batchImportValidate", "POST", "机构批量导入校验"],
    ["/hnnxbank/sm/auth/branch/branch/func_batchImport", "POST", "机构批量导入执行"],
    ["/hnnxbank/sm/auth/branch/branch/func_downloadModel", "GET", "机构导入模板下载"],
    ["/hnnxbank/sm/auth/branch/branchAdmin/func_batchImportValidate", "POST", "管理员批量导入校验"],
    ["/hnnxbank/sm/auth/branch/branchAdmin/func_batchImportBranchAdmin", "POST", "管理员批量导入执行"],
    ["/hnnxbank/sm/auth/branch/branchAdmin/func_downloadModel", "GET", "管理员导入模板下载"],
    ["/hnnxbank/sm/auth/branch/branchAdmin/func_batchCopyRole", "POST", "管理员批量复制角色"]
]

# 定位"API 接口清单"标题后的表格
api_heading_idx = None
for i, para in enumerate(doc.paragraphs):
    if para.style.name.startswith('Heading') and 'API' in para.text and '接口' in para.text:
        api_heading_idx = i
        break

if api_heading_idx is not None:
    # 找到该标题后的第一个表格
    target_table = None
    heading_element = doc.paragraphs[api_heading_idx]._element
    # 在标题之后遍历找表格
    found_heading = False
    for element in doc.element.body:
        if element is heading_element:
            found_heading = True
            continue
        if found_heading and element.tag.endswith('tbl'):
            # 找到表格
            # 删除旧表格
            parent = element.getparent()
            parent.remove(element)
            break
    
    # 在API接口清单标题后插入新表格
    # 找到标题段落的下一个元素位置
    heading_para = doc.paragraphs[api_heading_idx]
    next_element = heading_para._element.getnext()
    
    # 创建新表格（不设置样式，避免兼容性问题）
    table = doc.add_table(rows=1 + len(api_data), cols=3)
    
    # 表头
    headers = ["接口路径", "方法", "说明"]
    for j, h in enumerate(headers):
        cell = table.rows[0].cells[j]
        cell.text = h
        for run in cell.paragraphs[0].runs:
            run.bold = True
            run.font.size = Pt(10)
    
    # 数据行
    for i, row_data in enumerate(api_data):
        for j, cell_text in enumerate(row_data):
            cell = table.rows[i + 1].cells[j]
            cell.text = cell_text
            for run in cell.paragraphs[0].runs:
                run.font.size = Pt(9)
    
    # 移动表格到正确位置（标题之后）
    tbl_element = table._tbl
    tbl_element.getparent().remove(tbl_element)
    if next_element is not None:
        next_element.addprevious(tbl_element)
    else:
        heading_para._element.addnext(tbl_element)
    
    print("✓ API接口表格已重建（7个接口）")
else:
    print("✗ 未找到API接口清单标题")

# === 修复2: 补充数据结构定义（BranchImportVo/BranchAdminImportVo/HnnxBatchCopyRoleReq） ===
# 找到"数据结构定义"标题
data_struct_idx = None
for i, para in enumerate(doc.paragraphs):
    if para.style.name.startswith('Heading') and '数据结构定义' in para.text:
        data_struct_idx = i
        break

if data_struct_idx is not None:
    heading_para = doc.paragraphs[data_struct_idx]
    next_element = heading_para._element.getnext()
    
    # 添加BranchImportVo描述
    p1_text = "BranchImportVo（机构批量导入VO）：包含上级机构号、机构号、机构名称、核算机构号、组织机构代码、票交所机构代码等字段，用于机构批量导入时的数据承载。"
    p2_text = "BranchAdminImportVo（管理员批量导入VO）：包含用户号、用户姓名、所属机构号、手机号、邮箱等字段，用于管理员批量导入时的数据承载。"
    p3_text = "HnnxBatchCopyRoleReq（批量复制角色请求DTO）：包含源机构号、目标机构号列表（支持多选）、目标用户号列表（支持多选），用于批量复制角色接口的请求参数封装。"
    
    # 在数据结构定义标题后插入段落
    insert_point = heading_para._element
    
    for text in [p1_text, p2_text, p3_text]:
        p = OxmlElement('w:p')
        r = OxmlElement('w:r')
        rPr = OxmlElement('w:rPr')
        sz = OxmlElement('w:sz')
        sz.set(qn('w:val'), '21')  # 10.5pt
        rPr.append(sz)
        r.append(rPr)
        t = OxmlElement('w:t')
        t.text = text
        r.append(t)
        p.append(r)
        insert_point.addnext(p)
        insert_point = p
    
    print("✓ 数据结构定义已补充（3个VO/DTO）")

# === 修复3: 补充栏位描述表格（附录D） ===
# 找到附录D标题后的表格
appendix_d_idx = None
for i, para in enumerate(doc.paragraphs):
    if para.style.name.startswith('Heading') and '附录D' in para.text:
        appendix_d_idx = i
        break

if appendix_d_idx is not None:
    heading_para = doc.paragraphs[appendix_d_idx]
    # 找到该标题后的第一个表格
    found = False
    for element in doc.element.body:
        if element is heading_para._element:
            found = True
            continue
        if found and element.tag.endswith('tbl'):
            # 清空并重建表格内容
            table_obj = None
            for t in doc.tables:
                if t._tbl is element:
                    table_obj = t
                    break
            
            if table_obj and len(table_obj.rows) > 0:
                # 机构管理栏位
                field_data = [
                    ["目标机构号", "输入", "弹出框", "M(必输)", "", "支持多选"],
                    ["目标机构名称", "输入", "文本框", "M(必输)", "", "灰显，不可修改"],
                    ["目标用户号", "输入", "弹出框", "M(必输)", "", "支持多选"],
                    ["目标用户姓名", "输入", "文本框", "M(必输)", "", "灰显，不可修改"]
                ]
                
                # 如果表格行数不够，添加行
                while len(table_obj.rows) < 1 + len(field_data):
                    table_obj.add_row()
                
                # 填充表头
                headers = ["数据名称", "输入/输出", "表现形式", "是否必输", "数据约束", "备注"]
                for j, h in enumerate(headers):
                    if j < len(table_obj.rows[0].cells):
                        table_obj.rows[0].cells[j].text = h
                
                # 填充数据
                for i, row_data in enumerate(field_data):
                    for j, cell_text in enumerate(row_data):
                        if j < len(table_obj.rows[i + 1].cells):
                            table_obj.rows[i + 1].cells[j].text = cell_text
                
                print("✓ 附录D栏位描述表格已更新")
            break

# === 修复4: 补充错误码定义 ===
for table in doc.tables:
    if len(table.rows) > 0:
        first_row_text = " ".join([cell.text.strip() for cell in table.rows[0].cells])
        if "错误码" in first_row_text and "错误信息" in first_row_text and "触发场景" in first_row_text:
            error_data = [
                ["E001", "校验不通过", "新增机构未选择上级机构", "返回错误提示，操作回滚"],
                ["E002", "数据重复", "机构号/机构名称/核算机构号/组织机构代码重复", "返回错误提示，操作回滚"],
                ["E003", "层级超限", "机构层级超过4级（含总行）", "返回错误提示，操作回滚"],
                ["E004", "权限不足", "非法人管理员尝试新增机构管理员", "返回错误提示，拒绝操作"],
                ["E005", "角色冲突", "要删除的关联角色已有柜员在使用", "提示'角色[xxx]为机构下用户使用不能去除与当前机构关系'"],
                ["E006", "角色不存在", "机构不具备待分配的角色", "提示'机构【XX】无角色【XX】，不能分配'"]
            ]
            
            while len(table.rows) < 1 + len(error_data):
                table.add_row()
            
            for i, row_data in enumerate(error_data):
                for j, cell_text in enumerate(row_data):
                    if j < len(table.rows[i + 1].cells):
                        table.rows[i + 1].cells[j].text = cell_text
            
            print("✓ 错误码表格已更新（6个业务错误码）")
            break

doc.save(v3_path)
print(f"\n✓ v3文档已生成: {v3_path}")
