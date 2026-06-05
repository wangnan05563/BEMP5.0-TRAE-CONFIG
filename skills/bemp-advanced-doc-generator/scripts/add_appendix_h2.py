"""给附录H1添加H2子节占位符"""
import paths
from docx import Document
from docx.oxml.ns import qn
from copy import deepcopy

path = str(paths.detail_design_template())
doc = Document(path)

# 找到"附录"H1，并定位H2模板
appendix_h1 = None
h2_template = None
for i, p in enumerate(doc.paragraphs):
    if p.style and p.style.name == 'Heading 1' and p.text.strip() == '附录':
        appendix_h1 = p
        # 找文档中已有的H2作为模板
        for j, p2 in enumerate(doc.paragraphs):
            if p2.style and p2.style.name == 'Heading 2' and p2.text.strip() == '4.1 功能描述':
                h2_template = p2
                break
        break

if appendix_h1 is None or h2_template is None:
    print(f"未找到附录H1或H2模板")
    exit(1)

# 5个附录H2
appendix_h2s = ["A.1 术语缩写", "A.2 参考资料", "A.3 异常处理", "A.4 性能指标", "A.5 版本历史"]

# 在附录H1后依次插入
h1_elem = appendix_h1._element
parent = h1_elem.getparent()
h1_idx = list(parent).index(h1_elem)

# 收集要插入的H2元素
new_h2_elems = []
for h2_title in appendix_h2s:
    h2_elem = deepcopy(h2_template._element)
    # 清除所有run
    for r in list(h2_elem.findall(qn('w:r'))):
        h2_elem.remove(r)
    # 添加新文本
    new_r = h2_elem.makeelement(qn('w:r'), {})
    new_t = new_r.makeelement(qn('w:t'), {})
    new_t.text = h2_title
    new_t.set(qn('xml:space'), 'preserve')
    new_r.append(new_t)
    h2_elem.append(new_r)
    new_h2_elems.append(h2_elem)

# 依次插入到H1后
insert_pos = h1_idx + 1
for elem in new_h2_elems:
    parent.insert(insert_pos, elem)
    insert_pos += 1

doc.save(path)
print(f"已添加 {len(new_h2_elems)} 个附录H2子节")

# 验证
doc2 = Document(path)
for p in doc2.paragraphs:
    if p.style and p.style.name.startswith('Heading') and ('A.' in p.text or '附录' in p.text):
        print(f"  {p.style.name}: {p.text.strip()}")
