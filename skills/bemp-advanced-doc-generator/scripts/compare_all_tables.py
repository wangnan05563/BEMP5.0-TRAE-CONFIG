"""详细对比v4和new的表格结构"""
import paths
from docx import Document

v4_path = str(paths.BANK_REQUIREMENTS_DIR / '承兑行额度管理-详细设计说明书-v4.docx')
new_path = str(paths.OUTPUT_DIR / '承兑行额度管理-详细设计文档-20260603.docx')

for label, path in [("v4", v4_path), ("new", new_path)]:
    doc = Document(path)
    print(f"\n=== {label} 文档表格（{len(doc.tables)}个）===")
    for i, tbl in enumerate(doc.tables):
        rows = len(tbl.rows)
        cols = len(tbl.columns) if tbl.rows else 0
        first_row = ' | '.join(c.text.strip()[:18] for c in tbl.rows[0].cells) if tbl.rows else ''
        print(f"  表格[{i:3d}]: {rows}行x{cols}列 | {first_row[:80]}")
