from pathlib import Path
SKILL_ROOT = Path(__file__).resolve().parent.parent
PROJECT_ROOT = SKILL_ROOT.parent.parent.parent

from pathlib import Path
"""完整数据流诊断：从 JSON 到文档生成"""
import json
import os
from docx import Document
from docx.oxml.ns import qn

# 1. 检查 JSON 文件
json_path = str(SKILL_ROOT / "output" / "_design-data-20260617.json")
print("=" * 70)
print("1. JSON 文件检查")
print("=" * 70)

if os.path.exists(json_path):
    with open(json_path, 'r', encoding='utf-8') as f:
        data = json.load(f)
    
    chapters = data.get('chapters', [])
    business_subs = data.get('businessSubmodules', [])
    
    print(f"✓ JSON 文件存在: {json_path}")
    print(f"  chapters 数量: {len(chapters)}")
    print(f"  businessSubmodules 数量: {len(business_subs)}")
    
    if chapters:
        print(f"\n  前3个 chapters:")
        for i, ch in enumerate(chapters[:3]):
            print(f"    [{i}] {ch.get('title', '')} | sections: {len(ch.get('sections', []))}")
    
    if business_subs:
        print(f"\n  businessSubmodules:")
        for i, sub in enumerate(business_subs):
            name = sub.get('name', '') if isinstance(sub, dict) else str(sub)
            print(f"    [{i}] {name}")
    
    # 检查 _preserve 标志
    preserve = data.get('_preserve', False)
    print(f"\n  _preserve 标志: {preserve}")
    print(f"  _PRESERVE_MODE 条件: preserve={preserve} or len(chapters)==0={len(chapters)==0}")
    print(f"  实际 _PRESERVE_MODE: {preserve or len(chapters)==0}")
    
else:
    print(f"✗ JSON 文件不存在: {json_path}")
    # 查找所有 JSON 文件
    output_dir = os.path.dirname(json_path)
    json_files = [f for f in os.listdir(output_dir) if f.endswith('.json')]
    print(f"  目录中的 JSON 文件: {json_files}")

# 2. 检查生成的文档
print("\n" + "=" * 70)
print("2. 生成的文档检查")
print("=" * 70)

doc_path = str(SKILL_ROOT / "output" / "机构管理和管理员管理功能优化-详细设计文档-20260617.docx")

if os.path.exists(doc_path):
    doc = Document(doc_path)
    print(f"✓ 文档存在: {doc_path}")
    print(f"  总段落数: {len(doc.paragraphs)}")
    print(f"  总表格数: {len(doc.tables)}")
    
    # 统计 H1 和 H2
    h1_count = 0
    h2_count = 0
    for para in doc.paragraphs:
        style_name = para.style.name if para.style else ""
        if style_name == 'Heading 1':
            h1_count += 1
        elif style_name == 'Heading 2':
            h2_count += 1
    
    print(f"  H1 标题数: {h1_count}")
    print(f"  H2 标题数: {h2_count}")
    
    # 检查关键章节
    print(f"\n  关键章节检查:")
    for para in doc.paragraphs:
        text = para.text.strip()
        style_name = para.style.name if para.style else ""
        
        if '组件' in text and '模块列表' in text:
            print(f"    ✓ 找到: {text} (样式: {style_name})")
            # 检查后续内容
            idx = doc.paragraphs.index(para)
            next_paras = doc.paragraphs[idx+1:idx+5]
            print(f"      后续 {len(next_paras)} 个段落:")
            for np in next_paras:
                np_style = np.style.name if np.style else ""
                print(f"        [{np_style}] {np.text[:50]}")
        
        if '模块1设计说明' in text:
            print(f"    ✓ 找到: {text} (样式: {style_name})")
            # 检查后续内容
            idx = doc.paragraphs.index(para)
            next_paras = doc.paragraphs[idx+1:idx+10]
            print(f"      后续 {len(next_paras)} 个段落:")
            for np in next_paras:
                np_style = np.style.name if np.style else ""
                np_text = np.text.strip()
                if np_text:
                    print(f"        [{np_style}] {np_text[:50]}")
    
else:
    print(f"✗ 文档不存在: {doc_path}")

print("\n" + "=" * 70)
print("诊断完成")
print("=" * 70)
