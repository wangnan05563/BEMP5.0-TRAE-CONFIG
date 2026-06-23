"""v4 最终质量审核"""
from docx import Document
import os

v4_path = r"d:\code\QJ\BEMP5.0DEV\.trae\skills\bemp-advanced-doc-generator\output\机构管理和管理员管理功能优化-详细设计文档-v4-20260617.docx"
doc = Document(v4_path)

print("=" * 60)
print("详细设计说明书 v4 最终质量审核报告")
print("=" * 60)

# 1. 文档基本信息
file_size = os.path.getsize(v4_path) / 1024 / 1024
print(f"\n【文档基本信息】")
print(f"  文件大小: {file_size:.2f} MB")
print(f"  标题数量: {len([p for p in doc.paragraphs if p.style.name.startswith('Heading')])}")
print(f"  段落数量: {len([p for p in doc.paragraphs if p.text.strip()])}")
print(f"  表格数量: {len(doc.tables)}")

# 2. PRD 需求覆盖度检查
print(f"\n【PRD 需求覆盖度】")
full_text = '\n'.join([p.text for p in doc.paragraphs])

requirements = {
    "批量导入机构功能": "批量导入机构" in full_text,
    "批量导入管理员功能": "批量导入管理员" in full_text,
    "模板下载功能": "模板下载" in full_text or "模版下载" in full_text,
    "批量复制角色功能": "批量复制角色" in full_text,
    "机构管理模块设计": "机构管理" in full_text,
    "管理员管理模块设计": "管理员管理" in full_text,
}

coverage_count = sum(1 for v in requirements.values() if v)
coverage_rate = coverage_count / len(requirements) * 100

for req, covered in requirements.items():
    status = "✓" if covered else "✗"
    print(f"  {status} {req}")

print(f"\n  需求覆盖率: {coverage_rate:.1f}% ({coverage_count}/{len(requirements)})")

# 3. 技术实现覆盖度
print(f"\n【技术实现覆盖度】")
tech_items = {
    "BranchImportVo": "BranchImportVo" in full_text,
    "BranchAdminImportVo": "BranchAdminImportVo" in full_text,
    "HnnxBatchCopyRoleReq": "HnnxBatchCopyRoleReq" in full_text,
    "HnnxBankBranchController": "HnnxBankBranchController" in full_text or "机构批量导入" in full_text,
    "HnnxBankBranchAdminController": "HnnxBankBranchAdminController" in full_text or "管理员批量导入" in full_text,
}

tech_count = sum(1 for v in tech_items.values() if v)
tech_rate = tech_count / len(tech_items) * 100

for item, covered in tech_items.items():
    status = "✓" if covered else "✗"
    print(f"  {status} {item}")

print(f"\n  技术覆盖率: {tech_rate:.1f}% ({tech_count}/{len(tech_items)})")

# 4. 接口定义完整性
print(f"\n【接口定义完整性】")
api_table = None
for table in doc.tables:
    if len(table.rows) > 0 and len(table.rows[0].cells) >= 3:
        first_row = [cell.text.strip() for cell in table.rows[0].cells]
        if '接口路径' in first_row[0]:
            api_table = table
            break

if api_table:
    api_count = len(api_table.rows) - 1  # 减去表头
    print(f"  ✓ API 接口表格存在")
    print(f"  ✓ 接口数量: {api_count} 个")
    
    # 检查关键接口
    api_texts = []
    for row in api_table.rows[1:]:
        api_texts.append(row.cells[0].text)
    
    key_apis = {
        "机构批量导入校验": any("batchImportValidate" in t and "branch/branch" in t for t in api_texts),
        "机构批量导入执行": any("batchImport" in t and "branch/branch" in t and "Validate" not in t for t in api_texts),
        "机构模板下载": any("downloadModel" in t and "branch/branch" in t for t in api_texts),
        "管理员批量导入校验": any("batchImportValidate" in t and "branchAdmin" in t for t in api_texts),
        "管理员批量导入执行": any("batchImportBranchAdmin" in t for t in api_texts),
        "管理员模板下载": any("downloadModel" in t and "branchAdmin" in t for t in api_texts),
        "管理员批量复制角色": any("batchCopyRole" in t for t in api_texts),
    }
    
    for api_name, found in key_apis.items():
        status = "✓" if found else "✗"
        print(f"    {status} {api_name}")
else:
    print(f"  ✗ API 接口表格不存在")

# 5. 错误码定义
print(f"\n【错误码定义】")
error_table = None
for table in doc.tables:
    if len(table.rows) > 0 and len(table.rows[0].cells) >= 4:
        first_row = [cell.text.strip() for cell in table.rows[0].cells]
        if '错误码' in str(first_row) and '错误信息' in str(first_row):
            error_table = table
            break

if error_table:
    error_count = len(error_table.rows) - 1
    print(f"  ✓ 错误码表格存在")
    print(f"  ✓ 错误码数量: {error_count} 个")
else:
    print(f"  ✗ 错误码表格不存在")

# 6. 栏位描述
print(f"\n【栏位描述】")
field_table = None
for table in doc.tables:
    if len(table.rows) > 0 and len(table.rows[0].cells) == 2:
        first_row = [cell.text.strip() for cell in table.rows[0].cells]
        if '数据名称' in str(first_row[0]):
            field_table = table
            break

if field_table:
    field_count = sum(1 for row in field_table.rows[1:] if row.cells[0].text.strip())
    print(f"  ✓ 栏位描述表格存在")
    print(f"  ✓ 栏位数量: {field_count} 个")
else:
    print(f"  ✗ 栏位描述表格不存在")

# 7. 文档结构完整性
print(f"\n【文档结构完整性】")
h1_headings = [p.text.strip() for p in doc.paragraphs if p.style.name == 'Heading 1']
required_sections = ["概述", "模块1设计说明", "模块2设计说明", "接口定义", "数据模型设计", "附录"]

for section in required_sections:
    found = any(section in h for h in h1_headings)
    status = "✓" if found else "✗"
    print(f"  {status} {section}")

# 8. 总体评估
print(f"\n" + "=" * 60)
print("【总体评估】")
print(f"  需求覆盖率: {coverage_rate:.1f}%")
print(f"  技术覆盖率: {tech_rate:.1f}%")
print(f"  接口完整性: {'✓' if api_table and len(api_table.rows) >= 8 else '✗'}")
print(f"  错误码定义: {'✓' if error_table else '✗'}")
print(f"  栏位描述: {'✓' if field_table else '✗'}")

overall_pass = coverage_rate >= 100 and tech_rate >= 100 and api_table and error_table and field_table
print(f"\n  质量门禁: {'✓ 通过' if overall_pass else '✗ 未通过'}")
print("=" * 60)

if overall_pass:
    print(f"\n✓ v4 文档质量审核通过，可以交付")
    print(f"  文件路径: {v4_path}")
else:
    print(f"\n✗ v4 文档质量审核未通过，需要继续优化")
