"""深度诊断：检查修订记录表格XML结构和蓝色run详情"""
import sys, os
sys.path.insert(0, os.path.dirname(__file__))
import paths
from docx import Document
from docx.oxml.ns import qn
from docx.shared import RGBColor

def is_blue(rgb):
    if rgb is None:
        return False
    r, g, b = int(rgb[0:2], 16), int(rgb[2:4], 16), int(rgb[4:6], 16)
    return b > 150 and b > r * 1.5 and b > g * 1.3

def check_revision_table(doc_path):
    """检查修订记录表格的XML结构"""
    doc = Document(str(doc_path))
    print(f'\n===== 修订记录表格诊断: {doc_path.name} =====')

    for i, table in enumerate(doc.tables):
        header_text = ' '.join(cell.text.strip() for cell in table.rows[0].cells) if table.rows else ''
        if not any(kw in header_text for kw in ['版本', '日期', '修改人', '修改内容', '修订']):
            continue

        print(f'\n找到修订记录表 (table#{i})')
        print(f'  行数: {len(table.rows)}')
        print(f'  列数 (rows[0]): {len(table.rows[0].cells)}')

        # 检查表头
        print(f'\n  --- 表头行 ---')
        for ci, cell in enumerate(table.rows[0].cells):
            # 检查是否合并单元格
            tc = cell._tc
            gridspan = tc.find('.//' + qn('w:gridSpan'))
            vmerge = tc.find('.//' + qn('w:vMerge'))
            gs_val = gridspan.get(qn('w:val'), '1') if gridspan is not None else '1'
            vm_val = vmerge.get(qn('w:val'), 'continue') if vmerge is not None else 'none'
            print(f'  [{ci}] text="{cell.text.strip()}" gridSpan={gs_val} vMerge={vm_val}')

        # 检查第一行数据
        if len(table.rows) > 1:
            print(f'\n  --- 第一行数据 ---')
            for ci, cell in enumerate(table.rows[1].cells):
                tc = cell._tc
                gridspan = tc.find('.//' + qn('w:gridSpan'))
                vmerge = tc.find('.//' + qn('w:vMerge'))
                gs_val = gridspan.get(qn('w:val'), '1') if gridspan is not None else '1'
                vm_val = vmerge.get(qn('w:val'), 'continue') if vmerge is not None else 'none'
                # 检查 runs
                runs_info = []
                for p in cell.paragraphs:
                    for r in p.runs:
                        color = str(r.font.color.rgb) if r.font.color and r.font.color.rgb else 'none'
                        runs_info.append(f'[{color}]"{r.text[:30]}"')
                print(f'  [{ci}] text="{cell.text.strip()}" gridSpan={gs_val} vMerge={vm_val} runs={runs_info}')

        # 检查 header_cells 和 data_row 的对应关系
        print(f'\n  --- header_cells vs data_row 对应 ---')
        header_cells = table.rows[0].cells
        data_row = table.rows[1].cells if len(table.rows) > 1 else []
        print(f'  header_cells count: {len(header_cells)}')
        print(f'  data_row count: {len(data_row)}')
        # 检查是否有共享的 tc 元素（合并单元格的标志）
        header_tc_ids = [id(c._tc) for c in header_cells]
        data_tc_ids = [id(c._tc) for c in data_row]
        dup_headers = [i for i, x in enumerate(header_tc_ids) if header_tc_ids.count(x) > 1]
        dup_data = [i for i, x in enumerate(data_tc_ids) if data_tc_ids.count(x) > 1]
        if dup_headers:
            print(f'  表头合并单元格索引: {dup_headers}')
        if dup_data:
            print(f'  数据行合并单元格索引: {dup_data}')
        break


def check_blue_runs(doc_path):
    """检查蓝色run的详细内容"""
    doc = Document(str(doc_path))
    print(f'\n===== 蓝色run诊断: {doc_path.name} =====')

    blue_in_paras = 0
    blue_in_tables = 0
    blue_with_text = 0
    blue_empty = 0
    samples = []

    for p in doc.paragraphs:
        for run in p.runs:
            if run.font.color and run.font.color.rgb:
                if is_blue(str(run.font.color.rgb)):
                    blue_in_paras += 1
                    if run.text.strip():
                        blue_with_text += 1
                        if len(samples) < 10:
                            samples.append(f'  [段落/{p.style.name}] text="{run.text[:60]}"')
                    else:
                        blue_empty += 1

    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for run in p.runs:
                        if run.font.color and run.font.color.rgb:
                            if is_blue(str(run.font.color.rgb)):
                                blue_in_tables += 1
                                if run.text.strip():
                                    blue_with_text += 1
                                    if len(samples) < 10:
                                        samples.append(f'  [表格] text="{run.text[:60]}"')
                                else:
                                    blue_empty += 1

    total = blue_in_paras + blue_in_tables
    print(f'  蓝色run总数: {total}')
    print(f'  段落中: {blue_in_paras}, 表格中: {blue_in_tables}')
    print(f'  有文本: {blue_with_text}, 空文本: {blue_empty}')
    print(f'  样本:')
    for s in samples:
        print(s)


def check_design_constraint_heading(doc_path):
    """检查设计约束标题的实际文本"""
    doc = Document(str(doc_path))
    print(f'\n===== 设计约束标题诊断: {doc_path.name} =====')

    body = doc.element.body
    for elem in list(body):
        if not elem.tag.endswith('}p'):
            continue
        pPr = elem.find('.//' + qn('w:pPr') + '/' + qn('w:pStyle'))
        if pPr is None:
            continue
        style_val = pPr.get(qn('w:val'), '') or ''
        if style_val not in {'1', '2', '3', '4', '5', '6',
                             'Heading1', 'Heading2', 'Heading3',
                             'heading1', 'heading2', 'heading3'}:
            continue
        texts = []
        for t in elem.findall('.//' + qn('w:t')):
            if t.text:
                texts.append(t.text)
        text = ''.join(texts).strip()
        if '约束' in text or '设计' in text:
            print(f'  style="{style_val}" text="{text}" len={len(text)}')


def check_component_summary(doc_path):
    """检查组件汇总表内容"""
    doc = Document(str(doc_path))
    print(f'\n===== 组件汇总表诊断: {doc_path.name} =====')

    for i, table in enumerate(doc.tables):
        header_text = ' '.join(cell.text.strip() for cell in table.rows[0].cells) if table.rows else ''
        if '组件' in header_text and ('汇总' in header_text or '名称' in header_text):
            print(f'  找到组件汇总表 (table#{i})')
            for ri, row in enumerate(table.rows):
                cells = [cell.text.strip()[:30] for cell in row.cells]
                print(f'  行{ri}: {cells}')
            break
    else:
        print('  未找到组件汇总表')


docs = [
    paths.OUTPUT_DIR / '河南农商-概要设计说明书-20260604.docx',
    paths.OUTPUT_DIR / '河南农商-详细设计文档-20260604.docx',
]

# 执行诊断，输出到文件避免 PSReadLine bug
import io
output_path = os.path.join(os.environ.get('TEMP', '/tmp'), 'deep_diag.txt')
with open(output_path, 'w', encoding='utf-8') as f:
    old_stdout = sys.stdout
    sys.stdout = f
    try:
        for doc_path in docs:
            if doc_path.exists():
                check_revision_table(doc_path)
                check_blue_runs(doc_path)
                check_design_constraint_heading(doc_path)
                check_component_summary(doc_path)
    finally:
        sys.stdout = old_stdout
print(f'Diagnostic output saved to: {output_path}')
