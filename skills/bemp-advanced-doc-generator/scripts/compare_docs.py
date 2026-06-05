"""对比两个文档的质量差异"""
import paths
from docx import Document
from lxml import etree
import os

# 技能重构前生成的v4文档
v4_path = str(paths.BANK_REQUIREMENTS_DIR / '承兑行额度管理-详细设计说明书-v4.docx')
# 新生成的文档
new_path = str(paths.OUTPUT_DIR / '承兑行额度管理-详细设计文档-20260603.docx')

for label, path in [('v4(重构前)', v4_path), ('new(重构后)', new_path)]:
    print(f'\n{"="*80}')
    print(f'文档: {label}')
    print(f'路径: {path}')
    print(f'大小: {os.path.getsize(path)} bytes')
    doc = Document(path)
    print(f'段落数: {len(doc.paragraphs)}')
    print(f'表格数: {len(doc.tables)}')

    # 统计标题
    h1_count = sum(1 for p in doc.paragraphs if p.style and p.style.name == 'Heading 1')
    h2_count = sum(1 for p in doc.paragraphs if p.style and p.style.name == 'Heading 2')
    h3_count = sum(1 for p in doc.paragraphs if p.style and p.style.name == 'Heading 3')
    print(f'Heading 1: {h1_count}, Heading 2: {h2_count}, Heading 3: {h3_count}')

    # 蓝色文本统计
    blue_count = 0
    for p in doc.paragraphs:
        for run in p.runs:
            if run.font.color and run.font.color.rgb:
                color = str(run.font.color.rgb)
                if color.startswith('00') and color not in ['000000']:
                    blue_count += 1
    print(f'非超链接蓝色文本: {blue_count}')

    # 前30个段落预览
    print(f'\n--- 段落预览 (前30) ---')
    for i, p in enumerate(doc.paragraphs[:30]):
        if p.text.strip():
            style = p.style.name if p.style else 'N/A'
            print(f'  [{i}] {style[:15]}: {p.text.strip()[:70]}')

    # 表格预览
    print(f'\n--- 表格预览 (前5) ---')
    for i, tbl in enumerate(doc.tables[:5]):
        rows = len(tbl.rows)
        cols = len(tbl.columns)
        # 第一行内容
        first_row = ' | '.join(c.text.strip()[:15] for c in tbl.rows[0].cells) if rows else ''
        print(f'  表格[{i}]: {rows}行 x {cols}列 | {first_row}')
