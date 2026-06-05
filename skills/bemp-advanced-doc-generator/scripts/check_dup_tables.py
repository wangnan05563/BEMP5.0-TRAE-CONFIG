"""检查新文档的重复表格问题"""
import paths
from docx import Document
import hashlib

new_path = str(paths.OUTPUT_DIR / '承兑行额度管理-详细设计文档-20260603.docx')
doc = Document(new_path)

# 给每个表格哈希指纹
seen = {}
for i, tbl in enumerate(doc.tables):
    # 表格指纹：前2行内容
    fingerprint = ''
    for row in tbl.rows[:2]:
        for cell in row.cells:
            fingerprint += cell.text.strip()[:30]
    h = hashlib.md5(fingerprint.encode()).hexdigest()[:8]
    if h in seen:
        seen[h]['count'] += 1
        seen[h]['indices'].append(i)
    else:
        seen[h] = {'count': 1, 'indices': [i], 'first_text': fingerprint[:50]}

# 输出重复表格
print("=== 重复表格统计 ===")
dup_count = 0
for h, info in seen.items():
    if info['count'] > 1:
        dup_count += 1
        print(f"  指纹{h}: 出现{info['count']}次 at {info['indices']} | {info['first_text']}")
print(f"\n总共 {dup_count} 种重复模式，{len(doc.tables)} 个表格")
