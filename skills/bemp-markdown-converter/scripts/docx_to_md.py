#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
Docx to Markdown Converter - BEMP技能版
将Word文档转换为Markdown格式，支持BEMP需求文档的格式优化

核心流程：
1. 预缓存图片 → 避免重复I/O
2. 遍历文档元素 → 段落和表格
3. 段落处理 → 图片提取、标题识别、隐式标题、编号列表、段落间距
4. 表格处理 → 合并行、按钮行、表单行、标题行识别、栏位描述去重
"""
import sys
import json
import re
import hashlib
from pathlib import Path
from collections import Counter
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

# 默认配置文件路径
DEFAULT_CONFIG_PATH = Path(__file__).parent / "config" / "converter_config.json"


def load_config(config_path=None):
    """
    加载配置，支持外部配置文件覆盖默认配置
    
    Args:
        config_path: 外部配置文件路径（可选）
    
    Returns:
        dict: 合并后的配置字典
    """
    config = {}
    
    # 加载默认配置
    if DEFAULT_CONFIG_PATH.exists():
        with open(DEFAULT_CONFIG_PATH, 'r', encoding='utf-8') as f:
            config = json.load(f)
    
    # 加载外部配置（覆盖默认配置）
    if config_path and Path(config_path).exists():
        with open(config_path, 'r', encoding='utf-8') as f:
            external_config = json.load(f)
            config.update(external_config)
    
    return config


def get_heading_level(paragraph):
    """获取段落的标题级别（1-6），0表示非标题"""
    style_name = paragraph.style.name.lower() if paragraph.style.name else ''
    if 'heading' in style_name:
        parts = style_name.split()
        if len(parts) >= 2:
            try:
                return int(parts[-1])
            except ValueError:
                pass
    return 0


def is_bold(paragraph):
    """检查段落是否整体加粗"""
    if not paragraph.runs:
        return False
    return all(run.bold for run in paragraph.runs if run.text.strip())


def is_list_item(paragraph, config):
    """检查段落是否为列表项"""
    text = paragraph.text.strip()
    if not text:
        return False
    
    list_cfg = config.get("list_prefix", {})
    
    # 符号前缀
    symbols = list_cfg.get("symbols", [])
    if any(text.startswith(s) for s in symbols):
        return True
    
    # 数字前缀
    pattern = list_cfg.get("number_pattern", r"^\d+[\.\、]")
    if re.match(pattern, text):
        return True
    
    return False


def get_indent_level(paragraph, config):
    """获取段落的缩进级别"""
    indent_cfg = config.get("indent", {})
    if not indent_cfg.get("enabled", True):
        return indent_cfg.get("default_level", 0)
    
    unit = indent_cfg.get("unit", 720)
    
    pPr = paragraph._element.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}pPr')
    if pPr is not None:
        indent = pPr.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}ind')
        if indent is not None:
            left = indent.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}left')
            if left:
                try:
                    return int(left) // unit
                except ValueError:
                    pass
    return indent_cfg.get("default_level", 0)


def convert_paragraph_to_markdown(paragraph, heading_level, heading_number, list_prefix, config):
    """将段落转换为Markdown格式"""
    text = paragraph.text.strip()
    if not text:
        return ''
    
    # 表格内的段落，使用内联格式转换
    parent = paragraph._element.getparent()
    if parent is not None and parent.tag.endswith('tc'):
        return convert_inline_formatting(paragraph)
    
    # 标题
    max_level = config.get("max_heading_level", 6)
    if heading_level > 0 and heading_level <= max_level:
        if heading_number:
            text = f"{heading_number} {text}"
        return f"{'#' * heading_level} {text}"
    
    # 编号列表
    if list_prefix:
        return f"{list_prefix}{text}"
    
    # 符号列表
    if is_list_item(paragraph, config):
        indent_level = get_indent_level(paragraph, config)
        prefix = '  ' * indent_level + '- '
        return f"{prefix}{text}"
    
    # 加粗段落
    if is_bold(paragraph):
        return f"**{text}**"
    
    return text


def convert_inline_formatting(paragraph):
    """转换段落内的内联格式（粗体、斜体）"""
    result = []
    for run in paragraph.runs:
        text = run.text
        if not text:
            continue
        if run.bold:
            result.append(f"**{text}**")
        elif run.italic:
            result.append(f"*{text}*")
        else:
            result.append(text)
    return ''.join(result)


# ==================== 图片处理 ====================

_IMAGE_CACHE = {}


def _build_image_cache(document):
    """预构建文档级图片缓存"""
    _IMAGE_CACHE.clear()
    for rel_id, rel in document.part.rels.items():
        if rel.reltype.endswith('image'):
            try:
                _IMAGE_CACHE[rel_id] = {
                    'blob': rel.target_part.blob,
                    'content_type': rel.target_part.content_type,
                }
            except Exception:
                pass


def extract_images_from_paragraph(paragraph, image_config, output_dir, image_counter):
    """
    从段落中提取图片并保存，返回Markdown图片引用
    
    Args:
        paragraph: python-docx Paragraph对象
        image_config: 图片配置
        output_dir: 图片输出目录
        image_counter: 图片计数器（可变对象）
    
    Returns:
        list: Markdown图片引用列表
    """
    if not image_config.get("enabled", True):
        return []
    
    images = []
    
    # 查找段落中的图片（blip元素）
    blip_ns = '{http://schemas.openxmlformats.org/drawingml/2006/main}blip'
    embed_ns = '{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed'
    
    for blip in paragraph._element.findall(f'.//{blip_ns}'):
        embed_id = blip.get(embed_ns)
        if not embed_id:
            continue
        
        # 从预构建缓存获取图片
        try:
            img_data = _IMAGE_CACHE.get(embed_id)
            if not img_data:
                continue
            
            blob = img_data['blob']
            if not blob:
                continue
            
            image_counter[0] += 1
            img_hash = hashlib.md5(blob).hexdigest()[:8]
            naming_pattern = image_config.get("naming_pattern", "img_{index}_{hash}")
            img_name = naming_pattern.format(index=image_counter[0], hash=img_hash)
            
            # 获取扩展名
            content_type = img_data['content_type']
            ext_map = image_config.get("extensions", {})
            ext = ext_map.get(content_type, image_config.get("default_extension", ".png"))
            
            img_filename = f"{img_name}{ext}"
            img_path = output_dir / img_filename
            
            # 保存图片
            output_dir.mkdir(parents=True, exist_ok=True)
            img_path.write_bytes(blob)
            
            # 生成Markdown引用
            rel_path = image_config.get("output_dir", "images")
            images.append(f"![图{image_counter[0]}]({rel_path}/{img_filename})")
        except Exception:
            pass
    
    return images


# ==================== 表格处理核心逻辑 ====================

def _is_merge_row(row_data, config):
    """
    检测是否为合并行（大部分单元格内容相同）
    
    Args:
        row_data: 行数据列表
        config: 配置字典
    
    Returns:
        bool: 是否为合并行
    """
    merge_cfg = config.get("merge_row", {})
    if not merge_cfg.get("enabled", True):
        return False
    
    min_cells = merge_cfg.get("min_cell_count", 2)
    threshold = merge_cfg.get("threshold_ratio", 0.8)
    
    if len(row_data) < min_cells:
        return False
    
    non_empty = [cell for cell in row_data if cell.strip()]
    if not non_empty:
        return False
    
    counter = Counter(non_empty)
    most_common_count = counter.most_common(1)[0][1]
    
    return most_common_count / len(non_empty) >= threshold


def _get_merged_text(row_data):
    """获取合并行的文本（取第一个非空单元格）"""
    for cell in row_data:
        if cell.strip():
            return cell.strip()
    return ""


def _is_form_row(row_data, config):
    """
    检测是否为表单行（包含输入框的合并行）
    
    Returns:
        tuple: (is_form, form_text)
    """
    form_cfg = config.get("form_row", {})
    if not form_cfg.get("enabled", True):
        return False, ""
    
    if not _is_merge_row(row_data, config):
        return False, ""
    
    merged_text = _get_merged_text(row_data)
    patterns = form_cfg.get("patterns", [])
    
    for pattern in patterns:
        if re.search(pattern, merged_text):
            return True, merged_text
    
    return False, ""


def _is_button_row(row_data, config):
    """
    检测是否为按钮行
    支持：1. 多列各自包含按钮关键字  2. 合并行包含按钮关键字
    
    Returns:
        tuple: (is_button, button_text)
    """
    button_cfg = config.get("button_row", {})
    if not button_cfg.get("enabled", True):
        return False, ""
    
    keywords = button_cfg.get("cell_keywords", [])
    min_count = button_cfg.get("min_button_count", 2)
    
    # 方法1：多列各自包含按钮关键字
    non_empty = [cell for cell in row_data if cell.strip()]
    button_cells = []
    for cell in non_empty:
        for kw in keywords:
            if kw in cell:
                button_cells.append(cell)
                break
    
    if len(button_cells) >= min_count:
        unique_cells = list(dict.fromkeys(button_cells))
        separator = button_cfg.get("output_separator", " ")
        return True, separator.join(unique_cells)
    
    # 方法2：合并行包含按钮关键字
    if not _is_merge_row(row_data, config):
        return False, ""
    
    merged_text = _get_merged_text(row_data)
    for kw in keywords:
        if kw in merged_text:
            return True, merged_text
    
    return False, ""


def _is_title_row(row_data, config):
    """检测是否为标题行（合并行且位于指定位置）"""
    title_cfg = config.get("title_row", {})
    if not title_cfg.get("enabled", True):
        return False
    
    return _is_merge_row(row_data, config)


def _deduplicate_consecutive_rows(rows_data):
    """去除连续重复的行"""
    if not rows_data:
        return rows_data
    
    result = [rows_data[0]]
    for i in range(1, len(rows_data)):
        if rows_data[i] != result[-1]:
            result.append(rows_data[i])
    return result


def _is_field_desc_table(rows_data, config):
    """检测是否为栏位描述表格"""
    field_cfg = config.get("field_desc_table", {})
    if not field_cfg.get("enabled", True):
        return False
    
    if not rows_data:
        return False
    
    header_keywords = field_cfg.get("header_keywords", [])
    if not header_keywords:
        return False
    
    first_row = rows_data[0]
    match_count = sum(1 for kw in header_keywords if any(kw in cell for cell in first_row))
    
    return match_count >= len(header_keywords) * 0.5


def _extract_row_data(row, config):
    """
    提取表格行数据，处理合并单元格（gridSpan）
    
    在python-docx中，合并单元格的每个底层cell都包含相同文本，
    导致"分页展示"在12列中重复12次。通过检测gridSpan，只对
    合并区域的第一个单元格提取文本，其余填充空值。
    
    Args:
        row: python-docx Row对象
        config: 配置字典
    
    Returns:
        list: 行数据列表（长度等于实际列数）
    """
    row_data = []
    max_cols = len(row.cells) if row.cells else 0
    
    for cell in row.cells:
        cell_text = cell.text.strip().replace('\n', '<br>')
        
        # 检测gridSpan（跨列合并）
        tcPr = cell._tc.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tcPr')
        grid_span = 1
        if tcPr is not None:
            gs_elem = tcPr.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}gridSpan')
            if gs_elem is not None:
                val = gs_elem.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val')
                if val:
                    try:
                        grid_span = int(val)
                    except ValueError:
                        pass
        
        # 第一个单元格放文本，其余跨列位置放空值
        row_data.append(cell_text)
        for _ in range(grid_span - 1):
            row_data.append('')
    
    # 确保行数不超过表格最大列数（处理可能的异常）
    if len(row_data) > max_cols:
        row_data = row_data[:max_cols]
    
    return row_data


def table_to_markdown(table, config):
    """
    将表格转换为Markdown格式
    
    处理逻辑：
    1. 提取表格数据（处理合并单元格gridSpan）
    2. 识别特殊行（按钮行、表单行、标题行、合并行）
    3. 处理栏位描述表格（去重）
    4. 构建Markdown输出
    """
    lines = []
    rows_data = []
    max_cols = 0
    
    # 1. 提取表格数据（使用合并单元格处理）
    for row in table.rows:
        row_data = _extract_row_data(row, config)
        rows_data.append(row_data)
        max_cols = max(max_cols, len(row_data))
    
    if not rows_data:
        return ''
    
    # 2. 处理特殊行
    processed_rows = []
    title_extracted = False
    
    for i, row_data in enumerate(rows_data):
        while len(row_data) < max_cols:
            row_data.append('')
        
        is_first = (i == 0)
        
        # 按钮行（优先检测）
        is_button, button_text = _is_button_row(row_data, config)
        if is_button:
            button_fmt = config.get("button_row", {}).get("output_format", "[{button_text}]")
            processed_rows.append({
                "type": "button",
                "text": button_text,
                "output": button_fmt.format(button_text=button_text)
            })
            continue
        
        # 表单行
        is_form, form_text = _is_form_row(row_data, config)
        if is_form:
            processed_rows.append({
                "type": "form",
                "text": form_text,
                "output": form_text
            })
            continue
        
        # 标题行（仅第一行）
        if is_first and not title_extracted:
            if _is_title_row(row_data, config):
                merged_text = _get_merged_text(row_data)
                processed_rows.append({
                    "type": "title",
                    "text": merged_text,
                    "output": f"**{merged_text}**"
                })
                title_extracted = True
                continue
        
        # 普通数据行
        processed_rows.append({
            "type": "data",
            "values": row_data
        })
    
    # 3. 构建Markdown输出
    title_rows = [r for r in processed_rows if r["type"] == "title"]
    if title_rows:
        for tr in title_rows:
            lines.append(tr["output"])
        lines.append('')
    
    data_rows = [r for r in processed_rows if r["type"] == "data"]
    
    # 栏位描述表格去重
    if data_rows and _is_field_desc_table([r["values"] for r in data_rows], config):
        dedup_cfg = config.get("field_desc_table", {})
        if dedup_cfg.get("deduplicate_consecutive", True):
            original_values = [r["values"] for r in data_rows]
            deduped_values = _deduplicate_consecutive_rows(original_values)
            data_rows = [{"type": "data", "values": v} for v in deduped_values]
    
    # 构建表格
    if data_rows:
        header = data_rows[0]["values"]
        lines.append('| ' + ' | '.join(header) + ' |')
        lines.append('| ' + ' | '.join(['---'] * max_cols) + ' |')
        for row in data_rows[1:]:
            lines.append('| ' + ' | '.join(row["values"]) + ' |')
    
    # 按钮行
    button_rows = [r for r in processed_rows if r["type"] == "button"]
    if button_rows:
        lines.append('')
        for br in button_rows:
            lines.append(br["output"])
    
    # 表单行
    form_rows = [r for r in processed_rows if r["type"] == "form"]
    if form_rows:
        lines.append('')
        for fr in form_rows:
            lines.append(fr["output"])
    
    return '\n'.join(lines)


# ==================== 主转换函数 ====================

def docx_to_markdown(docx_path, output_path=None, config_path=None):
    """
    将docx文件转换为markdown
    
    Args:
        docx_path: 输入docx文件路径
        output_path: 输出md文件路径（可选）
        config_path: 配置文件路径（可选）
    
    Returns:
        str: 转换后的markdown内容
    """
    config = load_config(config_path)
    
    doc = Document(docx_path)
    _build_image_cache(doc)
    
    md_lines = []
    
    # 标题序号跟踪器（6个级别）
    heading_counters = [0, 0, 0, 0, 0, 0]
    last_heading_level = 0
    
    # 编号列表跟踪器
    list_counters = {}
    
    # 图片计数器
    image_counter = [0]
    
    # 图片输出目录
    if output_path:
        output_dir = Path(output_path).parent / config.get("image", {}).get("output_dir", "images")
    else:
        output_dir = Path.cwd() / config.get("image", {}).get("output_dir", "images")
    
    # 遍历所有元素
    prev_element_type = None  # 'heading', 'normal', 'table'
    
    for element in doc.element.body:
        tag = element.tag.split('}')[-1] if '}' in element.tag else element.tag
        
        if tag == 'tbl':
            # 处理表格
            for table in doc.tables:
                if table._element == element:
                    # 添加间距
                    if config.get("spacing", {}).get("enabled", True):
                        rules = config["spacing"].get("rules", {})
                        if prev_element_type == "normal":
                            for _ in range(rules.get("table_after_normal", 1)):
                                md_lines.append('')
                        elif prev_element_type == "heading":
                            for _ in range(rules.get("table_after_heading", 1)):
                                md_lines.append('')
                    
                    md_lines.append(table_to_markdown(table, config))
                    md_lines.append('')
                    prev_element_type = 'table'
                    break
        
        elif tag == 'p':
            # 处理段落
            for para in doc.paragraphs:
                if para._element == element:
                    # 提取图片
                    image_refs = extract_images_from_paragraph(
                        para, config.get("image", {}), output_dir, image_counter
                    )
                    if image_refs:
                        md_lines.extend(image_refs)
                        if md_lines and md_lines[-1] != '':
                            md_lines.append('')
                    
                    heading_level = get_heading_level(para)
                    
                    list_prefix = ""
                    is_implicit_heading = False
                    implicit_heading_level = 0
                    
                    # 检测编号列表/隐式标题
                    num_pr = para._element.find(
                        './/{http://schemas.openxmlformats.org/wordprocessingml/2006/main}pPr/'
                        '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}numPr'
                    )
                    if num_pr is not None:
                        num_id_elem = num_pr.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}numId')
                        ilvl_elem = num_pr.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}ilvl')
                        
                        if num_id_elem is not None and ilvl_elem is not None:
                            num_id = num_id_elem.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val')
                            ilvl = ilvl_elem.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val', '0')
                            
                            if num_id:
                                key = (num_id, ilvl)
                                if key not in list_counters:
                                    list_counters[key] = 0
                                list_counters[key] += 1
                                
                                text_for_check = para.text.strip()
                                implicit_cfg = config.get("implicit_heading", {})
                                
                                if implicit_cfg.get("enabled", True):
                                    keywords = implicit_cfg.get("keywords", [])
                                    max_len = implicit_cfg.get("max_length", 20)
                                    exclude_punct = implicit_cfg.get("exclude_punctuation",
                                        ['，', '。', '；', '：', '、', '【', '】'])
                                    
                                    is_keyword_match = text_for_check in keywords
                                    is_short_text = (len(text_for_check) < max_len and
                                        not any(c in text_for_check for c in exclude_punct))
                                    
                                    if is_keyword_match or is_short_text:
                                        is_implicit_heading = True
                                        implicit_heading_level = last_heading_level + 1 if last_heading_level > 0 else 4
                                    else:
                                        list_cfg = config.get("list_number", {})
                                        if list_cfg.get("enabled", True):
                                            indent = "  " * int(ilvl)
                                            list_prefix = f"{indent}{list_counters[key]}. "
                                else:
                                    list_cfg = config.get("list_number", {})
                                    if list_cfg.get("enabled", True):
                                        indent = "  " * int(ilvl)
                                        list_prefix = f"{indent}{list_counters[key]}. "
                    
                    # 更新标题序号
                    heading_number = ""
                    max_level = config.get("max_heading_level", 6)
                    
                    # 显式标题
                    if heading_level > 0 and heading_level <= max_level:
                        heading_counters[heading_level - 1] += 1
                        for i in range(heading_level, max_level):
                            heading_counters[i] = 0
                        
                        heading_number = ".".join(
                            str(heading_counters[i])
                            for i in range(heading_level)
                            if heading_counters[i] > 0
                        )
                        
                        last_heading_level = heading_level
                        
                        list_cfg = config.get("list_number", {})
                        if list_cfg.get("reset_on_heading", True):
                            list_counters.clear()
                    
                    # 隐式标题
                    elif is_implicit_heading and implicit_heading_level > 0:
                        level = min(implicit_heading_level, max_level)
                        heading_counters[level - 1] += 1
                        for i in range(level, max_level):
                            heading_counters[i] = 0
                        
                        heading_number = ".".join(
                            str(heading_counters[i])
                            for i in range(level)
                            if heading_counters[i] > 0
                        )
                        heading_level = level
                    
                    md_line = convert_paragraph_to_markdown(
                        para, heading_level, heading_number, list_prefix, config
                    )
                    
                    # 段落间距
                    if config.get("spacing", {}).get("enabled", True):
                        rules = config["spacing"].get("rules", {})
                        
                        if heading_level > 0:
                            if prev_element_type in ("normal", "heading"):
                                for _ in range(rules.get("heading_after_normal", 1)):
                                    md_lines.append('')
                        else:
                            if prev_element_type == "heading":
                                for _ in range(rules.get("normal_after_heading", 0)):
                                    md_lines.append('')
                            elif prev_element_type == "normal":
                                for _ in range(rules.get("normal_after_normal", 0)):
                                    md_lines.append('')
                            elif prev_element_type == "table":
                                for _ in range(rules.get("normal_after_table", 1)):
                                    md_lines.append('')
                    
                    if md_line:
                        md_lines.append(md_line)
                    else:
                        md_lines.append('')
                    
                    prev_element_type = 'heading' if heading_level > 0 else 'normal'
                    break
    
    content = '\n'.join(md_lines)
    
    if output_path:
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(content)
        print(f"转换完成: {output_path}")
    else:
        print(content)
    
    return content


if __name__ == '__main__':
    import argparse
    
    parser = argparse.ArgumentParser(description='将docx文件转换为markdown格式')
    parser.add_argument('input', help='输入docx文件路径')
    parser.add_argument('-o', '--output', help='输出md文件路径')
    parser.add_argument('-c', '--config', help='配置文件路径（JSON格式）')
    
    args = parser.parse_args()
    docx_to_markdown(args.input, args.output, args.config)
